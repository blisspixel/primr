"""
DocumentBuilder for premium report generation.

Orchestrates the document construction process with F-Pattern optimization
and cognitive load management.
"""

from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from primr.output.chapter_config import CHAPTER_CONFIG
from primr.output.citation_processor import CitationProcessor, CitationStyle, SourceCitation
from primr.output.content_pattern_detector import ContentPatternDetector
from primr.output.executive_summary_generator import ExecutiveSummaryGenerator
from primr.output.markdown_parser import MarkdownParser
from primr.output.models import CompanySnapshot, ContentBlock, DocumentMetadata
from primr.output.polish_elements import (
    DataConfidenceIndicator,
    DocumentDisclaimer,
    FinancialDashboard,
    KeyImplicationsBox,
    OneLinerSummary,
    QuickWinsSection,
)
from primr.output.style_engine import DualCodingEnhancer, StyleEngine
from primr.output.table_builder import TableBuilder


class DocumentBuilder:
    """Builds the complete DOCX document with all components."""

    # Miller's Law: max items per section
    MAX_ITEMS_PER_SECTION = 7

    def __init__(
        self,
        company_name: str,
        section_results: dict[str, str],
        citations: list[dict[str, str]] | None = None,
        citation_style: CitationStyle = CitationStyle.NUMBERED,
    ):
        """
        Initialize DocumentBuilder.

        Args:
            company_name: Name of the company
            section_results: Dict mapping section_key to content string
            citations: Optional list of citation dicts with 'text', 'url', 'title' keys
            citation_style: Citation formatting style (NUMBERED, INLINE, SIDECAR)
        """
        self.company_name = company_name
        self.sections = section_results
        self.citations = citations or []
        self.citation_style = citation_style
        self.document = Document()
        self.style_engine = StyleEngine(self.document)
        self.parser = MarkdownParser()
        self.table_builder = TableBuilder(self.document)
        self.detector = ContentPatternDetector()
        self.dual_coding = DualCodingEnhancer()
        self.metadata = DocumentMetadata(company_name=company_name)

        # Citation processor for numbered references
        self.citation_processor = CitationProcessor(style=citation_style)

        # Process sections through citation processor (unless INLINE style)
        if citation_style != CitationStyle.INLINE:
            self._process_citations()

        # Polish elements
        self.disclaimer = DocumentDisclaimer()
        self.confidence_indicator = DataConfidenceIndicator()

        # Calculate overall confidence based on citation count
        self._citation_count = len(self.citations) + self.citation_processor.citation_count
        self._overall_confidence = self._calculate_overall_confidence()

    def _process_citations(self) -> None:
        """Process all section content through citation processor."""
        processed_sections = {}
        for key, content in self.sections.items():
            # Handle None or non-string content defensively
            if content is None:
                processed_sections[key] = ""
                continue
            if not isinstance(content, str):
                content = str(content)
            result = self.citation_processor.process_content(content)
            processed_sections[key] = result.transformed_content
        self.sections = processed_sections

    def build(self) -> Document:
        """
        Build the complete document.

        Returns:
            Completed Document object
        """
        self.style_engine.setup_styles()
        self._add_cover_page()
        self._add_company_snapshot()
        self._add_executive_summary()
        self._add_table_of_contents()
        self._add_chapters()
        self._add_sources_appendix()
        self._add_headers_footers()
        return self.document

    def _calculate_overall_confidence(self) -> str:
        """
        Calculate overall report confidence based on citation count.

        High: 10+ citations (well-sourced)
        Medium: 3-9 citations (adequately sourced)
        Low: 0-2 citations (limited sources)

        Returns:
            Confidence level string ('high', 'medium', 'low')
        """
        if self._citation_count >= 10:
            return "high"
        elif self._citation_count >= 3:
            return "medium"
        else:
            return "low"

    def _get_confidence_description(self) -> str:
        """Get human-readable confidence description."""
        descriptions = {
            "high": f"High confidence ({self._citation_count} sources)",
            "medium": f"Medium confidence ({self._citation_count} sources)",
            "low": f"Limited sources ({self._citation_count} citations)",
        }
        return descriptions.get(self._overall_confidence, "Unknown")

    def _add_headers_footers(self) -> None:
        """
        Add headers and footers with page numbers.

        Header: Company name and report title
        Footer: Centered page numbers (skipped on cover page via section break)
        """
        # Get the default section
        section = self.document.sections[0]

        # Configure different first page (cover page has no header/footer)
        section.different_first_page_header_footer = True

        # Header for subsequent pages
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        header_run = header_para.add_run(f"{self.company_name} | {self.metadata.report_title}")
        header_run.font.size = Pt(9)
        header_run.font.color.rgb = self.style_engine.MUTED_COLOR

        # Footer with page numbers
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add page number field
        self._add_page_number_field(footer_para)

    def _add_page_number_field(self, paragraph: Any) -> None:
        """Add a Word page number field to a paragraph."""
        run = paragraph.add_run()

        # Create PAGE field
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.text = "PAGE"

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")

        # Placeholder text
        text_elem = OxmlElement("w:t")
        text_elem.text = "1"

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(text_elem)
        run._r.append(fld_char_end)

        run.font.size = Pt(9)
        run.font.color.rgb = self.style_engine.MUTED_COLOR

    def _add_cover_page(self) -> None:
        """
        Add cover page with F-Pattern layout.

        F-Zone 1 (top-left): Company name + one-liner summary
        F-Zone 2 (top-right): Key metric (revenue or similar)
        """
        # Company name (large, centered)
        title = self.document.add_heading(self.company_name, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Subtitle
        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = subtitle.add_run(self.metadata.report_title)
        run.font.size = Pt(16)
        run.font.color.rgb = self.style_engine.SECONDARY_COLOR

        # One-liner summary using OneLinerSummary
        one_liner_gen = self._generate_one_liner()
        if one_liner_gen:
            one_liner = self.document.add_paragraph()
            one_liner.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = one_liner.add_run(one_liner_gen)
            run.font.size = Pt(11)
            run.font.italic = True

        # Generation date
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_para.add_run(f"Prepared: {self.metadata.generation_date}")

        # Confidentiality notice
        conf_para = self.document.add_paragraph()
        conf_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        conf_run = conf_para.add_run(self.metadata.confidentiality)
        conf_run.font.size = Pt(9)
        conf_run.font.color.rgb = self.style_engine.MUTED_COLOR

        # AI disclaimer
        self.disclaimer.add_to_document(self.document, self.style_engine)

        # Page break
        self.document.add_page_break()

    def _generate_one_liner(self) -> str:
        """Generate one-liner summary for cover page."""
        # Extract key info from sections
        all_content = "\n".join(self.sections.values())
        metrics = self.detector.extract_metrics(all_content)

        # Try to get industry from content
        industry = self.sections.get("industry", "")
        if not industry:
            # Try to extract from overview
            overview = self.sections.get("company_overview", "")
            if "technology" in overview.lower():
                industry = "technology"
            elif "retail" in overview.lower():
                industry = "retail"
            elif "financial" in overview.lower() or "bank" in overview.lower():
                industry = "financial services"

        # Get differentiator from competitive position
        differentiator = ""
        competitive = self.sections.get("competitive_position", "")
        if competitive:
            # Extract first sentence as differentiator
            sentences = competitive.split(".")
            if sentences:
                differentiator = sentences[0].strip()[:100]  # Limit length

        one_liner = OneLinerSummary(
            company_name=self.company_name,
            industry=industry,
            differentiator=differentiator,
            revenue=metrics.get("revenue", ""),
        )

        return one_liner.generate()

    def _add_company_snapshot(self) -> None:
        """Add company snapshot table with extracted metrics."""
        # Extract metrics from all content
        all_content = "\n".join(self.sections.values())
        metrics = self.detector.extract_metrics(all_content)

        # Build snapshot
        snapshot = CompanySnapshot(
            company_name=self.company_name,
            website=self.sections.get("company_website", ""),
            industry=self.sections.get("industry", ""),
            founded=metrics.get("founded"),
            headquarters=metrics.get("headquarters"),
            revenue=metrics.get("revenue"),
            employees=metrics.get("employees"),
            ticker=metrics.get("ticker"),
        )

        # Add heading
        self.document.add_heading("Company Snapshot", level=1)

        # Add table
        self.table_builder.create_company_snapshot(snapshot)

        # Add financial dashboard if we have financial metrics
        financial_metrics = {}
        if metrics.get("revenue"):
            financial_metrics["Revenue"] = metrics["revenue"]
        if metrics.get("profit_margin"):
            financial_metrics["Profit Margin"] = metrics["profit_margin"]
        if metrics.get("employees"):
            financial_metrics["Employees"] = metrics["employees"]
        if metrics.get("growth_rate"):
            financial_metrics["Growth Rate"] = metrics["growth_rate"]

        if financial_metrics:
            self.document.add_paragraph()
            dashboard = FinancialDashboard(self.document, self.style_engine)
            dashboard.create_dashboard(financial_metrics)

        self.document.add_page_break()

    def _add_executive_summary(self) -> None:
        """
        Add executive summary with dual-column layout.

        Structure:
        - THE BOTTOM LINE (F-Zone 1)
        - DATA CONFIDENCE indicator (based on citation count)
        - KEY INSIGHTS with trend indicators
        - WATCH OUTS with confidence indicators
        """
        summary_gen = ExecutiveSummaryGenerator(self.sections)
        exec_summary = summary_gen.generate()

        # Main heading
        self.document.add_heading("Executive Summary", level=1)

        # DATA CONFIDENCE indicator (based on citation count)
        if self._citation_count > 0:
            confidence_para = self.document.add_paragraph()
            confidence_para.add_run("Data Confidence: ")
            self.confidence_indicator.add_to_paragraph(
                confidence_para, self._get_confidence_description(), self._overall_confidence
            )
        # THE BOTTOM LINE
        if exec_summary.narrative:
            self.document.add_heading("The Bottom Line", level=2)

            # Add narrative paragraphs
            for para_text in exec_summary.narrative.split("\n\n"):
                if para_text.strip():
                    para = self.document.add_paragraph(para_text.strip())

        # KEY INSIGHTS
        if exec_summary.key_takeaways:
            self.document.add_heading("Key Insights", level=2)
            self.table_builder.create_executive_highlights(exec_summary.key_takeaways, title="")

        # WATCH OUTS (Risk Factors)
        if exec_summary.risk_factors:
            self.document.add_heading("Watch Outs", level=2)
            for risk in exec_summary.risk_factors:
                para = self.document.add_paragraph(style="List Bullet")
                # Add confidence indicator using DataConfidenceIndicator
                self.confidence_indicator.add_to_paragraph(para, risk, "medium")

        self.document.add_page_break()

    def _add_table_of_contents(self) -> None:
        """Add table of contents with Word TOC field."""
        self.document.add_heading("Table of Contents", level=1)

        # Add TOC field (Word will populate this)
        para = self.document.add_paragraph()
        run = para.add_run()

        # Create TOC field
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(fld_char_end)

        # Instruction to update TOC
        update_para = self.document.add_paragraph()
        update_para.add_run('(Right-click and select "Update Field" to populate)')
        update_para.runs[0].font.size = Pt(9)
        update_para.runs[0].font.italic = True
        update_para.runs[0].font.color.rgb = self.style_engine.MUTED_COLOR

        self.document.add_page_break()

    def _add_chapters(self) -> None:
        """
        Add all chapters with proper numbering.

        Iterates through CHAPTER_CONFIG and renders each chapter
        with its sections.
        """
        for chapter_num, (chapter_title, chapter_data) in enumerate(CHAPTER_CONFIG.items(), 1):
            self._add_chapter(chapter_num, chapter_title, chapter_data)

    def _add_chapter(self, chapter_num: int, title: str, chapter_data: dict) -> None:
        """Add a single chapter with its sections."""
        # Chapter heading
        icon = chapter_data.get("icon", "")
        chapter_heading = f"{chapter_num}. {title}"
        if icon:
            chapter_heading = f"{icon} {chapter_heading}"

        self.document.add_heading(chapter_heading, level=1)

        # Collect chapter content for implications
        chapter_content = []

        # Add sections
        sections = chapter_data.get("sections", [])
        for section_idx, (section_title, section_key) in enumerate(sections, 1):
            section_num = f"{chapter_num}.{section_idx}"
            self._add_section(section_num, section_title, section_key)

            # Collect content for implications
            content = self.sections.get(section_key, "")
            if content:
                chapter_content.append(content)

        # Add key implications box at chapter end
        if chapter_content:
            implications_box = KeyImplicationsBox(self.document, self.style_engine)
            all_content = "\n".join(chapter_content)
            implications = implications_box.extract_implications(all_content)
            if implications:
                implications_box.add_to_document(title, implications)

        # Add quick wins for Strategic Assessment chapter
        if "Strategic" in title and chapter_content:
            quick_wins_section = QuickWinsSection(self.document, self.style_engine)
            all_content = "\n".join(chapter_content)
            quick_wins = quick_wins_section.extract_quick_wins(all_content)
            if quick_wins:
                quick_wins_section.add_to_document(quick_wins)

        # Page break after chapter (except last)
        if chapter_num < len(CHAPTER_CONFIG):
            self.document.add_page_break()

    def _add_section(self, section_num: str, title: str, section_key: str) -> None:
        """Add a single section with its content."""
        content = self.sections.get(section_key, "")

        if not content:
            return  # Skip empty sections

        # Section heading
        self.document.add_heading(f"{section_num} {title}", level=2)

        # Render content
        self._render_section_content(content)

    def _render_section_content(self, content: str) -> None:
        """
        Render section content with proper formatting.

        Uses the clean markdown converter for direct, simple rendering.
        """
        from primr.output.markdown_converter import render_section_content

        render_section_content(self.document, content)

    def _render_block(self, block: ContentBlock) -> None:
        """Render a single content block."""
        if block.type == "heading":
            for line in block.lines:
                level = min(line.level + 2, 4)  # Offset for document hierarchy
                self.document.add_heading(line.content, level=level)

        elif block.type == "bullet_list":
            for line in block.lines:
                para = self.document.add_paragraph(style="List Bullet")
                if line.level > 0:
                    para.paragraph_format.left_indent = Inches(0.25 * line.level)
                self.parser.apply_inline_formatting(para, line.content)

        elif block.type == "numbered_list":
            for line in block.lines:
                para = self.document.add_paragraph(style="List Number")
                if line.level > 0:
                    para.paragraph_format.left_indent = Inches(0.25 * line.level)
                self.parser.apply_inline_formatting(para, line.content)

        elif block.type == "paragraph":
            # Combine text lines into paragraph
            text_parts: list[str] = []
            for line in block.lines:
                if line.type == "inline_header":
                    # Render inline header with bold label
                    if text_parts:
                        para = self.document.add_paragraph()
                        self.parser.apply_inline_formatting(para, " ".join(text_parts))
                        text_parts = []

                    para = self.document.add_paragraph()
                    header = line.metadata.get("header_text", "")
                    self.style_engine.apply_inline_header_formatting(para, header, line.content)
                else:
                    text_parts.append(line.content)

            # Flush remaining text
            if text_parts:
                para = self.document.add_paragraph()
                self.parser.apply_inline_formatting(para, " ".join(text_parts))

        elif block.type == "table":
            # Render markdown table (common in Deep Research output)
            table_data = self.parser.parse_table_block(block)
            if table_data["headers"] or table_data["rows"]:
                self.table_builder.create_from_markdown(table_data["headers"], table_data["rows"])
                self.document.add_paragraph()  # Spacer after table

    def _add_sources_appendix(self) -> None:
        """
        Add sources/citations appendix.

        Lists all sources used in the research, extracted from:
        - CitationProcessor (numbered references from content)
        - Deep Research citations
        - Inline markdown links in content
        """
        # Get citations from processor first (these have reference numbers)
        processor_citations = self.citation_processor.citations

        # Collect additional sources not captured by processor
        additional_sources = self._collect_all_sources()

        total_sources = len(processor_citations) + len(additional_sources)

        if total_sources == 0:
            return  # No sources to add

        # Page break before appendix
        self.document.add_page_break()

        # Appendix heading
        self.document.add_heading("Sources", level=1)

        # Introduction paragraph
        intro = self.document.add_paragraph()
        intro.add_run(
            f"This report references {total_sources} source(s). "
            "All information should be verified independently."
        )
        intro.runs[0].font.size = Pt(10)
        intro.runs[0].font.italic = True

        self.document.add_paragraph()

        # List processor citations first (with their reference numbers)
        for citation in processor_citations:
            self._add_numbered_source_entry(citation)

        # List additional sources (continue numbering)
        next_num = len(processor_citations) + 1
        for source in additional_sources:
            self._add_source_entry(next_num, source)
            next_num += 1

    def _add_numbered_source_entry(self, citation: SourceCitation) -> None:
        """Add a source entry from CitationProcessor with its reference number."""
        para = self.document.add_paragraph()

        # Source number
        num_run = para.add_run(f"[{citation.reference_number}] ")
        num_run.font.bold = True
        num_run.font.size = Pt(10)

        # Title
        title_run = para.add_run(citation.title)
        title_run.font.size = Pt(10)

        # URL on new line
        if citation.url:
            para.add_run("\n")
            url_run = para.add_run(citation.url)
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = self.style_engine.SECONDARY_COLOR

    def _collect_all_sources(self) -> list[dict[str, str]]:
        """
        Collect additional sources not already captured by CitationProcessor.

        Returns:
            List of source dicts with 'title', 'url' keys
        """
        sources = []

        # Get URLs already captured by processor
        seen_urls = {c.url for c in self.citation_processor.citations}

        # Add explicit citations that weren't in content
        for citation in self.citations:
            url = citation.get("url", "")
            if url and url not in seen_urls:
                sources.append(
                    {"title": citation.get("title") or citation.get("text", url), "url": url}
                )
                seen_urls.add(url)

        return sources

    def _add_source_entry(self, idx: int, source: dict[str, str]) -> None:
        """Add a single source entry to the appendix."""
        para = self.document.add_paragraph()

        # Source number
        num_run = para.add_run(f"[{idx}] ")
        num_run.font.bold = True
        num_run.font.size = Pt(10)

        # Title
        title = source.get("title", "Unknown Source")
        title_run = para.add_run(title)
        title_run.font.size = Pt(10)

        # URL on new line
        url = source.get("url", "")
        if url:
            para.add_run("\n")
            url_run = para.add_run(url)
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = self.style_engine.SECONDARY_COLOR
