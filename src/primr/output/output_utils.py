"""
Report generation and output formatting utilities.

This module provides the main entry points for generating premium
consultant-grade reports in TXT, DOCX, and PDF formats.
"""

import os
import re
import shutil
import zipfile
from contextlib import suppress
from datetime import datetime
from io import StringIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from primr.config.config import OUTPUT_DIR, WORKING_DIR
from primr.config.sections_config import SECTION_KEY_MAP

# Import new premium report components
from primr.output.document_builder import DocumentBuilder
from primr.utils.console import console

os.makedirs(OUTPUT_DIR, exist_ok=True)


def parse_markdown_line(line):
    """
    Parse a line and return (type, content, indent_level).
    Types: 'heading', 'bullet', 'numbered', 'text', 'empty'
    """
    stripped = line.rstrip()

    if not stripped:
        return ("empty", "", 0)

    # Section heading (## )
    if stripped.startswith("## "):
        return ("heading", stripped[3:], 0)

    # Sub-heading (### )
    if stripped.startswith("### "):
        return ("subheading", stripped[4:], 0)

    # Calculate indent level (4 spaces = 1 level)
    indent = len(line) - len(line.lstrip())
    indent_level = indent // 4

    content = stripped.lstrip()

    # Bullet point (* or - at start)
    if content.startswith(("*   ", "* ")):
        bullet_content = content[4:] if content.startswith("*   ") else content[2:]
        return ("bullet", bullet_content, indent_level)

    if content.startswith(("-   ", "- ")):
        bullet_content = content[4:] if content.startswith("-   ") else content[2:]
        return ("bullet", bullet_content, indent_level)

    # Numbered list
    num_match = re.match(r"^(\d+)\.\s+(.+)$", content)
    if num_match:
        return ("numbered", num_match.group(2), indent_level)

    return ("text", content, indent_level)


def apply_inline_formatting(paragraph, text):
    """
    Apply inline markdown formatting (bold, italic) to a paragraph.
    """
    # Pattern for bold (**text** or __text__)
    pattern = r"\*\*(.+?)\*\*|__(.+?)__"

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the match
        if match.start() > last_end:
            paragraph.add_run(text[last_end : match.start()])

        # Add bold text
        bold_text = match.group(1) or match.group(2)
        bold_run = paragraph.add_run(bold_text)
        bold_run.bold = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def load_section_results(company_name: str) -> dict[str, str]:
    """Loads section data from the most recent working/{company_name}/ run folder."""
    section_results: dict[str, str] = {}
    working_dir = os.path.join(WORKING_DIR, company_name.replace(" ", "_"))

    if not os.path.exists(working_dir):
        console.error(f"Working directory not found: {working_dir}")
        return section_results

    # Find the most recent timestamped subfolder (e.g. 2026-02-18_2300)
    # Fall back to the flat directory if no subfolders exist
    search_dir = working_dir
    try:
        subdirs = sorted(
            [d for d in os.listdir(working_dir) if os.path.isdir(os.path.join(working_dir, d))],
            reverse=True,
        )
        if subdirs:
            search_dir = os.path.join(working_dir, subdirs[0])
    except OSError as e:
        console.error(f"Failed to list working directory {working_dir}: {e}")

    console.info(f"Loading section data from: {search_dir}")

    for _section_title, section_key in SECTION_KEY_MAP.items():
        file_path = os.path.join(search_dir, f"{section_key}.txt")
        if os.path.exists(file_path):
            with open(file_path, encoding="utf-8") as f:
                section_results[section_key] = f.read().strip()

    return section_results


def strip_markdown_artifacts(text: str) -> str:
    """
    Remove markdown formatting artifacts from text for clean plain-text output.

    Handles:
    - Bold markers (**text** -> text)
    - Heading markers (## -> removed)
    - Preserves bullet structure with clean formatting

    Note: We only strip double markers (**bold** and __bold__) to avoid
    false positives with single underscores in normal text.

    Args:
        text: Text potentially containing markdown

    Returns:
        Clean plain text
    """
    # Remove bold markers (double asterisks)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)

    # Remove bold markers (double underscores) - must have content between them
    # Use word boundaries to avoid matching things like __init__ or file__name
    text = re.sub(r"(?<!\w)__([^_]+)__(?!\w)", r"\1", text)

    # Remove heading markers but keep the text
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)

    return text


def save_report_as_txt(section_results, company_name, output_dir: str | Path | None = None):
    """
    Save AI-generated report as a clean TXT file.

    Produces well-formatted plain text with:
    - Clear section headers
    - Proper spacing between sections
    - Markdown artifacts removed
    - Consistent bullet formatting
    """
    date_str = datetime.now().strftime("%m-%d-%Y")
    file_name = f"{company_name}_Company_Overview_{date_str}.txt"
    destination_dir = Path(output_dir) if output_dir is not None else Path(OUTPUT_DIR)
    destination_dir.mkdir(parents=True, exist_ok=True)
    txt_path = destination_dir / file_name

    try:
        with open(txt_path, "w", encoding="utf-8") as f:
            # Title block
            f.write("=" * 60 + "\n")
            f.write(f"{company_name}: Strategic Company Overview\n")
            f.write(f"Prepared: {date_str}\n")
            f.write("=" * 60 + "\n\n")

            for section_title, section_key in SECTION_KEY_MAP.items():
                content = section_results.get(section_key, "").strip()
                if content:
                    # Section header
                    f.write("-" * 40 + "\n")
                    f.write(f"{section_title.upper()}\n")
                    f.write("-" * 40 + "\n\n")

                    # Clean content of markdown artifacts
                    clean_content = strip_markdown_artifacts(content)

                    # Normalize bullet formatting for plain text
                    lines = clean_content.split("\n")
                    for line in lines:
                        stripped = line.strip()
                        if stripped:
                            # Convert various bullet styles to consistent format
                            if stripped.startswith(("* ", "- ", "• ")):
                                # Preserve indentation for nested bullets
                                indent = len(line) - len(line.lstrip())
                                indent_str = "  " * (indent // 4)
                                bullet_content = stripped[2:].strip()
                                f.write(f"{indent_str}* {bullet_content}\n")
                            else:
                                f.write(f"{stripped}\n")
                        else:
                            f.write("\n")

                    f.write("\n")

            # Footer
            f.write("=" * 60 + "\n")
            f.write("This report was generated using AI-assisted research.\n")
            f.write("Please verify key information independently.\n")
            f.write("=" * 60 + "\n")

        console.ok("TXT saved")
        return str(txt_path)

    except Exception as e:
        console.error(f"Failed to save TXT report: {e}")
        return None


def save_report_as_docx_premium(
    section_results,
    company_name,
    citation_style="numbered",
    output_dir: str | Path | None = None,
):
    """
    Generate premium consultant-grade DOCX using DocumentBuilder.

    This is the new premium report generation that includes:
    - Professional cover page with one-liner summary
    - Company snapshot table
    - Executive summary with key insights and risk factors
    - Table of contents
    - 5 logical chapters with proper hierarchy
    - Professional styling and typography
    - Clean numbered citations [1] style (configurable)

    Args:
        section_results: Dict mapping section_key to content string
        company_name: Name of the company
        citation_style: Citation formatting - "numbered" (default), "inline", or "sidecar"

    Returns:
        Path to generated DOCX file, or None on failure
    """
    from primr.output.citation_processor import CitationProcessor, CitationStyle

    date_str = datetime.now().strftime("%m-%d-%Y")
    file_name = f"{company_name}_Strategic_Overview_{date_str}.docx"
    destination_dir = Path(output_dir) if output_dir is not None else Path(OUTPUT_DIR)
    destination_dir.mkdir(parents=True, exist_ok=True)
    docx_path = destination_dir / file_name

    try:
        # Map string to enum
        style_map = {
            "numbered": CitationStyle.NUMBERED,
            "inline": CitationStyle.INLINE,
            "sidecar": CitationStyle.SIDECAR,
        }
        style = style_map.get(citation_style, CitationStyle.NUMBERED)

        # Process citations if not inline style
        processed_sections = section_results
        citations = []

        if style != CitationStyle.INLINE:
            processor = CitationProcessor(style=style)
            processed_sections = {}
            for key, content in section_results.items():
                result = processor.process_content(content)
                processed_sections[key] = result.transformed_content
            citations = [{"title": c.title, "url": c.url} for c in processor.citations]

            # Generate sidecar file if requested
            if style == CitationStyle.SIDECAR:
                sidecar_filename, sidecar_content = processor.generate_sidecar_file(company_name)
                sidecar_path = destination_dir / sidecar_filename
                with open(sidecar_path, "w", encoding="utf-8") as f:
                    f.write(sidecar_content)
                console.ok(f"Sources file: {sidecar_filename}")

        builder = DocumentBuilder(company_name, processed_sections, citations=citations)
        document = builder.build()
        document.save(docx_path)
        console.ok("Premium DOCX saved")
        return str(docx_path)
    except Exception as e:
        console.error(f"Failed to generate premium DOCX: {e}")
        # Fall back to legacy generation
        return None


def save_report_as_docx(txt_path, company_name, output_dir: str | Path | None = None):
    """
    Convert markdown TXT content into a formatted DOCX document.

    This is the legacy function maintained for backward compatibility.
    For premium reports, use save_report_as_docx_premium() instead.
    """
    date_str = datetime.now().strftime("%m-%d-%Y")
    file_name = f"{company_name}_Company_Overview_{date_str}.docx"
    destination_dir = Path(output_dir) if output_dir is not None else Path(OUTPUT_DIR)
    destination_dir.mkdir(parents=True, exist_ok=True)
    docx_path = destination_dir / file_name

    try:
        document = Document()

        # Set up document styles
        style = document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        with open(txt_path, encoding="utf-8") as f:
            lines = f.readlines()

        # Title
        title = document.add_heading(lines[0].strip(), level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Date
        date_para = document.add_paragraph(lines[1].strip())
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Process content
        current_text_buffer: list[str] = []

        for line in lines[2:]:
            line_type, content, indent = parse_markdown_line(line)

            # Flush text buffer before structural elements
            if line_type in ("heading", "subheading", "bullet", "numbered") and current_text_buffer:
                para = document.add_paragraph()
                apply_inline_formatting(para, " ".join(current_text_buffer))
                current_text_buffer = []

            if line_type == "empty":
                if current_text_buffer:
                    para = document.add_paragraph()
                    apply_inline_formatting(para, " ".join(current_text_buffer))
                    current_text_buffer = []

            elif line_type == "heading":
                document.add_heading(content, level=1)

            elif line_type == "subheading":
                document.add_heading(content, level=2)

            elif line_type == "bullet":
                para = document.add_paragraph(style="List Bullet")
                if indent > 0:
                    para.paragraph_format.left_indent = Inches(0.25 * indent)
                apply_inline_formatting(para, content)

            elif line_type == "numbered":
                para = document.add_paragraph(style="List Number")
                if indent > 0:
                    para.paragraph_format.left_indent = Inches(0.25 * indent)
                apply_inline_formatting(para, content)

            elif line_type == "text":
                # Accumulate text for paragraph wrapping
                current_text_buffer.append(content)

        # Flush remaining buffer
        if current_text_buffer:
            para = document.add_paragraph()
            apply_inline_formatting(para, " ".join(current_text_buffer))

        document.save(docx_path)
        console.ok("DOCX saved")
        return str(docx_path)

    except Exception as e:
        console.error(f"Failed to convert TXT to DOCX: {e}")
        return None


def convert_docx_to_pdf(docx_path: str | Path) -> str | None:
    """
    Convert DOCX to PDF and return the PDF path on success.

    Strategy:
    1) Try ``docx2pdf`` (best on Windows/macOS with Office available)
    2) Fallback to LibreOffice ``soffice --headless --convert-to pdf`` when available
    """
    import subprocess
    import sys

    docx_file = Path(docx_path)
    if not docx_file.exists():
        console.warn(f"PDF conversion skipped: DOCX not found ({docx_file})")
        return None

    pdf_file = docx_file.with_suffix(".pdf")

    try:
        from docx2pdf import convert

        old_stderr = sys.stderr
        sys.stderr = StringIO()
        try:
            convert(str(docx_file))
        finally:
            sys.stderr = old_stderr

        if pdf_file.exists():
            console.ok("PDF saved")
            return str(pdf_file)
    except Exception as e:
        console.warn(f"docx2pdf unavailable, trying LibreOffice fallback: {e}")

    soffice = shutil.which("soffice")
    if not soffice:
        console.warn("PDF conversion unavailable: no supported converter found")
        return None

    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(docx_file.parent),
                str(docx_file),
            ],
            check=True,
            capture_output=True,
        )
        if pdf_file.exists():
            console.ok("PDF saved")
            return str(pdf_file)
        console.warn("PDF conversion skipped: converter returned without creating PDF")
        return None
    except Exception as e:
        console.warn(f"PDF conversion unavailable: {e}")
        # Best-effort cleanup if converter produced partial output
        with suppress(OSError):
            if pdf_file.exists() and pdf_file.stat().st_size == 0:
                pdf_file.unlink()
        return None


def zip_research_files(company_name):
    """Zips the research files for the company and moves to the output directory."""
    try:
        date_str = datetime.now().strftime("%m-%d-%Y")
        zip_filename = f"{company_name}_research_{date_str}.zip"
        zip_filepath = os.path.join(OUTPUT_DIR, zip_filename)

        company_folder = os.path.join(WORKING_DIR, company_name.replace(" ", "_"))
        if os.path.exists(company_folder):
            with zipfile.ZipFile(zip_filepath, "w", zipfile.ZIP_DEFLATED) as zipf:
                for root, _, files in os.walk(company_folder):
                    for file in files:
                        file_path = os.path.join(root, file)
                        archive_name = os.path.relpath(file_path, WORKING_DIR)
                        zipf.write(file_path, archive_name)

            console.ok("Research files archived")

    except Exception as e:
        console.warn(f"Archive creation skipped: {e}")


def cleanup(company_name):
    """Handles cleanup tasks: generates ZIP archive and removes temporary working files."""
    try:
        zip_research_files(company_name)
        company_folder = os.path.join(WORKING_DIR, company_name.replace(" ", "_"))
        if os.path.exists(company_folder):
            shutil.rmtree(company_folder)

    except Exception as e:
        console.warn(f"Cleanup incomplete: {e}")


def generate_final_report(
    company_name: str,
    premium: bool = True,
    citation_style: str = "numbered",
    output_dir: str | Path | None = None,
    diagnostics_dir: str | Path | None = None,
    write_txt: bool = True,
) -> str | None:
    """
    Generates the final structured report in TXT, DOCX, and ZIP archive formats.

    Args:
        company_name: Name of the company
        premium: If True, generate premium consultant-grade report (default)
                 If False, use legacy simple format
        citation_style: Citation formatting style - "numbered" (default), "inline", or "sidecar"

    Returns:
        Path to generated DOCX file, or None on failure
    """
    section_results = load_section_results(company_name)
    if not section_results:
        console.error("No section data available for report")
        return None

    txt_path = None
    if write_txt or diagnostics_dir is not None:
        txt_output_dir = output_dir if write_txt else diagnostics_dir
        txt_path = save_report_as_txt(section_results, company_name, output_dir=txt_output_dir)

    # Generate DOCX - try premium first, fall back to legacy
    docx_path = None
    if premium:
        docx_path = save_report_as_docx_premium(
            section_results, company_name, citation_style, output_dir=output_dir
        )

    # Fall back to legacy if premium fails or not requested
    if not docx_path and txt_path:
        docx_path = save_report_as_docx(txt_path, company_name, output_dir=output_dir)

    if docx_path:
        convert_docx_to_pdf(docx_path)

    cleanup(company_name)

    return docx_path
