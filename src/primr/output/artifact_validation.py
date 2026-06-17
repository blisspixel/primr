"""Final-artifact validation and salvage helpers.

Extracted from `primr.core.research_agent` for isolated unit testing.

This module holds the single source of truth for forbidden internal markers
that must never reach a shipped artifact, plus the scanner / auto-strip /
DOCX text-extractor pair that the pipeline uses to fail closed when the
shipping prep stage leaves residue behind.
"""

from __future__ import annotations

import logging
import os
import re
from pathlib import Path
from typing import Any, TypedDict

logger = logging.getLogger(__name__)

# Configurable warning ceiling for leaked internal-scaffolding markers in a
# shipped report (bare [workbook] / [cross-ref ...] refs, bold-wrapped "What to
# validate:" lines, informal [cite: label] markers). Default 0 = zero tolerance
# for warning visibility only. A leak that survived the upstream canonicalization
# seam is logged and eval-tracked, but it no longer blocks the polished DOCX.
_SCAFFOLDING_LEAK_THRESHOLD_ENV = "PRIMR_MAX_SCAFFOLDING_LEAKS"


def _scaffolding_leak_threshold() -> int:
    """Resolve the max tolerated scaffolding-leak count from the environment.

    Defaults to 0 (zero warning tolerance). A malformed or negative value falls
    back to 0 so the visibility signal can never be silently disabled by a bad
    env value.
    """
    raw = os.environ.get(_SCAFFOLDING_LEAK_THRESHOLD_ENV)
    if raw is None:
        return 0
    try:
        return max(0, int(raw))
    except (TypeError, ValueError):
        logger.warning(
            "Invalid %s=%r; falling back to zero-tolerance (0)",
            _SCAFFOLDING_LEAK_THRESHOLD_ENV,
            raw,
        )
        return 0


# Configurable ceiling for dangling inline citations (`[cite: N]` with no
# matching Sources-appendix entry) in a shipped report. Default 0 = zero
# tolerance: a citation that did not resolve after the upstream LLM repair
# blocks the polished DOCX (MD/TXT + sidecar validation report still written).
# Malformed/negative values fall back to 0 so the gate can't be silently
# disabled. Operators can relax it via PRIMR_MAX_DANGLING_CITATIONS.
_DANGLING_CITATION_THRESHOLD_ENV = "PRIMR_MAX_DANGLING_CITATIONS"


def _dangling_citation_threshold() -> int:
    """Resolve the max tolerated dangling-citation count from the environment."""
    raw = os.environ.get(_DANGLING_CITATION_THRESHOLD_ENV)
    if raw is None:
        return 0
    try:
        return max(0, int(raw))
    except (TypeError, ValueError):
        logger.warning(
            "Invalid %s=%r; falling back to zero-tolerance (0)",
            _DANGLING_CITATION_THRESHOLD_ENV,
            raw,
        )
        return 0


def _scan_citation_integrity_issues(markdown_content: str, threshold: int) -> list[str]:
    """Return shipping-gate issue strings when dangling citations exceed threshold.

    Empty list when integrity is within the configured threshold. Detection
    lives in ``primr.qa.report_analyzer.scan_citation_integrity`` (single source
    of truth, shared with the QA layer's citation concept).
    """
    from primr.qa.report_analyzer import scan_citation_integrity

    result = scan_citation_integrity(markdown_content)
    if result["missing_count"] <= threshold:
        return []

    missing_preview = ", ".join(str(n) for n in result["missing_citations"][:10])
    detail = (
        "no Sources appendix"
        if not result["has_bibliography"]
        else f"unresolved: {missing_preview}"
    )
    return [
        f"citation_integrity:dangling={result['missing_count']} (threshold {threshold}; {detail})"
    ]


# Configurable ceiling for unambiguous structural defects in a shipped report
# (duplicate top-level `##` headings, empty sections). Default 0 = zero
# tolerance. Malformed/negative values fall back to 0. Relax via
# PRIMR_MAX_STRUCTURE_DEFECTS. Only catches always-broken structure — required-
# section presence is deliberately NOT gated (report-type-dependent / heuristic).
_STRUCTURE_DEFECT_THRESHOLD_ENV = "PRIMR_MAX_STRUCTURE_DEFECTS"


def _structure_defect_threshold() -> int:
    """Resolve the max tolerated structural-defect count from the environment."""
    raw = os.environ.get(_STRUCTURE_DEFECT_THRESHOLD_ENV)
    if raw is None:
        return 0
    try:
        return max(0, int(raw))
    except (TypeError, ValueError):
        logger.warning(
            "Invalid %s=%r; falling back to zero-tolerance (0)",
            _STRUCTURE_DEFECT_THRESHOLD_ENV,
            raw,
        )
        return 0


def _scan_section_structure_issues(markdown_content: str, threshold: int) -> list[str]:
    """Return shipping-gate issue strings when structural defects exceed threshold.

    Empty list when structure is within the configured threshold. Detection
    lives in ``primr.qa.report_analyzer.scan_section_structure`` (single source
    of truth). Catches duplicate `##` headings and empty sections only.
    """
    from primr.qa.report_analyzer import scan_section_structure

    result = scan_section_structure(markdown_content)
    if result["total_defects"] <= threshold:
        return []

    issues: list[str] = [
        f"section_structure:defects={result['total_defects']} (threshold {threshold})"
    ]
    if result["duplicate_headings"]:
        issues.append("section_structure:duplicate=" + ", ".join(result["duplicate_headings"][:10]))
    if result["empty_sections"]:
        issues.append("section_structure:empty=" + ", ".join(result["empty_sections"][:10]))
    return issues


def _scan_scaffolding_leakage_issues(markdown_content: str, threshold: int) -> list[str]:
    """Return shipping-gate issue strings when scaffolding leaks exceed threshold.

    Empty list when the leak count is within the configured threshold. The
    detection logic lives in ``primr.qa.report_analyzer.scan_scaffolding_leakage``
    (single source of truth, shared with the QA scorecard).
    """
    from primr.qa.report_analyzer import scan_scaffolding_leakage

    leak = scan_scaffolding_leakage(markdown_content)
    if leak["total_leaked"] <= threshold:
        return []

    issues = [f"scaffolding_leak:total={leak['total_leaked']} (threshold {threshold})"]
    for key, label in (
        ("workbook_markers", "workbook_markers"),
        ("cross_ref_markers", "cross_ref_markers"),
        ("bare_bold_validate", "bold_validate_lines"),
        ("informal_cite_markers", "informal_cite_markers"),
    ):
        if leak[key]:
            issues.append(f"scaffolding_leak:{label}={leak[key]}")
    return issues


# Bracketed/filename internal tokens. These are matched case-insensitively
# because the delimiters ([], the .txt suffix) never occur in legitimate prose,
# so IGNORECASE cannot false-block real content. Detection is partial-match (no
# closing-bracket requirement) so the scanner catches truncated tokens too.
_FORBIDDEN_OUTPUT_PATTERNS: tuple[tuple[str, str], ...] = (
    ("raw_source_tag", r"\[Source:\s*(?:https?://)?[^\]\s]+"),
    ("section_cross_ref", r"\[\s*(?:see|cross-?ref|xref)\s+##\s+[^\]]+\]"),
    ("workbook_ref", r"\[Workbook:[^\]]*\]"),
    ("workbook_section_ref", r"\[workbook section[^\]]*\]"),
    ("workbook_section_symbol", r"\[Workbook §[^\]]*\]"),
    ("analysis_workbook_ref", r"\[Analysis Workbook[^\]]*\]"),
    ("analysis_ref", r"\[Analysis:[^\]]*\]"),
    ("external_sources_ref", r"\[External Sources\]"),
    ("citation_inventory", r"\[citation inventory[^\]]*\]"),
    ("vendor_research_file", r"vendor-research-[\w.-]+\.txt"),
)

# Cleaner patterns require the full closing-bracket form so substitution
# removes only the well-formed token (not arbitrary trailing text).
_FORBIDDEN_OUTPUT_CLEANERS: tuple[tuple[str, str], ...] = (
    ("raw_source_tag", r"\[Source:[^\]]*\]"),
    ("section_cross_ref", r"\[\s*(?:see|cross-?ref|xref)\s+##\s+[^\]]+\]"),
    ("workbook_ref", r"\[Workbook:[^\]]*\]"),
    ("workbook_section_ref", r"\[workbook section[^\]]*\]"),
    ("workbook_section_symbol", r"\[Workbook §[^\]]*\]"),
    ("analysis_workbook_ref", r"\[Analysis Workbook[^\]]*\]"),
    ("analysis_ref", r"\[Analysis:[^\]]*\]"),
    ("external_sources_ref", r"\[External Sources\]"),
    ("citation_inventory", r"\[citation inventory[^\]]*\]"),
    ("vendor_research_file", r"vendor-research-[\w.-]+\.txt"),
)

# Internal workbook labels that leak as Title-Case headers/labels. Matched
# CASE-SENSITIVELY (unlike the bracketed tokens above) so legitimate lowercase
# prose never false-blocks shipping: a report may say "based on our internal
# analysis" or "in the analysis context of X" - that is real content, not a
# leak, and gating it is exactly the brittle trap (agentic-balance.md). Only the
# exact Title-Case label form (the way the analysis workbook emits it) is caught.
_FORBIDDEN_LEAKED_LABELS: tuple[tuple[str, str], ...] = (
    ("internal_roi_model", r"\bInternal ROI Model\b"),
    ("internal_analysis", r"\bInternal Analysis\b"),
    ("analysis_context", r"\bAnalysis Context\b"),
)

# Bare internal terms (no bracket form) that must never leak.
_FORBIDDEN_INTERNAL_TERMS: tuple[str, ...] = ("vendor-research",)


class _ArtifactValidation(TypedDict):
    """Result of an artifact validation pass.

    ``issues`` are *blocking* — structural/referential validity and unambiguous
    internal-token leaks that withhold the polished DOCX. ``warnings`` are
    *non-blocking* content signals (the scaffolding-leak scan) that are surfaced
    and eval-tracked but never withhold the deliverable: a regex cannot be the
    quality moat (see docs/design/agentic-balance.md), so content-shape findings
    are reported, not gated.
    """

    passed: bool
    issues: list[str]
    warnings: list[str]
    errors: list[str]


def _auto_strip_forbidden_patterns(text: str) -> str:
    """Last-resort defensive sweep: strip anything the artifact scanner would flag."""
    if not text.strip():
        return text

    for _label, pattern in _FORBIDDEN_OUTPUT_CLEANERS:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)

    # Leaked Title-Case labels: case-sensitive so lowercase prose is preserved.
    for _label, pattern in _FORBIDDEN_LEAKED_LABELS:
        text = re.sub(pattern, "", text)

    lower = text.lower()
    for term in _FORBIDDEN_INTERNAL_TERMS:
        if term in lower:
            text = re.sub(re.escape(term), "", text, flags=re.IGNORECASE)

    text = re.sub(r"  +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text


def _scan_forbidden_output_patterns(text: str) -> list[str]:
    """Return a list of human-readable issue strings, one per detected pattern."""
    issues: list[str] = []
    for label, pattern in _FORBIDDEN_OUTPUT_PATTERNS:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            issues.append(f"{label}: {match.group(0)[:120]}")

    # Leaked Title-Case labels: case-sensitive so legitimate lowercase prose
    # ("our internal analysis") never false-blocks shipping.
    for label, pattern in _FORBIDDEN_LEAKED_LABELS:
        match = re.search(pattern, text)
        if match:
            issues.append(f"{label}: {match.group(0)[:120]}")

    lower = text.lower()
    for term in _FORBIDDEN_INTERNAL_TERMS:
        if term in lower:
            issues.append(f"internal_term: {term}")

    return issues


def _validate_output_markdown(
    markdown_content: str,
    *,
    scaffolding_threshold: int | None = None,
    citation_threshold: int | None = None,
    structure_threshold: int | None = None,
) -> _ArtifactValidation:
    """Validate that a markdown artifact is ship-ready. All checks fail-closed:

    - zero-tolerance forbidden-marker scan (raw [Source:], [Workbook:], etc.);
    - a configurable scaffolding-leak warning (bare [workbook]/[cross-ref],
      bold "What to validate:" lines, informal [cite: label]) - default 0,
      override via ``PRIMR_MAX_SCAFFOLDING_LEAKS``;
    - a configurable citation-integrity gate (inline [cite: N] with no matching
      Sources-appendix entry) — default 0, override via
      ``PRIMR_MAX_DANGLING_CITATIONS``;
    - a configurable section-structure gate (duplicate ``##`` headings, empty
      sections) — default 0, override via ``PRIMR_MAX_STRUCTURE_DEFECTS``.

    Blocking issues (withhold the polished DOCX; MD/TXT + sidecar still written):
    the zero-tolerance forbidden-token scan, citation-integrity, and
    section-structure. The scaffolding-leak scan is a NON-blocking *warning*: a
    regex can't be the quality moat (see docs/design/agentic-balance.md), so a
    leaked marker is surfaced and eval-tracked (`## Artifact Drift`) but does not
    withhold the deliverable. Quality is enforced upstream (the writer prompt) and
    measured by eval, not gated here.
    """
    if scaffolding_threshold is None:
        scaffolding_threshold = _scaffolding_leak_threshold()
    if citation_threshold is None:
        citation_threshold = _dangling_citation_threshold()
    if structure_threshold is None:
        structure_threshold = _structure_defect_threshold()
    try:
        # Blocking: unambiguous internal-token leaks + prose-invariant structure.
        issues = _scan_forbidden_output_patterns(markdown_content)
        issues.extend(_scan_citation_integrity_issues(markdown_content, citation_threshold))
        issues.extend(_scan_section_structure_issues(markdown_content, structure_threshold))
        # Non-blocking content signal: surfaced + eval-tracked, never withholds.
        warnings = _scan_scaffolding_leakage_issues(markdown_content, scaffolding_threshold)
        if warnings:
            logger.warning(
                "Scaffolding-leak warning (non-blocking; shipped): %s", "; ".join(warnings)
            )
        return {"passed": len(issues) == 0, "issues": issues, "warnings": warnings, "errors": []}
    except Exception as exc:
        # Fail closed: an exception inside the scanner means we could not
        # confirm the artifact is clean. Downstream code writes a sidecar
        # validation report and blocks DOCX shipping when this returns False.
        logger.warning("Markdown artifact validation failed: %s", exc)
        return {"passed": False, "issues": [], "warnings": [], "errors": [str(exc)]}


def _extract_docx_text(document: Any) -> str:
    """Flatten a python-docx Document into a single text string."""
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _validate_output_docx(docx_path: Path) -> _ArtifactValidation:
    """Validate that a generated DOCX contains no markdown artifacts or forbidden tokens."""
    try:
        from docx import Document

        from primr.output.markdown_parser import ArtifactDetector

        document = Document(str(docx_path))
        detector = ArtifactDetector()
        artifacts = detector.scan_document(document)
        issues = [
            f"markdown_artifact:{artifact['type']}:{artifact['match']}"
            for artifact in artifacts[:10]
        ]
        issues.extend(_scan_forbidden_output_patterns(_extract_docx_text(document)))
        return {"passed": len(issues) == 0, "issues": issues, "warnings": [], "errors": []}
    except Exception as exc:
        # Fail closed — see _validate_output_markdown for the rationale.
        logger.warning("DOCX artifact validation failed: %s", exc)
        return {"passed": False, "issues": [], "warnings": [], "errors": [str(exc)]}


def _write_output_validation_report(
    base_path: Path,
    phase: str,
    issues: list[str],
    errors: list[str],
    diagnostics_dir: str | Path | None = None,
) -> Path | None:
    """Write a sidecar text report describing why an artifact failed validation."""
    if not issues and not errors:
        return None

    if diagnostics_dir is not None:
        diagnostics_path = Path(diagnostics_dir)
        diagnostics_path.mkdir(parents=True, exist_ok=True)
        report_path = diagnostics_path / f"{base_path.stem}_{phase}_validation.txt"
    else:
        report_path = base_path.with_name(f"{base_path.stem}_{phase}_validation.txt")
    lines = [f"Artifact validation report ({phase})", ""]
    if issues:
        lines.append("Issues:")
        lines.extend(f"- {item}" for item in issues)
        lines.append("")
    if errors:
        lines.append("Validator errors:")
        lines.extend(f"- {item}" for item in errors)
        lines.append("")
    report_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return report_path
