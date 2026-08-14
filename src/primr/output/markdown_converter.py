"""
Clean markdown to DOCX converter.

Simple, direct conversion without excessive abstraction.
Handles: headings, bullets, bold, italic, links, tables.

Design decision: Converter owns the header.
- Title and subtitle are passed as parameters, rendered by the converter
- Markdown content should NOT include a top-level header block
- If markdown has a header block, it gets stripped automatically
"""

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def sanitize_text(text: str) -> str:
    """
    Remove characters that are invalid in XML/DOCX.

    DOCX files are XML-based and cannot contain:
    - NULL bytes (\\x00)
    - Control characters (\\x01-\\x08, \\x0B, \\x0C, \\x0E-\\x1F)

    Args:
        text: Input text that may contain invalid characters

    Returns:
        Sanitized text safe for DOCX
    """
    if not text:
        return text
    # Remove NULL and control characters (except tab, newline, carriage return)
    # XML 1.0 allows: #x9 (tab), #xA (newline), #xD (carriage return), #x20-#xD7FF, etc.
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    """Add a hyperlink to a paragraph."""
    # Sanitize text and URL
    text = sanitize_text(text)
    url = sanitize_text(url)
    if not text or not url:
        return
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def parse_inline_markdown(paragraph: Any, text: str) -> None:
    """Parse inline markdown (bold, italic, code, links) and add to paragraph."""
    # Sanitize text first to remove invalid XML characters
    text = sanitize_text(text)
    if not text:
        return
    # Pattern: **bold**, *italic*, `code`, [text](url)
    # italic: guarded so "5*3" / "a*b" math is not mis-italicized (the "*" must
    #   not be flanked by word chars or another "*").
    # link: URL group allows one level of balanced parens so "..._(company)"
    #   Wikipedia-style links are not truncated at the first ")".
    pattern = r"(\*\*.*?\*\*|(?<![\w*])\*[^*]+\*(?![\w*])|`[^`]+`|\[[^\]]+\]\([^()]*(?:\([^()]*\)[^()]*)*\))"
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        # Bold
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        # Italic (single asterisk, not double)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        # Code
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        # Link
        elif part.startswith("[") and "](" in part:
            match = re.match(r"\[(.*?)\]\((.*)\)", part)
            if match:
                add_hyperlink(paragraph, match.group(1), match.group(2))
        # Regular text
        else:
            paragraph.add_run(part)


def render_table(doc: Document, table_lines: list[str]) -> None:
    """Render markdown table as DOCX table."""
    if not table_lines or len(table_lines) < 2:
        return
    rows = []
    for line in table_lines:
        # Skip separator lines (|---|---|). The character class must include
        # `|` so multi-column separators (which carry internal pipes between
        # columns) are matched and not rendered as a spurious data row.
        if re.match(r"^\s*\|[\s\-:|]+\|\s*$", line):
            continue
        cells = [cell.strip() for cell in line.split("|")]
        if cells and not cells[0]:
            cells = cells[1:]
        if cells and not cells[-1]:
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    if not rows:
        return
    # Create table
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                parse_inline_markdown(cell.paragraphs[0], cell_text)
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True


def _is_table_separator(line: str) -> bool:
    """True for a markdown table separator row like ``|---|---|`` (only pipes,
    dashes, colons, and spaces, with at least one dash)."""
    s = line.strip()
    return bool(s) and "-" in s and set(s) <= {"|", "-", ":", " "}


def _flush_table_block(doc: Document, table_lines: list[str]) -> None:
    """Render buffered potential-table lines.

    Only render a DOCX table when the block is an actual markdown table (it has
    a ``|---|`` separator row). Otherwise render each line as a normal
    paragraph: ``render_table`` no-ops on blocks with <2 lines or no parsed
    rows, and the callers would otherwise silently drop prose that merely
    contains a ``|`` (e.g. "Strengths | Weaknesses").
    """
    if not table_lines:
        return
    if len(table_lines) >= 2 and any(_is_table_separator(ln) for ln in table_lines):
        render_table(doc, table_lines)
        return
    for ln in table_lines:
        if ln.strip():
            parse_inline_markdown(doc.add_paragraph(), ln.strip())


def strip_heading_markers(text: str) -> str:
    """Remove markdown formatting from heading text."""
    text = sanitize_text(text)
    if not text:
        return ""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    return text.strip()


def _is_header_metadata_line(line: str) -> bool:
    if line.startswith("**") and ":**" in line:
        return True
    if re.match(r"^\*[^*\n]+\*(?:\s*\|.*)?$", line):
        return True
    return bool(re.match(r"^(?:Date|Website|Generated):\s+\S", line, re.IGNORECASE))


def strip_markdown_header_block(markdown_text: str) -> str:
    """
    Strip the header block from markdown content.

    The header block typically looks like:
        # Title

        **Status:** Draft
        **Date:** December 18, 2024

        ---

    Since the converter owns the header (via title/subtitle params),
    we strip this to avoid duplicate headers in the DOCX.

    Args:
        markdown_text: Raw markdown content

    Returns:
        Markdown with header block removed
    """
    if not markdown_text:
        return markdown_text

    lines = markdown_text.split("\n")

    # Find where the header block ends
    # Look for: # Title, then metadata lines, then ---
    i = 0

    # Skip leading empty lines
    while i < len(lines) and not lines[i].strip():
        i += 1

    # Check if first non-empty line is a top-level heading
    if i < len(lines) and lines[i].strip().startswith("# "):
        i += 1

        # Skip empty lines after heading
        while i < len(lines) and not lines[i].strip():
            i += 1

        # Skip metadata lines.
        while i < len(lines):
            line = lines[i].strip()
            if _is_header_metadata_line(line) or not line:
                i += 1
            elif line == "---":
                i += 1
                break
            else:
                break

    # Return remaining content
    return "\n".join(lines[i:])


def _is_blockquote_line(line: str) -> bool:
    return line.strip().startswith(">")


def _blockquote_text(line: str) -> str:
    stripped = line.strip()
    if stripped.startswith(">"):
        return stripped[1:].lstrip()
    return ""


def _render_code_block(doc: Document, code_lines: list[str]) -> None:
    """Render a fenced code block as a monospaced paragraph."""
    text = sanitize_text("\n".join(code_lines))
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _render_body_lines(doc: Document, lines: list[str], *, allow_h1: bool) -> None:
    """Render markdown body lines into ``doc``. Shared by full and section renders."""
    i = 0
    in_table = False
    table_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        if not line.strip() and not in_table:
            i += 1
            continue
        stripped = line.strip()
        if not in_table and stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _render_code_block(doc, code_lines)
            if i < len(lines):
                i += 1
            continue
        if "|" in line and not in_table:
            next_line = lines[i + 1] if i + 1 < len(lines) else ""
            if _is_table_separator(line) or _is_table_separator(next_line):
                in_table = True
                table_lines = [line]
                i += 1
                continue
        if in_table:
            if "|" in line:
                table_lines.append(line)
                i += 1
                continue
            _flush_table_block(doc, table_lines)
            in_table = False
            table_lines = []
            continue
        if stripped.startswith("#### "):
            doc.add_heading(strip_heading_markers(stripped[5:]), level=4)
        elif stripped.startswith("### "):
            doc.add_heading(strip_heading_markers(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(strip_heading_markers(stripped[3:]), level=2)
        elif allow_h1 and stripped.startswith("# "):
            doc.add_heading(strip_heading_markers(stripped[2:]), level=1)
        elif _is_blockquote_line(line):
            quote_parts = [_blockquote_text(line)]
            j = i + 1
            while j < len(lines) and _is_blockquote_line(lines[j]):
                quote_parts.append(_blockquote_text(lines[j]))
                j += 1
            paragraph = doc.add_paragraph(style="Quote")
            parse_inline_markdown(paragraph, "\n".join(quote_parts))
            i = j
            continue
        elif stripped.startswith(("- ", "* ")):
            paragraph = doc.add_paragraph(style="List Bullet")
            parse_inline_markdown(paragraph, stripped[2:].strip())
        elif re.match(r"^\d+[.)]\s", stripped):
            match = re.match(r"^\d+[.)]\s+(.*)", stripped)
            if match:
                paragraph = doc.add_paragraph(style="List Number")
                parse_inline_markdown(paragraph, match.group(1))
        elif stripped.startswith("---"):
            pass
        elif stripped:
            paragraph = doc.add_paragraph()
            parse_inline_markdown(paragraph, stripped)
        i += 1
    if in_table and table_lines:
        _flush_table_block(doc, table_lines)


def setup_document_styles(doc: Document) -> None:
    """
    Set up consistent document styles.

    Explicitly defines styles to avoid environment-dependent defaults.
    Uses reduced heading spacing for a cleaner, less "Word-y" look.
    """
    styles = doc.styles

    # Normal text style
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    # Heading styles with graduated spacing (less chunky than Word defaults)
    # Format: (style_name, font_size, bold, space_before, space_after)
    heading_configs = [
        ("Heading 1", 18, True, 10, 6),  # Major sections
        ("Heading 2", 14, True, 10, 4),  # Sub-sections
        ("Heading 3", 12, True, 8, 4),  # Sub-sub-sections
        ("Heading 4", 11, True, 6, 2),  # Minor headings
    ]

    for style_name, size, bold, before, after in heading_configs:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = bold
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)


def markdown_to_docx(
    markdown_text: str, output_path: Path, title: str | None = None, subtitle: str | None = None
) -> Path:
    """
    Convert markdown to DOCX with clean formatting.

    Design: Converter owns the header.
    - If title/subtitle provided, they are rendered by the converter
    - Any markdown header block (# Title, metadata fields, etc.) is stripped
    - This prevents duplicate headers in the output

    Args:
        markdown_text: Markdown content
        output_path: Where to save the DOCX
        title: Optional document title (rendered as centered heading)
        subtitle: Optional subtitle (rendered as gray centered text)

    Returns:
        Path to created DOCX
    """
    doc = Document()

    # Set up consistent styles
    setup_document_styles(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add title if provided (converter owns the header)
    # Using 20pt for cleaner "memo" feel vs 24pt "report cover page" look
    if title:
        p = doc.add_heading(title, level=0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.size = Pt(20)
            run.font.name = "Calibri"

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(12)  # Smaller than body, muted
        run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)
        run.font.name = "Calibri"

    # Strip markdown header block if converter is adding title
    # This prevents duplicate headers
    content = markdown_text
    if title:
        content = strip_markdown_header_block(markdown_text)

    _render_body_lines(doc, content.split("\n"), allow_h1=True)
    doc.save(output_path)
    return output_path


def render_section_content(doc: Document, content: str) -> None:
    """
    Render markdown section content directly to a document.

    This is for adding content to an existing document without
    creating a new file.

    Args:
        doc: Existing Document object
        content: Markdown content to render
    """
    _render_body_lines(doc, content.split("\n"), allow_h1=False)
