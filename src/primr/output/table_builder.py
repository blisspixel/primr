"""
TableBuilder for premium report generation.

Creates professionally styled tables for company snapshots, metrics,
and callout boxes.
"""

from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

from primr.output.models import CompanySnapshot


class TableBuilder:
    """Builds professionally styled tables."""

    # Color constants (matching StyleEngine)
    TABLE_HEADER_BG = RGBColor(0, 51, 102)  # Navy - table headers
    TABLE_ALT_ROW = RGBColor(245, 247, 250)  # Light gray-blue - alternating rows
    TEXT_COLOR = RGBColor(51, 51, 51)  # Dark gray - body text
    WHITE = RGBColor(255, 255, 255)  # White - header text
    CALLOUT_BG = RGBColor(240, 248, 255)  # Alice blue - callout backgrounds

    def __init__(self, document: Document):
        """
        Initialize TableBuilder with a Document.

        Args:
            document: A python-docx Document object
        """
        self.document = document

    def create_company_snapshot(self, snapshot: CompanySnapshot) -> Table:
        """
        Create the company snapshot box (2-column key-value table).

        Args:
            snapshot: CompanySnapshot dataclass with company info

        Returns:
            Styled Table object
        """
        # Build data rows from snapshot
        rows_data = []

        if snapshot.industry:
            rows_data.append(("Industry", snapshot.industry))
        if snapshot.founded:
            rows_data.append(("Founded", snapshot.founded))
        if snapshot.headquarters:
            rows_data.append(("Headquarters", snapshot.headquarters))
        if snapshot.revenue:
            rows_data.append(("Revenue", snapshot.revenue))
        if snapshot.employees:
            rows_data.append(("Employees", snapshot.employees))
        if snapshot.ticker:
            rows_data.append(("Ticker", snapshot.ticker))
        if snapshot.website:
            rows_data.append(("Website", snapshot.website))

        if not rows_data:
            rows_data = [("Company", snapshot.company_name)]

        # Create table
        table = self.document.add_table(rows=len(rows_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Populate and style
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]

            # Label cell (bold)
            label_cell = row.cells[0]
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(10)

            # Value cell
            value_cell = row.cells[1]
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(str(value) if value else "")
            value_run.font.size = Pt(10)

        self.apply_table_style(table, "snapshot")
        return table

    def create_executive_highlights(
        self, highlights: list[str], title: str = "KEY TAKEAWAYS"
    ) -> Table:
        """
        Create a styled callout box for key takeaways.

        Args:
            highlights: List of highlight strings (bullet points)
            title: Title for the callout box

        Returns:
            Styled Table object (single-cell table with background)
        """
        # Create single-cell table for callout effect
        table = self.document.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # Add title
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = self.TABLE_HEADER_BG

        # Add highlights as bullet points
        for highlight in highlights:
            para = cell.add_paragraph()
            para.add_run("• " + highlight)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.left_indent = Inches(0.15)

        self.apply_table_style(table, "callout")
        return table

    def create_key_metrics_table(
        self, metrics: dict[str, str], title: str = "KEY METRICS"
    ) -> Table:
        """
        Create a key metrics summary table with professional styling.

        Args:
            metrics: Dict of metric_name -> value
            title: Title for the table

        Returns:
            Styled Table object
        """
        if not metrics:
            return None

        # Create table with header row
        table = self.document.add_table(rows=len(metrics) + 1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_row = table.rows[0]
        header_row.cells[0].paragraphs[0].add_run("Metric").bold = True
        header_row.cells[1].paragraphs[0].add_run("Value").bold = True

        # Data rows
        for i, (metric, value) in enumerate(metrics.items(), 1):
            row = table.rows[i]
            row.cells[0].paragraphs[0].add_run(metric.replace("_", " ").title())
            row.cells[1].paragraphs[0].add_run(str(value))

        self.apply_table_style(table, "professional")
        return table

    def apply_table_style(self, table: Table, style: str = "professional") -> None:
        """
        Apply professional styling to table.

        Args:
            table: Table object to style
            style: Style variant - 'professional', 'snapshot', 'callout'
        """
        if style == "professional":
            self._apply_professional_style(table)
        elif style == "snapshot":
            self._apply_snapshot_style(table)
        elif style == "callout":
            self._apply_callout_style(table)

    def _apply_professional_style(self, table: Table) -> None:
        """Apply professional style: blue header, alternating rows, subtle borders."""
        # Set column widths
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(2.5)

        # Style header row (first row)
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                self._set_cell_shading(cell, self.TABLE_HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = self.WHITE
                        run.font.bold = True

        # Alternating row colors for data rows
        for i, row in enumerate(table.rows[1:], 1):
            if i % 2 == 0:
                for cell in row.cells:
                    self._set_cell_shading(cell, self.TABLE_ALT_ROW)

    def _apply_snapshot_style(self, table: Table) -> None:
        """Apply snapshot style: clean key-value pairs, no header."""
        # Set column widths
        if table.columns:
            table.columns[0].width = Inches(1.5)  # Label column
            table.columns[1].width = Inches(3.0)  # Value column

        # Alternating row colors
        for i, row in enumerate(table.rows):
            if i % 2 == 1:
                for cell in row.cells:
                    self._set_cell_shading(cell, self.TABLE_ALT_ROW)

    def _apply_callout_style(self, table: Table) -> None:
        """Apply callout style: shaded background, no borders."""
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_shading(cell, self.CALLOUT_BG)
                cell.width = Inches(6.0)

    def _set_cell_shading(self, cell: Any, color: RGBColor) -> None:
        """Set background shading for a cell."""
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color[0]:02X}{color[1]:02X}{color[2]:02X}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def create_from_markdown(self, headers: list[str], rows: list[list[str]]) -> Table:
        """
        Create a table from parsed markdown table data.

        Used to render tables from Deep Research output which often
        includes markdown tables for competitive analysis, comparisons, etc.

        Args:
            headers: List of column header strings
            rows: List of row data (each row is a list of cell strings)

        Returns:
            Styled Table object
        """
        if not headers and not rows:
            return None

        # Determine column count
        num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
        if num_cols == 0:
            return None

        # Create table with header row + data rows
        num_rows = 1 + len(rows) if headers else len(rows)
        table = self.document.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        row_idx = 0

        # Add header row
        if headers:
            header_row = table.rows[0]
            for col_idx, header_text in enumerate(headers[:num_cols]):
                cell = header_row.cells[col_idx]
                para = cell.paragraphs[0]
                run = para.add_run(header_text)
                run.bold = True
                run.font.size = Pt(10)
            row_idx = 1

        # Add data rows
        for data_row in rows:
            if row_idx >= num_rows:
                break
            table_row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(data_row[:num_cols]):
                cell = table_row.cells[col_idx]
                para = cell.paragraphs[0]
                run = para.add_run(cell_text)
                run.font.size = Pt(10)
            row_idx += 1

        self.apply_table_style(table, "professional")
        return table
