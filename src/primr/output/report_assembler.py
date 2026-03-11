"""
Report Assembler for consulting-tier reports.

Combines sections into a complete report with sources appendix
and export capabilities.
"""

import contextlib
import os
import shutil
import tempfile
from datetime import datetime
from pathlib import Path

from primr.core.report_models import Insight, Report, ReportMetadata, SectionContent, SourceCitation
from primr.utils.logging_config import get_logger

logger = get_logger("report_assembler")


class ReportAssembler:
    """Assembles report sections into a complete document."""

    def assemble(
        self,
        company_name: str,
        website: str,
        industry: str,
        executive_summary: SectionContent,
        sections: list[SectionContent],
        insights: list[Insight],
        research_duration: float,
    ) -> Report:
        """
        Assemble sections into a complete report.

        Args:
            company_name: Name of the company
            website: Company website
            industry: Company industry
            executive_summary: The executive summary section
            sections: List of report sections
            insights: List of insights
            research_duration: Time taken for research in seconds

        Returns:
            Complete Report object
        """
        # Collect all sources from all sections
        all_sources = self._collect_sources(executive_summary, sections)

        metadata = ReportMetadata(
            company_name=company_name,
            website=website,
            industry=industry,
            generated_at=datetime.now(),
            research_duration_seconds=research_duration,
            sources_count=len(all_sources),
        )

        return Report(
            metadata=metadata,
            executive_summary=executive_summary,
            sections=sections,
            sources_appendix=all_sources,
            insights=insights,
        )

    def _collect_sources(
        self, executive_summary: SectionContent, sections: list[SectionContent]
    ) -> list[SourceCitation]:
        """Collect all unique sources from sections."""
        sources = []
        seen_urls = set()

        # Add sources from executive summary
        for source in executive_summary.sources:
            if source.url not in seen_urls:
                sources.append(source)
                seen_urls.add(source.url)

        # Add sources from other sections
        for section in sections:
            for source in section.sources:
                if source.url not in seen_urls:
                    sources.append(source)
                    seen_urls.add(source.url)

        return sources

    def generate_toc(self, sections: list[SectionContent]) -> str:
        """
        Generate a table of contents.

        Args:
            sections: List of sections

        Returns:
            Table of contents as string
        """
        toc_lines = ["Table of Contents", ""]
        toc_lines.append("Executive Summary")

        for section in sections:
            toc_lines.append(section.title)

        toc_lines.append("Sources")

        return "\n".join(toc_lines)

    def generate_sources_appendix(self, sources: list[SourceCitation]) -> str:
        """
        Generate a sources appendix.

        Args:
            sources: List of source citations

        Returns:
            Formatted sources appendix
        """
        if not sources:
            return "No sources cited."

        lines = ["Sources", ""]

        for _i, source in enumerate(sources, 1):
            accessed = source.accessed_at.strftime("%Y-%m-%d")
            source_type = source.source_type.value.replace("_", " ").title()

            line = f"{source.title}"
            line += f"\n  URL: {source.url}"
            line += f"\n  Type: {source_type}"
            line += f"\n  Accessed: {accessed}"

            if source.excerpt:
                excerpt = (
                    source.excerpt[:150] + "..." if len(source.excerpt) > 150 else source.excerpt
                )
                line += f"\n  Excerpt: {excerpt}"

            lines.append(line)
            lines.append("")

        return "\n".join(lines)

    def to_markdown(self, report: Report) -> str:
        """
        Convert report to Markdown format.

        Args:
            report: The report to convert

        Returns:
            Markdown string
        """
        lines = []

        # Title
        lines.append(f"# {report.metadata.company_name} Company Research Report")
        lines.append("")
        lines.append(f"Generated: {report.metadata.generated_at.strftime('%B %d, %Y')}")
        lines.append(f"Industry: {report.metadata.industry}")
        lines.append(f"Website: {report.metadata.website}")
        lines.append("")

        # Table of Contents
        lines.append("---")
        lines.append("")
        lines.append(self.generate_toc(report.sections))
        lines.append("")
        lines.append("---")
        lines.append("")

        # Executive Summary
        lines.append(f"## {report.executive_summary.title}")
        lines.append("")
        lines.append(report.executive_summary.content)
        lines.append("")

        # Sections
        for section in report.sections:
            lines.append(f"## {section.title}")
            lines.append("")
            lines.append(section.content)
            lines.append("")

            # Add confidence notes if any
            if section.confidence_notes:
                lines.append("*Confidence Notes:*")
                for note in section.confidence_notes:
                    lines.append(f"- {note.statement}: {note.confidence.value} ({note.basis})")
                lines.append("")

        # Sources Appendix
        lines.append("---")
        lines.append("")
        lines.append("## Sources")
        lines.append("")
        lines.append(self.generate_sources_appendix(report.sources_appendix))

        return "\n".join(lines)

    def export_docx(self, report: Report, path: str) -> bool:
        """
        Export report to DOCX format.

        Args:
            report: The report to export
            path: Output file path

        Returns:
            True if successful, False otherwise
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt  # noqa: F401
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return False

        try:
            doc = Document()

            # Title
            title = doc.add_heading(f"{report.metadata.company_name} Company Research Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            meta = doc.add_paragraph()
            meta.add_run(f"Generated: {report.metadata.generated_at.strftime('%B %d, %Y')}\n")
            meta.add_run(f"Industry: {report.metadata.industry}\n")
            meta.add_run(f"Website: {report.metadata.website}")

            doc.add_page_break()

            # Table of Contents
            doc.add_heading("Table of Contents", level=1)
            toc = doc.add_paragraph()
            toc.add_run("Executive Summary\n")
            for section in report.sections:
                toc.add_run(f"{section.title}\n")
            toc.add_run("Sources")

            doc.add_page_break()

            # Executive Summary
            doc.add_heading(report.executive_summary.title, level=1)
            doc.add_paragraph(report.executive_summary.content)

            # Sections
            for section in report.sections:
                doc.add_heading(section.title, level=1)
                doc.add_paragraph(section.content)

            # Sources
            doc.add_page_break()
            doc.add_heading("Sources", level=1)

            for source in report.sources_appendix:
                p = doc.add_paragraph()
                p.add_run(f"{source.title}\n").bold = True
                p.add_run(f"URL: {source.url}\n")
                p.add_run(f"Type: {source.source_type.value.replace('_', ' ').title()}\n")
                p.add_run(f"Accessed: {source.accessed_at.strftime('%Y-%m-%d')}")

            doc.save(path)
            logger.info(f"Report exported to {path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export DOCX: {e}")
            return False

    def export_pdf(self, report: Report, path: str) -> bool:
        """
        Export report to PDF format.

        Args:
            report: The report to export
            path: Output file path

        Returns:
            True if successful, False otherwise
        """
        # First export to DOCX, then convert
        # This is a simplified approach - could use reportlab for direct PDF
        try:
            import subprocess

            output_pdf = Path(path)
            output_dir = output_pdf.parent if output_pdf.parent != Path("") else Path(".")
            output_dir.mkdir(parents=True, exist_ok=True)

            # Create temp DOCX in target directory so soffice can emit a sibling PDF.
            fd, temp_docx_str = tempfile.mkstemp(
                suffix=".docx", prefix="primr_pdf_", dir=str(output_dir)
            )
            os.close(fd)
            temp_docx = Path(temp_docx_str)

            if not self.export_docx(report, str(temp_docx)):
                with contextlib.suppress(OSError):
                    temp_docx.unlink(missing_ok=True)
                return False

            # Try to convert using LibreOffice (if available)
            try:
                subprocess.run(
                    [
                        "soffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(output_dir),
                        str(temp_docx),
                    ],
                    check=True,
                    capture_output=True,
                )
                generated_pdf = output_dir / f"{temp_docx.stem}.pdf"
                if generated_pdf.exists():
                    if generated_pdf.resolve() != output_pdf.resolve():
                        output_pdf.parent.mkdir(parents=True, exist_ok=True)
                        generated_pdf.replace(output_pdf)
                    temp_docx.unlink(missing_ok=True)
                    return True
                logger.warning("LibreOffice conversion reported success but produced no PDF")
                return False

            except (subprocess.CalledProcessError, FileNotFoundError):
                logger.warning("LibreOffice not available for PDF conversion")
                fallback_docx = output_pdf.with_suffix(".docx")
                try:
                    if fallback_docx.exists():
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        fallback_docx = output_pdf.with_name(f"{output_pdf.stem}_{timestamp}.docx")
                    shutil.move(str(temp_docx), str(fallback_docx))
                except Exception:
                    logger.debug("Failed to preserve DOCX fallback", exc_info=True)
                    with contextlib.suppress(OSError):
                        temp_docx.unlink(missing_ok=True)
                return False

        except Exception as e:
            logger.error(f"Failed to export PDF: {e}")
            return False

    def validate_report(self, report: Report) -> tuple[bool, list[str]]:
        """
        Validate a report for completeness.

        Args:
            report: The report to validate

        Returns:
            Tuple of (is_valid, list of issues)
        """
        issues = []

        # Check metadata
        if not report.metadata.company_name:
            issues.append("Missing company name")

        # Check executive summary
        if not report.executive_summary.content:
            issues.append("Empty executive summary")

        # Check sections
        if not report.sections:
            issues.append("No sections in report")

        for section in report.sections:
            if not section.content:
                issues.append(f"Empty section: {section.title}")

        # Check sources
        if not report.sources_appendix:
            issues.append("No sources cited")

        return len(issues) == 0, issues
