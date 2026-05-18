"""
Professional polish elements for premium reports.

Provides additional formatting and content enhancement components
for consultant-grade document output.
"""

import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


@dataclass
class OneLinerSummary:
    """
    One-sentence company summary for cover page.

    Template: "{company} is a {industry} company that {differentiator},
    generating {revenue}"

    Must pass the "dinner test" - explainable in 2 sentences.
    """

    company_name: str
    industry: str = ""
    differentiator: str = ""
    revenue: str = ""

    def generate(self) -> str:
        """Generate the one-liner summary."""
        parts = [f"{self.company_name}"]

        if self.industry:
            parts.append(f"is a {self.industry} company")

        if self.differentiator:
            parts.append(f"that {self.differentiator}")

        if self.revenue:
            parts.append(f"generating {self.revenue}")

        if len(parts) == 1:
            return f"{self.company_name} is a company under research."

        return " ".join(parts) + "."


@dataclass
class DocumentDisclaimer:
    """
    AI-generated content disclaimer and data currency notice.
    """

    generation_date: str = ""

    def __post_init__(self):
        if not self.generation_date:
            self.generation_date = datetime.now().strftime("%B %d, %Y")

    def get_disclaimer_text(self) -> str:
        """Get the full disclaimer text."""
        return (
            f"This report was generated using AI-assisted research on {self.generation_date}. "
            "Information should be verified independently before making business decisions. "
            "Data reflects publicly available information at the time of generation."
        )

    def add_to_document(self, document: Document, style_engine: Any) -> None:
        """Add disclaimer to document."""
        para = document.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = para.add_run(self.get_disclaimer_text())
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = style_engine.MUTED_COLOR


class KeyImplicationsBox:
    """
    Styled callout box for key implications at chapter endings.

    Extracts 2-3 key implications from chapter content.
    """

    def __init__(self, document: Document, style_engine: Any):
        self.document = document
        self.style_engine = style_engine

    def extract_implications(self, content: str, max_items: int = 3) -> list[str]:
        """
        Extract key implications from content.

        Looks for sentences containing implication keywords.
        """
        implications = []

        # Keywords that suggest implications
        keywords = [
            "suggests",
            "indicates",
            "implies",
            "means",
            "therefore",
            "consequently",
            "as a result",
            "opportunity",
            "risk",
            "should consider",
            "may need",
            "could benefit",
        ]

        sentences = re.split(r"[.!?]+", content)

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence or len(sentence) < 20:
                continue

            lower = sentence.lower()
            if any(kw in lower for kw in keywords):
                implications.append(sentence)
                if len(implications) >= max_items:
                    break

        return implications

    def add_to_document(self, title: str, implications: list[str]) -> None:
        """Add implications box to document."""
        if not implications:
            return

        # Create a table for the callout box
        table = self.document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.rows[0].cells[0]

        # Title
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(f"Key Implications: {title}")
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = self.style_engine.PRIMARY_COLOR

        # Implications as bullets
        for impl in implications:
            para = cell.add_paragraph()
            para.add_run(f"  {impl}")
            para.paragraph_format.space_before = Pt(4)

        # Style the cell
        self._style_callout_cell(cell)

        # Add spacing after
        self.document.add_paragraph()

    def _style_callout_cell(self, cell: Any) -> None:
        """Apply callout box styling to cell."""
        # Light background
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F5F5F5")
        cell._tc.get_or_add_tcPr().append(shading)


class StrategicRecommendationFormatter:
    """
    Formats recommendations with priority indicators.

    Adds High/Medium/Low priority with color coding.
    """

    PRIORITY_COLORS = {
        "high": RGBColor(0xC0, 0x39, 0x2B),  # Red
        "medium": RGBColor(0xF3, 0x9C, 0x12),  # Orange
        "low": RGBColor(0x27, 0xAE, 0x60),  # Green
    }

    PRIORITY_INDICATORS = {
        "high": "[HIGH]",
        "medium": "[MED]",
        "low": "[LOW]",
    }

    def __init__(self, document: Document):
        self.document = document

    def detect_priority(self, text: str) -> str:
        """Detect priority level from text content."""
        lower = text.lower()

        high_keywords = ["critical", "urgent", "immediate", "essential", "must"]
        medium_keywords = ["should", "important", "consider", "recommend"]

        if any(kw in lower for kw in high_keywords):
            return "high"
        elif any(kw in lower for kw in medium_keywords):
            return "medium"
        return "low"

    def format_recommendation(self, number: int, text: str, priority: str | None = None) -> None:
        """Add a formatted recommendation to the document."""
        if priority is None:
            priority = self.detect_priority(text)

        para = self.document.add_paragraph()

        # Number
        num_run = para.add_run(f"{number}. ")
        num_run.bold = True

        # Priority indicator
        indicator = self.PRIORITY_INDICATORS.get(priority, "")
        if indicator:
            priority_run = para.add_run(f"{indicator} ")
            priority_run.font.color.rgb = self.PRIORITY_COLORS.get(priority, RGBColor(0, 0, 0))
            priority_run.bold = True

        # Recommendation text
        para.add_run(text)


class QuickWinsSection:
    """
    Extracts actionable items implementable in 30 days.

    Looks for low-cost, quick implementation, immediate impact items.
    """

    QUICK_WIN_KEYWORDS = [
        "quick",
        "easy",
        "simple",
        "immediate",
        "low-cost",
        "straightforward",
        "readily",
        "quickly",
        "soon",
    ]

    def __init__(self, document: Document, style_engine: Any):
        self.document = document
        self.style_engine = style_engine

    def extract_quick_wins(self, content: str, max_items: int = 5) -> list[str]:
        """Extract quick win items from content."""
        quick_wins = []

        sentences = re.split(r"[.!?]+", content)

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence or len(sentence) < 15:
                continue

            lower = sentence.lower()
            if any(kw in lower for kw in self.QUICK_WIN_KEYWORDS):
                quick_wins.append(sentence)
                if len(quick_wins) >= max_items:
                    break

        return quick_wins

    def add_to_document(self, quick_wins: list[str]) -> None:
        """Add quick wins section as callout box."""
        if not quick_wins:
            return

        # Heading
        self.document.add_heading("Quick Wins (30-Day Actions)", level=3)

        # Create callout table
        table = self.document.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        for i, win in enumerate(quick_wins, 1):
            para = cell.add_paragraph() if i > 1 else cell.paragraphs[0]
            para.add_run(f"{i}. {win}")

        # Style
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8F5E9")  # Light green
        cell._tc.get_or_add_tcPr().append(shading)

        self.document.add_paragraph()


class FinancialDashboard:
    """
    Creates 2x3 grid layout for key financial metrics.

    Adds trend indicators where data supports.
    """

    def __init__(self, document: Document, style_engine: Any):
        self.document = document
        self.style_engine = style_engine

    def create_dashboard(self, metrics: dict[str, str]) -> None:
        """
        Create financial dashboard table.

        Args:
            metrics: Dict of metric_name -> value (e.g., {"Revenue": "$5.2B"})
        """
        if not metrics:
            return

        # Heading
        self.document.add_heading("Financial Dashboard", level=3)

        # Create 2x3 table (or adjust based on metric count)
        metric_items = list(metrics.items())
        rows = (len(metric_items) + 2) // 3  # Ceiling division

        table = self.document.add_table(rows=rows, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (name, value) in enumerate(metric_items):
            row_idx = i // 3
            col_idx = i % 3

            if row_idx < len(table.rows):
                cell = table.rows[row_idx].cells[col_idx]

                # Metric name
                name_para = cell.paragraphs[0]
                name_run = name_para.add_run(name)
                name_run.font.size = Pt(9)
                name_run.font.color.rgb = self.style_engine.MUTED_COLOR

                # Metric value
                value_para = cell.add_paragraph()
                value_run = value_para.add_run(value)
                value_run.bold = True
                value_run.font.size = Pt(14)

        self.document.add_paragraph()


class DataConfidenceIndicator:
    """
    Adds confidence levels to key claims.

    High = multiple sources (3 dots)
    Medium = single source (2 dots)
    Low = inference (1 dot)
    """

    INDICATORS = {
        "high": "●●●",
        "medium": "●●○",
        "low": "●○○",
    }

    COLORS = {
        "high": RGBColor(0x27, 0xAE, 0x60),  # Green
        "medium": RGBColor(0xF3, 0x9C, 0x12),  # Orange
        "low": RGBColor(0xC0, 0x39, 0x2B),  # Red
    }

    def get_indicator(self, level: str) -> str:
        """Get the confidence indicator string."""
        return self.INDICATORS.get(level, self.INDICATORS["medium"])

    def format_with_confidence(self, text: str, level: str) -> str:
        """Format text with confidence indicator."""
        indicator = self.get_indicator(level)
        return f"{text} {indicator}"

    def add_to_paragraph(self, paragraph: Any, text: str, level: str) -> None:
        """Add text with confidence indicator to paragraph."""
        paragraph.add_run(text + " ")

        indicator_run = paragraph.add_run(self.get_indicator(level))
        indicator_run.font.color.rgb = self.COLORS.get(level, self.COLORS["medium"])
        indicator_run.font.size = Pt(10)
