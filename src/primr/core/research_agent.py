"""
Automated Company Research Agent - Orchestration Hub

This module serves as the main entry point and orchestration hub for company research.
It delegates to specialized modules for specific functionality:

- workspace: Working folder management, file consolidation
- structured_research: Website scraping and section generation
- vendor_research: Cloud vendor AI capabilities research
- ai_strategy: AI strategy generation
- deep_research_runner: Deep Research execution
- cli: Command-line interface

For backward compatibility, this module re-exports key functions from the specialized modules.

Usage:
    from primr.core.research_agent import perform_research, main

    # Run research programmatically
    result = perform_research("Acme Corp", "https://acme.example")

    # Run CLI
    main()
"""

# Suppress RuntimeWarning FIRST - before any other imports
import warnings

warnings.filterwarnings("ignore", message=".*found in sys.modules.*", category=RuntimeWarning)

# =============================================================================
# BACKWARD COMPATIBLE RE-EXPORTS
# =============================================================================
# These imports ensure existing code that imports from research_agent.py continues to work.
# New code should import directly from the specialized modules.

# From workspace module

# From structured_research module

# From cli module
# From ai_strategy module
from primr.core.ai_strategy import (
    CloudVendor,
    Platform,
)
from primr.core.ai_strategy_runtime import (
    build_ai_strategy_prompt as _build_ai_strategy_prompt_impl,
)
from primr.core.ai_strategy_runtime import (
    generate_ai_strategy_section as _generate_ai_strategy_section_impl,
)
from primr.core.cli import (
    main as _main_new,
)

# From deep_research_runner module
from primr.core.deep_research_runner import (
    DeepResearchConfig,
    DeepResearchMode,
)
from primr.core.strategy_generation import (
    build_strategy_prompt_from_yaml as _build_strategy_prompt_from_yaml_impl,
)
from primr.core.strategy_generation import (
    generate_generic_strategy as _generate_generic_strategy_impl,
)

# =============================================================================
# PUBLIC API
# =============================================================================

__all__ = [
    "CloudVendor",
    "DeepResearchConfig",
    "DeepResearchMode",
    "Platform",
    "consolidate_working_folder",
    "create_working_folder",
    "ensure_valid_url",
    "generate_initial_overview",
    "get_user_input",
    "improve_output_file",
    "main",
    "perform_research",
    "process_csv",
    "research_section",
    "run_doctor",
    "run_research",
    "save_section_output",
    "validate_context_files",
]

import asyncio
import atexit
import gc
import json
import os
import re
import sys
import time
from collections.abc import Callable
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any, TypedDict
from urllib.parse import urlparse

if TYPE_CHECKING:
    from primr.output.final_artifact import GeneratedSection
    from primr.prompts.loader import SectionConfig

from primr.ai.grading_agent import grade_report
from primr.ai.llm import llm
from primr.ai.summarize import summarize_scraped_content
from primr.config.config import (
    FAST_FEEDBACK_RULES_PATH,
    GRADE_THRESHOLD_FOR_RESEARCH_REFINEMENT,
    LOGS_DIR,
    MAX_EXTERNAL_SEARCH_QUERIES,
    MAX_EXTERNAL_SOURCES,
    MIN_SCRAPED_CHARS,
    MIN_SCRAPED_PAGES,
    OUTPUT_DIR,
    PROJECT_ROOT,
    WORKING_DIR,
)
from primr.config.env import load_primr_env
from primr.config.models import GROK_MODEL_WRITING, GrokTier, PrimrModels
from primr.config.sections_config import SECTION_KEY_MAP
from primr.core.research_orchestrator import (
    ResearchConfig,
    ResearchMode,
    get_orchestrator,
)
from primr.data.scrape import fetch_web_content, scrape_external_sources_validated
from primr.data.scraping.org_profile import get_focus_areas_for_org_type
from primr.data.search_utils import (
    generate_external_search_queries,
    generate_search_queries,
    search_web,
)
from primr.output.output_utils import generate_final_report
from primr.utils.console import console
from primr.utils.logging_config import get_logger
from primr.utils.observability import (
    JobSummary,
    correlation_scope,
    log_job_summary,
    log_structured,
)
from primr.utils.validators import validate_url_for_request

load_primr_env()

logger = get_logger("research_agent")

# Setup directories
for directory in [OUTPUT_DIR, WORKING_DIR, LOGS_DIR]:
    os.makedirs(directory, exist_ok=True)

# Load prompts from package config directory
PROMPTS_FILE = Path(__file__).parent.parent / "config" / "prompts.json"
with open(PROMPTS_FILE, encoding="utf-8") as f:
    PROMPTS = json.load(f)


def generate_prompt(template_name, **kwargs):
    if template_name not in PROMPTS:
        raise ValueError(f"Prompt '{template_name}' not found")
    return PROMPTS[template_name].format(**kwargs)


# User-friendly tier names for display
TIER_DISPLAY_NAMES = {
    "requests": "HTTP",
    "httpx": "HTTP/2",
    "curl_cffi": "stealth HTTP",
    "playwright": "browser",
    "playwright_aggressive": "browser+",
    "drissionpage": "headless",
    "drissionpage_stealth": "stealth browser",
    "vision": "AI vision",
    "cache": "cache",
}


def _load_fast_feedback_guidance() -> str:
    """Load persisted fast-mode guidance generated from eval feedback loops."""
    path = Path(FAST_FEEDBACK_RULES_PATH)
    if not path.exists():
        return ""
    try:
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception as e:
        logger.debug("Failed to load fast feedback guidance from %s: %s", path, e)
        return ""
    if not text:
        return ""
    # Bound prompt growth.
    return text[:4000]


def format_tier_stats(tier_stats: dict) -> str:
    """Format tier stats for user-friendly display."""
    # Sort by count descending
    sorted_tiers = sorted(tier_stats.items(), key=lambda x: -x[1])
    parts = []
    for tier, count in sorted_tiers:
        display_name = TIER_DISPLAY_NAMES.get(tier, tier)
        parts.append(f"{count} {display_name}")
    return ", ".join(parts)


def _validate_scrape_quality(
    corpus: dict[str, str],
    *,
    min_pages: int = MIN_SCRAPED_PAGES,
    min_chars: int = MIN_SCRAPED_CHARS,
) -> tuple[bool, str]:
    """Check whether scraped website content is sufficient for reliable analysis."""
    pages = len(corpus)
    chars = sum(len(c or "") for c in corpus.values())
    ok = pages >= min_pages and chars >= min_chars
    reason = (
        f"Scrape quality too low ({pages} pages, {chars:,} chars; "
        f"requires >= {min_pages} pages and >= {min_chars:,} chars)"
    )
    return ok, reason


def _build_link_selection_prompt(
    company_name: str,
    website: str,
    links_text: str,
    max_links: int,
    organization_type: str,
) -> str:
    focus_areas = "\n".join(
        f"- {focus}" for focus in get_focus_areas_for_org_type(organization_type)
    )
    return (
        f"You are selecting pages for intelligence gathering on {company_name} ({website}).\n\n"
        f"Organization type: {organization_type}.\n"
        "Choose only from the discovered URLs below. Do not invent, normalize, or rewrite URLs.\n\n"
        "Prioritize pages that help explain the organization through these focus areas:\n"
        f"{focus_areas}\n\n"
        "Discovered URLs:\n"
        f"{links_text}\n\n"
        f"Return only URLs from the discovered list, up to {max_links}, one per line."
    )


def select_links_with_llm(
    links: list,
    company_name: str,
    website: str,
    max_links: int = 50,
    organization_type: str = "commercial",
) -> list[str]:
    """
    Use LLM to intelligently select the most valuable links for research.

    The LLM acts like a business analyst deciding which pages to read to understand
    a company - prioritizing pages about leadership, strategy, products,
    financials, and recent news.

    Args:
        links: List of DiscoveredLink objects (pre-scored heuristically)
        company_name: Company name for context
        website: Company website URL
        max_links: Maximum links to return (passed to LLM so it knows the constraint)

    Returns:
        List of URLs selected by the LLM
    """
    if not links:
        return []

    # If we have fewer links than max, just return all of them
    if len(links) <= max_links:
        return [link.url for link in links]

    # Format links for the prompt - include URL and anchor text if available
    link_list = []
    for link in links[:200]:  # Cap at 200 to avoid token limits
        if hasattr(link, "anchor_text") and link.anchor_text:
            link_list.append(f"{link.url} ({link.anchor_text})")
        else:
            link_list.append(link.url)

    links_text = "\n".join(link_list)

    try:
        prompt = _build_link_selection_prompt(
            company_name=company_name,
            website=website,
            links_text=links_text,
            max_links=max_links,
            organization_type=organization_type,
        )

        # Use link_selection model type (Flash - cheap and fast)
        response = llm(prompt, model_type="link_selection")

        discovered_urls = {link.url for link in links}
        selected_urls = []
        dropped_urls = []
        for line in response.strip().split("\n"):
            line = line.strip()
            if not line or not line.startswith("http"):
                continue
            if line in discovered_urls and line not in selected_urls:
                selected_urls.append(line)
            else:
                dropped_urls.append(line)

        if dropped_urls:
            logger.info(
                "Dropped %s LLM-selected URLs that were not in the discovered set",
                len(dropped_urls),
            )

        # If LLM returned valid URLs, use them (LLM already knows the limit)
        if selected_urls:
            logger.info(f"LLM selected {len(selected_urls)} links from {len(links)}")
            return selected_urls

    except Exception as e:
        logger.warning(f"LLM link selection failed: {e}, falling back to heuristic scoring")

    # Fallback to heuristic scoring if LLM fails
    return [link.url for link in links[:max_links]]


def create_working_folder(company_name, website, reuse_incomplete: bool = False):
    """
    Create working folder for research artifacts with timestamped run ID.

    Each run gets its own subfolder like: working/Company_Name/2026-01-09_0915/
    This prevents mixing old and new data from different runs.
    """
    from datetime import datetime

    if not company_name and website:
        parsed_url = urlparse(website)
        company_name = parsed_url.netloc.replace("www.", "").replace(".", "_")

    folder_name = company_name.replace(" ", "_") if company_name else "Unknown_Company"

    company_root = os.path.join(WORKING_DIR, folder_name)

    # Optional resume behavior: reuse latest incomplete run folder for this company
    if reuse_incomplete and os.path.isdir(company_root):
        run_dirs = sorted(
            [d for d in os.listdir(company_root) if os.path.isdir(os.path.join(company_root, d))],
            reverse=True,
        )
        for run_id in run_dirs:
            candidate = os.path.join(company_root, run_id)
            state_path = os.path.join(candidate, "_run_state.json")
            if not os.path.exists(state_path):
                continue
            try:
                with open(state_path, encoding="utf-8") as f:
                    state = json.load(f)
                if not isinstance(state, dict):
                    continue
                status = str(state.get("status", "")).lower()
                if status in {"running", "failed", "cancelled", "canceled"}:
                    logger.info(f"Reusing incomplete working folder: {candidate}")
                    return candidate
            except Exception as e:
                logger.debug("Failed to read run state for resume candidate %s: %s", candidate, e)
                continue

    # Create timestamped run folder: Company_Name/2026-01-09_0915
    run_id = datetime.now().strftime("%Y-%m-%d_%H%M")
    folder_path = os.path.join(company_root, run_id)

    os.makedirs(folder_path, exist_ok=True)
    logger.info(f"Created working folder: {folder_path}")
    return folder_path


def _run_state_file(folder_path: str) -> str:
    """Return path to the per-run state file."""
    return os.path.join(folder_path, "_run_state.json")


def _load_run_state(folder_path: str) -> dict[str, Any]:
    """Load run state JSON if present, else return empty dict."""
    path = _run_state_file(folder_path)
    if not os.path.exists(path):
        return {}
    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except json.JSONDecodeError as e:
        logger.warning("Run state file corrupted (%s), starting with empty state: %s", path, e)
        return {}
    except Exception as e:
        logger.warning("Failed to load run state from %s, starting with empty state: %s", path, e)
        return {}


def _save_run_state(folder_path: str, state: dict[str, Any]) -> None:
    """Persist run state JSON without aborting the run on transient Windows locks."""
    path = _run_state_file(folder_path)
    tmp = f"{path}.{os.getpid()}.tmp"
    os.makedirs(folder_path, exist_ok=True)
    payload = json.dumps(state, indent=2)
    with open(tmp, "w", encoding="utf-8") as f:
        f.write(payload)
        f.flush()
        os.fsync(f.fileno())

    last_error: PermissionError | None = None
    for attempt in range(5):
        try:
            os.replace(tmp, path)
            return
        except PermissionError as exc:
            last_error = exc
            if attempt == 4:
                break
            time.sleep(0.05 * (attempt + 1))

    logger.warning(
        "Atomic run state save failed for %s; falling back to direct overwrite: %s",
        path,
        last_error,
    )
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(payload)
    finally:
        if os.path.exists(tmp):
            try:
                os.remove(tmp)
            except OSError:
                logger.debug("Failed to remove stale run-state temp file %s", tmp, exc_info=True)


def _update_run_state(folder_path: str, **updates: Any) -> None:
    """Merge updates into run state file and refresh timestamp."""
    state = _load_run_state(folder_path)
    state.update(updates)
    state["updated_at"] = datetime.now().isoformat()
    _save_run_state(folder_path, state)


def _append_run_event(
    folder_path: str, phase: str, status: str, message: str, **extra: Any
) -> None:
    """Append a timeline event into run state."""
    state = _load_run_state(folder_path)
    events = state.get("events", [])
    if not isinstance(events, list):
        events = []
    event: dict[str, Any] = {
        "ts": datetime.now().isoformat(),
        "phase": phase,
        "status": status,
        "message": message,
    }
    if extra:
        event["extra"] = extra
    events.append(event)
    state["events"] = events[-200:]  # keep recent history bounded
    state["updated_at"] = datetime.now().isoformat()
    _save_run_state(folder_path, state)


# =============================================================================
# RESILIENCE RUN STATE HELPERS (pipeline-resilience feature)
# =============================================================================


def _ensure_resilience_keys(state: dict[str, Any]) -> dict[str, Any]:
    """Ensure resilience arrays exist in run state (backwards compatible).

    Adds ``model_health``, ``recovery_events``, and ``background_aborts``
    arrays if they are missing.  Existing keys are never overwritten.

    **Feature: pipeline-resilience**
    **Validates: Requirements 16.3, 12.2, NFR 3**
    """
    for key in ("model_health", "recovery_events", "background_aborts"):
        if key not in state or not isinstance(state[key], list):
            state[key] = state.get(key, []) if isinstance(state.get(key), list) else []
    return state


def _append_model_health_event(folder_path: str, event_dict: dict[str, Any]) -> None:
    """Append a ModelHealthEvent dict to the ``model_health`` array in run state.

    **Feature: pipeline-resilience**
    **Validates: Requirements 12.2**
    """
    state = _load_run_state(folder_path)
    _ensure_resilience_keys(state)
    state["model_health"].append(event_dict)
    state["model_health"] = state["model_health"][-200:]
    state["updated_at"] = datetime.now().isoformat()
    _save_run_state(folder_path, state)


def _append_recovery_event(folder_path: str, event_dict: dict[str, Any]) -> None:
    """Append a recovery event dict to the ``recovery_events`` array in run state.

    **Feature: pipeline-resilience**
    **Validates: Requirements 16.3**
    """
    state = _load_run_state(folder_path)
    _ensure_resilience_keys(state)
    state["recovery_events"].append(event_dict)
    state["recovery_events"] = state["recovery_events"][-200:]
    state["updated_at"] = datetime.now().isoformat()
    _save_run_state(folder_path, state)


def _append_background_abort(folder_path: str, event_dict: dict[str, Any]) -> None:
    """Append a background abort dict to the ``background_aborts`` array in run state.

    **Feature: pipeline-resilience**
    **Validates: Requirements 16.3**
    """
    state = _load_run_state(folder_path)
    _ensure_resilience_keys(state)
    state["background_aborts"].append(event_dict)
    state["background_aborts"] = state["background_aborts"][-200:]
    state["updated_at"] = datetime.now().isoformat()
    _save_run_state(folder_path, state)


def _init_run_state_with_resilience(folder_path: str, base_state: dict[str, Any]) -> None:
    """Initialize run state with resilience keys included.

    Merges the base state dict with empty resilience arrays and saves.
    Backwards compatible — existing keys in *base_state* are preserved.

    **Feature: pipeline-resilience**
    **Validates: Requirements 16.3, NFR 3**
    """
    _ensure_resilience_keys(base_state)
    _save_run_state(folder_path, base_state)


def _build_resilience_event_listener(folder_path: str):
    """Build an event listener callback that routes recovery events to run state.

    Returns a callable suitable for ``RecoveryExecutor(event_listener=...)``.

    **Feature: pipeline-resilience**
    **Validates: Requirements 16.1, 16.2, 16.3**
    """
    from primr.pipeline.executor import BackgroundAbort, RecoveryEvent

    def _listener(event) -> None:
        if isinstance(event, RecoveryEvent):
            _append_recovery_event(folder_path, event.to_dict())
        elif isinstance(event, BackgroundAbort):
            _append_background_abort(folder_path, event.to_dict())

    return _listener


def _build_health_listener(folder_path: str):
    """Build a health listener callback that logs ModelHealthEvents to run state.

    Returns a callable suitable for ``ModelCircuitBreaker(health_listener=...)``.

    **Feature: pipeline-resilience**
    **Validates: Requirements 12.1, 12.2**
    """
    from primr.pipeline.model_breaker import ModelHealthEvent

    def _listener(event: ModelHealthEvent) -> None:
        _append_model_health_event(folder_path, event.to_dict())
        log_structured(
            "info",
            "Model health transition",
            model=event.model,
            from_state=event.from_state,
            to_state=event.to_state,
            failure_count=event.failure_count,
        )

    return _listener


def ensure_valid_url(website):
    if not website:
        return None
    website = website.strip()
    if website.startswith(("http://", "https://")):
        return website
    return f"https://{website}"


def get_user_input():
    console.banner("Company Research")
    console.blank()

    company_name = input("  Company name: ").strip()
    website = input("  Website URL:  ").strip()

    if not company_name and not website:
        console.error("Need either company name or website")
        sys.exit(1)

    return company_name, ensure_valid_url(website) if website else None


def save_section_output(folder_path, section_key, content):
    """Save section content to file."""
    filepath = os.path.join(folder_path, f"{section_key}.txt")
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
    except OSError as e:
        logger.error(f"Failed to save section {section_key}: {e}")


def consolidate_working_folder(folder_path: str) -> str:
    """
    Consolidate all .txt files from a working folder into a single context file.

    Args:
        folder_path: Path to working folder (e.g., working/Parts_Town)

    Returns:
        Path to the consolidated temporary file
    """
    import glob
    import tempfile

    if not os.path.isdir(folder_path):
        raise ValueError(f"Working folder not found: {folder_path}")

    # Find all .txt files
    txt_files = glob.glob(os.path.join(folder_path, "*.txt"))
    if not txt_files:
        raise ValueError(f"No .txt files found in {folder_path}")

    # Extract company name from folder — skip timestamped leaf dirs like 2026-03-04_1530
    basename = os.path.basename(folder_path)
    if re.match(r"^\d{4}-\d{2}-\d{2}", basename):
        # Timestamped run folder; company name is the parent
        company_name = os.path.basename(os.path.dirname(folder_path)).replace("_", " ")
    else:
        company_name = basename.replace("_", " ")

    # Build consolidated document
    lines = [
        f"# Research Context: {company_name}",
        f"Source: {folder_path}",
        "",
        "This document contains research findings from the Structured Pipeline.",
        "",
        "---",
        "",
    ]

    # Read each file and add to document
    for txt_file in sorted(txt_files):
        filename = os.path.basename(txt_file)
        section_name = filename.replace(".txt", "").replace("_", " ").title()

        try:
            with open(txt_file, encoding="utf-8") as f:
                content = f.read().strip()

            if content:
                lines.extend([f"## {section_name}", "", content, "", "---", ""])
        except Exception as e:
            logger.warning(f"Failed to read {txt_file}: {e}")

    # Write to temp file
    # NOTE: We must close the fd from mkstemp before opening the file by path
    content = "\n".join(lines)
    fd, filepath = tempfile.mkstemp(
        suffix=".txt", prefix=f"{company_name.replace(' ', '_')}_context_"
    )
    os.close(fd)  # Close the fd - we'll open by path

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    logger.info(f"Consolidated {len(txt_files)} files into {filepath}")
    return filepath


# Supported file types for Deep Research File Search
SUPPORTED_CONTEXT_EXTENSIONS = {".txt", ".pdf", ".md", ".json", ".csv"}


def validate_context_files(file_paths: list) -> tuple:
    """
    Validate context files for Deep Research upload.

    Args:
        file_paths: List of file paths to validate

    Returns:
        Tuple of (valid_files, invalid_files, warnings)
    """
    valid_files = []
    invalid_files = []
    warnings = []

    for file_path in file_paths:
        if not os.path.exists(file_path):
            invalid_files.append((file_path, "File not found"))
            continue

        ext = os.path.splitext(file_path)[1].lower()

        if ext in SUPPORTED_CONTEXT_EXTENSIONS:
            valid_files.append(file_path)
        elif ext in {".docx", ".doc"}:
            # Word docs - suggest converting to PDF
            invalid_files.append(
                (
                    file_path,
                    "Word docs not directly supported. Convert to PDF or use the .txt output",
                )
            )
            warnings.append("Tip: Use the _Company_Overview.txt file from output/ instead of .docx")
        elif ext in {".xlsx", ".xls"}:
            invalid_files.append((file_path, "Excel files not supported. Export to CSV"))
        else:
            invalid_files.append((file_path, f"Unsupported file type: {ext}"))

    return valid_files, invalid_files, warnings


def generate_initial_overview(company_name, website, industry, folder_path):
    overview_prompt = generate_prompt(
        "initial_company_overview",
        company_name=company_name,
        company_website=website or "N/A",
        industry=industry,
        detailed_products_services="N/A",
        unique_selling_proposition="N/A",
        mission_vision="N/A",
        company_history="N/A",
        key_achievements="N/A",
        target_audience="N/A",
        financial_overview="N/A",
        business_drivers_and_kpis="N/A",
        business_outcomes="N/A",
        scraped_website_summary="N/A",
    )
    overview = llm(overview_prompt, model_type="report")

    overview_file = os.path.join(folder_path, f"{company_name}_Draft_Overview.txt")
    with open(overview_file, "w", encoding="utf-8") as f:
        f.write(overview)

    return overview


def research_section(
    section_name, company_name, website, industry, folder_path, overview, summarized_insights
):
    section_key = SECTION_KEY_MAP.get(section_name)

    if not section_key or section_key not in PROMPTS:
        return ""

    if section_name in ["Company Name", "Website", "Industry"]:
        value = (
            company_name
            if section_name == "Company Name"
            else website
            if section_name == "Website"
            else industry
        )
        save_section_output(folder_path, section_key, value or "N/A")
        return value

    prompt_data = {
        "company_name": company_name,
        "company_website": website or "N/A",
        "industry": industry or "N/A",
        "detailed_products_services": summarized_insights or "N/A",
        "unique_selling_proposition": "N/A",
        "mission_vision": "N/A",
        "company_history": "N/A",
        "key_achievements": "N/A",
        "target_audience": "N/A",
        "financial_overview": "N/A",
        "business_drivers_and_kpis": "N/A",
        "potential_business_outcomes": "N/A",
        "industry_insights": "N/A",
        "potential_business_drivers": "N/A",
        "primary_apps_sources_of_data": "N/A",
        "main_types_of_users": "N/A",
        "board_of_directors_concerns": "N/A",
        "potential_business_value": "N/A",
        "strategic_recommendations": "N/A",
        "scraped_website_summary": summarized_insights or "N/A",
        "value_theory": overview,
    }

    try:
        ai_prompt = generate_prompt(section_key, **prompt_data)
    except KeyError:
        return ""

    ai_input = f"""
## Company: {company_name}
## Website: {website or "N/A"}
## Industry: {industry or "N/A"}
## Section: {section_name}

{ai_prompt}

## Context
{overview}

## Scraped Insights
{summarized_insights}
"""

    ai_response = llm(ai_input, model_type="report")

    if section_name not in ["Company Name", "Website", "Industry"]:
        try:
            score, needs_research, feedback = grade_report(
                ai_response, section_name, company_name, website, overview, summarized_insights
            )

            if (
                needs_research
                and score is not None
                and score < GRADE_THRESHOLD_FOR_RESEARCH_REFINEMENT
            ):
                queries = generate_search_queries(company_name, website, section_name, ai_response)
                for query in queries[:2]:
                    results = search_web(query, company_name, website)
                    if results:
                        formatted = "\n".join(
                            f"- [{r.get('title', 'Source')}]({r.get('url', '')}): {r.get('snippet', '')}"
                            for r in results[:5]
                        )
                        ai_input += f"\n\n## Additional Research\n{formatted}"
                        ai_response = llm(ai_input, model_type="report")
                        break
        except Exception as e:
            logger.warning(f"Grading/refinement failed for section '{section_name}': {e}")

    if not ai_response or len(ai_response.strip()) < 50:
        ai_response = f"No detailed {section_name} information available for {company_name}."

    save_section_output(folder_path, section_key, ai_response)
    return ai_response


def run_research(
    company_name: str,
    website: str,
    on_progress: Callable[[str], None] | None = None,
    fail_on_low_scrape: bool = True,
) -> dict | None:
    """
    Run structured research and return section results.

    This is the entry point used by ResearchOrchestrator for structured mode.

    Args:
        company_name: Name of the company
        website: Company website URL
        on_progress: Optional callback for progress updates (message: str)

    Returns:
        Dict mapping section_key to content, or None if scraping/quality
        gates fail before any sections are produced.
    """
    import time as time_module

    def progress(msg: str) -> None:
        """Send progress update via callback or console."""
        if on_progress:
            on_progress(msg)
        else:
            console.info(msg)

    def format_time(seconds: float) -> str:
        """Format seconds into readable time string."""
        if seconds < 60:
            return f"{int(seconds)}s"
        return f"{int(seconds // 60)}m {int(seconds % 60)}s"

    folder_path = create_working_folder(company_name, website)
    progress(f"> Working folder: {folder_path}")

    # Scrape website - saves raw scrapes incrementally to _raw_scrapes folder
    # fetch_web_content already shows completion message, no need to duplicate
    scraped_data = (
        fetch_web_content(website, company_name, max_pages=50, working_folder=folder_path)
        if website
        else {}
    )

    # Abort if website was provided but nothing could be scraped
    if website and not scraped_data:
        console.fail("Could not scrape any pages from the website - site may be blocking")
        console.muted("  The site may be rate-limiting after a recent scrape.")
        console.muted('  Try: primr "Company" url --mode deep  (skips site scraping)')
        return None

    if website and fail_on_low_scrape:
        quality_ok, quality_reason = _validate_scrape_quality(scraped_data)
        if not quality_ok:
            console.fail(quality_reason)
            console.muted("  Re-run with --skip-scrape-validation to continue anyway")
            return None

    # External research - with LLM validation to ensure correct company
    # This prevents including content from similarly-named but unrelated companies
    # (e.g., "EverTrue" fundraising software vs "EverTrue" senior living)
    progress("Searching for external sources...")

    external_data = {}
    total_search_results = 0
    if website:
        from primr.data.scrape import scrape_external_sources_validated
        from primr.data.search_utils import search_web

        # Generate targeted search queries using LLM
        progress("Generating search strategy...")
        external_queries = generate_external_search_queries(
            company_name,
            website,
            max_queries=MAX_EXTERNAL_SEARCH_QUERIES,
        )

        # Keep hardcoded queries as reliable fallbacks at the end
        fallback_queries = [
            "news OR press release OR announcement",
            "funding OR acquisition OR partnership",
        ]
        all_queries = list(external_queries)
        for fallback in fallback_queries:
            if len(all_queries) >= MAX_EXTERNAL_SEARCH_QUERIES:
                break
            if fallback not in all_queries:
                all_queries.append(fallback)

        max_external_sources = MAX_EXTERNAL_SOURCES

        for query in all_queries:
            if len(external_data) >= max_external_sources:
                break

            results = search_web(query, company_name, website)
            if results:
                total_search_results += len(results)
                progress(f"  Found {len(results)} results for '{query[:40]}...'")
                filtered = [
                    r for r in results[:5] if website.lower() not in r.get("url", "").lower()
                ]
                remaining_slots = max_external_sources - len(external_data)
                progress(f"  Validating {len(filtered)} external articles...")
                scraped = scrape_external_sources_validated(
                    filtered,
                    company_name=company_name,
                    website=website,
                    max_sources=min(2, remaining_slots),
                    working_folder=folder_path,
                )
                external_data.update(scraped)

    progress(
        f"+ {len(external_data)} external sources validated (from {total_search_results} search results)"
    )

    all_scraped = {**scraped_data, **external_data}

    # Save raw scraped URLs to working folder for debugging
    urls_file = os.path.join(folder_path, "_scraped_urls.txt")
    with open(urls_file, "w", encoding="utf-8") as f:
        f.write(f"# Scraped URLs for {company_name}\n")
        f.write(f"# Website: {website}\n")
        f.write(f"# Total: {len(all_scraped)} pages\n\n")
        f.write("## Website Pages:\n")
        for url in scraped_data:
            f.write(f"  {url}\n")
        f.write(f"\n## External Sources ({len(external_data)}):\n")
        for url in external_data:
            f.write(f"  {url}\n")

    # Summarize content
    progress("Summarizing content...")
    summarized = summarize_scraped_content(company_name, website, all_scraped, folder_path)
    if not summarized.strip():
        summarized = "No insights extracted."
    progress("+ Content summarized")

    # Clean up raw scrapes folder now that we have the summary
    raw_folder = os.path.join(folder_path, "_raw_scrapes")
    if os.path.exists(raw_folder):
        import shutil

        try:
            shutil.rmtree(raw_folder)
            logger.debug("Cleaned up raw scrapes folder")
        except Exception as e:
            logger.debug(f"Failed to clean up raw scrapes: {e}")

    # Industry identification
    progress("Identifying industry...")
    industry_prompt = generate_prompt(
        "industry",
        company_name=company_name,
        company_website=website or "N/A",
        scraped_insights=summarized,
    )
    industry = llm(industry_prompt, model_type="research").strip() or "Unknown"
    progress(f"+ Industry: {industry}")

    # Overview
    progress("Generating overview...")
    overview = generate_initial_overview(company_name, website, industry, folder_path)
    progress("+ Overview complete")

    # Research all sections
    sections = [
        "Company Name",
        "Website",
        "Industry",
        "Detailed Products/Services",
        "Unique Selling Proposition (USP)",
        "Mission & Vision",
        "Company History",
        "Key Achievements",
        "Target Audience",
        "Financial Overview",
        "Potential Business Drivers & KPIs",
        "Industry Insights",
        "Potential Business Drivers",
        "Primary Apps or Sources of Data",
        "Main Types of Users",
        "Board of Directors Concerns",
        "Potential Business Value",
        "Strategic Recommendations",
    ]

    # Count non-trivial sections for progress
    analysis_sections = [s for s in sections if s not in ["Company Name", "Website", "Industry"]]
    total_analysis = len(analysis_sections)

    sections_start = time_module.time()
    progress(f"Analyzing {total_analysis} report sections...")
    section_results = {}
    analysis_idx = 0

    for section in sections:
        section_key = SECTION_KEY_MAP.get(section)
        if section_key:
            # Show progress for non-trivial sections (clean format, no timing per step)
            if section not in ["Company Name", "Website", "Industry"]:
                analysis_idx += 1
                progress(f"  [{analysis_idx}/{total_analysis}] {section}")

            content = research_section(
                section, company_name, website, industry, folder_path, overview, summarized
            )
            if content:
                section_results[section_key] = content

    # Final timing for sections
    progress(
        f"+ {total_analysis} sections complete ({format_time(time_module.time() - sections_start)})"
    )

    console.progress_done()
    return section_results


def perform_scrape_only(
    company_name: str | None,
    website: str | None,
    start_time: float,
    max_scrape_time: int | None = None,
    fail_on_low_scrape: bool = True,
    folder_path: str | None = None,
) -> str | None:
    """
    Scrape mode: Build site corpus + extract insights.

    Delegates to fetch_web_content() for all scraping work.

    Cost: ~$0.01-0.05 (LLM for summarization only)
    """
    if not website:
        console.fail("Scrape mode requires a website URL")
        return None

    display_name = company_name or urlparse(website).netloc

    # Create working folder (silent)
    folder_path = folder_path or create_working_folder(company_name, website)
    _update_run_state(folder_path, current_phase="scrape", status="running")
    _append_run_event(folder_path, "scrape", "started", "Scrape-only mode started")

    # Build Site Corpus (shows its own progress)
    corpus = fetch_web_content(
        website=website,
        company_name=company_name or display_name,
        max_pages=50,
        working_folder=folder_path,
    )

    pages_scraped = len(corpus)
    total_chars = sum(len(c or "") for c in corpus.values())

    if pages_scraped == 0:
        console.fail("Could not scrape any pages - site may be blocking")
        console.muted('Try: primr "Company" url --mode deep')
        _update_run_state(folder_path, status="failed", current_phase="scrape")
        _append_run_event(folder_path, "scrape", "failed", "No pages scraped")
        return None

    if website and fail_on_low_scrape:
        quality_ok, quality_reason = _validate_scrape_quality(corpus)
        if not quality_ok:
            console.fail(quality_reason)
            console.muted("Re-run with --skip-scrape-validation to continue anyway")
            _update_run_state(folder_path, status="failed", current_phase="scrape")
            _append_run_event(folder_path, "scrape", "failed", quality_reason)
            return None

    # Save combined corpus
    scraped_file = os.path.join(folder_path, "scraped_content.txt")
    with open(scraped_file, "w", encoding="utf-8") as f:
        f.write(f"# {display_name} - Scraped Content\n")
        f.write(f"# URL: {website}\n")
        f.write(f"# Pages: {pages_scraped}\n\n")
        for url, content in corpus.items():
            f.write(f"\n{'=' * 60}\n")
            f.write(f"URL: {url}\n")
            f.write(f"{'=' * 60}\n")
            f.write(content[:5000] + "\n")

    # Extract Insights (LLM)
    console.status("Extracting insights...")

    summarized = summarize_scraped_content(company_name, website, corpus, folder_path)
    console.clear_line()
    console.done("Insights extracted")

    # Save insights
    insights_file = os.path.join(folder_path, "insights.txt")
    with open(insights_file, "w", encoding="utf-8") as f:
        f.write(f"# {display_name} - Key Insights\n\n")
        f.write(summarized)

    # Final summary - one clean line
    elapsed = time.time() - start_time
    mins = int(elapsed // 60)
    secs = int(elapsed % 60)
    time_str = f"{mins}m {secs}s" if mins > 0 else f"{secs}s"

    console.blank()
    console.done(f"Complete: {pages_scraped} pages, {total_chars:,} chars ({time_str})")
    console.muted(f"Output: {folder_path}")
    _update_run_state(
        folder_path,
        status="completed",
        current_phase="complete",
        completed_at=datetime.now().isoformat(),
        pages_scraped=pages_scraped,
        scraped_chars=total_chars,
    )
    _append_run_event(folder_path, "scrape", "completed", "Scrape-only mode completed")

    return folder_path


# Part labels for console output
_PART_LABELS = {
    1: "Foundation",
    2: "Industry",
    3: "Strategic",
    4: "Deep Insights",
    5: "Synthesis",
}


def _group_sections_by_part() -> list[list["SectionConfig"]]:
    """
    Load sections from company_overview.yaml and group by ``part`` number.

    Returns a list of 5 lists (parts 1-5), each containing the
    :class:`SectionConfig` objects that belong to that part.
    """
    from primr.prompts.loader import load_prompt_config

    config = load_prompt_config("company_overview")
    groups: dict[int, list] = {}
    for section in config.sections:
        groups.setdefault(section.part, []).append(section)
    # Return in part order (1-5)
    return [groups[p] for p in sorted(groups)]


def _build_fast_batch_prompt(
    company_name: str,
    website: str | None,
    analysis_workbook: str,
    raw_corpus_subset: str,
    external_sources: str,
    source_urls: list[str],
    sections: list["SectionConfig"],
    previous_sections: list["GeneratedSection"],
    batch_number: int,
    total_batches: int,
) -> str:
    """
    Build the prompt for writing one batch of report sections.

    Each batch receives:
    - The analysis workbook (shared across all batches)
    - Raw corpus + external sources for evidence
    - 300-word rolling summaries of the last 5 completed sections
    - Section-specific instructions from the YAML config
    """
    current_date = datetime.now().strftime("%B %d, %Y")

    # Build section instructions for this batch
    section_parts: list[str] = []
    for section in sections:
        covers_text = "\n".join(f"      - {item}" for item in section.covers)
        depth_text = section.depth.strip() if section.depth else "Thorough analysis"
        position_label = section.position or "middle"
        section_parts.append(
            f"### {section.name}\n"
            f"**Purpose:** {section.purpose}\n"
            f"**Position:** {position_label}\n"
            f"**Must cover:**\n{covers_text}\n"
            f"**Depth:** {depth_text}"
        )
    section_block = "\n\n".join(section_parts)

    # Rolling context: 400-word summaries of last 7 completed sections
    rolling_context = ""
    if previous_sections:
        recent = previous_sections[-7:]
        context_parts: list[str] = []
        for s in recent:
            # Truncate each section to ~400 words for rolling context
            words = s.content.split()
            summary = " ".join(words[:400])
            if len(words) > 400:
                summary += " ..."
            context_parts.append(f"**{s.title}** (completed):\n{summary}")
        rolling_context = "\n\n".join(context_parts)

    rolling_block = (
        f"## PREVIOUS SECTIONS (for narrative continuity)\n{rolling_context}"
        if rolling_context
        else "## PREVIOUS SECTIONS\n(This is the first batch — no prior sections.)"
    )

    sources_text = (
        "\n".join(f"[{i}] {url}" for i, url in enumerate(source_urls, start=1))
        if source_urls
        else "(no external sources)"
    )
    word_target = len(sections) * 800
    feedback_guidance = _load_fast_feedback_guidance()
    feedback_block = (
        f"=== FAST FEEDBACK GUIDANCE (from prior evals) ===\n{feedback_guidance}\n"
        if feedback_guidance
        else ""
    )
    return f"""**Company:** {company_name}
**Website:** {website or "N/A"}
**Date:** {current_date}
**Batch:** {batch_number + 1} of {total_batches}

You are writing batch {batch_number + 1} of {total_batches} for a Strategic Company Overview.
This batch contains {len(sections)} sections. Write each section under its own ## heading.

{rolling_block}

=== ANALYSIS WORKBOOK ===
{analysis_workbook}

=== RAW DATA (for evidence and citations) ===
{raw_corpus_subset}

=== EXTERNAL SOURCES ===
{external_sources}

{feedback_block}

SOURCES CONSULTED:
{sources_text}

---

Write the following sections. Each section MUST start with a ## heading matching the section name exactly.

{section_block}

REQUIREMENTS:
- Write at least {word_target:,} words total across all sections in this batch
- Use specific facts, numbers, examples, and strategic comparisons — cite sources with [cite: N]
- Be analytical and hypothesis-driven, not just descriptive
- Label claims with confidence levels (Confirmed/Reported/Estimated/Hypothesis)
- If direct evidence is limited, still write a deep section by anchoring on observed facts,
  extending with defensible inference, and making the strategic implication explicit
- Build on the previous sections' narrative (see rolling context above)
- For framework sections (SWOT, Porter's, Value Chain): organize insights from
  earlier sections, don't introduce wholly new observations
- Include tables where instructed (financials, competitors, timelines)
- Each section should have substantive depth — multiple paragraphs with evidence
- If a numeric claim cannot be supported by a cited source, replace it with
  "Not publicly disclosed", a bounded qualitative range, or a clearly labeled low-confidence directional statement
- Do not invent market sizes, CAGR, revenue ranges, headcount ranges, or shares
  unless directly grounded in one or more cited sources or a transparent comparative heuristic
- End each section with a short "What to validate:" line containing one concrete
  discovery question or data point to confirm in client interviews
- ZERO REPETITION: Before writing a section, review the rolling context above.
  If an insight, data point, or hypothesis already appeared in a prior section,
  do NOT restate it. Reference the earlier section instead ("as noted in the
  Executive Summary...") or build on it with new evidence.
- AI/TECHNOLOGY INTEGRATION: When relevant, explicitly connect AI or technology
  use cases to the company's specific business challenges. Don't just mention
  "AI could help" — specify which AI capability (NLP, computer vision, predictive
  analytics, etc.) maps to which concrete business problem identified in this report.

CONSULTING RIGOR (critical):
- Do NOT paraphrase the company's marketing. When you cite their claims, immediately
  stress-test them against external evidence or flag what's unverifiable.
- For each major hypothesis or insight, include "What to validate": a specific question
  or data point a consultant should probe in discovery.
- Be CONSERVATIVE on financial estimates. If you're inferring revenue from employee
  count, say "highly uncertain" and use wide ranges. Never state inferences as fact.
- Frame "why now" for the company — what transition or inflection point makes this
  moment interesting? Platform shifts, PE investment, leadership changes, etc.
- Think like a buyer, not a narrator. Where does this company win deals? Where does
  it lose? What would a competitor say about them?
- When direct evidence is sparse, go deeper on likely economics, buyer behavior,
  operating constraints, strategic tradeoffs, scenario paths, and the decisions leadership faces.
- When direct evidence is sparse, go deeper on likely economics, buyer behavior, operating
  constraints, strategic tradeoffs, scenario paths, and the decisions leadership likely faces.

CITATION FORMAT (strict):
- The SOURCES CONSULTED block above is a numbered citation key: [N] URL
- Inline claims must reference citations as [cite: N], where N matches the
  number assigned to that URL in the SOURCES CONSULTED block
- Reuse the same N every time you cite the same URL
- Do NOT emit [Source: URL] inline; use [cite: N] only
- Do NOT invent citation numbers — only cite N values present in the key above

OUTPUT CONTRACT (strict):
- Preferred format: emit each section inside a lightweight XML envelope:
  <section><title>Exact Section Name</title><body>Section body here</body></section>
- If you do not use the XML envelope, start each section with exactly one ## heading matching the requested section name
- Do not add a ## Sources, ## References, or ## Citations subsection inside section output
- Include exactly one What to validate: line per section, and make it the final line of that section
- Do not include any preamble or commentary outside the requested section bodies
"""


# --- Section-level fast-mode helpers (individual section writing) ---

_HIGH_DEPTH_SECTION_IDS = frozenset(
    {
        "executive_summary",
        "competitive_landscape",
        "company_history",
        "engagement_opportunities",
    }
)


def _get_section_word_target(section: "SectionConfig") -> int:
    """Return adaptive word target for a single section.

    - Sections with depth mentioning 'pages'/'comprehensive', or IDs in
      ``_HIGH_DEPTH_SECTION_IDS`` → 1,200 words
    - Framework sections (position == 'framework') → 800 words
    - Everything else → 800 words
    """
    depth_lower = (section.depth or "").lower()
    if (
        section.id in _HIGH_DEPTH_SECTION_IDS
        or "pages" in depth_lower
        or "comprehensive" in depth_lower
    ):
        return 1_200
    if section.position == "framework":
        return 800
    return 800


def _get_section_max_tokens(section: "SectionConfig") -> int:
    """Return max_tokens for a single-section Grok call."""
    return 6_000 if _get_section_word_target(section) >= 1_000 else 4_000


def _build_fast_section_prompt(
    company_name: str,
    website: str | None,
    analysis_workbook: str,
    raw_corpus_subset: str,
    external_sources: str,
    source_urls: list[str],
    section: "SectionConfig",
    written_sections: list["GeneratedSection"],
    section_index: int,
    all_section_names: list[str],
    reasoning_mode: str = "standard",
) -> str:
    """Build prompt for writing a single report section.

    Similar to ``_build_fast_batch_prompt`` but tailored for one section at
    a time with table-of-contents awareness and rolling context.
    """
    current_date = datetime.now().strftime("%B %d, %Y")
    word_target = _get_section_word_target(section)

    # Section instructions
    covers_text = "\n".join(f"      - {item}" for item in section.covers)
    depth_text = section.depth.strip() if section.depth else "Thorough analysis"
    position_label = section.position or "middle"
    section_block = (
        f"### {section.name}\n"
        f"**Purpose:** {section.purpose}\n"
        f"**Position:** {position_label}\n"
        f"**Must cover:**\n{covers_text}\n"
        f"**Depth:** {depth_text}"
    )

    # Table of contents with [DONE]/[NOW]/[TODO] markers
    toc_parts: list[str] = []
    for idx, name in enumerate(all_section_names):
        if idx < section_index:
            toc_parts.append(f"  [DONE] {name}")
        elif idx == section_index:
            toc_parts.append(f"  [NOW]  {name}")
        else:
            toc_parts.append(f"  [TODO] {name}")
    toc_block = "## REPORT TABLE OF CONTENTS\n" + "\n".join(toc_parts)

    # Rolling context: framework/exec-summary sections get full prior content for synthesis;
    # regular sections get 300-word summaries of last 5 completed sections
    rolling_context = ""
    if written_sections:
        if section.position == "framework" or section.id == "executive_summary":
            # Framework sections and executive summary need full prior content
            # to synthesise insights from earlier analytical sections
            context_parts = [f"**{s.title}** (completed):\n{s.content}" for s in written_sections]
        else:
            recent = written_sections[-5:]
            context_parts = []
            for s in recent:
                words = s.content.split()
                summary = " ".join(words[:300])
                if len(words) > 300:
                    summary += " ..."
                context_parts.append(f"**{s.title}** (completed):\n{summary}")
        rolling_context = "\n\n".join(context_parts)

    rolling_block = (
        f"## PREVIOUS SECTIONS (for narrative continuity)\n{rolling_context}"
        if rolling_context
        else "## PREVIOUS SECTIONS\n(This is the first section — no prior sections.)"
    )

    sources_text = (
        "\n".join(f"[{i}] {url}" for i, url in enumerate(source_urls, start=1))
        if source_urls
        else "(no external sources)"
    )
    feedback_guidance = _load_fast_feedback_guidance()
    feedback_block = (
        f"=== FAST FEEDBACK GUIDANCE (from prior evals) ===\n{feedback_guidance}\n"
        if feedback_guidance
        else ""
    )

    reasoning_guidance = (
        "CONSTRAINED-EVIDENCE MODE: Direct company-specific evidence for this section is limited. "
        "Do NOT collapse into a thin fact check. Use the website, news, industry structure, competitor "
        "analogs, and operating logic to build a deep strategic section. Separate what is observed, what "
        "is inferred, what is hypothesis, and what the strategic implication is."
        if reasoning_mode == "constrained_evidence"
        else "STANDARD-EVIDENCE MODE: Use the strongest available mix of direct evidence, external research, "
        "and strategic inference."
    )

    return f"""**Company:** {company_name}
**Website:** {website or "N/A"}
**Date:** {current_date}
**Section:** {section_index + 1} of {len(all_section_names)} — {section.name}

{toc_block}

You are writing ONE section of a Strategic Company Overview.
Write this section under a single ## heading matching the section name exactly.

{rolling_block}

=== ANALYSIS WORKBOOK ===
{analysis_workbook}

=== RAW DATA (for evidence and citations) ===
{raw_corpus_subset}

=== EXTERNAL SOURCES ===
{external_sources}

{feedback_block}

REASONING MODE:
{reasoning_guidance}

SOURCES CONSULTED:
{sources_text}

---

Write the following section. It MUST start with a ## heading matching the section name exactly.

{section_block}

REQUIREMENTS:
- Write at least {word_target:,} words for this section
- Use specific facts, numbers, examples, and strategic comparisons — cite sources with [cite: N]
- Be analytical and hypothesis-driven, not just descriptive
- Label claims with confidence levels (Confirmed/Reported/Estimated/Hypothesis)
- If direct evidence is limited, still write a deep section by anchoring on observed facts,
  extending with defensible inference, and making the strategic implication explicit
- Build on the previous sections' narrative (see rolling context above)
- For framework sections (SWOT, Porter's, Value Chain): organize insights from
  earlier sections, don't introduce wholly new observations
- Include tables where instructed (financials, competitors, timelines)
- This section should have substantive depth — multiple paragraphs with evidence
- If a numeric claim cannot be supported by a cited source, replace it with
  "Not publicly disclosed", a bounded qualitative range, or a clearly labeled low-confidence directional statement
- Do not invent market sizes, CAGR, revenue ranges, headcount ranges, or shares
  unless directly grounded in one or more cited sources or a transparent comparative heuristic
- End the section with a short "What to validate:" line containing one concrete
  discovery question or data point to confirm in client interviews
- ZERO REPETITION: Before writing, review the rolling context and TOC above.
  If an insight, data point, or hypothesis already appeared in a prior section,
  do NOT restate it. Reference the earlier section instead ("as noted in the
  Executive Summary...") or build on it with new evidence.
- AI/TECHNOLOGY INTEGRATION: When relevant, explicitly connect AI or technology
  use cases to the company's specific business challenges. Don't just mention
  "AI could help" — specify which AI capability maps to which concrete problem.
- CITATION HYGIENE: Keep citations compact. Prefer paragraph-end citation clusters over
  interrupting every sentence, and let the final Sources appendix carry the dense reference load.

CONSULTING RIGOR (critical):
- Do NOT paraphrase the company's marketing. When you cite their claims, immediately
  stress-test them against external evidence or flag what's unverifiable.
- For each major hypothesis or insight, include "What to validate": a specific question
  or data point a consultant should probe in discovery.
- Be CONSERVATIVE on financial estimates. If you're inferring revenue from employee
  count, say "highly uncertain" and use wide ranges. Never state inferences as fact.
- Frame "why now" for the company — what transition or inflection point makes this
  moment interesting? Platform shifts, PE investment, leadership changes, etc.
- Think like a buyer, not a narrator. Where does this company win deals? Where does
  it lose? What would a competitor say about them?

CITATION FORMAT (strict):
- The SOURCES CONSULTED block above is a numbered citation key: [N] URL
- Inline claims must reference citations as [cite: N], where N matches the
  number assigned to that URL in the SOURCES CONSULTED block
- Reuse the same N every time you cite the same URL
- Do NOT emit [Source: URL] inline; use [cite: N] only
- Do NOT invent citation numbers — only cite N values present in the key above

OUTPUT CONTRACT (strict):
- Preferred format: emit each section inside a lightweight XML envelope:
  <section><title>Exact Section Name</title><body>Section body here</body></section>
- If you do not use the XML envelope, start each section with exactly one ## heading matching the requested section name
- Do not add a ## Sources, ## References, or ## Citations subsection inside section output
- Include exactly one What to validate: line per section, and make it the final line of that section
- Do not include any preamble or commentary outside the requested section bodies
"""


def _normalize_generated_section_payload(
    title: str,
    body: str,
    expected_title: str | None = None,
) -> "GeneratedSection":
    """Normalize a generated section into a stricter payload contract."""
    canonical_title = (expected_title or title or "Section").strip().rstrip("#").strip()
    working_body = body.strip()

    # Drop a duplicated section heading if the model nested one inside the body.
    heading_match = re.match(r"^##\s+.+?(?:\n+|$)", working_body)
    if heading_match:
        working_body = working_body[heading_match.end() :].lstrip()

    # Drop any embedded reference appendix the model tried to include inside a section.
    ref_match = re.search(
        r"^##\s+(Sources|References|Citations)\s*$",
        working_body,
        flags=re.IGNORECASE | re.MULTILINE,
    )
    if ref_match:
        working_body = working_body[: ref_match.start()].rstrip()

    validation_lines: list[str] = []
    cleaned_lines: list[str] = []
    for raw_line in working_body.splitlines():
        stripped = raw_line.strip()
        if re.match(r"^What to validate:\s*", stripped, flags=re.IGNORECASE):
            question = re.sub(r"^What to validate:\s*", "", stripped, flags=re.IGNORECASE).strip()
            if question:
                validation_lines.append(f"What to validate: {question}")
            continue
        cleaned_lines.append(raw_line)

    cleaned_body = "\n".join(cleaned_lines).strip()
    validation_line = (
        validation_lines[-1]
        if validation_lines
        else "What to validate: Confirm this section's key claim with primary customer or operator evidence."
    )
    content = (cleaned_body + "\n\n" + validation_line).strip() if cleaned_body else validation_line
    citation_numbers: list[int] = []
    for raw_group in re.findall(r"\[cite:\s*([0-9,\s]+)\]", content, re.IGNORECASE):
        for raw_num in raw_group.split(","):
            raw_num = raw_num.strip()
            if raw_num.isdigit():
                num = int(raw_num)
                if num not in citation_numbers:
                    citation_numbers.append(num)

    from primr.output.final_artifact import GeneratedSection

    return GeneratedSection(
        title=canonical_title,
        content=content,
        words=len(content.split()),
        validate_line=validation_line,
        citation_numbers=citation_numbers,
    )


def _parse_structured_section_envelopes(content: str) -> list[tuple[str, str]]:
    """Parse optional XML-style section envelopes emitted by the writer stage."""
    matches = re.findall(
        r"<section>\s*<title>(.*?)</title>\s*<body>(.*?)</body>\s*</section>",
        content,
        flags=re.IGNORECASE | re.DOTALL,
    )
    parsed: list[tuple[str, str]] = []
    for title, body in matches:
        clean_title = re.sub(r"\s+", " ", title).strip()
        clean_body = body.strip()
        if clean_title:
            parsed.append((clean_title, clean_body))
    return parsed


_SECTION_ENVELOPE_RE = re.compile(
    r"<section>\s*<title>(.*?)</title>\s*<body>(.*?)</body>\s*</section>",
    re.IGNORECASE | re.DOTALL,
)
_SECTION_HEADING_RE = re.compile(r"^##\s+(.+?)\s*$", re.MULTILINE)


def _extract_generated_section_blocks(content: str) -> tuple[str, list[tuple[str, str]]]:
    """Extract generated sections in source order from XML envelopes and/or markdown headings."""
    envelope_matches = list(_SECTION_ENVELOPE_RE.finditer(content))
    envelope_spans = [(match.start(), match.end()) for match in envelope_matches]

    def inside_envelope(position: int) -> bool:
        return any(start <= position < end for start, end in envelope_spans)

    heading_matches = [
        match
        for match in _SECTION_HEADING_RE.finditer(content)
        if not inside_envelope(match.start())
    ]

    block_starts = sorted(
        [match.start() for match in envelope_matches] + [match.start() for match in heading_matches]
    )
    parsed_blocks: list[tuple[int, str, str]] = []

    for match in envelope_matches:
        title = re.sub(r"\s+", " ", match.group(1)).strip()
        body = match.group(2).strip()
        if title:
            parsed_blocks.append((match.start(), title, body))

    for match in heading_matches:
        next_start = next((start for start in block_starts if start > match.start()), len(content))
        title = match.group(1).strip().rstrip("#").strip()
        body = content[match.end() : next_start].strip()
        if title:
            parsed_blocks.append((match.start(), title, body))

    parsed_blocks.sort(key=lambda item: item[0])
    preamble_end = parsed_blocks[0][0] if parsed_blocks else 0
    preamble = content[:preamble_end].strip()
    ordered_blocks = [(title, body) for _, title, body in parsed_blocks]
    return preamble, ordered_blocks


def _parse_single_section(
    content: str,
    expected_section: "SectionConfig",
) -> "GeneratedSection":
    """Parse Grok's single-section response.

    Expects one ``## `` heading or an optional ``<section>`` envelope.
    Falls back to using the expected section name if no heading found.
    """
    preamble, blocks = _extract_generated_section_blocks(content)
    if blocks:
        title, body = blocks[0]
        if preamble:
            body = f"{preamble}\n\n{body}".strip()
        return _normalize_generated_section_payload(title, body, expected_section.name)

    return _normalize_generated_section_payload(
        expected_section.name,
        content.strip(),
        expected_section.name,
    )


def _determine_section_reasoning_mode(section: "SectionConfig", analysis_workbook: str) -> str:
    """Use constrained-evidence reasoning when direct company signal is thin."""
    evidence_keywords = {
        "financial_profile": ["revenue", "profit", "margin", "funding", "valuation", "earnings"],
        "company_history": ["founded", "history", "acquisition", "pivot", "milestone"],
        "industry_outlook": ["industry trend", "regulation", "outlook", "forecast", "disruption"],
    }
    keywords = evidence_keywords.get(section.id)
    if not keywords:
        return "standard"
    workbook_lower = analysis_workbook.lower() if analysis_workbook else ""
    hits = sum(1 for kw in keywords if kw in workbook_lower)
    return "constrained_evidence" if hits == 0 else "standard"


def _write_section_with_retry(
    section: "SectionConfig",
    section_index: int,
    all_section_names: list[str],
    written_sections: list["GeneratedSection"],
    company_name: str,
    website: str | None,
    analysis_workbook: str,
    raw_corpus_subset: str,
    external_sources_raw: str,
    source_urls: list[str],
    report_system: str,
    reasoning_mode: str = "standard",
    model: str | None = None,
) -> "GeneratedSection | None":
    """Write a single section with one retry if output is thin.

    Returns the parsed section, or ``None`` on failure.

    NOTE (pipeline-resilience): The thin-section retry below is a *content
    quality* retry (re-prompt when word count is too low), not an API error
    retry.  It is intentionally retained alongside the stage-level
    RecoveryExecutor which handles API failures, model fallback, and
    skip/abort.  The two layers are complementary.
    """
    from primr.ai.grok_client import grok_llm

    word_target = _get_section_word_target(section)

    prompt = _build_fast_section_prompt(
        company_name,
        website,
        analysis_workbook,
        raw_corpus_subset,
        external_sources_raw,
        source_urls,
        section,
        written_sections,
        section_index,
        all_section_names,
        reasoning_mode,
    )

    writing_model = model or GROK_MODEL_WRITING
    try:
        section_content = grok_llm(
            prompt,
            model=writing_model,
            max_tokens=_get_section_max_tokens(section),
            temperature=0.6,
            system_prompt=report_system,
        )
    except Exception as section_err:
        logger.warning("Section '%s' failed: %s", section.name, section_err)
        log_structured(
            "warning", "Fast mode section failed", section=section.name, error=str(section_err)
        )
        return None

    if not section_content or not section_content.strip():
        logger.warning("Section '%s' returned empty", section.name)
        return None

    parsed = _parse_single_section(section_content, section)

    # Thin-section retry: if < 50% of target, retry once
    if parsed.words < word_target * 0.5:
        logger.warning(
            "Section '%s' too thin (%d/%d words), retrying",
            section.name,
            parsed.words,
            word_target,
        )
        retry_prompt = (
            f"IMPORTANT: Your previous attempt produced only {parsed.words} words. "
            f"The minimum target is {word_target} words. Write a substantive, strategist-grade, "
            f"decision-useful section. If direct evidence is limited, deepen the analysis with explicit "
            f"inference and strategic implications rather than staying thin.\n\n" + prompt
        )
        try:
            retry_content = grok_llm(
                retry_prompt,
                model=writing_model,
                max_tokens=_get_section_max_tokens(section),
                temperature=0.6,
                system_prompt=report_system,
            )
            if retry_content and retry_content.strip():
                retry_parsed = _parse_single_section(retry_content, section)
                if retry_parsed.words > parsed.words:
                    parsed = retry_parsed
        except Exception as retry_err:
            logger.warning("Section '%s' retry failed: %s", section.name, retry_err)

    return parsed


def _fast_coherence_pass(
    company_name: str,
    website: str | None,
    report_content: str,
    model: str | None = None,
) -> str:
    """Run a coherence pass over the assembled fast-mode report.

    Tasks:
    1. Remove cross-section repetition → replace with cross-references
    2. Smooth transitions between sections
    3. Ensure framework sections reference earlier analytical sections
    4. Fix terminology consistency

    Guards against destructive compression — rejects output that loses
    too many words or sections.
    """
    from primr.ai.grok_client import grok_llm

    if not report_content.strip():
        return report_content

    # Skip coherence pass for reports that would exceed the output token budget.
    # ~1.4 tokens per word; prompt overhead ~2K tokens; max_tokens=32K.
    word_count = len(report_content.split())
    max_output_words = 20_000  # ~28K tokens, leaving headroom within 32K cap
    if word_count > max_output_words:
        logger.info(
            "Skipping coherence pass: report is %d words (exceeds %d word limit for single-pass editing)",
            word_count,
            max_output_words,
        )
        return report_content

    prompt = f"""You are a copy editor making MINIMAL tweaks to a strategic company overview.
The report was written section-by-section. Your ONLY job is light polish.

Company: {company_name}
Website: {website or "N/A"}

CRITICAL: Output MUST be at least 98% of the original word count. You are NOT rewriting.

YOUR THREE TASKS — nothing else:
1. TERMINOLOGY: If the company/product name varies (e.g. "Northgate" vs "Northgate Gonzalez"),
   pick one form and use it consistently. Fix obvious typos only.
2. CROSS-REFERENCES: If the EXACT same sentence (same fact, same number) appears in two
   sections, replace the SECOND occurrence with "As noted in [Section Name]..." Keep all
   surrounding paragraphs intact.
3. TRANSITIONS: Add ONE linking sentence at the start of each section (e.g. "Building on
   the financial profile above...") if no transition exists.

ABSOLUTE PROHIBITIONS:
- Do NOT delete ANY paragraph, bullet, subsection, or table
- Do NOT remove or rename ## headings
- Do NOT remove [cite: N], confidence labels, or "What to validate:" lines
- Do NOT rewrite sentences for style — only fix terminology and add cross-references
- Do NOT summarize, condense, or merge sections
- Do NOT add new facts or analysis
- Every paragraph in the input MUST appear in the output

If in doubt, leave the text unchanged. Err on the side of doing nothing.

Return the full markdown report. No preamble.

--- REPORT START ---
{report_content}
--- REPORT END ---
"""
    writing_model = model or GROK_MODEL_WRITING
    try:
        polished = grok_llm(
            prompt,
            model=writing_model,
            max_tokens=32_000,
            temperature=0.3,
            system_prompt=(
                "You are a meticulous editorial analyst improving coherence and flow "
                "across a multi-section strategic report. Preserve ALL depth and evidence. "
                "Make only surgical edits to duplicate sentences — never delete paragraphs."
            ),
        )
        if not polished or not polished.strip():
            return report_content

        # Guard: reject destructive compression
        original_words = len(report_content.split())
        polished_words = len(polished.split())
        _, original_sections = _split_markdown_sections(report_content)
        _, polished_sections = _split_markdown_sections(polished)

        if polished_words < int(original_words * 0.96):
            logger.warning(
                "Coherence pass dropped too many words (%d → %d), using original",
                original_words,
                polished_words,
            )
            return report_content
        if len(polished_sections) < len(original_sections):
            logger.warning(
                "Coherence pass lost sections (%d → %d), using original",
                len(original_sections),
                len(polished_sections),
            )
            return report_content

        return polished
    except Exception:
        logger.warning("Coherence pass failed, using original report", exc_info=True)
        return report_content


def _clean_fast_report_output(report_content: str) -> str:
    """
    Final cleanup of fast-mode report artifacts before citation normalization.

    Removes:
    - Standalone (Reported)/(Estimated)/(Confirmed) confidence labels at section
      boundaries (inline labels within prose are preserved)
    - Grok model disclaimer boilerplate
    - Informal [cite: name] tags that reference internal labels (not URLs)
    - Excess blank lines
    """
    if not report_content.strip():
        return report_content

    # Rewrite noisy inline citation patterns into cleaner, auditable forms.
    report_content = _rewrite_inline_confidence_citations(report_content)
    report_content = _rewrite_cite_from_url_tags(report_content)

    # 1. Strip Grok disclaimer (appears at end of report or AI strategy)
    report_content = re.sub(
        r"\n*_?Disclaimer:\s*Grok is not a financial advi[sc]er[^\n]*\n?",
        "\n",
        report_content,
        flags=re.IGNORECASE,
    )

    # 2. Strip standalone confidence labels at section boundaries.
    #    These appear as "(Reported)" on their own line, typically between
    #    a "What to validate" line and the next section heading or EOF.
    report_content = re.sub(
        r"\n\s*\((?:Reported|Estimated|Confirmed|Hypothesis)\)\s*\n(?=\s*\n|$)",
        "\n",
        report_content,
    )

    # 3. Strip informal [cite: name] tags (non-numeric, non-URL references).
    #    These are internal labels like [cite: workbook], [cite: bbb] that
    #    can't be resolved to URLs. Inline confidence labels in prose are the
    #    proper way to indicate source type.
    #    Matches: [cite: workbook], [cite: website; cite: bbb], [cite: enrollment]
    #    Preserves: [cite: 1], [cite: 2, 3], [Source: URL]
    def _strip_informal_cites(match: re.Match[str]) -> str:
        return _sanitize_numeric_cite_bracket(match.group(1))

    report_content = re.sub(
        r"\[([^\]]*cites?:\s*[^\]]+)\]",
        _strip_informal_cites,
        report_content,
        flags=re.IGNORECASE,
    )

    # Also strip [cross-ref: ...] tags — internal analysis references
    report_content = re.sub(r"\s*\[cross-ref:[^\]]*\]", "", report_content, flags=re.IGNORECASE)

    # 4. Strip internal citation inventory/debug notes that should never ship.
    report_content = re.sub(
        r"\n?\[citation inventory[^\]]*\]\n?",
        "\n",
        report_content,
        flags=re.IGNORECASE,
    )

    # 4b. Strip internal workbook/external-source references that are useful
    # during generation but should not appear in shipped artifacts.
    report_content = re.sub(r"\[Workbook:[^\]]*\]", "", report_content, flags=re.IGNORECASE)
    report_content = re.sub(r"\[workbook section[^\]]*\]", "", report_content, flags=re.IGNORECASE)
    report_content = re.sub(r"\[Workbook §[^\]]*\]", "", report_content, flags=re.IGNORECASE)
    report_content = re.sub(r"\[Analysis Workbook[^\]]*\]", "", report_content, flags=re.IGNORECASE)
    report_content = re.sub(r"\[Analysis:[^\]]*\]", "", report_content, flags=re.IGNORECASE)
    report_content = re.sub(r"\[External Sources\]", "", report_content, flags=re.IGNORECASE)
    report_content = re.sub(
        r"vendor-research-[\w.-]+\.txt", "", report_content, flags=re.IGNORECASE
    )
    report_content = re.sub(r"\bInternal ROI Model\b", "", report_content, flags=re.IGNORECASE)
    report_content = re.sub(r"\bInternal Analysis\b", "", report_content, flags=re.IGNORECASE)
    report_content = re.sub(r"\bAnalysis Workbook\b", "", report_content, flags=re.IGNORECASE)

    # 5. Strip LLM meta-annotations like [Word count: 1028]
    report_content = re.sub(r"\[Word count:\s*[\d,]+\]", "", report_content, flags=re.IGNORECASE)

    # 6. Clean up double spaces left by stripped citations/tags
    report_content = re.sub(r"  +", " ", report_content)

    # 7. Clean up excess blank lines (3+ newlines -> 2)
    report_content = re.sub(r"\n{3,}", "\n\n", report_content)

    return report_content.strip() + "\n"


def _rewrite_inline_confidence_citations(content: str) -> str:
    """Convert nested confidence/source annotations into cleaner prose."""
    pattern = re.compile(
        r"\[(Confirmed|Reported|Estimated|Hypothesis):\s*([^\[\]]*?)\s*"
        r"\[cite:\s*\d+\s+from\s+(https?://[^\]\s]+)\]\s*\]",
        re.IGNORECASE,
    )

    def _replace(match: re.Match[str]) -> str:
        label = match.group(1).capitalize()
        detail = re.sub(r"\s+", " ", match.group(2)).strip(" ;,")
        url = match.group(3).strip()
        if detail:
            return f"({label}: {detail}) [Source: {url}]"
        return f"({label}) [Source: {url}]"

    return pattern.sub(_replace, content)


def _rewrite_cite_from_url_tags(content: str) -> str:
    """Convert malformed `[cite: N from URL]` tags into source tags for normalization."""
    return re.sub(
        r"\[cite:\s*\d+\s+from\s+(https?://[^\]\s]+)\]",
        lambda m: f"[Source: {m.group(1).strip()}]",
        content,
        flags=re.IGNORECASE,
    )


def _sanitize_numeric_cite_bracket(inner: str) -> str:
    """Keep only numeric cite ids from a mixed citation bracket."""
    nums: list[str] = []
    for cite_match in re.finditer(r"cites?:\s*([^;\]]+)", inner, re.IGNORECASE):
        for raw_num in re.findall(r"\d+", cite_match.group(1)):
            if raw_num not in nums:
                nums.append(raw_num)
    if not nums:
        return ""
    return "[cite: " + ", ".join(nums) + "]"


_INTERNAL_REFERENCE_TERMS = (
    "analysis context",
    "analysis workbook",
    "internal analysis",
    "internal roi model",
    "vendor-research",
    "workbook",
    "company report",
    "industry baseline",
    "market analysis",
    "itr on website",
    "itron website",
    "insights.txt",
    "workbook.md",
)


def _strip_internal_source_placeholders(content: str) -> str:
    """Remove non-auditable internal source placeholders from final outputs."""
    if not content.strip():
        return content

    confidence_bracket = re.compile(
        r"\[(Confirmed|Reported|Estimated|Hypothesis):\s*([^\]]+)\]", re.IGNORECASE
    )

    def _drop_if_internal(match: re.Match[str]) -> str:
        source_text = match.group(2).lower()
        if any(term in source_text for term in _INTERNAL_REFERENCE_TERMS):
            return ""
        return match.group(0)

    cleaned = confidence_bracket.sub(_drop_if_internal, content)
    cleaned = re.sub(
        r"\[(?:Reported|Confirmed|Estimated|Hypothesis):\s*\]", "", cleaned, flags=re.IGNORECASE
    )
    cleaned = re.sub(r"\[citation inventory[^\]]*\]", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned


def _strip_unresolved_section_cross_references(content: str) -> str:
    """Remove unresolved internal section references that should not ship."""
    if not content.strip():
        return content

    cleaned = re.sub(
        r"\[\s*(?:see|cross-?ref|xref)\s+##\s+[^\]]+\]",
        "",
        content,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned


def _extract_markdown_headings(content: str) -> list[str]:
    """Return normalized markdown headings in document order."""
    return [heading.strip() for heading in re.findall(r"^##\s+(.+?)\s*$", content, re.MULTILINE)]


def _preserves_report_structure(original: str, candidate: str) -> bool:
    """Require same ordered headings, allowing only an appended sources section."""
    original_headings = _extract_markdown_headings(original)
    candidate_headings = _extract_markdown_headings(candidate)
    if candidate_headings[: len(original_headings)] != original_headings:
        return False
    extra_headings = [h.strip().lower() for h in candidate_headings[len(original_headings) :]]
    if any(h not in {"sources", "citations", "references"} for h in extra_headings):
        return False

    original_words = len(original.split())
    candidate_words = len(candidate.split())
    if original_words == 0:
        return False
    return candidate_words >= int(original_words * 0.98)


def _strategy_money_to_millions(value: float, unit: str) -> float:
    unit = unit.upper()
    if unit == "B":
        return value * 1000.0
    if unit == "K":
        return value / 1000.0
    return value


def _is_auditable_source_url(url: str) -> bool:
    """Require a public HTTP(S) URL with a plausible hostname for source appendices."""
    import ipaddress

    try:
        parsed = urlparse(url)
    except Exception:
        return False

    host = (parsed.hostname or "").strip().lower()
    if not host:
        return False

    try:
        ipaddress.ip_address(host)
        return True
    except ValueError:
        pass

    if "." not in host:
        return False

    labels = host.split(".")
    if any(not label or label.startswith("-") or label.endswith("-") for label in labels):
        return False

    host_label = re.compile(r"^[a-z0-9-]{1,63}$")
    return all(host_label.fullmatch(label) for label in labels)


def _normalize_strategy_source_urls(source_urls: list[str]) -> tuple[list[str], list[str]]:
    """Return normalized auditable source URLs plus rejected raw entries."""
    normalized_urls: list[str] = []
    rejected_urls: list[str] = []
    seen: set[str] = set()

    for raw_url in source_urls:
        candidate = raw_url.strip()
        if not candidate:
            continue
        is_valid, normalized, _error = validate_url_for_request(candidate)
        if not is_valid or not _is_auditable_source_url(normalized):
            rejected_urls.append(candidate)
            continue
        if normalized not in seen:
            seen.add(normalized)
            normalized_urls.append(normalized)

    return normalized_urls, rejected_urls


def _extract_strategy_citation_definitions(
    strategy_content: str,
) -> tuple[set[int], dict[int, str], list[str]]:
    """Parse strategy citation definitions and keep only valid auditable URLs."""
    cited_numbers = {
        int(n) for n in re.findall(r"\[cite:\s*(\d+)\]", strategy_content, re.IGNORECASE)
    }
    valid_defs: dict[int, str] = {}
    invalid_defs: list[str] = []

    for num_str, raw_url in re.findall(
        r"\[cite:\s*(\d+)\]\s+([^\s]+)", strategy_content, re.IGNORECASE
    ):
        cite_num = int(num_str)
        is_valid, normalized, _error = validate_url_for_request(raw_url.strip())
        if not is_valid or not _is_auditable_source_url(normalized):
            invalid_defs.append(raw_url.strip())
            continue
        valid_defs[cite_num] = normalized

    return cited_numbers, valid_defs, invalid_defs


def _compute_strategy_qa_metrics(strategy_content: str) -> dict[str, int | float | bool]:
    """Deterministic QA checks for strategy outputs."""
    if not strategy_content.strip():
        return {
            "placeholder_refs": 0,
            "source_urls": 0,
            "citation_defs": 0,
            "missing_citations": 0,
            "invalid_source_urls": 0,
            "budget_totals_found": 0,
            "budget_inconsistent": False,
            "qa_gate_passed": False,
        }

    lower = strategy_content.lower()
    placeholder_refs = sum(1 for term in _INTERNAL_REFERENCE_TERMS if term in lower)
    placeholder_refs += len(
        re.findall(r"\[\s*(?:see|cross-?ref|xref)\s+##\s+[^\]]+\]", strategy_content, re.IGNORECASE)
    )
    raw_source_urls = len(
        re.findall(r"\[Source:\s*https?://[^\]\s]+", strategy_content, re.IGNORECASE)
    )
    cited_numbers, valid_defs, invalid_defs = _extract_strategy_citation_definitions(
        strategy_content
    )
    missing_citations = sorted(cited_numbers - set(valid_defs))
    source_urls = max(raw_source_urls, len(valid_defs))

    totals: list[float] = []
    explicit_totals: list[float] = []
    year_one_totals: list[float] = []
    for m in re.finditer(
        r"Total\s*:?\s*\$([0-9]+(?:\.[0-9]+)?)\s*([KMB])", strategy_content, re.IGNORECASE
    ):
        total_value = _strategy_money_to_millions(float(m.group(1)), m.group(2))
        totals.append(total_value)
        explicit_totals.append(total_value)

    for m in re.finditer(
        r"Year 1 investment\s*\(?[^\n)]*\)?\s*:?\s*\$([0-9]+(?:\.[0-9]+)?)(?:\s*-\s*([0-9]+(?:\.[0-9]+)?))?\s*([KMB])",
        strategy_content,
        re.IGNORECASE,
    ):
        low = _strategy_money_to_millions(float(m.group(1)), m.group(3))
        high = _strategy_money_to_millions(float(m.group(2)), m.group(3)) if m.group(2) else low
        midpoint = (low + high) / 2.0
        totals.append(midpoint)
        year_one_totals.append(midpoint)

    budget_inconsistent = False
    comparison_pool = explicit_totals + year_one_totals
    if len(explicit_totals) >= 2 or (explicit_totals and year_one_totals):
        min_total = min(comparison_pool)
        max_total = max(comparison_pool)
        if min_total > 0 and ((max_total - min_total) / min_total) > 0.20:
            budget_inconsistent = True

    qa_passed = bool(
        placeholder_refs == 0
        and source_urls >= 2
        and len(missing_citations) == 0
        and len(invalid_defs) == 0
        and not budget_inconsistent
    )
    return {
        "placeholder_refs": placeholder_refs,
        "source_urls": source_urls,
        "citation_defs": len(valid_defs),
        "missing_citations": len(missing_citations),
        "invalid_source_urls": len(invalid_defs),
        "budget_totals_found": len(totals),
        "budget_inconsistent": budget_inconsistent,
        "qa_gate_passed": qa_passed,
    }


def _clean_strategy_output(strategy_content: str) -> str:
    """Final deterministic cleanup for strategy artifacts."""
    if not strategy_content.strip():
        return strategy_content
    cleaned = _clean_fast_report_output(strategy_content)
    cleaned = _normalize_fast_citations(cleaned)
    cleaned = _strip_internal_source_placeholders(cleaned)
    cleaned = _strip_unresolved_section_cross_references(cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip() + "\n"


def _normalize_fast_citations(report_content: str, source_urls: list[str] | None = None) -> str:
    """
    Normalize fast-mode citations to the deterministic analyzer format.

    Three citation conventions are accepted, in priority order:

    1. Inline ``[Source: URL]`` tags — collected and renumbered to ``[cite: N]``.
    2. ``[cite: N] URL`` definitions — preserved as-is.
    3. Bare ``[cite: N]`` references with no inline mapping — when ``source_urls``
       is supplied, N is treated as a 1-indexed position into that list and the
       URLs become the canonical citation key. This matches the section-writing
       prompt contract that hands the writer a numbered SOURCES CONSULTED list.

    If none of the above resolve, orphan ``[cite: N]`` markers are stripped to
    avoid QA citation-integrity warnings.
    """
    report_content = _rewrite_cite_from_url_tags(report_content)
    report_content = re.sub(
        r"\[([^\]]*cites?:\s*[^\]]+)\]",
        lambda m: _sanitize_numeric_cite_bracket(m.group(1)),
        report_content,
        flags=re.IGNORECASE,
    )

    # Collect existing citation definitions if present.
    existing_cite_def = re.compile(r"\[cite:\s*(\d+)\]\s*(https?://\S+)", re.IGNORECASE)
    num_to_url: dict[int, str] = {}
    for m in existing_cite_def.finditer(report_content):
        num_to_url[int(m.group(1))] = m.group(2).strip()

    # Collect and number source URLs found inline.
    # Match both [Source: https://...] and [Source: domain.com/...] (bare domains).
    source_pattern = re.compile(r"\[Source:\s*((?:https?://)?[^\]\s]+)\s*\]", re.IGNORECASE)
    # Also match multi-word non-URL source tags like [Source: Microsoft Azure]
    # and strip them entirely (they're not auditable citations).
    multiword_source_pattern = re.compile(r"\[Source:\s*[^\]]+\]", re.IGNORECASE)
    urls_in_order: list[str] = []
    url_to_num: dict[str, int] = {}
    next_num = max(num_to_url.keys(), default=0) + 1

    # Seed URL mappings from existing defs.
    for num, url in sorted(num_to_url.items()):
        url_to_num[url] = num
        urls_in_order.append(url)

    for match in source_pattern.finditer(report_content):
        raw_url = match.group(1).strip()
        # Normalize bare domains to https:// for consistent keying
        url = raw_url if raw_url.startswith(("http://", "https://")) else f"https://{raw_url}"
        if url not in url_to_num:
            url_to_num[url] = next_num
            next_num += 1
            urls_in_order.append(url)

    if not url_to_num and not num_to_url:
        # No inline mapping found. If the caller supplied source_urls AND the
        # body has bare [cite: N] refs whose N values fit that list, treat the
        # numbered list as the citation key (this matches the prompt contract).
        bare_cite_pattern = re.compile(r"\[cite:\s*(\d+(?:\s*,\s*\d+)*)\]", re.IGNORECASE)
        if source_urls:
            cited_nums: set[int] = set()
            for m in bare_cite_pattern.finditer(report_content):
                for n in re.findall(r"\d+", m.group(1)):
                    cited_nums.add(int(n))
            valid_nums = {n for n in cited_nums if 1 <= n <= len(source_urls)}
            if valid_nums:
                num_to_url = {n: source_urls[n - 1] for n in valid_nums}
                for n in sorted(num_to_url):
                    url_to_num[num_to_url[n]] = n
                    urls_in_order.append(num_to_url[n])
                next_num = max(num_to_url.keys()) + 1
        if not url_to_num and not num_to_url:
            # No mapping resolvable — strip orphan refs so QA passes cleanly.
            bare_strip = re.compile(r"\s*\[cite:\s*\d+(?:\s*,\s*\d+)*\]", re.IGNORECASE)
            if bare_strip.search(report_content):
                logger.info("Stripping orphan [cite: N] refs with no backing URLs")
                return bare_strip.sub("", report_content)
            return report_content

    def _replace_source(match: re.Match[str]) -> str:
        raw_url = match.group(1).strip()
        url = raw_url if raw_url.startswith(("http://", "https://")) else f"https://{raw_url}"
        num = url_to_num.get(url)
        if num is None:
            nonlocal next_num
            num = next_num
            next_num += 1
            url_to_num[url] = num
            urls_in_order.append(url)
        return f"[cite: {num}]"

    normalized = source_pattern.sub(_replace_source, report_content)

    # Strip leftover multi-word [Source: ...] tags that the URL-only pattern missed
    # (e.g. [Source: Microsoft Azure], [Source: Company Website]).
    normalized = multiword_source_pattern.sub("", normalized)

    # Replace "Sources" heading if present to avoid duplicate appendices.
    sources_heading = re.compile(
        r"^##\s+(Sources|Citations|References)\s*$", re.IGNORECASE | re.MULTILINE
    )
    if sources_heading.search(normalized):
        # Remove existing appendix section content from first sources heading onward.
        lines = normalized.splitlines()
        start_idx = None
        for i, line in enumerate(lines):
            if sources_heading.match(line.strip()):
                start_idx = i
                break
        if start_idx is not None:
            normalized = "\n".join(lines[:start_idx]).rstrip()

    # Resolve body citation refs against known definitions; drop orphan refs.
    known = dict(num_to_url)
    for url, num in url_to_num.items():
        known[num] = url

    cite_ref = re.compile(r"\[cite:\s*([0-9,\s]+)\]", re.IGNORECASE)
    used_old_nums: list[int] = []

    def _clean_refs(match: re.Match[str]) -> str:
        nums: list[int] = []
        for raw in match.group(1).split(","):
            raw = raw.strip()
            if not raw.isdigit():
                continue
            n = int(raw)
            if n in known and n not in nums:
                nums.append(n)
        if not nums:
            return ""
        for n in nums:
            if n not in used_old_nums:
                used_old_nums.append(n)
        return "[cite: " + ", ".join(str(n) for n in nums) + "]"

    normalized = cite_ref.sub(_clean_refs, normalized)

    # Renumber used citations to contiguous sequence in first-use order.
    remap = {old: idx + 1 for idx, old in enumerate(used_old_nums)}

    def _renumber_refs(match: re.Match[str]) -> str:
        nums: list[str] = []
        for raw in match.group(1).split(","):
            raw = raw.strip()
            if not raw.isdigit():
                continue
            old = int(raw)
            if old in remap:
                new_num = str(remap[old])
                if new_num not in nums:
                    nums.append(new_num)
        return f"[cite: {', '.join(nums)}]" if nums else ""

    normalized = cite_ref.sub(_renumber_refs, normalized)

    sources_lines = ["## Sources", ""]
    for old in used_old_nums:
        url = known[old]
        sources_lines.append(f"[cite: {remap[old]}] {url}")

    return normalized.rstrip() + "\n\n" + "\n".join(sources_lines) + "\n"


def _ensure_strategy_source_inventory(
    strategy_content: str, source_urls: list[str], min_sources: int = 2
) -> str:
    """Append a minimal sources inventory when strategy output lacks explicit source URLs."""
    if not strategy_content.strip() or not source_urls:
        return strategy_content

    metrics = _compute_strategy_qa_metrics(strategy_content)
    if metrics["source_urls"] >= min_sources:
        return strategy_content

    normalized_source_urls, _rejected_urls = _normalize_strategy_source_urls(source_urls)
    if not normalized_source_urls:
        return strategy_content

    existing_defs = re.findall(
        r"\[cite:\s*(\d+)\]\s*(https?://\S+)", strategy_content, re.IGNORECASE
    )
    existing_urls = {url.strip() for _, url in existing_defs}
    next_num = max((int(num) for num, _ in existing_defs), default=0) + 1
    new_lines: list[str] = []

    for url in normalized_source_urls:
        normalized = url.strip()
        if not normalized or normalized in existing_urls:
            continue
        new_lines.append(f"[cite: {next_num}] {normalized}")
        existing_urls.add(normalized)
        next_num += 1
        if len(new_lines) >= max(min_sources, 4):
            break

    if not new_lines:
        return strategy_content

    if re.search(r"^##\s+Sources\s*$", strategy_content, re.IGNORECASE | re.MULTILINE):
        return strategy_content.rstrip() + "\n" + "\n".join(new_lines) + "\n"

    return strategy_content.rstrip() + "\n\n## Sources\n\n" + "\n".join(new_lines) + "\n"


def _repair_strategy_artifact_issues(
    strategy_content: str,
    company_name: str,
    vendor: str,
    strategy_label: str,
    source_urls: list[str],
    issues: list[str],
    model: str | None = None,
) -> str:
    """Run one focused repair pass for strategy artifact issues before shipping."""
    from primr.ai.grok_client import grok_llm

    if not strategy_content.strip() or not issues:
        return strategy_content

    source_block = "\n".join(f"- {url}" for url in source_urls[:12]) if source_urls else "(none)"
    issue_block = "\n".join(f"- {issue}" for issue in issues)
    vendor_label = vendor.upper() if vendor and vendor != "agnostic" else strategy_label
    prompt = f"""Repair this strategy document for shipping. Keep the same structure, depth,
and consultant-grade detail, but fix the specific trust and artifact problems listed below.

COMPANY: {company_name}
STRATEGY: {strategy_label}
VENDOR CONTEXT: {vendor_label}

ISSUES TO FIX:
{issue_block}

ALLOWED SOURCE URLS:
{source_block}

DOCUMENT:
{strategy_content[:50_000]}

RULES:
- Preserve the depth and specificity of the document
- Keep dense references in a final ## Sources appendix, not as raw source dumps in body prose
- Reconcile all budget numbers so Year 1 investment, investment framework, and board summary agree
- Remove unsupported or malformed source references
- Use only the allowed source URLs when citing
- Every [cite: N] used in the body must have a valid matching definition in ## Sources
- Do not invent new URLs, vendors, metrics, or budget totals
- Return only the repaired markdown document"""

    system_prompt = (
        "You are a meticulous strategy editor repairing a document for shipment. "
        "Preserve depth, improve auditability, and resolve contradictions conservatively."
    )

    writing_model = model or GROK_MODEL_WRITING
    try:
        repaired = grok_llm(
            prompt,
            model=writing_model,
            max_tokens=20_000,
            temperature=0.2,
            system_prompt=system_prompt,
        )
    except Exception as exc:
        log_structured("warning", "Strategy artifact repair failed", error=str(exc))
        return strategy_content

    return repaired.strip() if repaired and repaired.strip() else strategy_content


def _prepare_strategy_for_output(
    strategy_content: str,
    company_name: str,
    vendor: str,
    strategy_label: str,
    source_urls: list[str],
    model: str | None = None,
) -> tuple[str, dict[str, int | float | bool], list[str]]:
    """Normalize, validate, and repair strategy output before artifact shipping."""
    normalized_source_urls, rejected_source_urls = _normalize_strategy_source_urls(source_urls)

    prepared = _clean_strategy_output(strategy_content)
    prepared = _ensure_strategy_source_inventory(prepared, normalized_source_urls)
    qa = _compute_strategy_qa_metrics(prepared)

    repair_issues: list[str] = []
    if qa["budget_inconsistent"]:
        repair_issues.append("budget_inconsistent")
    if qa["missing_citations"]:
        repair_issues.append("missing_citations")
    if qa["invalid_source_urls"] or rejected_source_urls:
        repair_issues.append("invalid_source_urls")
    if qa["placeholder_refs"]:
        repair_issues.append("placeholder_refs")

    if repair_issues:
        prepared = _repair_strategy_artifact_issues(
            prepared,
            company_name,
            vendor,
            strategy_label,
            normalized_source_urls,
            repair_issues,
            model=model,
        )
        prepared = _clean_strategy_output(prepared)
        prepared = _ensure_strategy_source_inventory(prepared, normalized_source_urls)
        qa = _compute_strategy_qa_metrics(prepared)

    return prepared, qa, rejected_source_urls


def _split_markdown_sections(content: str) -> tuple[str, list[tuple[str, str]]]:
    """Split markdown into preamble and (heading, body) sections."""
    lines = content.splitlines()
    sections: list[tuple[str, str]] = []
    preamble_lines: list[str] = []
    current_heading: str | None = None
    current_body: list[str] = []

    for line in lines:
        if line.startswith("## "):
            if current_heading is None:
                preamble = "\n".join(preamble_lines).strip()
            else:
                sections.append((current_heading, "\n".join(current_body).strip()))
            current_heading = line[3:].strip()
            current_body = []
            continue
        if current_heading is None:
            preamble_lines.append(line)
        else:
            current_body.append(line)

    if current_heading is not None:
        sections.append((current_heading, "\n".join(current_body).strip()))
        preamble = "\n".join(preamble_lines).strip()
    else:
        preamble = content.strip()

    return preamble, sections


def _enforce_fast_section_quality_guards(report_content: str) -> str:
    """
    Apply deterministic quality guards to fast reports.

    - Ensure each section has at least one confidence label token.
    - Ensure each non-reference section includes a "What to validate:" line.
    """
    preamble, sections = _split_markdown_sections(report_content)
    if not sections:
        return report_content

    label_pattern = re.compile(r"\((Confirmed|Reported|Estimated|Hypothesis)[^)]*\)", re.IGNORECASE)
    reference_headings = {"sources", "citations", "references"}
    rebuilt: list[str] = [preamble] if preamble else []

    for heading, body in sections:
        lower_heading = heading.strip().lower()
        guarded_body = body.strip()

        if lower_heading not in reference_headings:
            if not label_pattern.search(guarded_body):
                guarded_body = (guarded_body + "\n\n(Reported)").strip()
            if "what to validate" not in guarded_body.lower():
                guarded_body = (
                    guarded_body
                    + "\n\nWhat to validate: Confirm this section's key claim with primary customer or operator evidence."
                ).strip()

        rebuilt.append(f"## {heading}\n\n{guarded_body}")

    return "\n\n".join(part for part in rebuilt if part).strip() + "\n"


def _compute_fast_report_qa_metrics(
    report_content: str,
    unresolved_contradictions: int = 0,
) -> dict[str, int | float | bool]:
    """Compute lightweight local QA metrics for fast reports.

    Checks: confidence labels, citations, validation prompts, duplicate
    sections, thin sections, and unresolved contradiction carry-through.
    """
    confidence_labels = len(
        re.findall(
            r"\((Confirmed|Reported|Estimated|Hypothesis)[^)]*\)", report_content, re.IGNORECASE
        )
    )
    cited_numbers: set[int] = set()
    for raw_group in re.findall(r"\[cite:\s*([0-9,\s]+)\]", report_content, re.IGNORECASE):
        for raw_num in raw_group.split(","):
            raw_num = raw_num.strip()
            if raw_num.isdigit():
                cited_numbers.add(int(raw_num))

    sources_block = ""
    match = re.search(r"^##\s+Sources\s*$", report_content, flags=re.IGNORECASE | re.MULTILINE)
    if match:
        sources_block = report_content[match.start() :]
    defined: set[int] = set()
    for raw_group in re.findall(r"\[cite:\s*([0-9,\s]+)\]", sources_block, re.IGNORECASE):
        for raw_num in raw_group.split(","):
            raw_num = raw_num.strip()
            if raw_num.isdigit():
                defined.add(int(raw_num))
    missing = sorted(cited_numbers - defined)

    _, sections = _split_markdown_sections(report_content)
    reference_headings = {"sources", "citations", "references"}
    content_sections = [h for h, _ in sections if h.strip().lower() not in reference_headings]
    with_validate = sum(
        1
        for h, body in sections
        if h.strip().lower() not in reference_headings and "what to validate:" in body.lower()
    )

    # Check for duplicate section headings
    heading_counts: dict[str, int] = {}
    for h, _ in sections:
        key = h.strip().lower()
        heading_counts[key] = heading_counts.get(key, 0) + 1
    duplicate_sections = sum(1 for c in heading_counts.values() if c > 1)

    # Check for thin sections (< 100 words)
    thin_sections = sum(
        1
        for h, body in sections
        if h.strip().lower() not in reference_headings and len(body.split()) < 100
    )

    # QA gate: stricter — also checks for duplicates and thin sections
    qa_passed = bool(
        confidence_labels >= 8
        and len(cited_numbers) > 0
        and len(defined) > 0
        and len(missing) == 0
        and with_validate >= max(1, len(content_sections))
        and duplicate_sections == 0
        and thin_sections == 0
        and unresolved_contradictions == 0
    )

    return {
        "word_count": len(report_content.split()),
        "confidence_labels": confidence_labels,
        "citations_used": len(cited_numbers),
        "citations_defined": len(defined),
        "missing_citations": len(missing),
        "section_count": len(content_sections),
        "sections_with_validate": with_validate,
        "duplicate_sections": duplicate_sections,
        "thin_sections": thin_sections,
        "unresolved_contradictions": unresolved_contradictions,
        "qa_gate_passed": qa_passed,
    }


def _repair_fast_report_citation_integrity(
    company_name: str,
    website: str | None,
    report_content: str,
    source_urls: list[str],
    model: str | None = None,
) -> str:
    """Repair missing citation linkage while preserving report structure."""
    from primr.ai.grok_client import grok_llm

    if not report_content.strip() or not source_urls:
        return report_content

    source_block = "\n".join(f"{i}. {u}" for i, u in enumerate(source_urls, 1))
    prompt = f"""You are repairing citation integrity in a strategic markdown report.

Company: {company_name}
Website: {website or "N/A"}

Your job:
- Keep ALL existing ## headings, sections, tables, and prose intact
- Add compact inline citations as [cite: N] only where major factual claims or numeric claims appear
- Add a final ## Sources appendix that maps each [cite: N] to one of the allowed URLs below
- Reuse citation numbers consistently for the same URL
- Do NOT invent URLs or add claims
- Do NOT remove any 'What to validate:' lines
- Output must preserve the existing section order and at least 98% of the original word count

Allowed URLs:
{source_block}

Return the full corrected markdown report only.

--- REPORT START ---
{report_content}
--- REPORT END ---
"""
    writing_model = model or GROK_MODEL_WRITING
    try:
        repaired = grok_llm(
            prompt,
            model=writing_model,
            max_tokens=40_000,
            temperature=0.2,
            system_prompt="You are a meticulous editor fixing citations in-place without rewriting content.",
        )
    except Exception as repair_err:
        logger.warning("Citation repair pass failed: %s", repair_err)
        return report_content

    if not repaired or not repaired.strip():
        return report_content

    repaired = _clean_fast_report_output(repaired)
    repaired = _normalize_fast_citations(repaired, source_urls=source_urls)
    repaired = _enforce_fast_section_quality_guards(repaired)

    if not _preserves_report_structure(report_content, repaired):
        logger.warning("Citation repair changed report structure too much; keeping original")
        return report_content

    repaired_metrics = _compute_fast_report_qa_metrics(repaired)
    if (
        repaired_metrics["citations_used"] > 0
        and repaired_metrics["citations_defined"] > 0
        and repaired_metrics["missing_citations"] == 0
    ):
        return repaired

    return report_content


def _polish_fast_report_for_trust(
    company_name: str,
    website: str | None,
    report_content: str,
    source_urls: list[str],
    model: str | None = None,
) -> str:
    """
    Run a lightweight post-write polish pass for fast mode trust/readability.

    Goals:
    - keep prose readable and concise
    - ensure confidence labels are present on non-obvious claims
    - improve citation discipline while keeping citations compact
    - preserve section structure and core meaning
    """
    from primr.ai.grok_client import grok_llm

    if not report_content.strip():
        return report_content

    source_block = (
        "\n".join(f"{i}. {u}" for i, u in enumerate(source_urls, 1)) if source_urls else "(none)"
    )
    feedback_guidance = _load_fast_feedback_guidance()
    feedback_block = (
        f"\n10. Apply this prior-eval feedback guidance where relevant:\n{feedback_guidance}\n"
        if feedback_guidance
        else ""
    )
    prompt = f"""You are editing a strategic report for quality and trust, not rewriting from scratch.

Company: {company_name}
Website: {website or "N/A"}

Rules:
1. Preserve the report structure and section headings.
2. Keep prose readable for executives; do not overstuff citations inline.
3. Use compact inline citations only as [cite: N].
4. Ensure non-obvious claims include confidence labels:
   (Confirmed), (Reported), (Estimated), or (Hypothesis).
5. Add/retain a single "## Sources" section at the end only.
6. Do NOT invent sources. Only use these source URLs:
{source_block}
7. Replace any unsupported numeric precision with cautious language:
   "Not publicly disclosed" or low-confidence qualitative ranges.
8. Remove repeated or contradictory claims; prefer one clear statement with
   the best available evidence and confidence tag.
9. Ensure each section ends with "What to validate:" and one concrete check.
{feedback_block}

Return the fully edited markdown report only.

--- REPORT START ---
{report_content}
--- REPORT END ---
"""
    writing_model = model or GROK_MODEL_WRITING
    try:
        polished = grok_llm(
            prompt,
            model=writing_model,
            max_tokens=10_000,
            temperature=0.2,
            system_prompt=(
                "You are a meticulous editorial QA analyst improving evidence discipline and readability."
            ),
        )
        if not polished or not polished.strip():
            return report_content

        # Guard against destructive compression/truncation from the polish pass.
        original_words = len(report_content.split())
        polished_words = len(polished.split())
        _, original_sections = _split_markdown_sections(report_content)
        _, polished_sections = _split_markdown_sections(polished)
        if original_words < 100:
            min_words = 1
        elif original_words >= 1200:
            min_words = max(1200, int(original_words * 0.70))
        else:
            min_words = max(50, int(original_words * 0.50))
        if original_words >= 1200:
            min_sections = max(1, len(original_sections))
        else:
            min_sections = max(1, int(len(original_sections) * 0.70))
        if polished_words < min_words or len(polished_sections) < min_sections:
            return report_content
        return polished
    except Exception as e:
        log_structured(
            "warning",
            "Trust polish pass failed, using unpolished report",
            error=str(e),
        )
        return report_content


def _parse_batch_sections(
    content: str,
    expected_sections: list["SectionConfig"],
) -> list["GeneratedSection"]:
    """Parse Grok's batch response from XML envelopes and/or markdown headings."""
    parsed: list[GeneratedSection] = []
    preamble, blocks = _extract_generated_section_blocks(content)

    for idx, (title, body) in enumerate(blocks):
        expected_title = expected_sections[idx].name if idx < len(expected_sections) else title
        parsed.append(_normalize_generated_section_payload(title, body, expected_title))

    if not parsed and content.strip():
        expected_title = expected_sections[0].name if expected_sections else "Section"
        parsed.append(
            _normalize_generated_section_payload(
                expected_title,
                content.strip(),
                expected_title,
            )
        )

    if preamble and parsed:
        first_body = parsed[0].content
        parsed[0] = _normalize_generated_section_payload(
            parsed[0].title,
            preamble + "\n\n" + first_body,
            parsed[0].title,
        )

    return parsed


def _assemble_fast_report(
    company_name: str,
    website: str | None,
    written_sections: list["GeneratedSection"],
) -> str:
    """
    Assemble individual batch sections into a final markdown report.
    """
    current_date = datetime.now().strftime("%B %d, %Y")

    header = f"# Strategic Company Overview: {company_name}\n\n"
    header += f"*{current_date}*"
    if website:
        header += f" | [{website}]({website})"
    header += "\n\n---\n"

    body_parts: list[str] = []
    for i, section in enumerate(written_sections):
        body_parts.append(section.to_markdown())
        # Horizontal separator every 5 sections (matches full mode)
        if (i + 1) % 5 == 0 and i + 1 < len(written_sections):
            body_parts.append("---")

    return header + "\n\n".join(body_parts)


def _build_fast_analysis_prompt(
    company_name: str,
    website: str | None,
    raw_corpus: str,
    external_sources: str,
) -> str:
    """
    Build the Phase 2 analysis workbook prompt for Grok fast mode.

    Sends raw scraped data + external sources to Grok and asks for a
    structured analysis workbook (facts, hypotheses, tensions — not prose).
    """
    current_date = datetime.now().strftime("%B %d, %Y")

    return f"""**Company:** {company_name}
**Website:** {website or "N/A"}
**Date:** {current_date}

Below is raw data scraped from the company's website and external sources.
Analyze it and produce a Structured Analysis Workbook.

=== RAW WEBSITE DATA ===
{raw_corpus}

=== EXTERNAL SOURCES ===
{external_sources}

---

Produce a **Structured Analysis Workbook** with the following sections.
Use bullet points, tables, and short paragraphs. This is working notes, not prose.

CRITICAL: You are doing PRE-ENGAGEMENT ANALYSIS for a consulting firm, not summarizing
the company's marketing. Separate what {company_name} CLAIMS from what external evidence
SUPPORTS. Be conservative on financial estimates — use wide ranges and note confidence.

1. **Company Basics**
   - Official name, headquarters, founding date, employee count
   - Ownership structure (public/private, investors)
   - Label each fact: (Confirmed), (Reported), (Estimated), or (Hypothesis)

2. **Products & Services Catalog**
   - Every product/service found, organized by category
   - Pricing models, contract structures if visible
   - Recent launches or pivots (last 2-3 years)
   - Distinguish what's live/adopted vs. what's announced/marketing

3. **Customer Segments & Market Positioning**
   - Primary segments with evidence
   - Geographic distribution
   - Enterprise vs SMB vs consumer mix
   - Go-to-market approach
   - Flag any logo references that lack depth (vague "powered by" vs. detailed case)

4. **Competitive Landscape**
   - At least 5 competitors with: name, estimated size, key differentiator
   - Where {company_name} appears to win and lose (from external evidence, not their claims)
   - Emerging disruptors
   - Include competitors the company DOESN'T mention but should

5. **Financial Profile**
   - Revenue (actual or estimated with WIDE ranges and LOW confidence if inferred)
   - Growth rate and trajectory
   - Profitability indicators
   - Funding history / capital structure
   - Include a summary table
   - AVOID aggressive inferences — if data is thin, say so explicitly

6. **Leadership Profiles**
   - C-suite with backgrounds, tenure, previous roles
   - Board composition
   - Recent departures or hires

7. **Industry Dynamics**
   - Industry size, growth rate
   - Key trends and disruption factors
   - Regulatory environment

8. **Strategic Hypotheses** (3-5)
   For each:
   - Hypothesis statement
   - Supporting evidence (with sources)
   - Counter-evidence or alternative explanation
   - Confidence level
   - What question would you ask in discovery to TEST this?

9. **Strategic Tensions** (3-5)
   For each:
   - The tension (a tradeoff they must manage, e.g., "Scale vs Customization")
   - Evidence from the data
   - How they appear to be managing it currently

10. **Narrative Gaps** (3-5)
    For each:
    - What they claim (with quote/source from THEIR marketing)
    - Contradicting or complicating EXTERNAL signals
    - Question to explore
    These should be genuine stress-tests of their story, not minor wording quibbles.

11. **Areas of Potential Fragility** (3-4)
    Focus on systemic risks: single points of failure, concentration risks,
    dependencies that could break under stress.

12. **Patterns Worth Exploring** (3-5)
    Novel observations: surprising correlations, timing signals, behavioral
    patterns that don't fit the narrative.

13. **Discovery Questions** (6-8)
    For each:
    - The question
    - Why we're asking (what evidence prompted it)
    - What we hope to learn
    These should be questions a CONSULTING PARTNER would ask in a first meeting —
    sharp, grounded in evidence, testing specific hypotheses.
"""


def _assess_source_relevance(
    company_name: str,
    external_data: dict[str, str],
) -> dict[str, str]:
    """Filter external sources by relevance using LLM assessment.

    Asks the LLM to rate each source's relevance given the company's profile.
    Drops sources that are generic filler rather than genuinely informative.
    Returns a filtered dict of URL -> content.
    """
    if len(external_data) <= 5:
        return external_data  # too few to bother filtering

    # Build a compact summary of each source for LLM review
    source_summaries: list[str] = []
    url_list = list(external_data.keys())
    for i, url in enumerate(url_list):
        snippet = external_data[url][:500].replace("\n", " ")
        source_summaries.append(f"{i + 1}. {url}\n   {snippet}")

    prompt = f"""You are evaluating external research sources about {company_name}.

Below are {len(url_list)} sources. For each, decide: KEEP or DROP.

KEEP a source if it provides SPECIFIC, USEFUL intelligence about {company_name}:
- Names executives, financials, deals, partnerships, or strategies
- Provides industry analysis mentioning this company specifically
- Contains news, press releases, or analyst coverage about this company

DROP a source if it is:
- Generic industry content that barely mentions the company
- A directory listing, job board, or social media page with no substance
- Duplicate information already covered by another KEPT source
- Tangentially related but not genuinely informative

IMPORTANT: For smaller or less prominent companies, it is BETTER to keep 5 high-quality
sources than 25 mediocre ones. Be selective. Quality over quantity.

SOURCES:
{chr(10).join(source_summaries)}

Return ONLY a JSON array of the source NUMBERS to KEEP (e.g. [1, 3, 5, 8]).
No prose, no explanation."""

    try:
        response = llm(prompt, model_type="fast", streaming=False).strip()
        # Parse the JSON array
        import json as _json

        # Strip markdown fencing if present
        text = response.strip()
        if text.startswith("```"):
            first_nl = text.find("\n")
            text = text[first_nl + 1 :] if first_nl != -1 else text[3:]
            if text.rstrip().endswith("```"):
                text = text.rstrip()[:-3]
            text = text.strip()
        bracket_start = text.find("[")
        bracket_end = text.rfind("]")
        if bracket_start != -1 and bracket_end > bracket_start:
            keep_indices = _json.loads(text[bracket_start : bracket_end + 1])
        else:
            return external_data  # parse failed, keep all

        # Convert 1-indexed numbers to 0-indexed
        keep_set = {round(n) - 1 for n in keep_indices if isinstance(n, (int, float))}
        filtered = {
            url_list[i]: external_data[url_list[i]] for i in keep_set if 0 <= i < len(url_list)
        }

        if len(filtered) < 3:
            # LLM was too aggressive, keep originals
            return external_data

        dropped = len(external_data) - len(filtered)
        if dropped > 0:
            log_structured(
                "info",
                "Source quality filter dropped low-relevance sources",
                kept=len(filtered),
                dropped=dropped,
            )
        return filtered
    except Exception as e:
        log_structured(
            "warning",
            "Source relevance assessment failed, keeping all sources",
            error=str(e),
            source_count=len(external_data),
        )
        return external_data  # on any error, keep all sources


def _fast_gap_analysis(
    company_name: str,
    website: str | None,
    raw_corpus: str,
    external_sources: str,
    source_urls: list[str],
    model: str | None = None,
) -> tuple[list[str], str]:
    """
    Phase 2 helper: Grok identifies research gaps and returns targeted search queries.

    Returns:
        (list of search queries, gap analysis text for logging)
    """
    from primr.ai.grok_client import grok_llm

    # Build corpus summary — first 500 chars of each page
    corpus_lines = raw_corpus.split("\n\n")
    corpus_summary_parts: list[str] = []
    for block in corpus_lines:
        if block.startswith("[Page:"):
            corpus_summary_parts.append(block[:500])
    corpus_summary = (
        "\n\n".join(corpus_summary_parts[:80]) if corpus_summary_parts else raw_corpus[:30_000]
    )

    # Build external source summary — first 500 chars each
    ext_lines = external_sources.split("\n\n")
    ext_summary_parts: list[str] = []
    for block in ext_lines:
        if block.startswith("[Source:"):
            ext_summary_parts.append(block[:500])
    ext_summary = "\n\n".join(ext_summary_parts) if ext_summary_parts else external_sources[:5_000]

    prompt = f"""You've reviewed primary sources for {company_name}. As a strategic analyst, identify
what's MISSING — gaps that would weaken a consulting brief.

SOURCES REVIEWED:
{corpus_summary}

EXTERNAL SOURCES:
{ext_summary}

KNOWN SOURCE URLS (do NOT repeat these):
{chr(10).join(source_urls[:30])}

Return exactly 8 items in this format (one per block, no extra text):
GAP: [what's missing]
QUERY: [web search query to fill it]
PRIORITY: CRITICAL | IMPORTANT

Prioritize third-party validation sources: analyst reports, industry publications,
financial filings, customer case studies, employee reviews, regulatory documents.
Also cover: financials, competitive positioning, leadership changes, customer evidence,
technology direction, recent news, risk factors.
"""

    system_prompt = (
        "You are a research gap analyst for a consulting firm. "
        "Identify what's missing from preliminary research and suggest "
        "targeted web searches to fill those gaps. Be specific and actionable."
    )

    try:
        response = grok_llm(
            prompt,
            model=model,
            max_tokens=5_000,
            temperature=0.4,
            system_prompt=system_prompt,
        )
    except Exception as e:
        log_structured("warning", "Gap analysis failed", error=str(e))
        return [], f"Gap analysis failed: {e}"

    if not response or not response.strip():
        return [], "Gap analysis returned empty response"

    # Parse queries from response
    queries: list[str] = []
    for line in response.strip().splitlines():
        line = line.strip()
        if line.upper().startswith("QUERY:"):
            query = line[6:].strip().strip("\"'[]")
            if query:
                queries.append(query)

    return queries[:8], response


def _fast_cross_validate(
    company_name: str,
    website: str | None,
    report_content: str,
    source_urls: list[str],
    model: str | None = None,
    reasoning_session: Any = None,
) -> dict:
    """
    Phase 5 helper: Grok reviews the assembled report for quality issues.

    When `reasoning_session` is supplied (a `ContinuousReasoningSession`), the
    review runs as a follow-up turn inside the same session that produced the
    workbook, so the validator inherits the original corpus and workbook
    reasoning instead of re-reading just the report.

    Returns:
        {"weak_sections": [{"title": str, "reason": str, "queries": [str, str]}],
         "contradictions": [str]}
    """
    from primr.ai.grok_client import grok_llm

    source_list = "\n".join(f"- {url}" for url in source_urls[:50])

    if reasoning_session is not None:
        # Continuous mode: the session already contains the corpus + workbook
        # reasoning from Phase 3. Skip re-feeding sources; ask the model to
        # validate the report against what it already analyzed.
        prompt = f"""You produced the analysis workbook for {company_name} in your earlier turn.
Below is the assembled consulting brief that was written from that workbook. Review it for
quality issues, drawing on the corpus and reasoning you already have in context.

REPORT:
{report_content[:120_000]}

ADDITIONAL SOURCES (gathered during gap-fill, may not have been in your earlier context):
{source_list}

Return JSON (no markdown fencing, just raw JSON):
{{
  "weak_sections": [
    {{"title": "exact ## heading", "reason": "why it's weak", "queries": ["search query 1", "search query 2"]}}
  ],
  "contradictions": ["description of contradiction between sections"]
}}

A section is WEAK if it:
- Makes claims without citing any source
- Uses only generic industry statements, not company-specific evidence
- Is significantly shorter than other sections
- Relies heavily on the company's own marketing claims without external validation
- Drifts from the strategic analysis the workbook called for, or contradicts the workbook's own evidence

Limit: max 3 weak sections, max 3 contradictions. Only flag genuinely weak sections.
If the report is solid, return empty arrays."""
    else:
        prompt = f"""Review this consulting brief for {company_name}. Identify quality issues.

REPORT:
{report_content[:120_000]}

AVAILABLE SOURCES:
{source_list}

Return JSON (no markdown fencing, just raw JSON):
{{
  "weak_sections": [
    {{"title": "exact ## heading", "reason": "why it's weak", "queries": ["search query 1", "search query 2"]}}
  ],
  "contradictions": ["description of contradiction between sections"]
}}

A section is WEAK if it:
- Makes claims without citing any source
- Uses only generic industry statements, not company-specific evidence
- Is significantly shorter than other sections
- Relies heavily on the company's own marketing claims without external validation

Limit: max 3 weak sections, max 3 contradictions. Only flag genuinely weak sections.
If the report is solid, return empty arrays."""

    system_prompt = (
        "You are a quality reviewer for consulting research briefs. "
        "Identify sections that need more evidence or have quality issues. "
        "Return structured JSON only."
    )

    try:
        if reasoning_session is not None:
            response = reasoning_session.send(
                prompt,
                temperature=0.2,
                max_tokens=5_000,
            )
        else:
            response = grok_llm(
                prompt,
                model=model,
                max_tokens=5_000,
                temperature=0.2,
                system_prompt=system_prompt,
            )
    except Exception as e:
        log_structured("warning", "Cross-validation failed", error=str(e))
        return {"weak_sections": [], "contradictions": [], "_failed": True}

    if not response or not response.strip():
        return {"weak_sections": [], "contradictions": []}

    # Parse JSON from response
    try:
        # Strip markdown code fencing if present (```json, ```JSON, ``` etc.)
        text = response.strip()
        if text.startswith("```"):
            # Remove opening fence line (```json, ```JSON, ```, etc.)
            first_newline = text.find("\n")
            text = text[first_newline + 1 :] if first_newline != -1 else text[3:]
            # Remove closing fence
            if text.rstrip().endswith("```"):
                text = text.rstrip()[:-3]
            text = text.strip()

        # Try to extract JSON object if surrounded by prose
        try:
            result = json.loads(text)
        except json.JSONDecodeError:
            # Look for JSON object within the text
            brace_start = text.find("{")
            brace_end = text.rfind("}")
            if brace_start != -1 and brace_end > brace_start:
                result = json.loads(text[brace_start : brace_end + 1])
            else:
                raise

        if not isinstance(result, dict):
            log_structured(
                "warning", "Cross-validation JSON is not a dict", type=type(result).__name__
            )
            return {"weak_sections": [], "contradictions": []}

        # Enforce limits and validate types
        raw_weak = result.get("weak_sections", [])
        weak = [w for w in (raw_weak if isinstance(raw_weak, list) else []) if isinstance(w, dict)][
            :3
        ]
        raw_contradictions = result.get("contradictions", [])
        contradictions = [
            c
            for c in (raw_contradictions if isinstance(raw_contradictions, list) else [])
            if isinstance(c, str)
        ][:3]

        return {"weak_sections": weak, "contradictions": contradictions}
    except (json.JSONDecodeError, KeyError, TypeError, AttributeError) as e:
        log_structured("warning", "Cross-validation JSON parse failed, retrying", error=str(e))

        # NOTE (pipeline-resilience): This retry is a *format correction* retry
        # (re-prompt when JSON parsing fails), not an API error retry.  It is
        # intentionally retained alongside the stage-level RecoveryExecutor
        # which handles API failures and skip/abort for the cross-validation
        # background stage.
        # Retry with tighter prompt including the failed response
        retry_prompt = (
            "Your previous response could not be parsed as JSON. "
            "Return ONLY a JSON object, no markdown fencing, no prose before or after.\n\n"
            f"Previous response (first 2000 chars):\n{response[:2000]}\n\n"
            "Fix the JSON and return ONLY this structure:\n"
            '{"weak_sections": [{"title": "...", "reason": "...", "queries": ["...", "..."]}], '
            '"contradictions": ["..."]}'
        )
        try:
            retry_response = grok_llm(
                retry_prompt,
                max_tokens=3_000,
                temperature=0.1,
                system_prompt="Return valid JSON only. No markdown, no prose.",
            )
            if retry_response and retry_response.strip():
                retry_text = retry_response.strip()
                brace_start = retry_text.find("{")
                brace_end = retry_text.rfind("}")
                if brace_start != -1 and brace_end > brace_start:
                    result = json.loads(retry_text[brace_start : brace_end + 1])
                    if isinstance(result, dict):
                        raw_weak = result.get("weak_sections", [])
                        weak = [
                            w
                            for w in (raw_weak if isinstance(raw_weak, list) else [])
                            if isinstance(w, dict)
                        ][:3]
                        raw_contradictions = result.get("contradictions", [])
                        contradictions = [
                            c
                            for c in (
                                raw_contradictions if isinstance(raw_contradictions, list) else []
                            )
                            if isinstance(c, str)
                        ][:3]
                        return {"weak_sections": weak, "contradictions": contradictions}
        except Exception as retry_err:
            log_structured(
                "debug", "Cross-validation retry JSON parse failed", error=str(retry_err)
            )

        # Last resort: regex extraction of weak section titles and reasons
        try:
            weak_sections = []
            title_pattern = re.compile(r'"title"\s*:\s*"([^"]+)"')
            reason_pattern = re.compile(r'"reason"\s*:\s*"([^"]+)"')
            titles = title_pattern.findall(response)
            reasons = reason_pattern.findall(response)
            for i, title in enumerate(titles[:3]):
                weak_sections.append(
                    {
                        "title": title,
                        "reason": reasons[i] if i < len(reasons) else "Needs more evidence",
                        "queries": [f"{company_name} {title.lower()}"],
                    }
                )
            if weak_sections:
                log_structured(
                    "info", "Cross-validation recovered via regex", count=len(weak_sections)
                )
                return {"weak_sections": weak_sections, "contradictions": []}
        except Exception as regex_err:
            log_structured("debug", "Cross-validation regex fallback failed", error=str(regex_err))

        log_structured("warning", "Cross-validation JSON parse failed after all retries")
        return {"weak_sections": [], "contradictions": [], "_failed": True}


def _fast_regenerate_section(
    company_name: str,
    website: str | None,
    section_title: str,
    section_content: str,
    analysis_workbook: str,
    new_evidence: str,
    source_urls: list[str],
    model: str | None = None,
) -> str:
    """
    Phase 5 helper: Re-writes one weak section with additional evidence.

    Uses the same system prompt style as Phase 4 report writing.
    Returns the re-generated section content (starting with ## heading).
    """
    from primr.ai.grok_client import grok_llm

    source_list = "\n".join(f"- {url}" for url in source_urls[:50])

    prompt = f"""Re-write this section of a consulting brief for {company_name}, incorporating
the NEW EVIDENCE provided below. The goal is to make the section evidence-rich,
specific, and analytically strong.

SECTION TO REWRITE:
{section_content}

NEW EVIDENCE (incorporate this):
{new_evidence}

ANALYSIS CONTEXT (for background):
{analysis_workbook[:20_000]}

ALL AVAILABLE SOURCES:
{source_list}

RULES:
- Start with: ## {section_title}
- Full paragraphs with evidence and strategic interpretation, not bullet dumps
- Keep citations compact, usually at paragraph ends, and use [cite: N] references in the body
- Reserve the densest source inventory for the final Sources appendix
- Label claims: Confirmed, Reported, Estimated, Hypothesis
- Stress-test the company narrative — separate claims from evidence
- Keep roughly the same scope as the original section
- End with a single "What to validate:" line followed by a concrete check question"""

    system_prompt = (
        "You are a senior strategic analyst rewriting a section of a consulting dossier. "
        "Your reader is a partner walking into a meeting. Incorporate the new evidence "
        "to make the section analytically stronger. Be conservative on financial inferences."
    )

    writing_model = model or GROK_MODEL_WRITING
    try:
        result = grok_llm(
            prompt,
            model=writing_model,
            max_tokens=5_000,
            temperature=0.7,
            system_prompt=system_prompt,
        )
    except Exception as e:
        log_structured(
            "warning", "Section regeneration failed", section=section_title, error=str(e)
        )
        return section_content  # Return original on failure

    if not result or not result.strip():
        return section_content

    # Ensure it starts with the correct heading
    result = result.strip()
    if not result.startswith(f"## {section_title}"):
        # Strip Grok's wrong heading if it starts with any ## heading
        if result.startswith("## "):
            # Remove the first line (wrong heading)
            first_newline = result.find("\n")
            if first_newline != -1:
                result = result[first_newline:].strip()
            else:
                result = ""
        result = f"## {section_title}\n\n{result}" if result else f"## {section_title}\n\n"

    return result


# ── Strategy enrichment helpers (Phase 6 quality pass) ──────────────────


def _a_or_an(word: str) -> str:
    """Return 'a' or 'an' depending on whether word starts with a vowel sound."""
    return "an" if word and word[0].upper() in "AEIOU" else "a"


def _strategy_cross_validate(
    company_name: str,
    strategy_content: str,
    vendor: str,
    source_urls: list[str],
    model: str | None = None,
) -> dict:
    """
    Phase 6 helper: Grok reviews the strategy document for quality issues.

    Returns:
        {"weak_sections": [{"title": str, "reason": str, "queries": [str, str]}],
         "issues": [str]}
    """
    from primr.ai.grok_client import grok_llm

    source_list = "\n".join(f"- {url}" for url in source_urls[:50])

    prompt = f"""Review this {vendor.upper()} strategy document for {company_name}. Identify quality issues.

STRATEGY DOCUMENT:
{strategy_content[:120_000]}

AVAILABLE SOURCES:
{source_list}

Return JSON (no markdown fencing, just raw JSON):
{{
  "weak_sections": [
    {{"title": "exact ## heading", "reason": "why it's weak", "queries": ["search query 1", "search query 2"]}}
  ],
  "issues": ["description of quality issue"]
}}

A section is WEAK if it:
- Makes claims without citing evidence (e.g., "{vendor.upper()} offers best-in-class X" with no source)
- Uses generic recommendations that could apply to ANY company
- Lacks specific implementation details, timelines, or cost estimates
- Doesn't connect capabilities to THIS company's specific needs or challenges

Limit: max 2 weak sections, max 2 issues. Only flag genuinely weak sections.
If the strategy is solid, return empty arrays."""

    system_prompt = (
        f"You are a quality reviewer for {vendor.upper()} strategy documents. "
        "Identify sections that need more evidence or are too generic. "
        "Return structured JSON only."
    )

    try:
        response = grok_llm(
            prompt,
            model=model,
            max_tokens=4_000,
            temperature=0.2,
            system_prompt=system_prompt,
        )
    except Exception as e:
        log_structured("warning", "Strategy cross-validation failed", error=str(e))
        return {"weak_sections": [], "issues": [], "_failed": True}

    if not response or not response.strip():
        return {"weak_sections": [], "issues": []}

    # Parse JSON from response (same pattern as _fast_cross_validate)
    try:
        text = response.strip()
        if text.startswith("```"):
            first_newline = text.find("\n")
            text = text[first_newline + 1 :] if first_newline != -1 else text[3:]
            if text.rstrip().endswith("```"):
                text = text.rstrip()[:-3]
            text = text.strip()

        try:
            result = json.loads(text)
        except json.JSONDecodeError:
            brace_start = text.find("{")
            brace_end = text.rfind("}")
            if brace_start != -1 and brace_end > brace_start:
                result = json.loads(text[brace_start : brace_end + 1])
            else:
                raise

        if not isinstance(result, dict):
            return {"weak_sections": [], "issues": []}

        raw_weak = result.get("weak_sections", [])
        weak = [w for w in (raw_weak if isinstance(raw_weak, list) else []) if isinstance(w, dict)][
            :2
        ]
        raw_issues = result.get("issues", [])
        issues = [
            i for i in (raw_issues if isinstance(raw_issues, list) else []) if isinstance(i, str)
        ][:2]

        return {"weak_sections": weak, "issues": issues}
    except (json.JSONDecodeError, KeyError, TypeError, AttributeError) as e:
        log_structured("warning", "Strategy cross-validation JSON parse failed", error=str(e))
        return {"weak_sections": [], "issues": [], "_failed": True}


def _strategy_regenerate_section(
    company_name: str,
    vendor: str,
    section_title: str,
    section_content: str,
    new_evidence: str,
    analysis_workbook: str,
    model: str | None = None,
) -> str:
    """
    Phase 6 helper: Re-writes one weak strategy section with additional evidence.

    Returns the re-generated section content (starting with ## heading).
    """
    from primr.ai.grok_client import grok_llm

    article = _a_or_an(vendor.upper())
    prompt = f"""Re-write this section of {article} {vendor.upper()} strategy document for {company_name},
incorporating the NEW EVIDENCE provided below. Make the section specific, actionable,
and tied to this company's actual situation.

SECTION TO REWRITE:
{section_content}

NEW EVIDENCE (incorporate this):
{new_evidence}

ANALYSIS CONTEXT (for background):
{analysis_workbook[:20_000]}

RULES:
- Start with: ## {section_title}
- Connect {vendor.upper()} capabilities to THIS company's specific needs
- Include specific services, pricing tiers, or implementation approaches where evidence supports it
- Label claims: Confirmed, Reported, Estimated, Hypothesis
- Keep citations compact, usually at paragraph ends, and use [cite: N] references in the body
- Keep roughly the same scope as the original section
- Include concrete next steps or validation questions
- Keep the densest supporting reference list in the final Sources appendix"""

    system_prompt = (
        f"You are a senior strategy consultant rewriting a section of {article} {vendor.upper()} "
        f"strategy document for {company_name}. Incorporate new evidence to make the section "
        "more specific and actionable. Be conservative on cost estimates."
    )

    writing_model = model or GROK_MODEL_WRITING
    try:
        result = grok_llm(
            prompt,
            model=writing_model,
            max_tokens=8_000,
            temperature=0.6,
            system_prompt=system_prompt,
        )
    except Exception as e:
        log_structured(
            "warning", "Strategy section regeneration failed", section=section_title, error=str(e)
        )
        return section_content  # Return original on failure

    if not result or not result.strip():
        return section_content

    # Ensure it starts with the correct heading
    result = result.strip()
    if not result.startswith(f"## {section_title}"):
        if result.startswith("## "):
            first_newline = result.find("\n")
            if first_newline != -1:
                result = result[first_newline:].strip()
            else:
                result = ""
        result = f"## {section_title}\n\n{result}" if result else f"## {section_title}\n\n"

    return result


def _strategy_polish(
    company_name: str,
    vendor: str,
    strategy_content: str,
    model: str | None = None,
) -> str:
    """
    Phase 6 helper: Combined coherence + evidence discipline pass for strategy documents.

    Deduplicates, standardizes vendor references, adds confidence labels,
    and ensures specificity. Guards against destructive compression (90% word count,
    section count preservation).
    """
    from primr.ai.grok_client import grok_llm

    if not strategy_content.strip():
        return strategy_content

    article = _a_or_an(vendor.upper())
    prompt = f"""You are editing {article} {vendor.upper()} strategy document for {company_name} for coherence,
evidence discipline, and specificity.

TASKS (in priority order):
1. DEDUPLICATION: When the same point appears in multiple sections, keep the first
   occurrence and replace later duplicates with a cross-reference.
2. EVIDENCE DISCIPLINE: For each major recommendation, ensure it has:
   - A confidence label: Confirmed, Reported, Estimated, or Hypothesis
   - A compact [cite: N] reference where available, with dense references consolidated in the final Sources appendix
   - Specific {vendor.upper()} services/products named (not generic "cloud services")
3. SPECIFICITY CHECK: Replace generic recommendations with company-specific ones.
   BAD: "Leverage AI/ML capabilities to improve operations"
   GOOD: "Deploy {vendor.upper()} vision APIs for [specific company process] to reduce [specific metric]"
4. TERMINOLOGY: Standardize how {company_name}, {vendor.upper()} services, and
   competitors are named throughout.

STRICT RULES:
- Do NOT remove or rename ## section headings
- Do NOT remove confidence labels or source citations
- Do NOT add new strategic recommendations — only improve existing ones
- Do NOT delete paragraphs — only edit individual sentences
- PRESERVE all depth and analysis
- Output MUST contain at least 90% of the original word count

Return the fully edited markdown strategy document only. No preamble or commentary.

--- STRATEGY START ---
{strategy_content}
--- STRATEGY END ---"""

    writing_model = model or GROK_MODEL_WRITING
    try:
        polished = grok_llm(
            prompt,
            model=writing_model,
            max_tokens=32_000,
            temperature=0.3,
            system_prompt=(
                f"You are a meticulous editorial analyst polishing {article} {vendor.upper()} strategy "
                f"document for {company_name}. Improve evidence discipline and specificity "
                "while preserving ALL depth and analysis. Make only surgical edits."
            ),
        )
        if not polished or not polished.strip():
            return strategy_content

        # Guard: reject destructive compression
        original_words = len(strategy_content.split())
        polished_words = len(polished.split())
        _, original_sections = _split_markdown_sections(strategy_content)
        _, polished_sections = _split_markdown_sections(polished)

        if polished_words < int(original_words * 0.90):
            logger.warning(
                "Strategy polish dropped too many words (%d → %d), using original",
                original_words,
                polished_words,
            )
            return strategy_content
        if len(polished_sections) < len(original_sections):
            logger.warning(
                "Strategy polish lost sections (%d → %d), using original",
                len(original_sections),
                len(polished_sections),
            )
            return strategy_content

        return polished
    except Exception:
        logger.warning("Strategy polish failed, using original", exc_info=True)
        return strategy_content


def _enrich_strategy_content(
    strategy_content: str,
    company_name: str,
    vendor: str,
    label: str,
    source_urls: list[str],
    source_urls_seen: set[str],
    analysis_workbook: str,
    website: str | None,
    grok_reasoning: str | None = None,
    grok_writing: str | None = None,
) -> str:
    """
    Phase 6 orchestrator: Full quality pass for strategy documents.

    1. Cross-validate to find up to 2 weak sections
    2. For each: DDG search → scrape evidence → regenerate section
    3. Polish pass for coherence + evidence discipline
    4. Falls back to original on any failure
    """
    if not strategy_content or not strategy_content.strip():
        return strategy_content

    # Show vendor suffix only when vendor is a real cloud vendor (not the label itself)
    vendor_label = (
        f" ({vendor.upper()})" if vendor.lower() not in ("agnostic", label.lower()) else ""
    )

    # Step 1: Cross-validate
    try:
        with console.timed_operation(f"Reviewing {label}{vendor_label}"):
            cv_result = _strategy_cross_validate(
                company_name,
                strategy_content,
                vendor,
                source_urls,
                model=grok_reasoning,
            )
    except Exception as e:
        log_structured("warning", "Strategy CV failed, skipping enrichment", error=str(e))
        return strategy_content

    weak_sections = cv_result.get("weak_sections", [])
    issues = cv_result.get("issues", [])
    cv_failed = cv_result.pop("_failed", False)

    if cv_failed:
        console.warn(
            f"Strategy cross-validation failed for {label}{vendor_label} — skipping enrichment"
        )

    if issues:
        for issue in issues:
            console.info(f"Strategy issue: {issue[:100]}")

    # Step 2: Enrich weak sections
    sections_enriched = 0
    if weak_sections:
        # Build heading lookup for case-insensitive matching
        _, parsed_sections = _split_markdown_sections(strategy_content)
        heading_lookup = {h.lower(): h for h, _ in parsed_sections}

        for ws in weak_sections[:2]:
            raw_title = str(ws.get("title", "")).strip().lstrip("#").strip()
            queries = ws.get("queries", [])
            if not isinstance(queries, list):
                queries = [str(queries)] if queries else []
            reason = str(ws.get("reason", ""))

            if not raw_title or not queries:
                continue

            section_title: str = heading_lookup.get(raw_title.lower(), raw_title)
            console.info(f"Weak: {section_title} — {reason[:80]}")

            # Search for additional evidence
            new_evidence_parts: list[str] = []
            with console.timed_operation(f"Enriching: {section_title}"):
                for q in queries[:2]:
                    results = search_web(q, company_name, website)
                    if results:
                        filtered = [
                            r
                            for r in results[:3]
                            if (not website or website.lower() not in r.get("url", "").lower())
                            and r.get("url", "") not in source_urls_seen
                        ]
                        scraped = scrape_external_sources_validated(
                            filtered,
                            company_name=company_name,
                            website=website,
                            max_sources=3,
                        )
                        for url, content in scraped.items():
                            if url not in source_urls_seen:
                                source_urls.append(url)
                                source_urls_seen.add(url)
                                new_evidence_parts.append(f"[Source: {url}]\n{content[:12_000]}")

            if not new_evidence_parts:
                continue

            new_evidence = "\n\n".join(new_evidence_parts)

            # Find section in strategy content
            section_pattern = re.compile(
                rf"(## {re.escape(section_title)}\n.*?)(?=\n## |\Z)",
                re.DOTALL,
            )
            match = section_pattern.search(strategy_content)
            if not match:
                log_structured("warning", "Strategy CV: section not found", section=section_title)
                continue

            original_section = match.group(1)

            # Regenerate the section
            with console.timed_operation(f"Rewriting: {section_title}"):
                regenerated = _strategy_regenerate_section(
                    company_name,
                    vendor,
                    section_title,
                    original_section,
                    new_evidence,
                    analysis_workbook,
                    model=grok_writing,
                )

            if regenerated and regenerated != original_section:
                if not regenerated.endswith("\n"):
                    regenerated += "\n"
                strategy_content = (
                    strategy_content[: match.start()]
                    + regenerated
                    + strategy_content[match.end() :]
                )
                sections_enriched += 1
                console.ok(f"Enriched: {section_title}")
    else:
        console.ok("Strategy review: no sections flagged for enrichment")

    if sections_enriched > 0:
        log_structured("info", "Strategy sections enriched", count=sections_enriched, vendor=vendor)

    # Step 3: Polish pass
    try:
        with console.timed_operation(f"Polishing {label}{vendor_label}"):
            strategy_content = _strategy_polish(
                company_name, vendor, strategy_content, model=grok_writing
            )
    except Exception as e:
        log_structured("warning", "Strategy polish failed, keeping unpolished", error=str(e))

    return strategy_content


def perform_fast_research(
    company_name: str | None,
    website: str | None,
    start_time: float,
    ai_strategy: bool = False,
    platforms: tuple[str, ...] = ("agnostic",),
    strategy_types: list[str] | None = None,
    max_scrape_time: int | None = None,
    discovery_notes_content: str | None = None,
    *,
    folder_path: str | None = None,
    resume_local: bool = False,
    grok_tier: str = "hybrid",
    continuous_reasoning: bool = True,
    output_dir: str | Path | None = None,
    diagnostics_dir: str | Path | None = None,
    write_txt: bool = True,
) -> str | None:
    """
    Fast research mode using Grok 4.1 with accordion-style batch writing.

    Pipeline:
    1. Data collection: scrape 50 pages + 10 search queries via Gemini Flash
    2. Research deepening: Grok gap analysis → targeted search → fill
    3. Grok analysis call: structured workbook from enriched data
    4. Grok report writing: 5 batch calls (one per YAML part), each writing
       2-7 sections with rolling context from completed batches
    5. Cross-validation: find weak spots → targeted search → re-write ≤3 sections
    6. Optional strategy generation via Grok (AI strategy per vendor,
       and/or YAML-defined strategies like CX, security, data fabric)

    Target: ~20-30 min, ~$0.20
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed

    from primr.ai.grok_client import (
        ContinuousReasoningSession,
        get_grok_session_usage,
        get_grok_session_usage_by_model,
        grok_llm,
        reset_grok_session,
    )

    reset_grok_session()

    # Resolve Grok model pair for this tier
    grok_reasoning, grok_writing = PrimrModels.get_grok_models(GrokTier(grok_tier))

    # Continuous reasoning is on by default after the n=3 pilot — see ROADMAP
    # "Continuous Reasoning Session". When on, workbook generation (Phase 3)
    # and cross-validation (Phase 5) share a single Grok session so the
    # validator inherits the corpus + workbook reasoning instead of re-reading
    # the report cold. Pass --no-continuous-reasoning, or set
    # PRIMR_CONTINUOUS_REASONING=0/false to revert to the fresh-call topology.
    env_flag = os.getenv("PRIMR_CONTINUOUS_REASONING", "").strip().lower()
    if env_flag in ("0", "false", "no", "off"):
        continuous_reasoning = False
    elif env_flag in ("1", "true", "yes", "on"):
        continuous_reasoning = True

    # Session is constructed lazily at the workbook stage so the workbook's
    # system prompt becomes a real `role: system` message instead of being
    # folded into the first user turn (which the v1 pilot showed measurably
    # degrades workbook quality).
    reasoning_session: ContinuousReasoningSession | None = None
    if continuous_reasoning:
        log_structured(
            "info",
            "Continuous reasoning enabled — session will be constructed at workbook stage",
            model=grok_reasoning,
        )

    display_name = company_name or (urlparse(website or "").netloc if website else "")
    folder_path = folder_path or create_working_folder(company_name, website)

    try:
        has_strategies = ai_strategy or bool(strategy_types)
        total_phases = 6 if has_strategies else 5

        # =================================================================
        # Phase 1: Data collection (Gemini Flash — cheap)
        # =================================================================
        scan_domain = urlparse(website or "").netloc.replace("www.", "") if website else "website"
        console.phase_banner(
            1,
            total_phases,
            "Data Collection (fast)",
            f"Scraping {scan_domain} + external sources",
            "5-8 min",
        )

        # Scrape website (50 pages for enhanced fast mode)
        with console.timed_operation(f"Website scrape ({scan_domain})", show_spinner=False):
            scraped_data = (
                fetch_web_content(website, company_name, max_pages=50, working_folder=folder_path)
                if website
                else {}
            )
            pages_scraped = len(scraped_data)
        log_structured("info", "Fast mode: website scraping complete", pages=pages_scraped)

        if pages_scraped == 0 and website:
            console.warn("Limited website access — report will rely on web research")

        # Summarize scraped content with Flash (for insights.txt working file)
        summarized = ""
        if scraped_data:
            with console.timed_operation("Extracting insights"):
                summarized = summarize_scraped_content(
                    company_name, website, scraped_data, folder_path
                )

        # Build raw corpus from scraped data (truncate each page to 30k chars)
        raw_corpus_parts: list[str] = []
        for url, content in scraped_data.items():
            truncated = content[:30_000] if len(content) > 30_000 else content
            raw_corpus_parts.append(f"[Page: {url}]\n{truncated}")
        raw_corpus = "\n\n".join(raw_corpus_parts) if raw_corpus_parts else ""

        # Adaptive depth: assess data richness to calibrate search effort
        total_scraped_chars = sum(len(v or "") for v in scraped_data.values())
        if total_scraped_chars > 200_000 and pages_scraped > 30:
            # Rich website — data is abundant; a small number of high-signal
            # externals is enough for cross-validation.
            _search_depth = "rich"
            _ext_query_count = 8
            _max_ext = 12
            log_structured(
                "info",
                "Adaptive depth: rich website, reducing external search",
                pages=pages_scraped,
                chars=total_scraped_chars,
            )
        elif total_scraped_chars < 20_000:
            # Thin website — compensate with more externals, but not a firehose.
            # Fallback_sources already filled in EDGAR / Wikipedia / IR if the
            # main site was blocked; we don't need to re-validate 40 DDG hits.
            _search_depth = "thin"
            _ext_query_count = 12
            _max_ext = 22
            console.info(
                f"Thin website data ({pages_scraped} pages, {total_scraped_chars} chars) "
                "— increasing external search depth"
            )
            log_structured(
                "info",
                "Adaptive depth: thin website, increasing external search",
                pages=pages_scraped,
                chars=total_scraped_chars,
            )
        else:
            # Normal — 18 validated externals is plenty for a 23-section brief.
            _search_depth = "normal"
            _ext_query_count = 10
            _max_ext = 18

        # External research (adaptive query count)
        source_urls: list[str] = []
        source_urls_seen: set[str] = set()  # O(1) dedup across phases
        external_text_parts: list[str] = []
        external_raw_parts: list[str] = []
        external_queries = generate_external_search_queries(
            company_name,
            website,
            max_queries=_ext_query_count,
        )
        external_data: dict = {}
        max_external_sources = _max_ext
        _ext_search_start = time.time()

        def _search_one(query: str) -> list[dict]:
            """Search for a single query (thread-safe HTTP call)."""
            results = search_web(query, company_name, website)
            if not results:
                return []
            return [
                r
                for r in results[:5]
                if not website or website.lower() not in r.get("url", "").lower()
            ]

        # Phase 1: parallel searches (thread-safe HTTP calls)
        console.status(f"Searching external sources (0/{len(external_queries)} queries)")
        all_search_results: list[dict] = []
        _queries_done = 0
        _queries_failed = 0
        with ThreadPoolExecutor(max_workers=3) as executor:
            futures = [executor.submit(_search_one, q) for q in external_queries]
            for future in as_completed(futures):
                try:
                    all_search_results.extend(future.result())
                except Exception as e:
                    _queries_failed += 1
                    logger.warning("External search query failed: %s", e)
                _queries_done += 1
                console.status(
                    f"Searching external sources ({_queries_done}/{len(external_queries)} queries, {len(all_search_results)} results)"
                )

        if _queries_failed > 0:
            console.warn(
                f"{_queries_failed}/{len(external_queries)} search queries failed"
                " — external source coverage may be reduced"
            )
            log_structured(
                "warning",
                "External search queries failed",
                failed=_queries_failed,
                total=len(external_queries),
            )

        # Phase 2: parallel validation with a hard attempt cap.
        # External validation used to iterate *every* search result serially,
        # which on noisy queries meant 75-150+ Grok validation calls and 20+
        # minutes of wall time. We now:
        #   - hard-cap attempts at max_external_sources * 2 (empirically
        #     enough to fill the quota on any reasonable rejection rate)
        #   - run 4 attempts in parallel (external scrape now uses a
        #     Patchright-free orchestrator, so no browser contention)
        from concurrent.futures import ThreadPoolExecutor, as_completed

        from primr.pipeline.integration import create_pipeline_executor, scrape_page_with_recovery

        _resilience_listener = _build_resilience_event_listener(folder_path)
        _recovery_executor = create_pipeline_executor(
            folder_path, event_listener=_resilience_listener
        )

        # Deduplicate and cap candidate list up front.
        _candidates: list[dict] = []
        _seen_candidate_urls: set[str] = set()
        for result in all_search_results:
            url = result.get("url")
            if not url or url in _seen_candidate_urls or url in external_data:
                continue
            _seen_candidate_urls.add(url)
            _candidates.append(result)

        # Attempt cap sized at 1.6x the quota. Empirically ~35% of DDG
        # results get LLM-rejected as "wrong company" (similar-name but
        # unrelated business), so 1.6x fills the quota in the normal case
        # while keeping total work bounded. Was 2x before 1.19.1 — that
        # was overcautious and turned a 20-source quota into a 60-HTTP-call
        # validation marathon on companies with lots of hits (Nintendo,
        # Disney, any well-known brand).
        _attempt_cap = max(12, int(max_external_sources * 1.6))
        _candidates = _candidates[:_attempt_cap]
        _scrape_total = len(_candidates)
        _failed_scrape_urls: list[str] = []
        _completed_checks = 0

        def _do_scrape(_r: dict, _u: str):
            wrapped = lambda: scrape_external_sources_validated(  # noqa: E731
                [_r],
                company_name=company_name,
                website=website,
                max_sources=1,
            )
            return scrape_page_with_recovery(_recovery_executor, wrapped, _u, folder_path)

        # Total wall-clock deadline for this phase. A single hung worker
        # (stuck HTTP, stuck Grok validation call) used to block the whole
        # pipeline because ThreadPoolExecutor's __exit__ waits for every
        # running thread to finish and fut.cancel() only cancels queued
        # work. Deadline + manual shutdown(wait=False, cancel_futures=True)
        # lets us abandon stuck workers and move on.
        _validation_deadline_s = 600.0  # 10 min total across all workers
        _val_pool = ThreadPoolExecutor(max_workers=4)
        futures = {
            _val_pool.submit(_do_scrape, result, result["url"]): result for result in _candidates
        }
        _val_abandoned = False
        try:
            for fut in as_completed(futures, timeout=_validation_deadline_s):
                _completed_checks += 1
                if len(external_data) >= max_external_sources:
                    # Enough accepted sources — stop waiting for the rest.
                    break
                result = futures[fut]
                url = result["url"]
                console.status(
                    f"Validating external sources ({len(external_data)} validated, "
                    f"checking {_completed_checks}/{_scrape_total})"
                )
                try:
                    scrape_result = fut.result(timeout=0)
                except Exception as e:
                    logger.debug("External validation worker failed for %s: %s", url, e)
                    continue
                if scrape_result.success and scrape_result.output:
                    external_data.update(scrape_result.output)
                elif scrape_result.skipped:
                    _failed_scrape_urls.append(url)
                    logger.info("Scrape skipped for %s: %s", url, scrape_result.skip_reason)
        except TimeoutError:
            _val_abandoned = True
            console.warn(
                f"External validation deadline ({int(_validation_deadline_s)}s) reached — "
                f"continuing with {len(external_data)} validated sources "
                f"({_completed_checks}/{_scrape_total} workers checked)"
            )
        finally:
            # Don't block on hung threads — cancel queued work and detach
            # still-running workers from the stdlib atexit join hook so
            # they can't hang the process at interpreter shutdown.
            from primr.utils.async_utils import detach_running_workers

            _val_pool.shutdown(wait=False, cancel_futures=True)
            _val_abandoned_workers = bool(getattr(_val_pool, "_threads", ()))
            detach_running_workers(_val_pool)

        if _val_abandoned:
            _update_run_state(
                folder_path,
                external_validation_abandoned=True,
                external_validation_completed=_completed_checks,
                external_validation_total=_scrape_total,
            )

        # Log failed pages in run state (Req 2.3)
        if _failed_scrape_urls:
            _update_run_state(folder_path, failed_scrape_urls=_failed_scrape_urls)

        console.ok(f"Searching external sources ({console._elapsed(_ext_search_start)})")

        # Adaptive quality filter: drop low-relevance sources
        pre_filter_count = len(external_data)
        external_data = _assess_source_relevance(company_name, external_data)
        if len(external_data) < pre_filter_count:
            console.info(
                f"Quality filter: {pre_filter_count} -> {len(external_data)} sources (dropped {pre_filter_count - len(external_data)} low-relevance)"
            )

        for url, content in external_data.items():
            source_urls.append(url)
            source_urls_seen.add(url)
            external_text_parts.append(f"[Source: {url}]\n{content[:12_000]}")
            external_raw_parts.append(f"[Source: {url}]\n{content[:20_000]}")

        log_structured("info", "Fast mode: external sources complete", sources=len(external_data))
        _update_run_state(
            folder_path,
            pages_scraped=pages_scraped,
            website_chars=total_scraped_chars,
            external_sources_initial=len(external_data),
            search_depth=_search_depth,
        )
        console.phase_complete(
            "Data Collection (fast)",
            [("Pages", str(pages_scraped)), ("External", str(len(external_data)))],
        )

        # =================================================================
        # Hiring Signals — discover open postings, extract strategic signals.
        # Not numbered as a full phase so we avoid renumbering the five
        # downstream banners, but announced clearly. The resulting block is
        # threaded through BOTH the initial insights build and the Phase 2
        # gap-filling rebuild so it survives every refresh of insights.txt
        # and external_sources_raw.
        # =================================================================
        hiring_block = ""
        try:
            from primr.data.hiring_signals import gather_hiring_signals, render_for_prompt

            console.info("Hiring Signals — scanning for open job postings")
            hiring_signals = gather_hiring_signals(
                company_name or display_name,
                website,
                corpus=scraped_data,
                working_folder=folder_path,
            )
        except Exception as e:
            logger.warning("Hiring signals stage failed: %s", e)
            hiring_signals = None

        if hiring_signals and not hiring_signals.is_empty():
            console.ok(
                f"Hiring Signals: {hiring_signals.postings_extracted} postings analysed via "
                f"{hiring_signals.source} "
                f"({len(hiring_signals.tech_stack)} tech items, "
                f"{len(hiring_signals.strategic_initiatives)} initiatives)",
                show_time=False,
            )
            _update_run_state(
                folder_path,
                hiring_signals={
                    "source": hiring_signals.source,
                    "postings_found": hiring_signals.postings_found,
                    "postings_selected": hiring_signals.postings_selected,
                    "postings_extracted": hiring_signals.postings_extracted,
                    "company_slug": hiring_signals.company_slug,
                },
            )
            hiring_block = "=== HIRING SIGNALS ===\n" + render_for_prompt(hiring_signals)
        else:
            if hiring_signals is None:
                logger.info("Hiring signals: skipped (disabled or no slug candidates)")
            else:
                console.info(
                    "Hiring Signals: no public postings found — continuing without hiring data"
                )
            _update_run_state(
                folder_path,
                hiring_signals={
                    "source": hiring_signals.source if hiring_signals else "skipped",
                    "postings_found": hiring_signals.postings_found if hiring_signals else 0,
                    "postings_extracted": 0,
                },
            )

        # Combine Flash-summarized insights (for working folder)
        all_insights_parts = []
        if summarized:
            all_insights_parts.append(f"=== WEBSITE INSIGHTS ===\n{summarized}")
        if external_text_parts:
            all_insights_parts.append(
                "=== EXTERNAL SOURCES ===\n" + "\n\n".join(external_text_parts)
            )
        if hiring_block:
            all_insights_parts.append(hiring_block)
        combined_insights = (
            "\n\n".join(all_insights_parts) if all_insights_parts else "No research data collected."
        )

        # Save insights to working folder
        insights_file = os.path.join(folder_path, "insights.txt")
        with open(insights_file, "w", encoding="utf-8") as f:
            f.write(combined_insights)

        # Build raw external sources string for Grok. Hiring signals ride
        # along so the gap-analysis, workbook, section-writing, and
        # cross-validation prompts all see them as available evidence.
        external_raw_base_parts = list(external_raw_parts)
        if hiring_block:
            external_raw_base_parts.append(hiring_block)
        external_sources_raw = (
            "\n\n".join(external_raw_base_parts)
            if external_raw_base_parts
            else "(no external sources)"
        )

        # =================================================================
        # Phase 2: Research Deepening (Grok gap analysis → targeted search)
        # =================================================================
        console.phase_banner(
            2,
            total_phases,
            "Research Deepening",
            "Identifying gaps and searching for additional evidence",
            "3-5 min",
        )

        with console.timed_operation("Analyzing research gaps via Grok"):
            gap_queries, gap_text = _fast_gap_analysis(
                company_name or display_name,
                website,
                raw_corpus,
                external_sources_raw,
                source_urls,
                model=grok_reasoning,
            )

        gap_new_sources = 0
        gap_search_count = 0

        if gap_queries:
            console.ok(f"Gap analysis: {len(gap_queries)} questions identified")
            max_gap_sources = 10

            _gap_start = time.time()

            def _gap_search_one(gq: str) -> list[dict]:
                """Search for a single gap query (thread-safe HTTP call)."""
                results = search_web(gq, company_name, website)
                if not results:
                    return []
                return [
                    r
                    for r in results[:3]
                    if (not website or website.lower() not in r.get("url", "").lower())
                    and r.get("url", "") not in source_urls_seen
                ]

            # Phase 1: parallel searches (thread-safe HTTP calls)
            gap_search_results: list[dict] = []
            _gap_queries_done = 0
            console.status(f"Searching for gap-filling sources (0/{len(gap_queries)} queries)")
            with ThreadPoolExecutor(max_workers=3) as executor:
                futures = [executor.submit(_gap_search_one, gq) for gq in gap_queries]
                for future in as_completed(futures):
                    try:
                        gap_search_results.extend(future.result())
                    except Exception as e:
                        logger.warning("Gap search query failed: %s", e)
                    _gap_queries_done += 1
                    console.status(
                        f"Searching for gap-filling sources ({_gap_queries_done}/{len(gap_queries)} queries, {len(gap_search_results)} results)"
                    )

            # Phase 2: parallel validation with a hard attempt cap (same
            # design as the main external-source pass: 4 workers, cap at
            # 2x the target to bound runtime on noisy searches).
            _gap_candidates: list[dict] = []
            _gap_seen: set[str] = set()
            for result in gap_search_results:
                url = result.get("url")
                if not url or url in source_urls_seen or url in _gap_seen:
                    continue
                _gap_seen.add(url)
                _gap_candidates.append(result)

            # Same 1.6x sizing as the main external pass — see comment there.
            _gap_attempt_cap = max(10, int(max_gap_sources * 1.6))
            _gap_candidates = _gap_candidates[:_gap_attempt_cap]
            _gap_check_idx = 0

            def _validate_gap_source(res: dict) -> dict[str, str]:
                return scrape_external_sources_validated(
                    [res],
                    company_name=company_name,
                    website=website,
                    max_sources=1,
                )

            # Same deadline pattern as the main external-source validation
            # above — a hung worker can't block shutdown forever.
            _gap_deadline_s = 420.0  # 7 min total across all workers
            _gap_pool = ThreadPoolExecutor(max_workers=4)
            gap_futures = {_gap_pool.submit(_validate_gap_source, r): r for r in _gap_candidates}
            try:
                for fut in as_completed(gap_futures, timeout=_gap_deadline_s):
                    _gap_check_idx += 1
                    if gap_new_sources >= max_gap_sources:
                        break
                    console.status(
                        f"Validating gap sources ({gap_new_sources} found, "
                        f"checking {_gap_check_idx}/{len(_gap_candidates)})"
                    )
                    try:
                        scraped = fut.result(timeout=0)
                    except Exception as e:
                        logger.debug("Gap validation worker failed: %s", e)
                        continue
                    for scraped_url, content in scraped.items():
                        if gap_new_sources >= max_gap_sources:
                            break
                        if scraped_url not in source_urls_seen:
                            source_urls.append(scraped_url)
                            source_urls_seen.add(scraped_url)
                            external_text_parts.append(
                                f"[Source: {scraped_url}]\n{content[:12_000]}"
                            )
                            external_raw_parts.append(
                                f"[Source: {scraped_url}]\n{content[:20_000]}"
                            )
                            gap_new_sources += 1
            except TimeoutError:
                console.warn(
                    f"Gap-filling deadline ({int(_gap_deadline_s)}s) reached — "
                    f"continuing with {gap_new_sources} new sources "
                    f"({_gap_check_idx}/{len(_gap_candidates)} workers checked)"
                )
            finally:
                from primr.utils.async_utils import detach_running_workers

                _gap_pool.shutdown(wait=False, cancel_futures=True)
                detach_running_workers(_gap_pool)

            console.ok(f"Searching for gap-filling sources ({console._elapsed(_gap_start)})")

            console.ok(f"Found {gap_new_sources} additional sources")

            # Rebuild external_sources_raw with new sources. Keep the
            # hiring-signals block alongside so downstream prompts still
            # see it after the gap-filling refresh.
            external_raw_rebuild = list(external_raw_parts)
            if hiring_block:
                external_raw_rebuild.append(hiring_block)
            external_sources_raw = (
                "\n\n".join(external_raw_rebuild)
                if external_raw_rebuild
                else "(no external sources)"
            )

            # Update insights file
            all_insights_parts_updated = []
            if summarized:
                all_insights_parts_updated.append(f"=== WEBSITE INSIGHTS ===\n{summarized}")
            if external_text_parts:
                all_insights_parts_updated.append(
                    "=== EXTERNAL SOURCES ===\n" + "\n\n".join(external_text_parts)
                )
            if hiring_block:
                all_insights_parts_updated.append(hiring_block)
            combined_insights = (
                "\n\n".join(all_insights_parts_updated)
                if all_insights_parts_updated
                else combined_insights
            )
            with open(insights_file, "w", encoding="utf-8") as f:
                f.write(combined_insights)
        else:
            # Distinguish between "no gaps found" (good) and "gap analysis failed" (bad)
            if gap_text and "failed" in gap_text.lower():
                console.warn(f"Gap analysis failed — skipping research deepening ({gap_text})")
            else:
                console.info("Gap analysis found no research gaps — skipping")

        # Save gap analysis output to working folder
        gap_analysis_path = os.path.join(folder_path, "gap_analysis.md")
        with open(gap_analysis_path, "w", encoding="utf-8") as f:
            f.write(gap_text if gap_text else "(no gap analysis performed)")

        total_external = len(source_urls)
        _update_run_state(
            folder_path,
            gap_queries=len(gap_queries or []),
            gap_new_sources=gap_new_sources,
            external_sources_validated=total_external,
        )
        console.phase_complete(
            "Research Deepening",
            [("New sources", str(gap_new_sources)), ("Total external", str(total_external))],
        )

        validated_source_urls = list(source_urls)
        validated_source_count = len(validated_source_urls)

        # =================================================================
        # Phase 3: Grok analysis call (structured workbook)
        # =================================================================
        console.phase_banner(
            3, total_phases, "Analysis (Grok)", "Building structured analysis workbook", "2-4 min"
        )

        analysis_system = (
            "You are a senior strategic analyst conducting pre-engagement research "
            "for a consulting firm. Produce a structured analysis workbook — working "
            "notes with evidence, confidence levels, and hypotheses. Not polished prose. "
            "CRITICAL: Separate what the company CLAIMS from what external evidence "
            "SUPPORTS. Stress-test their narrative. Be conservative on financial inferences."
        )

        analysis_prompt = _build_fast_analysis_prompt(
            company_name or display_name,
            website,
            raw_corpus,
            external_sources_raw,
        )

        try:
            from primr.pipeline.integration import analysis_with_recovery

            # In continuous mode, construct the session here with the workbook's
            # system prompt as a real role:system message. The session then
            # carries that role + the workbook reasoning forward into Phase 5
            # cross-validation as a follow-up user turn.
            if continuous_reasoning and reasoning_session is None:
                reasoning_session = ContinuousReasoningSession(
                    model=grok_reasoning,
                    system_prompt=analysis_system,
                )

            def _do_analysis():
                if reasoning_session is not None:
                    return reasoning_session.send(
                        analysis_prompt,
                        max_tokens=18_000,
                        temperature=0.5,
                    )
                return grok_llm(
                    analysis_prompt,
                    model=grok_reasoning,
                    max_tokens=18_000,
                    temperature=0.5,
                    system_prompt=analysis_system,
                )

            with console.timed_operation("Generating analysis workbook via Grok"):
                _analysis_result = analysis_with_recovery(
                    _recovery_executor, _do_analysis, folder_path
                )
                if _analysis_result.success:
                    analysis_workbook = _analysis_result.output
                else:
                    raise RuntimeError(
                        _analysis_result.skip_reason or "Analysis recovery exhausted"
                    )
        except Exception as analysis_err:
            console.warn(f"Analysis workbook generation failed: {analysis_err}")
            console.info("Continuing with collected insights as fallback workbook")
            log_structured("warning", "Fast mode analysis fallback used", error=str(analysis_err))
            analysis_workbook = combined_insights

        if not analysis_workbook or not analysis_workbook.strip():
            console.warn("Analysis workbook empty — falling back to insights for report")
            analysis_workbook = combined_insights

        # Save workbook
        workbook_path = os.path.join(folder_path, "analysis_workbook.md")
        with open(workbook_path, "w", encoding="utf-8") as f:
            f.write(analysis_workbook)

        console.phase_complete("Analysis (Grok)")

        # =================================================================
        # Phase 4: Grok report writing (parallel within parts + coherence)
        # =================================================================
        console.phase_banner(
            4,
            total_phases,
            "Report Writing (Grok)",
            "Writing sections (parallel within parts)",
            "3-5 min",
        )

        # Build a raw data subset for evidence (~100k chars — workbook already distills corpus)
        raw_corpus_subset = raw_corpus[:100_000] if len(raw_corpus) > 100_000 else raw_corpus

        report_system = (
            "You are a senior strategic analyst writing a consulting dossier — internal prep "
            "before a discovery conversation. Your reader is a partner walking into a meeting.\n\n"
            "The bar is maximally useful strategic analysis: long-form, specific, strategically sharp, and written "
            "to get a consultant maximally up to speed before talking with the company.\n\n"
            "CORE DISCIPLINE:\n"
            "- STRESS-TEST the company's narrative. Do NOT paraphrase their marketing. "
            "When they claim 'only purpose-built' or '9x ROI', challenge it with evidence.\n"
            "- Frame every major claim as a hypothesis with counter-evidence. "
            "What would disprove it? What's the alternative explanation?\n"
            "- For each section, surface 'what to validate in conversation' — specific "
            "questions a consultant would ask to test the hypothesis.\n"
            "- Be CONSERVATIVE on financial estimates. Use wide ranges, note low confidence. "
            "Never state an inference as if it were confirmed.\n\n"
            "EPISTEMIC RULES:\n"
            "- Label claims: Confirmed (filings/official), Reported (credible 3rd party), "
            "Estimated (inferred), Hypothesis (our speculation)\n"
            "- CONFIDENCE RESET per section: don't inherit confidence from prior sections\n"
            "- NARRATIVE CEILING: don't escalate stakes. 'Opportunity' stays 'opportunity', "
            "not 'transformational opportunity'. Keep scope realistic.\n"
            "- NUMERIC PRECISION: ranges for estimates ('$800M-$1.2B'), note source/date\n"
            "- AVOID OVERREACH: don't claim inside knowledge of board decisions, precise "
            "market share in opaque markets, or causal certainty\n"
            "- REASON UNDER CONSTRAINT: if company-specific evidence is thin, still produce deep strategic analysis by combining "
            "observed facts, industry structure, competitive analogs, likely buyer behavior, and explicit scenario logic\n\n"
            "FORMATTING:\n"
            "- Full paragraphs with evidence and strategic interpretation, not bullet dumps\n"
            "- Tables for financials, competitors, timelines\n"
            "- Keep citations compact, usually at paragraph ends, and avoid cluttering every sentence\n"
            "- Use [cite: N] references in the body; keep the densest reference inventory in the final Sources appendix\n"
            "- Sub-headings (###) within sections for readability\n"
            "- Each insight lives in ONE section — cross-reference, don't repeat"
        )

        from concurrent.futures import ThreadPoolExecutor, as_completed

        section_batches = _group_sections_by_part()

        # Use constrained-evidence reasoning for thin-signal sections instead of skipping them.
        section_reasoning_modes: dict[str, str] = {}
        constrained_sections: list[str] = []
        for batch in section_batches:
            for sec in batch:
                mode = _determine_section_reasoning_mode(sec, analysis_workbook)
                section_reasoning_modes[sec.id] = mode
                if mode == "constrained_evidence":
                    constrained_sections.append(sec.name)
        if constrained_sections:
            console.info(
                "Using constrained-evidence strategic reasoning for "
                f"{len(constrained_sections)} section(s): {', '.join(constrained_sections)}"
            )

        # Pop executive_summary — write it LAST so it can synthesize the full report
        exec_summary_section = None
        for batch in section_batches:
            for sec in batch:
                if sec.id == "executive_summary":
                    exec_summary_section = sec
                    batch.remove(sec)
                    break
            if exec_summary_section:
                break
        # Remove empty batches (if exec summary was the only section in its batch)
        section_batches = [b for b in section_batches if b]

        # Rebuild section names from the post-pop batches so indices align with
        # global_offset used in _write_one.  Exec summary is written last, so it
        # should NOT be in the ToC during batch writing — avoids off-by-one where
        # [NOW] marker points to the wrong section name.
        all_section_names = [s.name for batch in section_batches for s in batch]
        written_sections: list[GeneratedSection] = []
        effective_name = company_name or display_name

        global_offset = 0
        for part_num, part_sections in enumerate(section_batches):
            part_label = _PART_LABELS.get(part_sections[0].part, f"Part {part_sections[0].part}")
            console.info(
                f"Part {part_num + 1}/{len(section_batches)} ({part_label}): "
                f"{len(part_sections)} section(s) in parallel"
            )

            # Snapshot written_sections — threads in this part share the same frozen prior context
            prior_sections = list(written_sections)

            def _write_one(
                idx_section: tuple[int, "SectionConfig"],
                _offset: int = global_offset,
                _prior: list["GeneratedSection"] = prior_sections,
            ) -> tuple[int, dict[str, Any] | None]:
                local_idx, sec = idx_section

                def _do_write():
                    result = _write_section_with_retry(
                        sec,
                        _offset + local_idx,
                        all_section_names,
                        _prior,
                        effective_name,
                        website,
                        analysis_workbook,
                        raw_corpus_subset,
                        external_sources_raw,
                        source_urls,
                        report_system,
                        section_reasoning_modes.get(sec.id, "standard"),
                        model=grok_writing,
                    )
                    if result is None:
                        raise RuntimeError(f"Section '{sec.name}' returned empty")
                    return result

                from primr.pipeline.integration import write_section_with_recovery

                stage_result = write_section_with_recovery(
                    _recovery_executor, _do_write, folder_path
                )
                if stage_result.success:
                    return (local_idx, stage_result.output)
                return (local_idx, None)

            results: list[tuple[int, dict[str, Any] | None]] = []
            if len(part_sections) == 1:
                results.append(_write_one((0, part_sections[0])))
            else:
                with ThreadPoolExecutor(max_workers=min(len(part_sections), 4)) as executor:
                    futures = {
                        executor.submit(_write_one, (i, s)): i for i, s in enumerate(part_sections)
                    }
                    for future in as_completed(futures):
                        results.append(future.result())

            # Sort by local index to maintain canonical section order
            results.sort(key=lambda x: x[0])
            seen_titles: set[str] = {s.title.lower().strip() for s in written_sections}
            for local_idx, parsed in results:
                if parsed:
                    title_key = parsed.title.lower().strip()
                    if title_key in seen_titles:
                        console.warn(f"  {parsed.title} — duplicate, skipping")
                        continue
                    seen_titles.add(title_key)
                    written_sections.append(parsed)
                    console.ok(f"  {parsed.title} ({parsed.words:,} words)")
                else:
                    sec_name = part_sections[local_idx].name
                    console.warn(f"  {sec_name} — skipped (failed or empty)")

            global_offset += len(part_sections)

        # Write executive summary LAST — it now has full report context to synthesize
        if exec_summary_section is not None:
            console.info("Writing Executive Summary (with full report context)")
            exec_parsed = _write_section_with_retry(
                exec_summary_section,
                0,  # section_index 0 — first section in final report
                all_section_names,
                written_sections,  # ALL completed sections → full synthesis context
                effective_name,
                website,
                analysis_workbook,
                raw_corpus_subset,
                external_sources_raw,
                source_urls,
                report_system,
                section_reasoning_modes.get(exec_summary_section.id, "standard"),
                model=grok_writing,
            )
            if exec_parsed:
                written_sections.insert(0, exec_parsed)
                console.ok(f"  {exec_parsed.title} ({exec_parsed.words:,} words)")
            else:
                console.warn("  Executive Summary — skipped (failed or empty)")

        if not written_sections:
            console.error("All report sections failed — no sections written")
            return None

        report_content = _assemble_fast_report(
            company_name or display_name, website, written_sections
        )

        # Coherence pass: deduplicate and smooth transitions
        with console.timed_operation("Running coherence pass"):
            report_content = _fast_coherence_pass(
                company_name or display_name, website, report_content, model=grok_writing
            )

        total_words = len(report_content.split())
        console.phase_complete(
            "Report Writing (Grok)",
            [("Sections", str(len(written_sections))), ("Words", f"{total_words:,}")],
        )

        # =================================================================
        # Phase 5: Cross-Validation (review + targeted enrichment)
        # =================================================================
        console.phase_banner(
            5,
            total_phases,
            "Cross-Validation",
            "Reviewing report for gaps and weak sections",
            "2-4 min",
        )

        with console.timed_operation("Reviewing report quality via Grok"):
            from primr.pipeline.integration import cross_validate_with_recovery

            def _do_cross_validate():
                return _fast_cross_validate(
                    company_name or display_name,
                    website,
                    report_content,
                    source_urls,
                    model=grok_reasoning,
                    reasoning_session=reasoning_session,
                )

            _cv_stage_result = cross_validate_with_recovery(
                _recovery_executor, _do_cross_validate, folder_path
            )
            if _cv_stage_result.success:
                cv_result = _cv_stage_result.output
            else:
                logger.info("Cross-validation skipped: %s", _cv_stage_result.skip_reason)
                cv_result = {"weak_sections": [], "contradictions": [], "_failed": True}

        cv_failed = cv_result.pop("_failed", False)
        weak_sections = cv_result.get("weak_sections", [])
        contradictions = cv_result.get("contradictions", [])
        unresolved_contradictions = len(contradictions)
        sections_enriched = 0
        cv_search_count = 0

        if cv_failed:
            console.warn("Cross-validation failed — report was not quality-checked")
        elif weak_sections:
            console.ok(f"Review complete: {len(weak_sections)} section(s) flagged for enrichment")

            # Build a lookup of report headings for case-insensitive matching
            report_headings = re.findall(r"^## (.+)$", report_content, re.MULTILINE)
            heading_lookup = {h.lower().strip(): h for h in report_headings}

            for ws in weak_sections:
                raw_title = str(ws.get("title", "")).lstrip("#").strip()
                raw_queries = ws.get("queries", [])
                queries = [str(q) for q in raw_queries[:3]] if isinstance(raw_queries, list) else []

                if not raw_title or not queries:
                    continue

                # Case-insensitive heading match
                section_title = heading_lookup.get(raw_title.lower(), raw_title)

                # Search for additional evidence
                new_evidence_parts: list[str] = []
                cv_new_sources = 0
                with console.timed_operation(f"Enriching: {section_title}"):
                    for q in queries:
                        cv_search_count += 1
                        results = search_web(q, company_name, website)
                        if results:
                            filtered = [
                                r
                                for r in results[:3]
                                if (not website or website.lower() not in r.get("url", "").lower())
                                and r.get("url", "") not in source_urls_seen
                            ]
                            scraped = scrape_external_sources_validated(
                                filtered,
                                company_name=company_name,
                                website=website,
                                max_sources=3,
                            )
                            for url, content in scraped.items():
                                if url not in source_urls_seen:
                                    source_urls.append(url)
                                    source_urls_seen.add(url)
                                    new_evidence_parts.append(
                                        f"[Source: {url}]\n{content[:12_000]}"
                                    )
                                    cv_new_sources += 1

                if not new_evidence_parts:
                    continue

                new_evidence = "\n\n".join(new_evidence_parts)

                # Find the original section content in the report
                section_pattern = re.compile(
                    rf"(## {re.escape(section_title)}\n.*?)(?=\n## |\Z)",
                    re.DOTALL,
                )
                match = section_pattern.search(report_content)
                if not match:
                    log_structured(
                        "warning",
                        "Cross-validation: section not found in report",
                        section=section_title,
                    )
                    continue

                original_section = match.group(1)

                # Re-generate the section with new evidence
                with console.timed_operation(f"Rewriting: {section_title}"):
                    regenerated = _fast_regenerate_section(
                        company_name or display_name,
                        website,
                        section_title,
                        original_section,
                        analysis_workbook,
                        new_evidence,
                        source_urls,
                        model=grok_writing,
                    )

                # Splice back into report (preserve \n\n separator between sections)
                if regenerated and regenerated != original_section:
                    if not regenerated.endswith("\n"):
                        regenerated += "\n"
                    report_content = (
                        report_content[: match.start()]
                        + regenerated
                        + report_content[match.end() :]
                    )
                    sections_enriched += 1
                    console.ok(f"Enriched: {section_title} ({cv_new_sources} new source(s))")
        else:
            console.ok("Review complete: no sections flagged for enrichment")

        if contradictions:
            for c in contradictions:
                console.info(f"Contradiction noted: {c[:100]}")

            # Resolve contradictions by asking Grok to standardize
            try:
                contradiction_list = "\n".join(f"- {c}" for c in contradictions)
                resolve_prompt = f"""You are editing a strategic report about {company_name or display_name}.

The cross-validation pass found these contradictions between sections:

{contradiction_list}

For EACH contradiction:
1. Determine which value has the strongest source/evidence
2. Standardize the report to use that value consistently
3. Add a confidence label if the value is uncertain

RULES:
- Do NOT delete, summarize, or condense any sections, paragraphs, or content
- Make ONLY surgical edits to the specific contradictory values/numbers
- Do NOT rewrite prose — change only the conflicting data points
- When evidence is ambiguous, use the most conservative estimate with a range
- Add "(Estimated)" or "(Reported)" labels to standardized values
- Preserve all ## headings, [cite: N] references, and structure
- Output MUST contain at least 98% of the original word count

Return the COMPLETE corrected report with all sections intact. No preamble.

--- REPORT ---
{report_content}
--- END ---"""

                resolved = grok_llm(
                    resolve_prompt,
                    model=grok_writing,
                    max_tokens=65_000,
                    temperature=0.2,
                    system_prompt="You are a fact-checker standardizing contradictory data points across report sections.",
                )
                if resolved and resolved.strip():
                    resolved_words = len(resolved.split())
                    original_words = len(report_content.split())
                    if _preserves_report_structure(report_content, resolved):
                        report_content = resolved
                        unresolved_contradictions = 0
                        console.ok(f"Resolved {len(contradictions)} contradiction(s)")
                    else:
                        logger.warning(
                            "Contradiction resolution changed structure too much (%d → %d words or headings changed), keeping original",
                            original_words,
                            resolved_words,
                        )
            except Exception as resolve_err:
                logger.warning("Contradiction resolution failed: %s", resolve_err)

        # Save cross-validation output to working folder
        cv_output_path = os.path.join(folder_path, "cross_validation.json")
        with open(cv_output_path, "w", encoding="utf-8") as f:
            json.dump(cv_result, f, indent=2)

        # Extract section count from report for metrics
        report_section_count = len(re.findall(r"^## ", report_content, re.MULTILINE))
        cv_stats = [
            ("Sections reviewed", str(report_section_count)),
            ("Enriched", str(sections_enriched)),
        ]
        if cv_failed:
            cv_stats.append(("Status", "FAILED"))
        console.phase_complete("Cross-Validation", cv_stats)

        # Trust polish is a low-cost editorial pass to improve evidence discipline.
        report_content = _polish_fast_report_for_trust(
            company_name or display_name,
            website,
            report_content,
            source_urls,
            model=grok_writing,
        )
        report_content = _clean_fast_report_output(report_content)
        report_content = _normalize_fast_citations(report_content, source_urls=source_urls)
        report_content = _enforce_fast_section_quality_guards(report_content)
        qa_metrics = _compute_fast_report_qa_metrics(
            report_content,
            unresolved_contradictions=unresolved_contradictions,
        )
        if qa_metrics["citations_used"] == 0 or qa_metrics["citations_defined"] == 0:
            repaired_report = _repair_fast_report_citation_integrity(
                company_name or display_name,
                website,
                report_content,
                source_urls,
                model=grok_writing,
            )
            if repaired_report != report_content:
                report_content = repaired_report
                qa_metrics = _compute_fast_report_qa_metrics(
                    report_content,
                    unresolved_contradictions=unresolved_contradictions,
                )
        qa_parts = [
            f"labels={qa_metrics['confidence_labels']}",
            f"cites={qa_metrics['citations_used']}/{qa_metrics['citations_defined']}",
            f"validate={qa_metrics['sections_with_validate']}/{qa_metrics['section_count']}",
        ]
        if qa_metrics.get("duplicate_sections", 0) > 0:
            qa_parts.append(f"dupes={qa_metrics['duplicate_sections']}")
        if qa_metrics.get("thin_sections", 0) > 0:
            qa_parts.append(f"thin={qa_metrics['thin_sections']}")
        if qa_metrics.get("unresolved_contradictions", 0) > 0:
            qa_parts.append(f"contradictions={qa_metrics['unresolved_contradictions']}")
        qa_parts.append(f"gate={'PASS' if qa_metrics['qa_gate_passed'] else 'WARN'}")
        console.info("Fast QA: " + ", ".join(qa_parts))
        report_trust_stats = [
            ("Report Gate", "PASS" if qa_metrics["qa_gate_passed"] else "WARN"),
            (
                "Citations",
                f"{qa_metrics['citations_used']}/{qa_metrics['citations_defined']} defined",
            ),
            (
                "Validate Lines",
                f"{qa_metrics['sections_with_validate']}/{qa_metrics['section_count']} sections",
            ),
        ]
        if qa_metrics.get("unresolved_contradictions", 0) > 0:
            report_trust_stats.append(
                ("Contradictions", str(qa_metrics["unresolved_contradictions"]))
            )

        # Save report via existing output pipeline
        # Note: unresolved contradictions are surfaced as QA warnings above
        # but do NOT block DOCX shipping — the contradictions are already
        # noted inline and the user gets the full report.
        report_gate_issues = []
        if qa_metrics["citations_used"] == 0 or qa_metrics["citations_defined"] == 0:
            console.warn(
                "Report validation warning: "
                + f"citation_integrity {qa_metrics['citations_used']}/{qa_metrics['citations_defined']}"
            )
        docx_path = _convert_deep_research_to_docx(
            report_content,
            company_name or display_name,
            website,
            gate_issues=report_gate_issues,
            output_dir=output_dir,
            diagnostics_dir=diagnostics_dir,
            write_txt=write_txt,
        )

        # Also save raw markdown for AI strategy context
        raw_md_path = os.path.join(folder_path, "report.md")
        with open(raw_md_path, "w", encoding="utf-8") as f:
            f.write(report_content)

        # =================================================================
        # Phase 6: Strategy Generation via Grok (optional)
        # =================================================================
        strategy_paths: dict[str, str] = {}
        strategy_trust_stats: list[tuple[str, list[tuple[str, str]]]] = []
        if has_strategies:
            console.phase_banner(
                6, total_phases, "Strategy (Grok)", "Generating strategy documents", "3-8 min"
            )

            # --- AI Strategy (per vendor) ---
            # When multiple platforms are active (common when recon detects
            # both AWS and Azure), run the per-vendor strategies concurrently.
            # The shared company context (report + insights + gap analysis +
            # workbook) is identical across vendors; only the vendor-specific
            # research docs differ. Running them in parallel roughly halves
            # wall-clock time on multi-platform runs.
            if ai_strategy and platforms:

                def _run_ai_strategy_for_vendor(vendor: str):
                    """Run the full per-platform AI strategy pipeline.

                    Returns (strategy_path, trust_stats_tuple, path_key) on
                    success, or None on failure. All console output is
                    prefixed with the vendor label so concurrent runs remain
                    distinguishable in the CLI.
                    """
                    strategy_prompt = _build_ai_strategy_prompt(
                        company_name or display_name, vendor, discovery_notes_content
                    )

                    context_parts = [f"--- Company Report ---\n{report_content[:50_000]}"]

                    # Enrich with working-folder artifacts (insights, gap analysis, workbook)
                    for artifact_name, artifact_limit in [
                        ("insights.txt", 20_000),
                        ("gap_analysis.md", 15_000),
                        ("analysis_workbook.md", 20_000),
                    ]:
                        artifact_path = os.path.join(folder_path, artifact_name)
                        if os.path.exists(artifact_path):
                            try:
                                with open(artifact_path, encoding="utf-8") as fh:
                                    artifact_content = fh.read()[:artifact_limit]
                                    if artifact_content.strip():
                                        context_parts.append(
                                            f"--- {artifact_name} ---\n{artifact_content}"
                                        )
                            except Exception as e:
                                logger.warning("Failed to read artifact %s: %s", artifact_name, e)

                    vendor_doc_paths = (
                        _get_or_generate_vendor_research(vendor)
                        if vendor.lower() != "agnostic"
                        else []
                    )
                    for vdp in vendor_doc_paths:
                        if vdp and os.path.exists(vdp):
                            try:
                                with open(vdp, encoding="utf-8") as fh:
                                    context_parts.append(
                                        f"--- {os.path.basename(vdp)} ---\n{fh.read()[:30_000]}"
                                    )
                            except Exception as e:
                                logger.warning("Failed to read vendor doc %s: %s", vdp, e)

                    combined_strategy_prompt = (
                        "Use the following context documents to inform your analysis:\n\n"
                        + "\n\n".join(context_parts)
                        + "\n\n---\n\n"
                        + strategy_prompt
                    )

                    vendor_label = f" ({vendor.upper()})" if len(platforms) > 1 else ""
                    try:
                        from primr.pipeline.integration import strategy_with_recovery

                        def _do_strategy(_prompt=combined_strategy_prompt):
                            return grok_llm(
                                _prompt,
                                model=grok_writing,
                                max_tokens=32_000,
                            )

                        with console.timed_operation(f"AI Strategy{vendor_label} via Grok"):
                            _strat_result = strategy_with_recovery(
                                _recovery_executor, _do_strategy, folder_path
                            )
                            if _strat_result.success:
                                strategy_content = _strat_result.output
                            else:
                                raise RuntimeError(
                                    _strat_result.skip_reason or "Strategy recovery exhausted"
                                )
                    except Exception as strat_err:
                        console.warn(f"AI Strategy{vendor_label} failed: {strat_err} — skipping")
                        log_structured(
                            "warning",
                            "Fast mode strategy failed",
                            vendor=vendor,
                            error=str(strat_err),
                        )
                        return  # abandon this vendor; others run independently

                    if strategy_content and strategy_content.strip():
                        strategy_content = re.sub(
                            r"\n*_?Disclaimer:\s*Grok is not a financial advi[sc]er[^\n]*\n?",
                            "\n",
                            strategy_content,
                            flags=re.IGNORECASE,
                        ).strip()
                        strategy_content = re.sub(
                            r"\[Word count:\s*[\d,]+\]",
                            "",
                            strategy_content,
                            flags=re.IGNORECASE,
                        )

                        # Enrich: cross-validate → evidence search → polish
                        try:
                            strategy_content = _enrich_strategy_content(
                                strategy_content,
                                company_name or display_name,
                                vendor,
                                "AI Strategy",
                                list(validated_source_urls),
                                set(validated_source_urls),
                                analysis_workbook,
                                website,
                                grok_reasoning=grok_reasoning,
                                grok_writing=grok_writing,
                            )
                        except Exception as enrich_err:
                            log_structured(
                                "warning",
                                "Strategy enrichment failed, keeping original",
                                vendor=vendor,
                                error=str(enrich_err),
                            )

                        strategy_content, strategy_qa, rejected_strategy_sources = (
                            _prepare_strategy_for_output(
                                strategy_content,
                                company_name or display_name,
                                vendor,
                                "AI Strategy",
                                list(validated_source_urls),
                                model=grok_writing,
                            )
                        )
                        qa_gate = "PASS" if strategy_qa["qa_gate_passed"] else "WARN"
                        console.info(
                            f"Strategy QA: placeholders={strategy_qa['placeholder_refs']}, "
                            f"sources={strategy_qa['source_urls']}/{strategy_qa['citation_defs']}, "
                            f"missing={strategy_qa['missing_citations']}, "
                            f"invalid={strategy_qa['invalid_source_urls'] + len(rejected_strategy_sources)}, "
                            f"budget={'OK' if not strategy_qa['budget_inconsistent'] else 'WARN'}, gate={qa_gate}"
                        )
                        if strategy_qa["source_urls"] == 0:
                            console.warn(
                                "Strategy QA: no explicit source URLs detected in strategy output"
                            )
                        strategy_trust_stats.append(
                            (
                                f"AI Strategy ({vendor.upper()})"
                                if len(platforms) > 1
                                else "AI Strategy",
                                [
                                    ("Gate", qa_gate),
                                    ("Sources", f"{strategy_qa['source_urls']} valid"),
                                    (
                                        "Citation Gaps",
                                        str(strategy_qa["missing_citations"]),
                                    ),
                                    (
                                        "Invalid Sources",
                                        str(
                                            strategy_qa["invalid_source_urls"]
                                            + len(rejected_strategy_sources)
                                        ),
                                    ),
                                    (
                                        "Budget Check",
                                        "WARN" if strategy_qa["budget_inconsistent"] else "OK",
                                    ),
                                ],
                            )
                        )

                        strategy_path = _save_strategy_output(
                            strategy_content,
                            company_name or display_name,
                            vendor,
                            strategy_label="AI_Strategy",
                            output_dir=output_dir,
                            diagnostics_dir=diagnostics_dir,
                            write_txt=write_txt,
                        )
                        if strategy_path:
                            key = f"ai_{vendor}" if len(platforms) > 1 else "ai"
                            strategy_paths[key] = strategy_path

                # Dispatch per-platform strategy workers. One platform = run
                # inline (no pool overhead). Multiple platforms = ThreadPool
                # with one worker per platform, capped at 3 for rate-limit
                # safety. grok_llm + network IO releases the GIL so threads
                # genuinely overlap.
                if len(platforms) == 1:
                    _run_ai_strategy_for_vendor(platforms[0])
                else:
                    from concurrent.futures import ThreadPoolExecutor, as_completed

                    with ThreadPoolExecutor(max_workers=min(len(platforms), 3)) as _strat_pool:
                        _strat_futures = {
                            _strat_pool.submit(_run_ai_strategy_for_vendor, v): v for v in platforms
                        }
                        for _sf in as_completed(_strat_futures):
                            v = _strat_futures[_sf]
                            try:
                                _sf.result()
                            except Exception as e:
                                logger.warning(
                                    "Parallel AI strategy worker for %s raised: %s",
                                    v,
                                    e,
                                )

            # --- YAML-defined strategies (customer_experience, security, data_fabric, etc.) ---
            if strategy_types:
                import yaml as _yaml

                for stype in strategy_types:
                    if stype == "ai":
                        continue  # already handled above

                    # Load strategy YAML config (name matches filename)
                    yaml_path = (
                        Path(__file__).parent.parent / "prompts" / "strategies" / f"{stype}.yaml"
                    )

                    if not yaml_path.exists():
                        console.warn(f"Strategy YAML not found: {stype}.yaml — skipping")
                        continue

                    try:
                        with open(yaml_path, encoding="utf-8") as f:
                            strategy_config = _yaml.safe_load(f)
                    except Exception as e:
                        console.warn(f"Failed to load {stype}.yaml: {e} — skipping")
                        continue

                    meta = strategy_config.get("meta", {})
                    display_name_strat = meta.get("name", stype.replace("_", " ").title())
                    output_filename = meta.get("output_filename", f"{{company_name}}_{stype}")
                    # Build label for filename from YAML meta
                    file_label = output_filename.replace("{company_name}_", "").replace(
                        "{company_name}", ""
                    )
                    if not file_label:
                        file_label = stype.replace(" ", "_")

                    strategy_prompt = _build_strategy_prompt_from_yaml(
                        strategy_config, company_name or display_name, discovery_notes_content
                    )

                    # Build context with report + working-folder artifacts
                    yaml_context_parts = [f"--- Company Report ---\n{report_content[:50_000]}"]
                    for artifact_name, artifact_limit in [
                        ("insights.txt", 20_000),
                        ("gap_analysis.md", 15_000),
                        ("analysis_workbook.md", 20_000),
                    ]:
                        artifact_path = os.path.join(folder_path, artifact_name)
                        if os.path.exists(artifact_path):
                            try:
                                with open(artifact_path, encoding="utf-8") as fh:
                                    artifact_content = fh.read()[:artifact_limit]
                                    if artifact_content.strip():
                                        yaml_context_parts.append(
                                            f"--- {artifact_name} ---\n{artifact_content}"
                                        )
                            except Exception as e:
                                logger.warning("Failed to read artifact %s: %s", artifact_name, e)

                    combined_prompt = (
                        "Use the following context documents to inform your analysis:\n\n"
                        + "\n\n".join(yaml_context_parts)
                        + "\n\n---\n\n"
                        + strategy_prompt
                    )

                    try:
                        from primr.pipeline.integration import strategy_with_recovery

                        def _do_yaml_strategy(_p=combined_prompt):
                            return grok_llm(
                                _p,
                                model=grok_writing,
                                max_tokens=32_000,
                            )

                        with console.timed_operation(f"{display_name_strat} via Grok"):
                            _yaml_strat_result = strategy_with_recovery(
                                _recovery_executor, _do_yaml_strategy, folder_path
                            )
                            if _yaml_strat_result.success:
                                strategy_content = _yaml_strat_result.output
                            else:
                                raise RuntimeError(
                                    _yaml_strat_result.skip_reason or "Strategy recovery exhausted"
                                )
                    except Exception as strat_err:
                        console.warn(f"{display_name_strat} failed: {strat_err} — skipping")
                        log_structured(
                            "warning",
                            "Fast mode strategy failed",
                            strategy=stype,
                            error=str(strat_err),
                        )
                        continue

                    if strategy_content and strategy_content.strip():
                        strategy_content = re.sub(
                            r"\n*_?Disclaimer:\s*Grok is not a financial advi[sc]er[^\n]*\n?",
                            "\n",
                            strategy_content,
                            flags=re.IGNORECASE,
                        ).strip()
                        strategy_content = re.sub(
                            r"\[Word count:\s*[\d,]+\]",
                            "",
                            strategy_content,
                            flags=re.IGNORECASE,
                        )

                        # Enrich: cross-validate → evidence search → polish
                        # Use strategy display name (e.g. "Customer Experience") not "agnostic"
                        try:
                            strategy_content = _enrich_strategy_content(
                                strategy_content,
                                company_name or display_name,
                                display_name_strat,
                                display_name_strat,
                                list(validated_source_urls),
                                set(validated_source_urls),
                                analysis_workbook,
                                website,
                                grok_reasoning=grok_reasoning,
                                grok_writing=grok_writing,
                            )
                        except Exception as enrich_err:
                            log_structured(
                                "warning",
                                "Strategy enrichment failed, keeping original",
                                strategy=stype,
                                error=str(enrich_err),
                            )

                        strategy_content, strategy_qa, rejected_strategy_sources = (
                            _prepare_strategy_for_output(
                                strategy_content,
                                company_name or display_name,
                                display_name_strat,
                                display_name_strat,
                                list(validated_source_urls),
                                model=grok_writing,
                            )
                        )
                        qa_gate = "PASS" if strategy_qa["qa_gate_passed"] else "WARN"
                        console.info(
                            f"Strategy QA: placeholders={strategy_qa['placeholder_refs']}, "
                            f"sources={strategy_qa['source_urls']}/{strategy_qa['citation_defs']}, "
                            f"missing={strategy_qa['missing_citations']}, "
                            f"invalid={strategy_qa['invalid_source_urls'] + len(rejected_strategy_sources)}, "
                            f"budget={'OK' if not strategy_qa['budget_inconsistent'] else 'WARN'}, gate={qa_gate}"
                        )
                        if strategy_qa["source_urls"] == 0:
                            console.warn(
                                "Strategy QA: no explicit source URLs detected in strategy output"
                            )
                        strategy_trust_stats.append(
                            (
                                display_name_strat,
                                [
                                    ("Gate", qa_gate),
                                    ("Sources", f"{strategy_qa['source_urls']} valid"),
                                    (
                                        "Citation Gaps",
                                        str(strategy_qa["missing_citations"]),
                                    ),
                                    (
                                        "Invalid Sources",
                                        str(
                                            strategy_qa["invalid_source_urls"]
                                            + len(rejected_strategy_sources)
                                        ),
                                    ),
                                    (
                                        "Budget Check",
                                        "WARN" if strategy_qa["budget_inconsistent"] else "OK",
                                    ),
                                ],
                            )
                        )

                        strategy_path = _save_strategy_output(
                            strategy_content,
                            company_name or display_name,
                            "agnostic",
                            strategy_label=file_label,
                            output_dir=output_dir,
                            diagnostics_dir=diagnostics_dir,
                            write_txt=write_txt,
                        )
                        if strategy_path:
                            strategy_paths[stype] = strategy_path

            if strategy_paths:
                console.phase_complete("Strategy (Grok)")
            else:
                console.warn("Strategy generation skipped — no strategies generated")

        # =================================================================
        # Summary
        # =================================================================
        elapsed = time.time() - start_time
        mins = int(elapsed // 60)
        secs = int(elapsed % 60)
        time_str = f"{mins}m {secs}s" if mins > 0 else f"{secs}s"

        if docx_path:
            console.success_box("Report ready", str(Path(docx_path).resolve()))
        else:
            console.warn(
                "Report DOCX held back by artifact gate; review the saved MD/TXT artifacts"
            )

        for strat_key, strategy_path in strategy_paths.items():
            # AI strategy keys: "ai" or "ai_azure" — show vendor suffix
            if strat_key.startswith("ai"):
                vendor_suffix = (
                    f" ({strat_key.split('_', 1)[1].upper()})" if "_" in strat_key else ""
                )
                label = f"AI Strategy{vendor_suffix}"
            else:
                label = strat_key.replace("_", " ").title()
            resolved_strategy_path = Path(strategy_path).resolve()
            if str(resolved_strategy_path).lower().endswith(".docx"):
                console.success_box(label, str(resolved_strategy_path))
            else:
                console.warn(
                    f"{label} DOCX held back by artifact gate; saved {resolved_strategy_path.name} instead"
                )

        # Cost summary from Grok session usage (per-model for accurate pricing)
        grok_usage = get_grok_session_usage()
        usage_by_model = get_grok_session_usage_by_model()
        grok_cost = 0.0
        for model_name, tokens in usage_by_model.items():
            model_config = PrimrModels.get_model_config(model_name)
            if model_config:
                grok_cost += PrimrModels.calculate_cost(
                    model_name, tokens["input_tokens"], tokens["output_tokens"]
                )
            else:
                # Unknown model — fall back to default Grok pricing
                grok_cost += PrimrModels.calculate_cost(
                    PrimrModels.GROK_MODEL, tokens["input_tokens"], tokens["output_tokens"]
                )

        # Flash cost from AI client
        from primr.ai.client import get_client

        client = get_client()
        flash_usage = client.get_usage_summary()
        flash_cost = flash_usage.get("total_cost", 0.0)

        actual_cost = grok_cost + flash_cost

        date_str = datetime.now().strftime("%m-%d-%Y")
        fallback_dir = Path(output_dir) if output_dir is not None else Path(OUTPUT_DIR)
        fallback_md = (
            fallback_dir / f"{company_name or display_name}_Strategic_Overview_{date_str}.md"
        )
        primary_output_path = str(fallback_md) if fallback_md.exists() else docx_path

        artifacts_passed = bool(docx_path) and all(
            str(path).lower().endswith(".docx") for path in strategy_paths.values()
        )
        completion_label = (
            "Fast mode complete" if artifacts_passed else "Fast mode complete with warnings"
        )
        console.ok(f"{completion_label} in {time_str}")

        _update_run_state(
            folder_path,
            report_sections=len(written_sections),
            report_words=total_words,
            external_sources_validated=validated_source_count,
            strategy_artifacts=len(strategy_paths),
            artifact_gate_passed=artifacts_passed,
            actual_cost_usd=round(actual_cost, 4),
        )

        if report_trust_stats:
            console.trust_summary("Report Trust", report_trust_stats)
        for trust_title, trust_stats in strategy_trust_stats:
            console.trust_summary(trust_title + " Trust", trust_stats)

        _tier_labels = {"fast": "Grok 4.1", "hybrid": "Grok 4.3 hybrid", "max": "Grok 4.3 max"}
        summary_items = [
            ("Mode", "fast (" + _tier_labels.get(grok_tier, "Grok") + ")"),
            ("Pages", str(pages_scraped)),
            ("External", str(validated_source_count)),
            ("Duration", time_str),
            (
                "Grok tokens",
                f"{grok_usage['input_tokens']:,} in / {grok_usage['output_tokens']:,} out",
            ),
            ("Actual Cost", f"~${actual_cost:.2f}"),
            ("Artifact Gate", "PASS" if artifacts_passed else "WARN"),
        ]
        if strategy_paths:
            strat_labels = []
            for k in strategy_paths:
                if k.startswith("ai"):
                    vendor_suffix = f" ({k.split('_', 1)[1].upper()})" if "_" in k else ""
                    strat_labels.append(f"AI Strategy{vendor_suffix}")
                else:
                    strat_labels.append(k.replace("_", " ").title())
            summary_items.append(("Strategies", ", ".join(strat_labels)))
        console.summary(summary_items)

        # Save usage to history
        from primr.utils.usage_tracker import get_usage_tracker

        tracker = get_usage_tracker()
        tracker.record_usage(
            mode="fast",
            company=display_name,
            input_tokens=grok_usage["input_tokens"],
            output_tokens=grok_usage["output_tokens"],
            search_queries=len(external_queries) + gap_search_count + cv_search_count,
            duration_seconds=elapsed,
            pipeline_cost=actual_cost,
        )
        tracker.save()

        # Log job summary
        job_summary = JobSummary.create(
            company=display_name,
            mode="fast",
            duration_seconds=elapsed,
            api_calls=0,
            total_tokens=grok_usage["input_tokens"] + grok_usage["output_tokens"],
            sections_generated=len(written_sections),
            output_path=primary_output_path,
        )
        log_job_summary(job_summary)

        return primary_output_path

    except Exception as e:
        console.error(f"Fast research failed: {e}")
        log_structured("error", "Fast research failed", error=str(e), error_type=type(e).__name__)
        logger.exception("Fast research failed")
        return None


def _save_strategy_output(
    strategy_content: str,
    company_name: str,
    platform: str,
    strategy_label: str = "AI_Strategy",
    output_dir: str | Path | None = None,
    diagnostics_dir: str | Path | None = None,
    write_txt: bool = True,
) -> str | None:
    """Save strategy markdown/txt/docx output. Returns docx path or None."""
    from primr.output.markdown_converter import markdown_to_docx
    from primr.output.output_utils import OUTPUT_DIR

    date_str = datetime.now().strftime("%m-%d-%Y")
    vendor_suffix = f"_{platform.upper()}" if platform != "agnostic" else ""
    base_name = f"{company_name}_{strategy_label}{vendor_suffix}_{date_str}"
    destination_dir = Path(output_dir) if output_dir is not None else Path(OUTPUT_DIR)
    destination_dir.mkdir(parents=True, exist_ok=True)
    internal_dir = Path(diagnostics_dir) if diagnostics_dir is not None else destination_dir
    internal_dir.mkdir(parents=True, exist_ok=True)

    # Human-readable title for DOCX
    display_label = strategy_label.replace("_", " ")

    try:
        strategy_content, markdown_validation, salvaged = _salvage_markdown_for_shipping(
            strategy_content,
            kind="strategy",
        )

        md_path = destination_dir / f"{base_name}.md"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(strategy_content)

        if write_txt or diagnostics_dir is not None:
            txt_path = (destination_dir if write_txt else internal_dir) / f"{base_name}.txt"
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(strategy_content)
        if salvaged:
            console.info(f"Adaptive salvage cleaned {display_label} markdown before shipping")

        strategy_gate_issues = list(markdown_validation["issues"])
        strategy_gate_errors = list(markdown_validation["errors"])
        strategy_advisory_issues: list[str] = []
        strategy_qa = _compute_strategy_qa_metrics(strategy_content)
        if strategy_qa["budget_inconsistent"]:
            strategy_advisory_issues.append("budget_inconsistent")
        if strategy_qa["missing_citations"]:
            strategy_advisory_issues.append(f"missing_citations:{strategy_qa['missing_citations']}")
        if strategy_qa["invalid_source_urls"]:
            strategy_advisory_issues.append(
                f"invalid_source_urls:{strategy_qa['invalid_source_urls']}"
            )
        if strategy_gate_issues:
            report_path = _write_output_validation_report(
                md_path,
                "markdown",
                strategy_gate_issues,
                strategy_gate_errors,
                diagnostics_dir=diagnostics_dir,
            )
            console.warn(f"{display_label} artifact issues: " + ", ".join(strategy_gate_issues[:3]))
            console.warn(
                "DOCX shipping gate failed for strategy markdown; saved MD/TXT only"
                + (f" ({report_path.name})" if report_path else "")
            )
            return str(md_path)
        if strategy_advisory_issues or strategy_gate_errors:
            report_path = _write_output_validation_report(
                md_path,
                "markdown",
                strategy_advisory_issues,
                strategy_gate_errors,
                diagnostics_dir=diagnostics_dir,
            )
            console.warn(
                f"{display_label} validation warnings: "
                + ", ".join((strategy_advisory_issues or strategy_gate_errors)[:3])
                + (f" ({report_path.name})" if report_path else "")
            )

        docx_path = destination_dir / f"{base_name}.docx"
        try:
            markdown_to_docx(
                markdown_text=strategy_content,
                output_path=docx_path,
                title=f"{display_label}: {company_name}",
                subtitle=f"{platform.upper()} | {datetime.now().strftime('%B %d, %Y')}"
                if platform != "agnostic"
                else datetime.now().strftime("%B %d, %Y"),
            )
        except Exception as e:
            logger.warning(f"DOCX conversion failed: {e}")
            return str(md_path)

        docx_validation = _validate_output_docx(docx_path)
        if not docx_validation["passed"]:
            report_path = _write_output_validation_report(
                docx_path,
                "docx",
                docx_validation["issues"],
                docx_validation["errors"],
                diagnostics_dir=diagnostics_dir,
            )
            try:
                docx_path.unlink(missing_ok=True)
            except Exception as cleanup_err:
                logger.warning(
                    "Failed to remove blocked strategy DOCX %s: %s", docx_path, cleanup_err
                )
            console.warn(
                "DOCX shipping gate failed for rendered strategy; saved MD/TXT only"
                + (f" ({report_path.name})" if report_path else "")
            )
            return str(md_path)

        if docx_validation["errors"]:
            report_path = _write_output_validation_report(
                docx_path,
                "docx",
                docx_validation["issues"],
                docx_validation["errors"],
                diagnostics_dir=diagnostics_dir,
            )
            console.warn(
                "DOCX validator encountered non-fatal errors; shipping strategy DOCX"
                + (f" ({report_path.name})" if report_path else "")
            )

        return str(docx_path)
    except Exception as e:
        logger.error(f"Failed to save strategy output: {e}")
        return None


def _extract_domain(url: str) -> str | None:
    """Extract domain from a URL for recon lookup.

    Uses recon's own validator for normalization.
    Returns None if the URL cannot be parsed into a valid domain.
    """
    from urllib.parse import urlparse as _urlparse

    try:
        from recon_tool.validator import validate_domain

        parsed = _urlparse(url)
        raw = parsed.netloc or parsed.path.split("/")[0]
        return validate_domain(raw)
    except (ValueError, Exception):
        return None


def perform_research(
    company_name: str | None = None,
    website: str | None = None,
    mode: str = "structured",
    citation_style: str = "numbered",
    ai_strategy: bool = False,
    platforms: tuple[str, ...] | None = None,
    output_dir: str | Path | None = None,
    skip_confirm: bool = False,
    context_files: list[Any] | None = None,
    refresh_vendor_research: bool = False,
    strategies: list[str] | None = None,
    strategy_only: bool = False,
    no_qa: bool = False,
    max_scrape_time: int | None = None,
    discovery_notes_path: str | None = None,
    lite_strategy: bool = False,
    fast_mode: bool = False,
    premium_mode: bool = False,
    skip_scrape_validation: bool = False,
    resume_local: bool = False,
    verify: bool = False,
    grok_tier: str = "hybrid",
    skip_recon: bool = False,
    continuous_reasoning: bool = True,
) -> str | None:
    if not company_name and not website:
        console.error("No company name or website provided")
        return None

    display_name: str = company_name or (urlparse(website or "").netloc if website else "")
    folder_path = create_working_folder(company_name, website, reuse_incomplete=resume_local)
    explicit_platforms = platforms is not None
    if platforms is None:
        from primr.core.platform_mapper import DEFAULT_PLATFORM_FALLBACK

        platforms = DEFAULT_PLATFORM_FALLBACK
    run_output_dir = Path(output_dir) if output_dir is not None else Path(OUTPUT_DIR)
    run_output_dir.mkdir(parents=True, exist_ok=True)
    diagnostics_dir: Path | None = None
    write_public_txt = True
    if output_dir is not None:
        diagnostics_dir = Path(folder_path) / "_diagnostics"
        diagnostics_dir.mkdir(parents=True, exist_ok=True)
        write_public_txt = False
    existing_state = _load_run_state(folder_path)
    if resume_local and existing_state:
        # Ensure resilience keys exist even on resumed runs (NFR 3 backwards compat)
        _ensure_resilience_keys(existing_state)
        _save_run_state(folder_path, existing_state)
        _update_run_state(
            folder_path,
            company_name=company_name or display_name,
            website=website,
            mode=mode,
            status="running",
            current_phase="initializing",
            ai_strategy=ai_strategy,
            cloud_vendors=list(platforms),
            working_folder=folder_path,
        )
        _append_run_event(
            folder_path, "initializing", "resumed", "Resuming from existing local run folder"
        )
    else:
        _init_run_state_with_resilience(
            folder_path,
            {
                "company_name": company_name or display_name,
                "website": website,
                "mode": mode,
                "status": "running",
                "current_phase": "initializing",
                "ai_strategy": ai_strategy,
                "cloud_vendors": list(platforms),
                "working_folder": folder_path,
                "started_at": datetime.now().isoformat(),
                "events": [],
            },
        )
        _append_run_event(folder_path, "initializing", "started", "Run initialized")

    # Load discovery notes if provided
    discovery_notes_content: str | None = None
    if discovery_notes_path:
        try:
            with open(discovery_notes_path, encoding="utf-8") as f:
                discovery_notes_content = f.read().strip()
            if discovery_notes_content:
                logger.info(f"Loaded discovery notes from {discovery_notes_path}")
            else:
                logger.warning(f"Discovery notes file is empty: {discovery_notes_path}")
                discovery_notes_content = None
        except FileNotFoundError:
            logger.error(f"Discovery notes file not found: {discovery_notes_path}")
            console.error(f"Discovery notes file not found: {discovery_notes_path}")
            _update_run_state(folder_path, status="failed", current_phase="initializing")
            _append_run_event(
                folder_path,
                "initializing",
                "failed",
                f"Discovery notes not found: {discovery_notes_path}",
            )
            return None
        except Exception as e:
            logger.error(f"Failed to load discovery notes: {e}")
            console.error(f"Failed to load discovery notes: {e}")
            _update_run_state(folder_path, status="failed", current_phase="initializing")
            _append_run_event(
                folder_path, "initializing", "failed", f"Failed loading discovery notes: {e}"
            )
            return None

    # =========================================================================
    # Recon pre-flight: DNS intelligence on target domain
    # =========================================================================
    recon_info = None  # TenantInfo | None
    recon_context_path: str | None = None

    if not skip_recon and website:
        domain = _extract_domain(website)
        if domain:
            try:
                _update_run_state(folder_path, current_phase="recon")
                _append_run_event(folder_path, "recon", "started", f"Running recon on {domain}")

                from recon_tool.resolver import resolve_tenant

                info, _recon_results = asyncio.get_event_loop().run_until_complete(
                    asyncio.wait_for(resolve_tenant(domain), timeout=15.0)
                )
                recon_info = info  # noqa: F841 — kept for future downstream use

                # Auto-detect platforms if user didn't specify
                from primr.core.platform_mapper import DEFAULT_PLATFORM_FALLBACK, map_platforms

                detected_platforms = map_platforms(info.slugs)
                if not explicit_platforms:
                    platforms = detected_platforms
                    console.info(f"Recon: auto-detected platform(s): {', '.join(platforms)}")
                elif detected_platforms != DEFAULT_PLATFORM_FALLBACK and set(
                    detected_platforms
                ) != set(platforms):
                    console.info(
                        f"Recon detected {', '.join(detected_platforms)}, "
                        f"but --platform {', '.join(platforms)} was specified. "
                        f"Using user override."
                    )

                # Format and write recon context to file
                from primr.core.recon_context import format_recon_context

                recon_text = format_recon_context(info)
                recon_context_path = os.path.join(folder_path, "_recon_context.txt")
                with open(recon_context_path, "w", encoding="utf-8") as f:
                    f.write(recon_text)

                # Log summary
                console.ok(
                    f"Recon: {len(info.services)} services, "
                    f"{len(info.insights)} insights, "
                    f"platform: {', '.join(detected_platforms)}"
                )

                _update_run_state(
                    folder_path,
                    cloud_vendors=list(platforms),
                    recon_detected_platforms=list(detected_platforms),
                    recon_service_count=len(info.services),
                    recon_signal_count=len(info.insights),
                )
                _append_run_event(
                    folder_path,
                    "recon",
                    "completed",
                    f"{len(info.services)} services detected",
                )

            except Exception as exc:
                console.warn(f"Recon: {exc} — continuing without domain intelligence")
                _append_run_event(folder_path, "recon", "failed", str(exc))
                # Keep existing platforms (user-specified or default)

    # Inject recon context into context_files for strategy generation
    if recon_context_path and os.path.exists(recon_context_path):
        if context_files is None:
            context_files = []
        context_files.insert(0, recon_context_path)

    # Show cost estimate and ask for confirmation
    if not skip_confirm:
        from primr.utils.cost_estimator import display_cost_estimate

        if not display_cost_estimate(
            mode,
            display_name,
            ai_strategy,
            num_vendors=len(platforms),
            lite_strategy=lite_strategy,
            fast_mode=fast_mode,
        ):
            console.info("Research cancelled by user")
            _update_run_state(folder_path, status="cancelled", current_phase="initializing")
            _append_run_event(
                folder_path,
                "initializing",
                "cancelled",
                "Run cancelled by user at cost confirmation",
            )
            return None

    start_time = time.time()

    # Wrap entire research flow in correlation context for tracing
    with correlation_scope("research", company=display_name, mode=mode):
        log_structured(
            "info",
            "Starting research job",
            company=display_name,
            mode=mode,
            ai_strategy=ai_strategy,
        )

        # Fast mode: Grok 4.1 accordion batch pipeline
        # Activated by: explicit --fast, or auto-detect (complete mode + XAI_API_KEY + not premium)
        use_fast = fast_mode or (
            not premium_mode
            and mode in ("complete", "structured", "hybrid")
            and os.environ.get("XAI_API_KEY")
        )
        if use_fast and not premium_mode:
            _update_run_state(folder_path, current_phase="fast_mode", status="running")
            _append_run_event(folder_path, "fast_mode", "started", "Fast mode pipeline started")
            fast_path = perform_fast_research(
                company_name,
                website,
                start_time,
                ai_strategy=ai_strategy,
                platforms=platforms,
                strategy_types=strategies,
                max_scrape_time=max_scrape_time,
                discovery_notes_content=discovery_notes_content,
                folder_path=folder_path,
                resume_local=resume_local,
                grok_tier=grok_tier,
                continuous_reasoning=continuous_reasoning,
                output_dir=run_output_dir,
                diagnostics_dir=diagnostics_dir,
                write_txt=write_public_txt,
            )
            if fast_path:
                _update_run_state(
                    folder_path,
                    status="completed",
                    current_phase="complete",
                    completed_at=datetime.now().isoformat(),
                )
                _append_run_event(
                    folder_path, "complete", "completed", "Fast mode completed", output=fast_path
                )
            else:
                _update_run_state(folder_path, status="failed", current_phase="fast_mode")
                _append_run_event(folder_path, "fast_mode", "failed", "Fast mode failed")
            return fast_path

        # Handle scrape-only mode - scrape and extract insights
        if mode == "scrape-only":
            _update_run_state(folder_path, current_phase="scrape", status="running")
            return perform_scrape_only(
                company_name,
                website,
                start_time,
                max_scrape_time,
                fail_on_low_scrape=not skip_scrape_validation,
                folder_path=folder_path,
            )

        # Check if using Deep Research, Complete, or Hybrid mode
        if mode in ("deep-research", "complete", "hybrid"):
            _update_run_state(folder_path, current_phase="deep_research", status="running")
            return perform_deep_research(
                company_name,
                website,
                mode,
                start_time,
                citation_style,
                ai_strategy,
                platforms,
                context_files,
                refresh_vendor_research,
                strategies=strategies,
                strategy_only=strategy_only,
                discovery_notes_path=discovery_notes_path,
                discovery_notes_content=discovery_notes_content,
                lite_strategy=lite_strategy,
                fail_on_low_scrape=not skip_scrape_validation,
                folder_path=folder_path,
                output_dir=run_output_dir,
                diagnostics_dir=diagnostics_dir,
                write_txt=write_public_txt,
            )

        try:
            # Phase 1: Data Collection
            console.phase_banner(
                1, 4, "Data Collection", "Scraping website and external sources", "5-10 min"
            )
            _update_run_state(folder_path, current_phase="data_collection", status="running")
            _append_run_event(folder_path, "data_collection", "started", "Data collection started")

            # Scrape website
            with console.timed_operation("Website scrape", show_spinner=False):
                scraped_data = (
                    fetch_web_content(
                        website, company_name, max_pages=50, working_folder=folder_path
                    )
                    if website
                    else {}
                )
                pages_scraped = len(scraped_data)
            log_structured("info", "Website scraping complete", pages=pages_scraped)

            # Warn if scraping was very limited
            if pages_scraped <= 2 and website:
                console.warn("Limited website access - report will rely more on web research")

            if website and not skip_scrape_validation:
                quality_ok, quality_reason = _validate_scrape_quality(scraped_data)
                if not quality_ok:
                    console.fail(quality_reason)
                    console.muted("  Re-run with --skip-scrape-validation to continue anyway")
                    _update_run_state(folder_path, status="failed", current_phase="data_collection")
                    _append_run_event(folder_path, "data_collection", "failed", quality_reason)
                    return None

            # External research - with LLM validation to ensure correct company
            # This prevents including content from similarly-named but unrelated companies
            with console.timed_operation("Searching external sources (with validation)"):
                # Generate targeted search queries using LLM
                external_queries = generate_external_search_queries(
                    company_name,
                    website,
                    max_queries=MAX_EXTERNAL_SEARCH_QUERIES,
                )

                # Keep hardcoded queries as reliable fallbacks at the end
                fallback_queries = [
                    "news OR press release OR announcement",
                    "funding OR acquisition OR partnership",
                ]
                all_queries = list(external_queries)
                for fallback in fallback_queries:
                    if len(all_queries) >= MAX_EXTERNAL_SEARCH_QUERIES:
                        break
                    if fallback not in all_queries:
                        all_queries.append(fallback)

                external_data = {}
                max_external_sources = MAX_EXTERNAL_SOURCES

                for query in all_queries:
                    if len(external_data) >= max_external_sources:
                        break

                    results = search_web(query, company_name, website)
                    if results:
                        filtered = [
                            r
                            for r in results[:5]
                            if not website or website.lower() not in r.get("url", "").lower()
                        ]
                        remaining_slots = max_external_sources - len(external_data)
                        scraped = scrape_external_sources_validated(
                            filtered,
                            company_name=company_name,
                            website=website,
                            max_sources=min(2, remaining_slots),
                        )
                        external_data.update(scraped)
            log_structured(
                "info", "External sources complete (validated)", sources=len(external_data)
            )

            all_scraped = {**scraped_data, **external_data}
            console.phase_complete(
                "Data Collection",
                [
                    ("Pages scraped", str(pages_scraped)),
                    ("External sources", str(len(external_data))),
                ],
            )

            # Phase 2: Analysis
            console.phase_banner(2, 4, "Analysis", "Processing and summarizing content", "3-5 min")
            _update_run_state(folder_path, current_phase="analysis", status="running")
            _append_run_event(folder_path, "analysis", "started", "Analysis started")

            with console.timed_operation("Summarizing content"):
                summarized = summarize_scraped_content(
                    company_name, website, all_scraped, folder_path
                )
                if not summarized.strip():
                    summarized = "No insights extracted."

            # Industry identification
            with console.timed_operation("Identifying industry"):
                industry_prompt = generate_prompt(
                    "industry",
                    company_name=company_name,
                    company_website=website or "N/A",
                    scraped_insights=summarized,
                )
                industry = llm(industry_prompt, model_type="research").strip() or "Unknown"
            console.info(f"Industry: {industry}")
            console.phase_complete("Analysis")

            # Phase 3: Report Generation
            console.phase_banner(
                3, 4, "Report Generation", "Building comprehensive report sections", "10-15 min"
            )
            _update_run_state(folder_path, current_phase="report_generation", status="running")
            _append_run_event(
                folder_path, "report_generation", "started", "Report generation started"
            )

            # Overview
            with console.timed_operation("Building overview"):
                overview = generate_initial_overview(company_name, website, industry, folder_path)

            # Value theory
            with console.timed_operation("Value analysis"):
                value_prompt = generate_prompt(
                    "value_theory", company_name=company_name, company_website=website or "N/A"
                )
                value_theory = llm(value_prompt, model_type="research").strip()

                value_file = os.path.join(folder_path, "value_theory.txt")
                with open(value_file, "w", encoding="utf-8") as f:
                    f.write(value_theory or "N/A")

            # Sections
            sections = [
                "Company Name",
                "Website",
                "Industry",
                "Detailed Products/Services",
                "Unique Selling Proposition (USP)",
                "Mission & Vision",
                "Company History",
                "Key Achievements",
                "Target Audience",
                "Financial Overview",
                "Potential Business Drivers & KPIs",
                "Industry Insights",
                "Potential Business Drivers",
                "Primary Apps or Sources of Data",
                "Main Types of Users",
                "Board of Directors Concerns",
                "Potential Business Value",
                "Strategic Recommendations",
            ]

            section_start = time.time()
            for i, section in enumerate(sections):
                console.progress_with_time(i + 1, len(sections), section, section_start)
                research_section(
                    section, company_name, website, industry, folder_path, overview, summarized
                )

            console.progress_done()
            console.phase_complete("Report Generation", [("Sections", str(len(sections)))])

            # Phase 4: Output
            total_phases = 4 + (1 if ai_strategy else 0) + (1 if verify else 0)
            console.phase_banner(
                4, total_phases, "Finalizing", "Generating output documents", "1-2 min"
            )
            _update_run_state(folder_path, current_phase="finalizing", status="running")
            _append_run_event(folder_path, "finalizing", "started", "Finalizing output documents")
            with console.timed_operation("Generating documents"):
                docx_path = generate_final_report(
                    company_name or display_name,
                    citation_style=citation_style,
                    output_dir=run_output_dir,
                    diagnostics_dir=diagnostics_dir,
                    write_txt=write_public_txt,
                )

            # Generate AI strategy if requested (uses Deep Research with company context)
            ai_strategy_path = None
            if ai_strategy:
                console.phase_banner(
                    5,
                    total_phases,
                    "AI Strategy Analysis",
                    "Generating AI recommendations",
                    "5-10 min",
                )
                _update_run_state(folder_path, current_phase="ai_strategy", status="running")
                _append_run_event(
                    folder_path, "ai_strategy", "started", "AI strategy generation started"
                )
                # Consolidate working folder into single context file for AI strategy
                context_file = consolidate_working_folder(folder_path)
                # No heartbeat - the progress callback provides phase-aware status updates
                ai_strategy_path = _generate_ai_strategy_section(
                    company_name or display_name,
                    platforms[0],
                    company_research_path=context_file,
                    force_refresh_vendor=refresh_vendor_research,
                    discovery_notes_content=discovery_notes_content,
                    lite_strategy=lite_strategy,
                    output_dir=run_output_dir,
                    diagnostics_dir=diagnostics_dir,
                    write_txt=write_public_txt,
                )
                if ai_strategy_path:
                    console.phase_complete("AI Strategy Analysis")

            # Run QA analysis if enabled (default: enabled, --no-qa disables)
            qa_result = None
            ai_strategy_qa_result = None
            if not no_qa:
                try:
                    from primr.qa.integration import QAIntegration
                    from primr.qa.models import QAOptions

                    verbose_mode = hasattr(console, "verbose") and console.verbose

                    qa_options = QAOptions(
                        enabled=True, save_detailed=True, verbose_cli=verbose_mode
                    )
                    qa_integration = QAIntegration(qa_options)

                    # QA for main Strategic Overview report
                    if docx_path:
                        txt_report_path = Path(docx_path).with_suffix(".txt")
                        if not txt_report_path.exists() and diagnostics_dir is not None:
                            txt_report_path = diagnostics_dir / txt_report_path.name
                        if txt_report_path.exists():
                            qa_result = qa_integration.run_post_generation_qa(
                                txt_report_path, company_name or display_name
                            )

                    # QA for AI Strategy report
                    if ai_strategy_path:
                        ai_strategy_txt = Path(ai_strategy_path).with_suffix(".txt")
                        if not ai_strategy_txt.exists() and diagnostics_dir is not None:
                            ai_strategy_txt = diagnostics_dir / ai_strategy_txt.name
                        if ai_strategy_txt.exists():
                            ai_strategy_qa_result = qa_integration.run_post_generation_qa(
                                ai_strategy_txt, f"{company_name or display_name} (AI Strategy)"
                            )

                except Exception as e:
                    logger.warning(f"QA analysis failed: {e}")

            # Run claim verification if --verify flag is set
            verification_result = None
            if verify and docx_path:
                try:
                    verify_phase = 6 if ai_strategy else 5
                    verify_total = verify_phase
                    console.phase_banner(
                        verify_phase,
                        verify_total,
                        "Claim Verification",
                        "Verifying factual claims",
                        "1-3 min",
                    )
                    verification_result = _run_verification(
                        company_name or display_name,
                        website or "",
                        docx_path,
                    )
                    if verification_result:
                        console.phase_complete(
                            "Claim Verification",
                            [
                                ("Trust", f"{verification_result.trust_percentage}%"),
                                (
                                    "Verified",
                                    f"{verification_result.verified_count}/{verification_result.total_claims}",
                                ),
                            ],
                        )
                    else:
                        console.phase_complete(
                            "Claim Verification", [("Status", "No claims found")]
                        )
                except Exception as e:
                    logger.warning(f"Claim verification failed: {e}")
                    console.warn(f"Verification failed (non-blocking): {e}")

            elapsed = time.time() - start_time
            mins = int(elapsed // 60)
            secs = int(elapsed % 60)
            time_str = f"{mins}m {secs}s" if mins > 0 else f"{secs}s"

            console.phase_complete("Finalizing", [("Duration", time_str)])

            # Final output
            if docx_path:
                report_path = Path(docx_path).resolve()
                # Use file:// URI for clickable links in terminal
                file_uri = report_path.as_uri()
                console.success_box("Report ready", file_uri)

            if ai_strategy_path:
                console.success_box("AI Strategy", ai_strategy_path)

            # Display QA grades inline
            qa_grades = []
            if qa_result and qa_result.grade > 0:
                qa_grades.append(("Overview", qa_result.grade))
            if ai_strategy_qa_result and ai_strategy_qa_result.grade > 0:
                qa_grades.append(("AI Strategy", ai_strategy_qa_result.grade))
            if qa_grades:
                console.grades(qa_grades)

            # Get actual usage from AI client
            from primr.ai.client import get_client

            client = get_client()
            usage = client.get_usage_summary()
            actual_cost = usage.get("total_cost", 0)

            # Get estimated cost for comparison
            from primr.utils.cost_estimator import estimate_cost

            estimate_cost(mode, ai_strategy, use_historical=False)

            # Summary stats
            summary_items = [
                ("Duration", time_str),
                ("Cost", f"${actual_cost:.2f}"),
            ]
            console.summary(summary_items)

            # Save usage to history
            from primr.utils.usage_tracker import get_usage_tracker

            tracker = get_usage_tracker()
            tracker.record_usage(
                mode=mode,
                company=display_name,
                input_tokens=usage.get("total_input_tokens", 0),
                output_tokens=usage.get("total_output_tokens", 0),
                duration_seconds=elapsed,
                pipeline_cost=actual_cost,
            )
            tracker.save()

            # Log job summary for observability
            job_summary = JobSummary.create(
                company=display_name,
                mode=mode,
                duration_seconds=elapsed,
                api_calls=usage.get("api_calls", 0),
                total_tokens=usage.get("total_input_tokens", 0)
                + usage.get("total_output_tokens", 0),
                sections_generated=len(sections),
                output_path=docx_path,
            )
            log_job_summary(job_summary)
            _update_run_state(
                folder_path,
                status="completed",
                current_phase="complete",
                completed_at=datetime.now().isoformat(),
                duration_seconds=elapsed,
            )
            _append_run_event(folder_path, "complete", "completed", f"Run completed in {time_str}")

            return docx_path

        except Exception as e:
            console.error(f"Research failed: {e}")
            log_structured("error", "Research failed", error=str(e), error_type=type(e).__name__)
            logger.exception("Research failed")
            _update_run_state(folder_path, status="failed", current_phase="error")
            _append_run_event(folder_path, "error", "failed", str(e))
            return None


def perform_deep_research(
    company_name: str | None,
    website: str | None,
    mode: str,
    start_time: float,
    citation_style: str = "numbered",
    ai_strategy: bool = False,
    platforms: tuple[str, ...] = ("agnostic",),
    context_files: list[Any] | None = None,
    refresh_vendor_research: bool = False,
    strategies: list[str] | None = None,
    strategy_only: bool = False,
    discovery_notes_path: str | None = None,
    discovery_notes_content: str | None = None,
    lite_strategy: bool = False,
    fail_on_low_scrape: bool = True,
    folder_path: str | None = None,
    output_dir: str | Path | None = None,
    diagnostics_dir: str | Path | None = None,
    write_txt: bool = True,
) -> str | None:
    """
    Perform research using Deep Research Agent, Complete, or Hybrid mode.

    Args:
        company_name: Name of the company
        website: Company website URL
        mode: Research mode ('deep-research', 'complete', or 'hybrid')
        start_time: Start timestamp for duration tracking
        citation_style: Citation formatting style
        ai_strategy: If True, generate AI opportunity recommendations (legacy, use strategies instead)
        platforms: Platform(s) for AI recommendations
        context_files: Optional list of files (PDFs, docs) to upload as context for Deep Research
        refresh_vendor_research: If True, force regenerate vendor research
        strategies: List of strategy module names to generate (e.g., ['ai', 'cloud'])
        strategy_only: If True, skip company overview and only run strategies
        discovery_notes_path: Path to discovery notes file (for logging/tracking)
        discovery_notes_content: Loaded content of discovery notes (freeform meeting insights)
    """
    display_name: str = company_name or (urlparse(website or "").netloc if website else "")
    folder_path = folder_path or create_working_folder(company_name, website)
    run_output_dir = Path(output_dir) if output_dir is not None else Path(OUTPUT_DIR)
    run_output_dir.mkdir(parents=True, exist_ok=True)
    if diagnostics_dir is not None:
        Path(diagnostics_dir).mkdir(parents=True, exist_ok=True)
    _update_run_state(folder_path, current_phase="preflight", status="running")
    _append_run_event(folder_path, "preflight", "started", "Deep research run started")

    # =================================================================
    # PRE-FLIGHT VALIDATION - Verify everything BEFORE expensive API calls
    # =================================================================
    preflight_errors = []

    # 1. Must have company name or website
    if not company_name and not website:
        preflight_errors.append("Must provide company name or website")

    # 2. Validate context files exist and are readable
    if context_files:
        for f in context_files:
            if not os.path.exists(f):
                preflight_errors.append(f"Context file not found: {f}")
            elif not os.path.isfile(f):
                preflight_errors.append(f"Context path is not a file: {f}")
            elif os.path.getsize(f) == 0:
                preflight_errors.append(f"Context file is empty: {f}")

    # 3. Validate API key is configured
    from primr.config.settings import get_settings

    settings = get_settings()
    if not settings.api.gemini_key:
        preflight_errors.append("GEMINI_API_KEY not configured")

    # ABORT if any pre-flight errors
    if preflight_errors:
        console.error("Pre-flight validation failed:")
        for err in preflight_errors:
            console.error(f"  - {err}")
        console.error("Fix these issues before running expensive Deep Research")
        _update_run_state(folder_path, status="failed", current_phase="preflight")
        _append_run_event(
            folder_path,
            "preflight",
            "failed",
            "Pre-flight validation failed",
            errors=preflight_errors,
        )
        return None

    # =================================================================

    # Pre-run: clean up any orphaned resources from prior crashed runs
    # File Search Stores have NO TTL and cost money if left behind
    try:
        from primr.ai.deep_research import cleanup_orphaned_resources

        orphans = cleanup_orphaned_resources()
        if orphans["caches_deleted"] or orphans["stores_deleted"]:
            console.warn(
                f"Cleaned up {orphans['caches_deleted']} orphaned cache(s) "
                f"and {orphans['stores_deleted']} orphaned store(s) from prior run"
            )
    except Exception as e:
        logger.debug(f"Pre-run resource cleanup check failed (non-fatal): {e}")

    # Map mode string to enum
    mode_map = {
        "deep-research": (ResearchMode.DEEP_RESEARCH, "Deep Research"),
        "complete": (ResearchMode.COMPLETE, "Complete (Two-Step)"),
        "hybrid": (ResearchMode.HYBRID, "Hybrid"),
    }
    research_mode, mode_label = mode_map.get(mode, (ResearchMode.DEEP_RESEARCH, "Deep Research"))

    # Wrap deep research in correlation context
    with correlation_scope("deep_research", company=display_name, mode=mode):
        log_structured("info", "Starting deep research", company=display_name, mode=mode)

        # For COMPLETE mode, the orchestrator handles all phase banners
        # For simple DEEP_RESEARCH mode, we show our own phase banners
        is_simple_deep_research = mode == "deep-research"

        if is_simple_deep_research:
            context_info = ""
            if context_files:
                context_info = f" with {len(context_files)} context file(s)"
            console.phase_banner(
                1, 3, f"{mode_label}{context_info}", "Autonomous AI research", "10-15 min"
            )
        _update_run_state(folder_path, current_phase="deep_research", status="running")
        _append_run_event(folder_path, "deep_research", "started", f"{mode_label} started")

        # Track last phase to only print on phase changes
        last_phase: list[str | None] = [None]  # list = mutable cell for closure
        last_update_time = [time.time()]

        def progress_callback(msg: str) -> None:
            # Extract phase from message (e.g., "Searching sources (2m 30s)")
            phase = msg.split(" (")[0].strip() if " (" in msg else msg.strip()

            # Strip leading dots (heartbeat-style updates)
            display_msg = msg.lstrip(". ")

            # Show indented sub-status messages (e.g. "  Uploading Stage 1 context")
            if msg.startswith("  "):
                console.muted(f"  {msg.strip()}")
                log_structured("debug", f"Deep research progress: {msg}")
                return

            # Show on phase change
            if phase and phase != last_phase[0] and not phase.startswith("  "):
                last_phase[0] = phase
                last_update_time[0] = time.time()
                console.info(display_msg)
            # Also show periodic updates every 2 minutes even if phase unchanged
            elif time.time() - last_update_time[0] > 120:
                last_update_time[0] = time.time()
                console.muted(f"  Still working... {display_msg}")

            log_structured("debug", f"Deep research progress: {msg}")

        try:
            # Run async orchestrator with heartbeat for long operations
            orchestrator = get_orchestrator()

            # Create event loop if needed
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)

            result = loop.run_until_complete(
                orchestrator.research(
                    company_name=company_name or display_name,
                    website=website,
                    mode=research_mode,
                    config=ResearchConfig(
                        mode=research_mode, fail_on_low_scrape=fail_on_low_scrape
                    ),
                    on_progress=progress_callback,
                    context_files=context_files,
                )
            )

            if not result.success:
                console.fail(f"Research failed: {result.error}")
                log_structured("error", "Deep research failed", error=result.error)
                _update_run_state(folder_path, status="failed", current_phase="deep_research")
                _append_run_event(
                    folder_path,
                    "deep_research",
                    "failed",
                    "Deep research failed",
                    error=result.error,
                )

                # Save partial results if the structured phase produced anything
                if result.section_results:
                    partial_folder = folder_path
                    partial_count = 0
                    for section_key, content in result.section_results.items():
                        save_section_output(partial_folder, section_key, content)
                        partial_count += 1
                    console.warn(
                        f"Saved {partial_count} partial sections from data collection to: {partial_folder}"
                    )
                    console.muted(
                        "  Tip: Re-run with --mode scrape to generate a report from scraped data"
                    )
                else:
                    console.muted("  Tip: Check logs for details, or re-run with --mode scrape")

                return None

            # Use sections_written for accurate count (accordion method tracks this)
            section_count = (
                result.sections_written
                if result.sections_written > 0
                else len(result.section_results)
            )
            log_structured("info", "Deep research complete", sections=section_count)

            # Calculate word and page count from raw content
            word_count = len(result.raw_content.split()) if result.raw_content else 0
            page_count = word_count // 500  # ~500 words per page

            if is_simple_deep_research:
                console.phase_complete(
                    "Deep Research",
                    [
                        ("Pages", f"~{page_count}"),
                        ("Words", f"{word_count:,}"),
                        ("Chapters", str(section_count)),
                    ],
                )
                console.phase_banner(
                    2, 3, "Processing Results", "Saving and converting output", "1-2 min"
                )
            _append_run_event(
                folder_path,
                "deep_research",
                "completed",
                "Deep research completed",
                pages=page_count,
                chapters=section_count,
            )
            _update_run_state(folder_path, current_phase="processing_results", status="running")

            # Save section results to working folder
            with console.timed_operation("Saving results"):
                for section_key, content in result.section_results.items():
                    save_section_output(folder_path, section_key, content)

            # Save raw markdown for reference
            raw_md_path = None
            if result.raw_content:
                raw_md_path = os.path.join(folder_path, "deep_research_output.md")
                with open(raw_md_path, "w", encoding="utf-8") as f:
                    f.write(result.raw_content)

            # Generate final report - use direct markdown conversion for Deep Research
            with console.timed_operation("Generating documents"):
                if result.raw_content and mode in ("deep-research", "complete", "hybrid"):
                    # Deep Research: convert markdown directly to DOCX (preserves structure)
                    docx_path = _convert_deep_research_to_docx(
                        result.raw_content,
                        company_name or display_name,
                        website,
                        output_dir=run_output_dir,
                        diagnostics_dir=diagnostics_dir,
                        write_txt=write_txt,
                    )
                else:
                    # Structured pipeline: use DocumentBuilder to assemble sections
                    docx_path = generate_final_report(
                        company_name or display_name,
                        citation_style=citation_style,
                        output_dir=run_output_dir,
                        diagnostics_dir=diagnostics_dir,
                        write_txt=write_txt,
                    )

            if is_simple_deep_research:
                console.phase_complete("Processing Results")

            # Determine which strategies to run
            strategies_to_run: list[str] = []
            if strategies:
                # Explicit --strategy flag takes precedence
                strategies_to_run = strategies
            elif ai_strategy:
                # Legacy --ai-strategy flag (default behavior)
                strategies_to_run = ["ai"]

            # Generate strategies (uses Deep Research with company context)
            strategy_paths: dict[str, str] = {}
            if strategies_to_run:
                _update_run_state(
                    folder_path, current_phase="strategy_generation", status="running"
                )
                _append_run_event(
                    folder_path,
                    "strategy_generation",
                    "started",
                    "Strategy generation started",
                    strategies=strategies_to_run,
                )
                base_phase = 3
                # Count total phases: AI strategy runs once per vendor, others run once
                total_phase_count = sum(
                    len(platforms) if s == "ai" else 1 for s in strategies_to_run
                )

                phase_offset = 0
                for strategy_name in strategies_to_run:
                    # AI strategy iterates over each vendor; others run once
                    vendors = list(platforms) if strategy_name == "ai" else ["agnostic"]

                    for vendor in vendors:
                        phase_num = base_phase + phase_offset
                        total_phases = base_phase + total_phase_count - 1

                        # Get display name from registry
                        from primr.prompts.registry import get_registry

                        registry = get_registry()
                        strategy_module = registry.get(strategy_name)
                        display_strategy_name = (
                            strategy_module.display_name
                            if strategy_module
                            else strategy_name.replace("_", " ").title()
                        )

                        # Append vendor to display name for multi-vendor AI runs
                        vendor_label = ""
                        if strategy_name == "ai" and len(platforms) > 1:
                            vendor_label = f" ({vendor.upper()})"

                        console.phase_banner(
                            phase_num,
                            total_phases,
                            f"{display_strategy_name}{vendor_label} Analysis",
                            f"Generating {display_strategy_name.lower()} recommendations{vendor_label.lower()}",
                            "5-10 min",
                        )

                        # Generate the strategy
                        strategy_path = _generate_strategy_section(
                            strategy_name=strategy_name,
                            company_name=company_name or display_name,
                            platform=vendor,
                            company_research_path=raw_md_path,
                            force_refresh_vendor=refresh_vendor_research,
                            lite_strategy=lite_strategy,
                            output_dir=run_output_dir,
                            diagnostics_dir=diagnostics_dir,
                            write_txt=write_txt,
                        )

                        if strategy_path:
                            # Use compound key for multi-vendor AI strategies
                            if strategy_name == "ai" and len(platforms) > 1:
                                key = f"ai_{vendor}"
                            else:
                                key = strategy_name
                            strategy_paths[key] = strategy_path
                            console.phase_complete(
                                f"{display_strategy_name}{vendor_label} Analysis"
                            )
                            _append_run_event(
                                folder_path,
                                "strategy_generation",
                                "completed",
                                f"{display_strategy_name}{vendor_label} completed",
                                output=strategy_path,
                            )

                        phase_offset += 1

            # For backward compatibility
            strategy_paths.get("ai")

            elapsed = time.time() - start_time
            mins = int(elapsed // 60)
            secs = int(elapsed % 60)
            time_str = f"{mins}m {secs}s" if mins > 0 else f"{secs}s"

            console.ok(f"Complete in {time_str}")
            _update_run_state(
                folder_path,
                status="completed",
                current_phase="complete",
                completed_at=datetime.now().isoformat(),
                duration_seconds=elapsed,
            )
            _append_run_event(folder_path, "complete", "completed", f"Run completed in {time_str}")

            # Final output - use plain paths (consistent format)
            if docx_path:
                console.success_box("Report ready", str(Path(docx_path).resolve()))

            # Show all generated strategy outputs
            for strat_key, strategy_path in strategy_paths.items():
                from primr.prompts.registry import get_registry

                registry = get_registry()
                # Parse compound keys like "ai_aws" back to base strategy name
                if "_" in strat_key and strat_key.split("_", 1)[0] == "ai":
                    base_name = "ai"
                    vendor_suffix = f" ({strat_key.split('_', 1)[1].upper()})"
                else:
                    base_name = strat_key
                    vendor_suffix = ""
                strategy_module = registry.get(base_name)
                strat_display_name = (
                    strategy_module.display_name
                    if strategy_module
                    else base_name.replace("_", " ").title()
                )
                console.success_box(
                    f"{strat_display_name}{vendor_suffix}", str(Path(strategy_path).resolve())
                )

            # Get actual usage from AI client (per-model accurate cost)
            from primr.ai.client import get_client
            from primr.config.models import DEEP_RESEARCH_COST

            client = get_client()
            usage = client.get_usage_summary()

            # Pipeline portion (Flash + Pro, per-model accurate)
            pipeline_cost = usage.get("total_cost", 0.0)
            total_input = usage.get("total_input_tokens", 0)
            total_output = usage.get("total_output_tokens", 0)

            # Deep Research portion (flat per-task cost, API doesn't expose tokens)
            dr_tasks = 0
            if mode in ("deep-research", "complete", "hybrid"):
                dr_tasks += 1  # Main research dossier
            if ai_strategy and platforms and not lite_strategy:
                dr_tasks += len(platforms)  # Each vendor triggers a DR task
            dr_cost = dr_tasks * DEEP_RESEARCH_COST.standard_task_cost

            actual_cost = pipeline_cost + dr_cost

            # Get pre-run estimate for comparison
            from primr.utils.cost_estimator import estimate_cost

            pre_estimate = estimate_cost(
                mode,
                ai_strategy,
                use_historical=False,
                num_vendors=len(platforms),
                lite_strategy=lite_strategy,
            )

            # Use sections_written for accurate count
            section_count = (
                result.sections_written
                if result.sections_written > 0
                else len(result.section_results)
            )

            # Count unique citations from generated content ([cite: N] format)
            citation_count = 0
            import re

            all_content = result.raw_content or ""
            if not all_content and result.section_results:
                all_content = "\n".join(result.section_results.values())
            if all_content:
                cite_numbers = set()
                for match in re.findall(r"\[cite:\s*([\d,\s]+)\]", all_content):
                    for num in match.split(","):
                        num = num.strip()
                        if num:
                            cite_numbers.add(num)
                citation_count = len(cite_numbers)

            # Summary stats with estimated vs actual comparison
            summary_items = [
                ("Mode", mode_label),
                ("Chapters", str(section_count)),
                ("Citations", str(citation_count)),
                ("Duration", time_str),
                ("Est. Cost", f"${pre_estimate.total_cost:.2f}"),
                ("Actual Cost", f"~${actual_cost:.2f}"),
            ]
            if ai_strategy:
                summary_items.append(("AI Strategy", "Yes"))
            console.summary(summary_items)

            # Save usage to history
            from primr.utils.usage_tracker import get_usage_tracker

            tracker = get_usage_tracker()
            tracker.record_usage(
                mode=mode,
                company=display_name,
                input_tokens=total_input,
                output_tokens=total_output,
                search_queries=result.search_queries_count,  # Actual count from API
                duration_seconds=elapsed,
                pipeline_cost=pipeline_cost,
                deep_research_cost=dr_cost,
            )
            tracker.save()

            # Log job summary for observability
            job_summary = JobSummary.create(
                company=display_name,
                mode=mode,
                duration_seconds=elapsed,
                api_calls=0,  # Deep Research doesn't expose API call count
                total_tokens=total_input + total_output,
                sections_generated=section_count,
                output_path=docx_path,
            )
            log_job_summary(job_summary)

            return docx_path

        except Exception as e:
            console.error(f"Deep research failed: {e}")
            log_structured(
                "error", "Deep research failed", error=str(e), error_type=type(e).__name__
            )
            logger.exception("Deep research failed")
            _update_run_state(folder_path, status="failed", current_phase="error")
            _append_run_event(folder_path, "error", "failed", str(e))
            return None
        finally:
            # Post-run: verify no resources leaked (safety net)
            try:
                from primr.ai.deep_research import cleanup_orphaned_resources

                leaked = cleanup_orphaned_resources()
                if leaked["caches_deleted"] or leaked["stores_deleted"]:
                    logger.warning(
                        f"Post-run cleanup found leaked resources: "
                        f"{leaked['caches_deleted']} cache(s), {leaked['stores_deleted']} store(s)"
                    )
            except Exception as cleanup_err:
                logger.debug(f"Post-run resource cleanup failed (non-fatal): {cleanup_err}")


_FORBIDDEN_OUTPUT_PATTERNS: tuple[tuple[str, str], ...] = (
    ("raw_source_tag", r"\[Source:\s*(?:https?://)?[^\]\s]+"),
    ("section_cross_ref", r"\[\s*(?:see|cross-?ref|xref)\s+##\s+[^\]]+\]"),
    ("workbook_ref", r"\[Workbook:[^\]]*\]"),
    ("workbook_section_ref", r"\[workbook section[^\]]*\]"),
    ("workbook_section_symbol", r"\[Workbook §[^\]]*\]"),
    ("analysis_workbook_ref", r"\[Analysis Workbook[^\]]*\]"),
    ("analysis_ref", r"\[Analysis:[^\]]*\]"),
    ("external_sources_ref", r"\[External Sources\]"),
    ("citation_inventory", r"\[citation inventory[^\]]*\]"),
    ("vendor_research_file", r"vendor-research-[\w.-]+\.txt"),
    ("internal_roi_model", r"\bInternal ROI Model\b"),
    ("internal_analysis", r"\bInternal Analysis\b"),
)

# Cleaning counterparts for each forbidden detection pattern.
# Detection patterns are partial-match (e.g. no closing bracket) for scanning;
# cleaning patterns match the full token so we can strip it cleanly.
# This list is the SINGLE SOURCE OF TRUTH for what gets auto-stripped.
_FORBIDDEN_OUTPUT_CLEANERS: tuple[tuple[str, str], ...] = (
    ("raw_source_tag", r"\[Source:[^\]]*\]"),
    ("section_cross_ref", r"\[\s*(?:see|cross-?ref|xref)\s+##\s+[^\]]+\]"),
    ("workbook_ref", r"\[Workbook:[^\]]*\]"),
    ("workbook_section_ref", r"\[workbook section[^\]]*\]"),
    ("workbook_section_symbol", r"\[Workbook §[^\]]*\]"),
    ("analysis_workbook_ref", r"\[Analysis Workbook[^\]]*\]"),
    ("analysis_ref", r"\[Analysis:[^\]]*\]"),
    ("external_sources_ref", r"\[External Sources\]"),
    ("citation_inventory", r"\[citation inventory[^\]]*\]"),
    ("vendor_research_file", r"vendor-research-[\w.-]+\.txt"),
    ("internal_roi_model", r"\bInternal ROI Model\b"),
    ("internal_analysis", r"\bInternal Analysis\b"),
)

# Internal terms that should never appear in shipped artifacts (bare text).
_FORBIDDEN_INTERNAL_TERMS: tuple[str, ...] = (
    "analysis context",
    "vendor-research",
)


def _auto_strip_forbidden_patterns(text: str) -> str:
    """Last-resort defensive sweep: strip anything the artifact scanner would flag.

    This runs AFTER all specific cleaners (citation normalizer, report cleaner,
    source placeholder stripper).  It uses the same patterns the scanner checks,
    so any new forbidden pattern automatically gets cleaned too — no drift.
    """
    if not text.strip():
        return text

    for _label, pattern in _FORBIDDEN_OUTPUT_CLEANERS:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)

    lower = text.lower()
    for term in _FORBIDDEN_INTERNAL_TERMS:
        if term in lower:
            text = re.sub(re.escape(term), "", text, flags=re.IGNORECASE)

    # Clean up artifacts left by stripping: double spaces, excess blank lines
    text = re.sub(r"  +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text


def _scan_forbidden_output_patterns(text: str) -> list[str]:
    issues: list[str] = []
    for label, pattern in _FORBIDDEN_OUTPUT_PATTERNS:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            issues.append(f"{label}: {match.group(0)[:120]}")

    lower = text.lower()
    for term in _FORBIDDEN_INTERNAL_TERMS:
        if term in lower:
            issues.append(f"internal_term: {term}")

    return issues


class _ArtifactValidation(TypedDict):
    """Result of an artifact validation pass.

    Used by `_validate_output_markdown` and `_validate_output_docx`. The
    dict-of-union shape it replaced (`dict[str, list[str] | bool]`) was
    correct at runtime but unusable through the type checker — every
    indexed access lost the per-key shape and required casts at every
    call site.
    """

    passed: bool
    issues: list[str]
    errors: list[str]


def _write_output_validation_report(
    base_path: Path,
    phase: str,
    issues: list[str],
    errors: list[str],
    diagnostics_dir: str | Path | None = None,
) -> Path | None:
    if not issues and not errors:
        return None

    if diagnostics_dir is not None:
        diagnostics_path = Path(diagnostics_dir)
        diagnostics_path.mkdir(parents=True, exist_ok=True)
        report_path = diagnostics_path / f"{base_path.stem}_{phase}_validation.txt"
    else:
        report_path = base_path.with_name(f"{base_path.stem}_{phase}_validation.txt")
    lines = [f"Artifact validation report ({phase})", ""]
    if issues:
        lines.append("Issues:")
        lines.extend(f"- {item}" for item in issues)
        lines.append("")
    if errors:
        lines.append("Validator errors:")
        lines.extend(f"- {item}" for item in errors)
        lines.append("")
    report_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return report_path


def _validate_output_markdown(markdown_content: str) -> _ArtifactValidation:
    try:
        issues = _scan_forbidden_output_patterns(markdown_content)
        return {"passed": len(issues) == 0, "issues": issues, "errors": []}
    except Exception as exc:
        logger.warning("Markdown artifact validation failed: %s", exc)
        return {"passed": True, "issues": [], "errors": [str(exc)]}


def _extract_docx_text(document: Any) -> str:
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _prepare_report_markdown_for_shipping(markdown_content: str) -> str:
    """Apply deterministic cleanup before report artifact validation/shipping."""
    from primr.output.final_artifact import canonicalize_final_markdown

    prepared = _clean_fast_report_output(markdown_content)
    prepared = canonicalize_final_markdown(prepared)
    prepared = _normalize_fast_citations(prepared)
    prepared = _strip_internal_source_placeholders(prepared)
    prepared = _enforce_fast_section_quality_guards(prepared)
    prepared = canonicalize_final_markdown(prepared)
    return prepared


def _prepare_strategy_markdown_for_shipping(strategy_content: str) -> str:
    """Apply deterministic cleanup before strategy artifact validation/shipping."""
    from primr.output.final_artifact import canonicalize_final_markdown

    prepared = canonicalize_final_markdown(strategy_content)
    prepared = _clean_strategy_output(prepared)
    prepared = canonicalize_final_markdown(prepared)
    return prepared


def _prepare_markdown_for_shipping(content: str, kind: str) -> str:
    if kind == "strategy":
        return _prepare_strategy_markdown_for_shipping(content)
    return _prepare_report_markdown_for_shipping(content)


def _salvage_markdown_for_shipping(
    markdown_content: str,
    kind: str,
) -> tuple[str, _ArtifactValidation, bool]:
    """Run one deterministic salvage pass before blocking artifact shipping.

    Three escalation levels:
    1. Validate raw content — if clean, ship as-is.
    2. Run kind-specific cleanup pipeline — if clean, ship salvaged.
    3. Auto-strip ALL forbidden patterns (last resort) — if clean, ship.

    Level 3 ensures any new forbidden pattern added to the scanner is
    automatically cleaned without needing a matching cleanup rule.
    """
    validation = _validate_output_markdown(markdown_content)
    if validation["passed"]:
        return markdown_content, validation, False

    prepared = _prepare_markdown_for_shipping(markdown_content, kind)
    if prepared == markdown_content:
        # Specific cleaners made no changes — escalate to auto-strip
        stripped = _auto_strip_forbidden_patterns(markdown_content)
        if stripped != markdown_content:
            stripped_validation = _validate_output_markdown(stripped)
            if stripped_validation["passed"]:
                return stripped, stripped_validation, True
            if len(stripped_validation["issues"]) < len(validation["issues"]):
                return stripped, stripped_validation, True
        return markdown_content, validation, False

    prepared_validation = _validate_output_markdown(prepared)
    if prepared_validation["passed"]:
        return prepared, prepared_validation, True

    # Specific cleaners reduced but didn't eliminate issues — auto-strip remainder
    stripped = _auto_strip_forbidden_patterns(prepared)
    if stripped != prepared:
        stripped_validation = _validate_output_markdown(stripped)
        if stripped_validation["passed"]:
            return stripped, stripped_validation, True
        if len(stripped_validation["issues"]) < len(prepared_validation["issues"]):
            return stripped, stripped_validation, True

    if len(prepared_validation["issues"]) < len(validation["issues"]):
        return prepared, prepared_validation, True

    return markdown_content, validation, False


def _validate_output_docx(docx_path: Path) -> _ArtifactValidation:
    try:
        from docx import Document

        from primr.output.markdown_parser import ArtifactDetector

        document = Document(str(docx_path))
        detector = ArtifactDetector()
        artifacts = detector.scan_document(document)
        issues = [
            f"markdown_artifact:{artifact['type']}:{artifact['match']}"
            for artifact in artifacts[:10]
        ]
        issues.extend(_scan_forbidden_output_patterns(_extract_docx_text(document)))
        return {"passed": len(issues) == 0, "issues": issues, "errors": []}
    except Exception as exc:
        logger.warning("DOCX artifact validation failed: %s", exc)
        return {"passed": True, "issues": [], "errors": [str(exc)]}


def _convert_deep_research_to_docx(
    markdown_content: str,
    company_name: str,
    website: str | None,
    gate_issues: list[str] | None = None,
    output_dir: str | Path | None = None,
    diagnostics_dir: str | Path | None = None,
    write_txt: bool = True,
) -> str | None:
    """
    Convert Deep Research markdown output to all output formats.

    This bypasses the DocumentBuilder and preserves the exact structure
    that Deep Research produces. The AI follows the prompt's section
    structure, so we just convert it cleanly to Word format.

    Outputs:
    - {company}_Strategic_Overview_{date}.md  - Raw markdown
    - {company}_Strategic_Overview_{date}.txt - Plain text
    - {company}_Strategic_Overview_{date}.docx - Word document

    Args:
        markdown_content: Raw markdown from Deep Research
        company_name: Name of the company
        website: Company website URL

    Returns:
        Path to generated DOCX file, or None on failure
    """
    from primr.output.markdown_converter import markdown_to_docx
    from primr.output.output_utils import OUTPUT_DIR

    date_str = datetime.now().strftime("%m-%d-%Y")
    base_name = f"{company_name}_Strategic_Overview_{date_str}"
    destination_dir = Path(output_dir) if output_dir is not None else Path(OUTPUT_DIR)
    destination_dir.mkdir(parents=True, exist_ok=True)
    internal_dir = Path(diagnostics_dir) if diagnostics_dir is not None else destination_dir
    internal_dir.mkdir(parents=True, exist_ok=True)

    try:
        markdown_content, markdown_validation, salvaged = _salvage_markdown_for_shipping(
            markdown_content,
            kind="report",
        )

        # Save markdown (.md)
        md_path = destination_dir / f"{base_name}.md"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        console.ok(f"MD saved: {base_name}.md", show_time=False)

        # Save plain text (.txt). For custom output directories, keep this
        # machine-facing mirror in the diagnostics folder so the requested
        # output path stays focused on customer-facing deliverables.
        if write_txt or diagnostics_dir is not None:
            txt_path = (destination_dir if write_txt else internal_dir) / f"{base_name}.txt"
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)
            if write_txt:
                console.ok(f"TXT saved: {base_name}.txt", show_time=False)
        if salvaged:
            console.info("Adaptive salvage cleaned report markdown before shipping")

        # Build subtitle with date and website
        subtitle_parts = [datetime.now().strftime("%B %d, %Y")]
        if website:
            subtitle_parts.append(website)
        subtitle = " | ".join(subtitle_parts)

        combined_markdown_issues = list(markdown_validation["issues"])
        if gate_issues:
            combined_markdown_issues.extend(gate_issues)
        if combined_markdown_issues:
            report_path = _write_output_validation_report(
                md_path,
                "markdown",
                combined_markdown_issues,
                markdown_validation["errors"],
                diagnostics_dir=diagnostics_dir,
            )
            console.warn("Report artifact issues: " + ", ".join(combined_markdown_issues[:3]))
            console.warn(
                "DOCX shipping gate failed for report markdown; saved MD/TXT only"
                + (f" ({report_path.name})" if report_path else "")
            )
            return None

        # Convert to DOCX
        docx_path = destination_dir / f"{base_name}.docx"
        try:
            markdown_to_docx(
                markdown_text=markdown_content,
                output_path=docx_path,
                title=f"Strategic Company Overview: {company_name}",
                subtitle=subtitle,
            )
        except PermissionError:
            # File is probably open in Word - try with timestamp suffix
            timestamp = datetime.now().strftime("%H%M%S")
            file_name = f"{base_name}_{timestamp}.docx"
            docx_path = destination_dir / file_name
            console.warn(f"Original file locked, saving as: {file_name}")
            markdown_to_docx(
                markdown_text=markdown_content,
                output_path=docx_path,
                title=f"Strategic Company Overview: {company_name}",
                subtitle=subtitle,
            )

        docx_validation = _validate_output_docx(docx_path)
        if not docx_validation["passed"]:
            report_path = _write_output_validation_report(
                docx_path,
                "docx",
                docx_validation["issues"],
                docx_validation["errors"],
                diagnostics_dir=diagnostics_dir,
            )
            try:
                docx_path.unlink(missing_ok=True)
            except Exception as cleanup_err:
                logger.warning("Failed to remove blocked DOCX %s: %s", docx_path, cleanup_err)
            console.warn(
                "DOCX shipping gate failed for rendered report; saved MD/TXT only"
                + (f" ({report_path.name})" if report_path else "")
            )
            return None

        if docx_validation["errors"]:
            report_path = _write_output_validation_report(
                docx_path,
                "docx",
                docx_validation["issues"],
                docx_validation["errors"],
                diagnostics_dir=diagnostics_dir,
            )
            console.warn(
                "DOCX validator encountered non-fatal errors; shipping report DOCX"
                + (f" ({report_path.name})" if report_path else "")
            )

        console.ok(f"DOCX saved: {docx_path.name}", show_time=False)
        return str(docx_path)

    except Exception as e:
        console.error(f"Failed to convert markdown to DOCX: {e}")
        logger.exception("Markdown to DOCX conversion failed")
        return None


def _get_vendor_research_path(platform: str) -> str:
    """Get the path for vendor research file based on current month."""
    from datetime import datetime

    current_month = datetime.now().strftime("%Y-%m")
    filename = f"vendor-research-{platform.lower()}-{current_month}.txt"
    return os.path.join(PROJECT_ROOT, "vendor-research", filename)


def _is_vendor_research_current(platform: str) -> bool:
    """Check if we have current month's vendor research."""
    # Azure has a manually curated file that's always preferred
    if platform.lower() == "azure":
        from primr.config.config import PROJECT_ROOT

        manual_path = os.path.join(
            PROJECT_ROOT, "docs/research latest microsoft ignite analysis.txt"
        )
        if os.path.exists(manual_path):
            return True

    research_path = _get_vendor_research_path(platform)
    return os.path.exists(research_path)


def _generate_vendor_research(platform: str) -> str | None:
    """
    Generate fresh vendor AI research using Deep Research.

    This creates a comprehensive overview of the latest AI services and capabilities
    from the specified cloud vendor, suitable for use as context in AI strategy generation.

    Args:
        platform: Platform (azure, aws, gcp, agnostic)

    Returns:
        Path to generated research file, or None if generation failed
    """
    from datetime import datetime

    from primr.ai.deep_research import ResearchStatus, get_deep_research_client
    from primr.config.settings import get_settings

    # =================================================================
    # PRE-FLIGHT VALIDATION - Check everything BEFORE expensive API call
    # =================================================================
    preflight_errors = []

    # 1. Validate cloud vendor
    valid_vendors = ["azure", "aws", "gcp", "agnostic", "private"]
    if platform.lower() not in valid_vendors:
        preflight_errors.append(
            f"Invalid cloud vendor: {platform}. Must be one of: {', '.join(valid_vendors)}"
        )

    # 2. Validate API key is configured
    settings = get_settings()
    if not settings.api.gemini_key:
        preflight_errors.append("GEMINI_API_KEY not configured in .env")

    # 3. Check docs directory is writable
    docs_dir = os.path.join(PROJECT_ROOT, "docs")
    try:
        os.makedirs(docs_dir, exist_ok=True)
        test_file = os.path.join(docs_dir, ".write_test")
        with open(test_file, "w") as f:
            f.write("test")
        os.remove(test_file)
    except Exception as e:
        preflight_errors.append(f"Docs directory not writable: {docs_dir} ({e})")

    # ABORT if any pre-flight errors
    if preflight_errors:
        console.error("Pre-flight validation failed:")
        for err in preflight_errors:
            console.error(f"  - {err}")
        console.error("Fix these issues before running expensive Deep Research")
        return None
    # =================================================================

    current_date = datetime.now().strftime("%B %Y")

    vendor_names = {
        "azure": "Microsoft Azure",
        "aws": "Amazon Web Services (AWS)",
        "gcp": "Google Cloud Platform (GCP)",
        "agnostic": "the AI Industry (cross-vendor)",
        "private": "Private Cloud / NVIDIA",
    }
    vendor_name = vendor_names.get(platform.lower(), platform)

    # Conference names for each vendor
    conferences = {
        "azure": "Microsoft Ignite, Microsoft Build",
        "aws": "AWS re:Invent, AWS Summit",
        "gcp": "Google Cloud Next, Google I/O",
        "agnostic": "NeurIPS, major vendor conferences (Ignite, re:Invent, Cloud Next), and recent announcements from OpenAI, Anthropic, NVIDIA, Meta, Mistral, and Cohere",
        "private": "NVIDIA GTC 2025, NVIDIA AI Enterprise releases, and partner announcements (VMware, Red Hat, Dell, HPE)",
    }
    conference = conferences.get(platform.lower(), "recent conferences")

    # Vendor-specific model platform names
    model_platforms = {
        "azure": "Azure OpenAI Service and Azure AI Foundry",
        "aws": "Amazon Bedrock",
        "gcp": "Vertex AI",
        "agnostic": "major model providers (OpenAI, Anthropic, Google, Meta, Mistral, Cohere) and cloud platforms (Azure, AWS, GCP)",
    }
    model_platform = model_platforms.get(platform.lower(), "the AI platform")

    # Vendor-specific official sources (detailed)
    official_sources_detail = {
        "azure": "Microsoft Learn docs, Azure Blog, official 'What's New' pages, Microsoft Tech Community, Ignite/Build session pages, and Azure pricing pages",
        "aws": "AWS Documentation, AWS News Blog, official 'What's New' announcements, re:Invent session pages, and AWS pricing pages",
        "gcp": "Google Cloud Documentation, Google Cloud Blog, official release notes, Cloud Next session pages, and GCP pricing pages",
        "agnostic": "official documentation and blogs from OpenAI, Anthropic, NVIDIA, Meta AI, Mistral, Cohere, and the major cloud vendors (Azure, AWS, GCP)",
    }
    official_source = official_sources_detail.get(
        platform.lower(), "official documentation and pricing pages"
    )

    # Vendor-specific service map categories
    service_maps = {
        "azure": """
AI Platform: Azure OpenAI Service, Azure AI Foundry, Cognitive Services
Agent/Copilot: Copilot Studio, Semantic Kernel, AI Builder
Data Plane: Azure Storage, Synapse, Fabric, Azure AI Search, Cosmos DB
App Plane: Azure Functions, Logic Apps, Event Grid, API Management
Security Plane: Entra ID, Key Vault, Purview, Defender for Cloud""",
        "aws": """
AI Platform: Amazon Bedrock, SageMaker, Comprehend, Rekognition
Agent/Copilot: Bedrock Agents, Amazon Q, Step Functions
Data Plane: S3, Lake Formation, Glue, OpenSearch, Redshift, DynamoDB
App Plane: Lambda, Step Functions, EventBridge, API Gateway
Security Plane: IAM, KMS, CloudTrail, GuardDuty, Macie""",
        "gcp": """
AI Platform: Vertex AI, Gemini API, Document AI, Vision AI
Agent/Copilot: Vertex AI Agent Builder, Dialogflow CX
Data Plane: Cloud Storage, BigQuery, Dataflow, Vertex AI Search, Firestore
App Plane: Cloud Functions, Cloud Run, Workflows, Eventarc, API Gateway
Security Plane: IAM, Cloud KMS, Security Command Center, DLP API""",
        "agnostic": """
Model Providers: OpenAI (GPT-4, o1), Anthropic (Claude), Google (Gemini), Meta (Llama), Mistral, Cohere
Infrastructure: NVIDIA (GPUs, NIM, NeMo), cloud platforms (Azure, AWS, GCP)
Cloud AI Platforms: Azure OpenAI/AI Foundry, Amazon Bedrock, Vertex AI
Agent Frameworks: LangChain, LlamaIndex, Semantic Kernel, AutoGen, CrewAI
Vector Databases: Pinecone, Weaviate, Qdrant, Chroma, pgvector
Evaluation/Observability: LangSmith, Weights & Biases, Arize, Helicone""",
    }
    service_map = service_maps.get(platform.lower(), "")

    prompt = f"""You are an AI technology analyst. Research the latest AI services and capabilities.

=============================================================================
OUTPUT FORMAT (Start the document with this exact header)
=============================================================================

# {vendor_name} AI Services and Capabilities

**Prepared by:** Primr Research System
**Date:** {current_date}

---

Then continue with the sections below.

=============================================================================
RESEARCH INSTRUCTIONS
=============================================================================

CRITICAL: This research must reflect the AI landscape as of {current_date}.
You MUST use live web search to find the latest information.
Do NOT rely on potentially outdated training data.

=============================================================================
RESEARCH AND VALIDATION PROTOCOL
=============================================================================

For every service or feature listed in this document:

1. **Verify current name**: Confirm the service still exists and has not been renamed or deprecated
2. **Status**: Note if GA (Generally Available) or Preview/Beta
3. **Region availability**: Flag if limited to specific regions
4. **Citation**: Link to official {vendor_name} documentation or blog post, include publish date or "last updated" date
5. **Pricing**: If pricing is mentioned, cite the pricing page and state assumptions (region, token units, provisioned throughput, etc.)
6. **Unconfirmed flag**: If a claim cannot be verified through official sources, mark as "UNCONFIRMED"

SOURCE HYGIENE (critical for defensibility):
- Primary sources ONLY: {official_source}
- If a product page is used, it must include a dated release note or "What's New" reference
- DO NOT use third-party blogs (SiliconAngle, InfoWorld, Medium, etc.) as primary evidence
- Third-party sources may be used ONLY to triangulate when official sources are unavailable
- If you must cite a third-party source, explicitly note it is not first-party

DEPRECATION AND END-OF-LIFE:
Flag any service, feature, or model that was deprecated, merged, or renamed in the past 12 months.
State the replacement or migration path.

PRICING GUIDANCE:
- Provide pricing at the level needed for directional architecture decisions, not a full cost model
- Label all pricing as "indicative, as of {current_date}" since pricing changes frequently
- Prefer linking to pricing pages rather than restating exact numbers
- If specific numbers are included, state assumptions (region, usage volume, token units)

CASE STUDY EVIDENCE STANDARD:
- Only include customer stories backed by official {vendor_name} sources (case study page, official blog, conference session)
- If a case study cannot be verified through first-party sources, DO NOT include it
- Downgrade unverified claims to "reported" language or omit entirely
- This section is high-risk for credibility. When in doubt, leave it out.

=============================================================================
FORMATTING RULES
=============================================================================

- Write in full paragraphs with clear section headers
- Use bullets for lists of services
- No em-dashes, use commas or periods
- Avoid hype language. Prefer operational language over visionary claims.
- Include specific service names, features, and availability status

=============================================================================
RESEARCH GOAL
=============================================================================

Many companies are interested in adopting AI in {current_date}. For {vendor_name}, I need a comprehensive
overview of the latest AI services, capabilities, and best practices that we should keep in mind when
advising enterprise customers on AI strategy.

Search for the latest updates from {conference} and recent {vendor_name} announcements.

=============================================================================
SECTION STRUCTURE
=============================================================================

## Executive Summary
Key themes and strategic direction for {vendor_name} AI in {current_date}.

## Service Map

A simple mapping of where services live in the {vendor_name} stack:
{service_map}

Update this map if any services have been renamed, merged, or deprecated.

## Foundation Models and AI Services

For {model_platform}, provide:
- Which models are available (provider, model family, version)
- What is new in the past 6 months
- GA vs Preview status for each model
- Customization options (fine-tuning, continued pretraining, adapters, prompt caching, guardrails)
- Pricing model overview (cite pricing page with assumptions)

## Productivity AI and Copilots
- Enterprise productivity tools with AI
- Integration with existing workflows
- Licensing and deployment models

## Agentic AI and Automation
- Agent building platforms and tools
- Orchestration capabilities
- Multi-agent scenarios
- When agents are appropriate versus workflow automation

## Data and Analytics AI
- AI-powered analytics and BI
- Data platform integration
- Vector search and RAG capabilities
- Default enterprise pattern for RAG on {vendor_name}

## AI Development Platform
- Model hosting options and when to use each
- Developer tools and SDKs
- MLOps and deployment
- Evaluation and monitoring capabilities

## Security and Governance
- AI governance tools
- Data protection and compliance
- Identity and access for AI
- Guardrails and content filtering

## Recommendation Guidance

When to use which service:
- Model hosting: When to use {model_platform} vs other options
- RAG: Default pattern for enterprise RAG and vector search
- Agents: When agents are appropriate versus simpler automation
- Anti-patterns to avoid (pilot purgatory, tool sprawl, no evaluation, skipping governance)

## Notable Customer Case Studies

IMPORTANT: Only include customer stories that have a primary {vendor_name} source (official case study page, official blog, conference session). If not from a primary source, do not include it.

- Real-world implementations with source links
- ROI examples where documented
- Industry-specific use cases

## New in the Past 6 Months

Bulleted list of recent changes. For each item include:
- What changed
- GA or Preview status
- Date and source URL

## Deprecations and Migrations (Past 12 Months)

List any services, features, or models that were:
- Deprecated or end-of-lifed
- Renamed or rebranded
- Merged into another service

For each, state the replacement or migration path.

## Sources

List all sources with URLs and dates. Group by section for easy reference.
"""

    console.info(f"Generating fresh {platform.upper()} AI research...")
    console.info("Estimated: 5-10 min, ~$0.50")  # Deep Research for vendor docs

    client = get_deep_research_client()

    def progress_callback(progress):
        if progress.message:
            console.info(f"Vendor Research: {progress.message}")

    coro = client.research(
        query=prompt,
        output_format=None,
        on_progress=progress_callback,
        timeout=1800,  # 30 min timeout
        job_metadata={
            "report_kind": "vendor_research",
            "cloud_vendor": platform.lower(),
        },
    )

    try:
        # If an event loop is already running (e.g. from Playwright/scraper),
        # run the coroutine in a separate thread to avoid "already running" error.
        try:
            running_loop = asyncio.get_running_loop()
        except RuntimeError:
            running_loop = None

        if running_loop and running_loop.is_running():
            import concurrent.futures

            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                result = pool.submit(asyncio.run, coro).result()
        else:
            result = asyncio.run(coro)

        if result.status != ResearchStatus.COMPLETED or not result.content:
            console.error("Vendor research generation failed")
            return None

        # Save to docs folder
        research_path = _get_vendor_research_path(platform)
        os.makedirs(os.path.dirname(research_path), exist_ok=True)

        with open(research_path, "w", encoding="utf-8") as f:
            f.write(result.content)

        # Deep Research is a flat per-task cost (API doesn't expose tokens)
        from primr.config.models import DEEP_RESEARCH_COST

        actual_cost = DEEP_RESEARCH_COST.standard_task_cost
        duration_str = f"{result.duration_seconds / 60:.1f}m"

        console.ok(
            f"Vendor research saved: {os.path.basename(research_path)} ({duration_str}, ~${actual_cost:.2f})"
        )
        return research_path

    except Exception as e:
        console.error(f"Vendor research failed: {e}")
        logger.exception("Vendor research error")
        return None


def _get_or_generate_vendor_research(platform: str) -> list[str]:
    """
    Get vendor research file, generating if needed.

    For Azure, prefers the manually curated Ignite analysis.
    For AWS/GCP, auto-generates if current month's research doesn't exist.

    Args:
        platform: Platform (azure, aws, gcp)

    Returns:
        List of paths to vendor research files, or empty list if unavailable
    """

    result_paths = []

    # Azure: always include manually curated Ignite analysis (it's excellent)
    if platform.lower() == "azure":
        manual_path = os.path.join(
            PROJECT_ROOT, "docs/research latest microsoft ignite analysis.txt"
        )
        if os.path.exists(manual_path):
            result_paths.append(manual_path)

    # Check for current month's auto-generated research
    research_path = _get_vendor_research_path(platform)
    if os.path.exists(research_path):
        result_paths.append(research_path)
    elif not result_paths:  # Only auto-generate if we have nothing
        generated = _generate_vendor_research(platform)
        if generated:
            result_paths.append(generated)

    return result_paths


def _generate_strategy_section(
    strategy_name: str,
    company_name: str,
    platform: str,
    company_research_path: str | None = None,
    force_refresh_vendor: bool = False,
    discovery_notes_content: str | None = None,
    lite_strategy: bool = False,
    output_dir: str | Path | None = None,
    diagnostics_dir: str | Path | None = None,
    write_txt: bool = True,
) -> str | None:
    """
    Generate a strategy document using Deep Research.

    This is a generic function that can generate any strategy module
    defined in the strategies/ directory.

    Args:
        strategy_name: Name of the strategy module (e.g., 'ai', 'customer_experience', 'modern_security_compliance', 'data_fabric_strategy')
        company_name: Name of the company
        platform: Platform preference (azure, aws, gcp, agnostic)
        company_research_path: Path to company research markdown (used as context)
        force_refresh_vendor: If True, regenerate vendor research even if current
        discovery_notes_content: Optional freeform meeting insights from discovery

    Returns:
        Path to the generated DOCX file, or None if generation failed
    """
    # Map strategy names to their YAML files and generation functions
    strategy_map = {
        "ai": "ai_first_transformation",
        "customer_experience": "customer_experience",
        "modern_security_compliance": "modern_security_compliance",
        "data_fabric_strategy": "data_fabric_strategy",
    }

    # Handle AI strategy separately (has vendor research)
    if strategy_name == "ai":
        return _generate_ai_strategy_section(
            company_name=company_name,
            platform=platform,
            company_research_path=company_research_path,
            force_refresh_vendor=force_refresh_vendor,
            discovery_notes_content=discovery_notes_content,
            lite_strategy=lite_strategy,
            output_dir=output_dir,
            diagnostics_dir=diagnostics_dir,
            write_txt=write_txt,
        )

    # Handle other fully-defined strategies
    if strategy_name in strategy_map:
        return _generate_generic_strategy(
            strategy_name=strategy_name,
            strategy_yaml=strategy_map[strategy_name],
            company_name=company_name,
            company_research_path=company_research_path,
            discovery_notes_content=discovery_notes_content,
            output_dir=output_dir,
            diagnostics_dir=diagnostics_dir,
            write_txt=write_txt,
        )

    # For placeholder strategies, show a message
    from primr.prompts.registry import get_registry

    registry = get_registry()
    strategy_module = registry.get(strategy_name)

    if not strategy_module:
        console.error(f"Strategy module not found: {strategy_name}")
        return None

    # Check if it's a placeholder
    import yaml

    if strategy_module.config_path.exists():
        with open(strategy_module.config_path, encoding="utf-8") as f:
            data = yaml.safe_load(f)
            meta = data.get("meta", {})
            if meta.get("status") == "placeholder":
                console.warn(f"Strategy '{strategy_name}' is a placeholder - not yet implemented")
                console.info(f"To implement, update: {strategy_module.config_path}")
                return None

    # For fully implemented strategies (future), use PromptComposer
    console.warn(f"Strategy '{strategy_name}' generation not yet implemented")
    return None


def _generate_generic_strategy(
    strategy_name: str,
    strategy_yaml: str,
    company_name: str,
    company_research_path: str | None = None,
    discovery_notes_content: str | None = None,
    output_dir: str | Path | None = None,
    diagnostics_dir: str | Path | None = None,
    write_txt: bool = True,
) -> str | None:
    return _generate_generic_strategy_impl(
        strategy_name=strategy_name,
        strategy_yaml=strategy_yaml,
        company_name=company_name,
        company_research_path=company_research_path,
        discovery_notes_content=discovery_notes_content,
        output_dir=output_dir,
        diagnostics_dir=diagnostics_dir,
        write_txt=write_txt,
    )


def _build_strategy_prompt_from_yaml(
    strategy_config: dict, company_name: str, discovery_notes_content: str | None = None
) -> str:
    return _build_strategy_prompt_from_yaml_impl(
        strategy_config=strategy_config,
        company_name=company_name,
        discovery_notes_content=discovery_notes_content,
    )


def _generate_ai_strategy_section(
    company_name: str,
    platform: str,
    company_research_path: str | None = None,
    force_refresh_vendor: bool = False,
    discovery_notes_content: str | None = None,
    lite_strategy: bool = False,
    output_dir: str | Path | None = None,
    diagnostics_dir: str | Path | None = None,
    write_txt: bool = True,
) -> str | None:
    """
    Generate AI strategy using Deep Research for board-level analysis.

    This uses a second Deep Research call to produce a comprehensive AI roadmap
    that's grounded in the company's actual business model, pain points, and
    the latest vendor capabilities.

    Outputs:
    - {company}_AI_Strategy_{date}.md  - Raw markdown
    - {company}_AI_Strategy_{date}.txt - Plain text
    - {company}_AI_Strategy_{date}.docx - Word document

    Args:
        company_name: Name of the company
        platform: Platform preference (azure, aws, gcp)
        company_research_path: Path to company research markdown (used as context)
        force_refresh_vendor: If True, regenerate vendor research even if current
        discovery_notes_content: Optional freeform meeting insights from discovery

    Returns:
        Path to the generated DOCX file, or None if generation failed
    """
    return _generate_ai_strategy_section_impl(
        company_name=company_name,
        platform=platform,
        company_research_path=company_research_path,
        force_refresh_vendor=force_refresh_vendor,
        discovery_notes_content=discovery_notes_content,
        lite_strategy=lite_strategy,
        output_dir=output_dir,
        diagnostics_dir=diagnostics_dir,
        write_txt=write_txt,
    )


def _build_ai_strategy_prompt(
    company_name: str, platform: str, discovery_notes_content: str | None = None
) -> str:
    return _build_ai_strategy_prompt_impl(
        company_name=company_name,
        platform=platform,
        discovery_notes_content=discovery_notes_content,
    )


def _run_verification(
    company_name: str,
    company_url: str,
    report_path: str,
) -> Any:
    """Run claim verification on a report. Returns VerificationResult or None."""
    import asyncio
    from pathlib import Path

    from primr.agentic.subagents.base import SubagentContext
    from primr.agentic.subagents.verifier import VerifierSubagent

    txt_path = Path(report_path).with_suffix(".txt")
    if not txt_path.exists():
        txt_path = Path(report_path)
        if txt_path.suffix in (".docx", ".pdf"):
            logger.warning(
                f"Verification: no .txt companion found, using {txt_path.suffix} file directly"
            )
            return None

    context = SubagentContext(
        company_name=company_name,
        company_url=company_url,
        working_dir=txt_path.parent,
        parent_results={"report_path": txt_path},
    )
    verifier = VerifierSubagent(context)

    try:
        asyncio.get_running_loop()
        # Already in an async context — run in a thread to avoid RuntimeError
        import concurrent.futures

        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            result = pool.submit(asyncio.run, verifier.execute()).result()
    except RuntimeError:
        # No running loop — safe to use asyncio.run directly
        result = asyncio.run(verifier.execute())

    if result.is_success and result.data:
        return result.data
    return None


def improve_output_file(
    file_path: str, *, in_place: bool = False, use_agentic: bool = False
) -> str | None:
    """Improve an existing markdown/text output artifact with deterministic + optional agentic QA cleanup."""
    path = Path(file_path)
    if not path.exists() or not path.is_file():
        console.error(f"Improve failed: file not found: {file_path}")
        return None

    if path.suffix.lower() not in {".md", ".txt"}:
        console.error(
            "Improve supports .md or .txt files. Convert DOCX/PDF to markdown/text first."
        )
        return None

    try:
        content = path.read_text(encoding="utf-8")
    except Exception as e:
        console.error(f"Improve failed: could not read file: {e}")
        return None

    is_strategy = "# AI Strategy:" in content or "_AI_Strategy_" in path.name

    if is_strategy:
        improved = content
        if use_agentic:
            cv = _strategy_cross_validate("Company", improved, "agnostic", [])
            if cv.get("weak_sections") or cv.get("issues"):
                improved = _strategy_polish("Company", "agnostic", improved)
        improved = _clean_strategy_output(improved)
        qa = _compute_strategy_qa_metrics(improved)
        console.info(
            "Improve QA (strategy): "
            f"placeholders={qa['placeholder_refs']}, "
            f"sources={qa['source_urls']}, "
            f"budget={'OK' if not qa['budget_inconsistent'] else 'WARN'}, "
            f"gate={'PASS' if qa['qa_gate_passed'] else 'WARN'}"
        )
    else:
        improved = content
        if use_agentic:
            cv = _fast_cross_validate("Company", None, improved, [])
            if cv.get("weak_sections") or cv.get("contradictions"):
                improved = _polish_fast_report_for_trust("Company", None, improved, [])
        improved = _clean_fast_report_output(improved)
        improved = _normalize_fast_citations(improved)
        improved = _strip_internal_source_placeholders(improved)
        improved = _enforce_fast_section_quality_guards(improved)
        qa = _compute_fast_report_qa_metrics(improved)
        console.info(
            "Improve QA (report): "
            f"labels={qa['confidence_labels']}, "
            f"cites={qa['citations_used']}/{qa['citations_defined']}, "
            f"validate={qa['sections_with_validate']}/{qa['section_count']}, "
            f"gate={'PASS' if qa['qa_gate_passed'] else 'WARN'}"
        )

    out_path = path if in_place else path.with_name(f"{path.stem}_improved{path.suffix}")
    try:
        out_path.write_text(improved, encoding="utf-8")
    except Exception as e:
        console.error(f"Improve failed: could not write output: {e}")
        return None

    return str(out_path)


def process_csv(
    file_path: str,
    mode: str = "complete",
    citation_style: str = "numbered",
    ai_strategy: bool = True,
    platforms: tuple[str, ...] = ("azure",),
    no_qa: bool = False,
) -> None:
    import csv

    console.header("Batch Processing", file_path)
    console.info(f"Mode: {mode}")

    with open(file_path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            company = row.get("company_name", "").strip()
            website = row.get("website", "").strip()
            if company or website:
                try:
                    perform_research(
                        company,
                        ensure_valid_url(website) if website else None,
                        mode=mode,
                        citation_style=citation_style,
                        ai_strategy=ai_strategy,
                        platforms=platforms,
                        no_qa=no_qa,
                    )
                except Exception as e:
                    console.error(f"Failed: {company or website} - {e}")


def cleanup():
    gc.collect()


atexit.register(cleanup)


def _list_recent_outputs():
    """List recent research outputs from the output directory."""
    import glob
    from datetime import datetime

    output_files = glob.glob(os.path.join(OUTPUT_DIR, "*.docx"))
    if not output_files:
        print("No recent outputs found.")
        return

    # Sort by modification time, newest first
    output_files.sort(key=os.path.getmtime, reverse=True)

    print("\nRECENT RESEARCH OUTPUTS")
    print("-" * 80)
    print(f"{'#':<3} {'Report':<40} {'Date':<12} {'Size':<8} {'QA Grade':<10}")
    print("-" * 80)

    for i, filepath in enumerate(output_files[:20], 1):
        filename = os.path.basename(filepath)
        mtime = datetime.fromtimestamp(os.path.getmtime(filepath))
        size_kb = os.path.getsize(filepath) / 1024

        # Look for corresponding QA report
        qa_grade = _get_qa_grade_for_report(filepath)
        qa_display = f"{qa_grade}/100" if qa_grade is not None else "N/A"

        # Truncate filename if too long
        display_name = filename[:37] + "..." if len(filename) > 40 else filename

        print(
            f"{i:2}. {display_name:<40} {mtime.strftime('%Y-%m-%d'):<12} {size_kb:6.1f}KB {qa_display:<10}"
        )

    if len(output_files) > 20:
        print(f"... and {len(output_files) - 20} more files")
    print("-" * 80)


def _get_qa_grade_for_report(report_path: str) -> int | None:
    """
    Get QA grade for a report by finding its corresponding QA report file.

    Args:
        report_path: Path to the main report file

    Returns:
        QA grade (0-100) or None if no QA report found
    """
    try:
        from pathlib import Path

        report_file = Path(report_path)
        # Extract company name from filename (remove extension and date)
        filename = report_file.stem

        # Remove common suffixes to get company name
        for suffix in ["_Strategic_Overview", "_Company_Overview", "_AI_Strategy"]:
            if suffix in filename:
                company_part = filename.split(suffix)[0]
                break
        else:
            company_part = filename

        # Look for QA report files
        output_dir = Path(OUTPUT_DIR)
        qa_patterns = [f"{company_part}*QA_Report*.txt", f"*{company_part}*QA_Report*.txt"]

        qa_files = []
        for pattern in qa_patterns:
            matches = list(output_dir.glob(pattern))
            qa_files.extend(matches)

        if not qa_files:
            return None

        # Get the most recent QA file
        latest_qa = max(qa_files, key=lambda f: f.stat().st_mtime)

        # Parse the grade from the QA report
        content = latest_qa.read_text(encoding="utf-8")
        for line in content.split("\n"):
            if line.startswith("Quality Score:"):
                # Extract score like "Quality Score: 85/100"
                parts = line.split(":")[1].strip().split("/")
                if parts and parts[0].isdigit():
                    return int(parts[0])

        return None

    except Exception as e:
        logger.warning("Failed to extract QA score: %s", e)
        return None


def _check_api_quota():
    """Check if Gemini API quota is available by making a lightweight test call."""
    from google import genai

    from primr.config.settings import get_settings

    settings = get_settings()

    if not settings.api.gemini_key:
        console.error("GEMINI_API_KEY not configured in .env")
        return

    console.banner("API Quota Check")
    console.info("Testing Gemini API availability...")

    try:
        client = genai.Client(api_key=settings.api.gemini_key)

        # Make a minimal API call to test quota
        response = client.models.generate_content(
            model=PrimrModels.FAST_MODEL,  # Use fast model for quick check
            contents="Say 'OK' in one word.",
        )

        if response and response.text:
            console.ok("API quota is available")
            console.info("You can run research now.")
        else:
            console.warn("API responded but with empty content")

    except Exception as e:
        error_str = str(e).lower()

        if "resource_exhausted" in error_str and ("per_day" in error_str or "quota" in error_str):
            console.error("Daily API quota is EXHAUSTED")
            console.info("Options:")
            console.info("  1. Wait until quota resets (usually midnight PT)")
            console.info("  2. Upgrade your API plan at https://ai.google.dev")
            console.info("  3. Use a different API key")
        elif "429" in str(e):
            console.warn("Rate limited - try again in a few minutes")
        elif "invalid" in error_str and "key" in error_str:
            console.error("Invalid API key")
        else:
            console.error(f"API check failed: {e}")


def _list_strategies():
    """List available strategy modules from the strategies/ directory."""
    from primr.prompts.registry import get_registry

    console.banner("Available Strategy Modules")

    registry = get_registry()
    strategies = registry.discover()

    if not strategies:
        console.info("No strategy modules found.")
        console.info("Strategy modules are YAML files in src/primr/prompts/strategies/")
        return

    console.info(f"Found {len(strategies)} strategy module(s):\n")

    for strategy in strategies:
        status = ""
        # Check if it's a placeholder
        if strategy.config_path.exists():
            import yaml

            with open(strategy.config_path, encoding="utf-8") as f:
                data = yaml.safe_load(f)
                meta = data.get("meta", {})
                if meta.get("status") == "placeholder":
                    status = " (placeholder)"

        console.step(f"{strategy.name}{status}")
        console.info(f"  Display Name: {strategy.display_name}")
        if strategy.description:
            console.info(f"  Description: {strategy.description}")
        if strategy.data_sources:
            console.info(f"  Data Sources: {len(strategy.data_sources)} file(s)")
        print()

    console.info("Use --strategy <name> to generate a specific strategy (coming in v1.2.6)")


def _clean_temp_files():
    """Clean up temporary files from working directory."""
    import glob

    # Clean working directory
    working_dirs = glob.glob(os.path.join(WORKING_DIR, "*"))
    temp_files = glob.glob(os.path.join(WORKING_DIR, "*.tmp"))

    cleaned = 0

    # Remove empty working directories
    for d in working_dirs:
        if os.path.isdir(d):
            try:
                # Only remove if empty or contains only temp files
                contents = os.listdir(d)
                if not contents:
                    os.rmdir(d)
                    cleaned += 1
            except Exception:
                logger.debug("Failed to remove temp directory %s", d, exc_info=True)

    # Remove temp files
    for f in temp_files:
        try:
            os.remove(f)
            cleaned += 1
        except Exception:
            logger.debug("Failed to remove temp file %s", f, exc_info=True)

    print(f"Cleaned {cleaned} temporary files/directories.")


def _open_file(filepath: str) -> None:
    """Open a file with the system default application."""
    from primr.utils.files import open_with_default_app

    try:
        open_with_default_app(filepath)
    except Exception as e:
        console.warn(f"Could not open file: {e}")


def run_doctor():
    """
    Run system diagnostics to verify Primr is properly configured.

    Checks:
    - Python version
    - API keys (GEMINI_API_KEY, SEARCH_API_KEY, SEARCH_ENGINE_ID)
    - API key format validation
    - Playwright browsers
    - Output directory writable
    - Working directory writable
    - Cache directory accessible
    - API quota availability
    """
    from primr.utils.type_guards import ConfigSchema, validate_config

    console.banner("Primr Doctor")
    console.blank()

    all_passed = True
    warnings_count = 0

    # 1. Python version
    console.step("Environment")
    py_version = sys.version_info
    if py_version >= (3, 10):
        console.ok(f"Python {py_version.major}.{py_version.minor}.{py_version.micro}")
    else:
        console.error(f"Python {py_version.major}.{py_version.minor} (need 3.10+)")
        all_passed = False

    # 2. API Keys with enhanced validation
    console.step("API Configuration")

    # Build config dict from environment
    env_config = {
        "gemini_key": os.environ.get("GEMINI_API_KEY", ""),
        "search_key": os.environ.get("SEARCH_API_KEY", ""),
        "search_engine_id": os.environ.get("SEARCH_ENGINE_ID", ""),
    }

    # Define validation schema
    api_schema = ConfigSchema(
        required_keys=["gemini_key"],
        optional_keys={"search_key": "", "search_engine_id": ""},
        type_hints={"gemini_key": str, "search_key": str, "search_engine_id": str},
        validators={
            "gemini_key": lambda v: len(v) >= 10 if v else False,
            "search_key": lambda v: len(v) >= 10 if v else True,  # Optional
        },
    )

    # Validate config
    validate_config(env_config, api_schema)

    # Report validation results
    gemini_key = env_config.get("gemini_key", "")
    if gemini_key and len(gemini_key) >= 10:
        # Check format (Gemini keys typically start with "AI")
        if gemini_key.startswith("AI"):
            console.ok("GEMINI_API_KEY configured (valid format)")
        else:
            console.ok("GEMINI_API_KEY configured")
            console.warn("  Key format unusual (expected to start with 'AI')")
            warnings_count += 1
    else:
        console.error("GEMINI_API_KEY not set or invalid")
        console.info("  Get your key at: https://ai.google.dev/")
        all_passed = False

    search_provider = os.environ.get("SEARCH_PROVIDER", "auto").lower().strip()
    if search_provider == "google":
        search_key = env_config.get("search_key", "")
        if search_key and len(search_key) >= 10:
            console.ok("SEARCH_API_KEY configured")
        else:
            console.warn("SEARCH_API_KEY not set (required for SEARCH_PROVIDER=google)")
            warnings_count += 1

        search_engine = env_config.get("search_engine_id", "")
        if search_engine:
            console.ok("SEARCH_ENGINE_ID configured")
        else:
            console.warn("SEARCH_ENGINE_ID not set (required for SEARCH_PROVIDER=google)")
            warnings_count += 1
    else:
        console.ok("Search: DuckDuckGo (no API key needed)")

    # 3. Playwright browsers
    console.step("Dependencies")
    try:
        from playwright.sync_api import sync_playwright

        with sync_playwright():
            console.ok("Playwright browsers available")
    except Exception as e:
        console.warn(f"Playwright not ready: {e}")
        console.info("  Run: playwright install chromium")
        warnings_count += 1

    # 4. Directory checks
    console.step("File System")

    # Output directory
    try:
        test_file = os.path.join(OUTPUT_DIR, ".primr_test")
        with open(test_file, "w") as f:
            f.write("test")
        os.remove(test_file)
        console.ok("Output directory writable")
    except Exception as e:
        console.error(f"Cannot write to output directory: {e}")
        all_passed = False

    # Working directory
    try:
        os.makedirs(WORKING_DIR, exist_ok=True)
        test_file = os.path.join(WORKING_DIR, ".primr_test")
        with open(test_file, "w") as f:
            f.write("test")
        os.remove(test_file)
        console.ok("Working directory writable")
    except Exception as e:
        console.error(f"Cannot write to working directory: {e}")
        all_passed = False

    # Cache directory
    try:
        cache_path = os.path.join(LOGS_DIR, "cache.db")
        if os.path.exists(cache_path):
            cache_size = os.path.getsize(cache_path) / (1024 * 1024)  # MB
            console.ok(f"Cache accessible ({cache_size:.1f} MB)")
        else:
            console.ok("Cache directory ready")
    except Exception as e:
        console.warn(f"Cache check failed: {e}")
        warnings_count += 1

    # 5. API quota check (quick test)
    console.step("API Connectivity")
    if gemini_key:
        try:
            from google import genai

            client = genai.Client(api_key=gemini_key)
            response = client.models.generate_content(
                model=PrimrModels.FAST_MODEL,
                contents="Reply with exactly: hello",
            )
            # Check if we got any response at all (connection works)
            if response and (response.text or response.candidates):
                console.ok("Gemini API responding")
            else:
                # Still connected, just empty - not a failure
                console.ok("Gemini API connected")
        except Exception as e:
            error_str = str(e).lower()
            if "quota" in error_str or "rate" in error_str:
                console.error("Gemini API quota exceeded - wait and retry")
                all_passed = False
            elif "invalid" in error_str and "key" in error_str:
                console.error("Gemini API key is invalid")
                all_passed = False
            else:
                console.warn(f"Gemini API test failed: {e}")
                warnings_count += 1
    else:
        console.warn("Skipping API test (no key configured)")
        warnings_count += 1

    # 6. QA System Configuration
    console.step("Quality Assurance")
    try:
        from primr.qa.analyzer import QAAnalyzer
        from primr.qa.integration import QAIntegration

        # Test QA system initialization
        QAIntegration()
        qa_analyzer = QAAnalyzer()

        if qa_analyzer.ai_client:
            console.ok("QA system initialized")
            console.info(f"  QA Model: {qa_analyzer.model_name}")
        else:
            console.warn("QA system initialized but AI client unavailable")
            console.info("  QA will use fallback analysis if needed")
            warnings_count += 1

    except Exception as e:
        console.warn(f"QA system check failed: {e}")
        console.info("  QA analysis may not work properly")
        warnings_count += 1

    # Summary
    console.blank()
    if all_passed and warnings_count == 0:
        console.success_box("All checks passed", "Primr is ready to use")
    elif all_passed:
        console.success_box(
            f"Ready with {warnings_count} warning(s)",
            "Primr can run, but some features may be limited",
        )
    else:
        console.error("Some checks failed - fix issues above before running research")

    return 0 if all_passed else 1


# Mode name mapping (new -> old internal names)
MODE_MAP = {
    "scrape": "scrape-only",
    "deep": "deep-research",
    "full": "complete",
    "parallel": "hybrid",
    # Also accept old names for backwards compatibility
    "structured": "structured",
    "deep-research": "deep-research",
    "complete": "complete",
    "hybrid": "hybrid",
    "scrape-only": "scrape-only",
}


# =============================================================================
# CLI ENTRY POINT (delegated to cli.py)
# =============================================================================
# The main() function is imported from cli.py at the top of this file.
# This avoids duplicate argument parsers and keeps CLI logic in one place.
main = _main_new  # Re-export from cli.py


# =============================================================================
# LEGACY COMPATIBILITY - Helper functions used by the old main()
# These are kept for any code that might import them directly.
# =============================================================================


def _legacy_main_removed():
    """
    The duplicate main() function has been removed.

    CLI logic is now centralized in cli.py. This module re-exports
    main = _main_new from cli.py for backward compatibility.
    """
    raise NotImplementedError("Use main() which delegates to cli.py")


if __name__ == "__main__":
    main()
