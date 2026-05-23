"""Coverage tests for primr.output.document_builder.DocumentBuilder."""

from __future__ import annotations

from docx.document import Document as DocxDocument

from primr.output.citation_processor import CitationStyle
from primr.output.document_builder import DocumentBuilder
from primr.output.markdown_parser import MarkdownParser
from primr.output.models import ContentBlock, ParsedLine


# --------------------------------------------------------------------------- #
# __init__ / citation processing
# --------------------------------------------------------------------------- #
def test_init_processes_citations_numbered():
    sections = {"financial_overview": "See [Acme](https://acme.example) for details."}
    builder = DocumentBuilder("Acme", sections)
    # NUMBERED style transforms inline links during init.
    assert "https://acme.example" not in builder.sections["financial_overview"]


def test_init_inline_style_skips_processing():
    sections = {"financial_overview": "See [Acme](https://acme.example)."}
    builder = DocumentBuilder("Acme", sections, citation_style=CitationStyle.INLINE)
    # INLINE style leaves the content untouched.
    assert "acme.example" in builder.sections["financial_overview"]


def test_process_citations_handles_none_and_nonstring():
    sections = {"a": None, "b": 12345}
    builder = DocumentBuilder("Acme", sections)
    assert builder.sections["a"] == ""
    assert builder.sections["b"] == "12345"


# --------------------------------------------------------------------------- #
# confidence calculation
# --------------------------------------------------------------------------- #
def test_calculate_overall_confidence_levels():
    builder = DocumentBuilder("Acme", {})
    builder._citation_count = 12
    assert builder._calculate_overall_confidence() == "high"
    builder._citation_count = 5
    assert builder._calculate_overall_confidence() == "medium"
    builder._citation_count = 1
    assert builder._calculate_overall_confidence() == "low"


def test_get_confidence_description():
    builder = DocumentBuilder("Acme", {})
    builder._citation_count = 12
    builder._overall_confidence = "high"
    assert "High confidence" in builder._get_confidence_description()
    builder._overall_confidence = "low"
    assert "Limited sources" in builder._get_confidence_description()
    builder._overall_confidence = "weird"
    assert builder._get_confidence_description() == "Unknown"


# --------------------------------------------------------------------------- #
# build() — full document
# --------------------------------------------------------------------------- #
def test_build_full_document():
    sections = {
        "company_website": "https://acme.example",
        "industry": "technology",
        "competitive_position": "Acme is the leading provider. It dominates the segment.",
        "company_overview": "A technology company.",
        "mission_vision": "## Mission\nTo build great things.",
        "company_history": "Founded: 2001 in a garage.",
        "financial_overview": (
            "Annual Revenue: $480M. Net Profit Margin: 15%. "
            "Strong growth potential ahead of competition."
        ),
        "unique_selling_proposition": "Acme is a unique pioneer with growth potential.",
        "strategic_recommendations": (
            "There is regulatory risk. Quick win: launch a pilot program. "
            "We recommend expansion into adjacent markets."
        ),
        "value_theory": "Value is created through scale.",
        "industry_insights": "The industry faces fierce competition and uncertainty.",
    }
    builder = DocumentBuilder("Acme", sections, citations=[{"title": "Ref", "url": "https://r.example"}])
    doc = builder.build()
    assert isinstance(doc, DocxDocument)
    # Document should contain headings we know are added.
    heading_texts = [p.text for p in doc.paragraphs]
    assert any("Executive Summary" in t for t in heading_texts)
    assert any("Company Snapshot" in t for t in heading_texts)
    assert any("Table of Contents" in t for t in heading_texts)
    assert any("Sources" in t for t in heading_texts)


def test_build_minimal_no_sources():
    builder = DocumentBuilder("Acme", {})
    doc = builder.build()
    assert isinstance(doc, DocxDocument)
    # No sources -> no Sources appendix heading.
    assert not any(p.text == "Sources" for p in doc.paragraphs)


# --------------------------------------------------------------------------- #
# _generate_one_liner industry inference branches
# --------------------------------------------------------------------------- #
def test_generate_one_liner_infers_industry_from_overview():
    builder = DocumentBuilder(
        "Acme",
        {"company_overview": "We are a retail business serving customers."},
    )
    one_liner = builder._generate_one_liner()
    assert isinstance(one_liner, str)
    assert one_liner


def test_generate_one_liner_financial_industry():
    builder = DocumentBuilder(
        "Acme",
        {"company_overview": "We are a bank with financial products."},
    )
    assert builder._generate_one_liner()


# --------------------------------------------------------------------------- #
# _collect_all_sources dedup
# --------------------------------------------------------------------------- #
def test_collect_all_sources_dedup_and_title_fallback():
    sections = {"financial_overview": "See [Acme](https://acme.example)."}
    builder = DocumentBuilder(
        "Acme",
        sections,
        citations=[
            {"url": "https://acme.example", "title": "Dup"},  # already in content -> skipped
            {"url": "https://new.example", "text": "fallback text"},  # title via text
        ],
    )
    extra = builder._collect_all_sources()
    urls = [s["url"] for s in extra]
    assert "https://new.example" in urls
    assert "https://acme.example" not in urls
    new_entry = next(s for s in extra if s["url"] == "https://new.example")
    assert new_entry["title"] == "fallback text"


# --------------------------------------------------------------------------- #
# _render_block branch coverage
# --------------------------------------------------------------------------- #
def test_render_block_heading():
    builder = DocumentBuilder("Acme", {})
    block = ContentBlock("heading", [ParsedLine("heading", "My Heading", 1, "# My Heading")])
    builder._render_block(block)
    assert any("My Heading" in p.text for p in builder.document.paragraphs)


def test_render_block_bullet_and_numbered_with_indent():
    builder = DocumentBuilder("Acme", {})
    bullet_block = ContentBlock(
        "bullet_list",
        [ParsedLine("bullet", "indented bullet", 1, "  - x")],
    )
    builder._render_block(bullet_block)
    num_block = ContentBlock(
        "numbered_list",
        [ParsedLine("numbered", "indented number", 1, "  1. x")],
    )
    builder._render_block(num_block)
    texts = [p.text for p in builder.document.paragraphs]
    assert any("indented bullet" in t for t in texts)
    assert any("indented number" in t for t in texts)


def test_render_block_paragraph_with_inline_header():
    builder = DocumentBuilder("Acme", {})
    block = ContentBlock(
        "paragraph",
        [
            ParsedLine("text", "leading text", 0, "leading text"),
            ParsedLine("inline_header", "the value", 0, "Label: the value", {"header_text": "Label"}),
            ParsedLine("text", "trailing text", 0, "trailing text"),
        ],
    )
    builder._render_block(block)
    full = " ".join(p.text for p in builder.document.paragraphs)
    assert "leading text" in full
    assert "the value" in full
    assert "trailing text" in full


def test_render_block_table():
    builder = DocumentBuilder("Acme", {})
    parser = MarkdownParser()
    blocks = parser.parse_content("| A | B |\n|---|---|\n| 1 | 2 |\n")
    table_block = next(b for b in blocks if b.type == "table")
    builder._render_block(table_block)
    assert len(builder.document.tables) == 1
