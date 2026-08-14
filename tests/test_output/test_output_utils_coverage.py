"""Coverage tests for primr.output.output_utils helpers and report writers."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

from docx import Document

from primr.output import output_utils
from primr.output.output_utils import (
    _safe_working_subdir,
    apply_inline_formatting,
    convert_docx_to_pdf,
    load_section_results,
    parse_markdown_line,
    save_report_as_docx,
    save_report_as_docx_premium,
    save_report_as_txt,
    strip_markdown_artifacts,
)


# --------------------------------------------------------------------------- #
# parse_markdown_line
# --------------------------------------------------------------------------- #
def test_parse_markdown_line_empty():
    assert parse_markdown_line("") == ("empty", "", 0)
    assert parse_markdown_line("    ") == ("empty", "", 0)


def test_parse_markdown_line_heading():
    assert parse_markdown_line("## Title") == ("heading", "Title", 0)


def test_parse_markdown_line_subheading():
    assert parse_markdown_line("### Sub") == ("subheading", "Sub", 0)


def test_parse_markdown_line_bullet_variants():
    assert parse_markdown_line("* item") == ("bullet", "item", 0)
    assert parse_markdown_line("*   item") == ("bullet", "item", 0)
    assert parse_markdown_line("- item") == ("bullet", "item", 0)
    assert parse_markdown_line("-   item") == ("bullet", "item", 0)


def test_parse_markdown_line_bullet_indent_level():
    line = "        * nested"
    assert parse_markdown_line(line) == ("bullet", "nested", 2)


def test_parse_markdown_line_numbered():
    assert parse_markdown_line("1. first") == ("numbered", "first", 0)


def test_parse_markdown_line_plain_text():
    assert parse_markdown_line("just text") == ("text", "just text", 0)


# --------------------------------------------------------------------------- #
# apply_inline_formatting
# --------------------------------------------------------------------------- #
def test_apply_inline_formatting_bold_and_plain():
    doc = Document()
    para = doc.add_paragraph()
    apply_inline_formatting(para, "before **bold** after")
    texts = [r.text for r in para.runs]
    assert "before " in texts
    assert "bold" in texts
    assert " after" in texts
    bold_run = next(r for r in para.runs if r.text == "bold")
    assert bold_run.bold is True


def test_apply_inline_formatting_underscore_bold():
    doc = Document()
    para = doc.add_paragraph()
    apply_inline_formatting(para, "__strong__")
    assert any(r.bold for r in para.runs)


def test_apply_inline_formatting_no_markup():
    doc = Document()
    para = doc.add_paragraph()
    apply_inline_formatting(para, "plain text only")
    assert para.runs[0].text == "plain text only"


# --------------------------------------------------------------------------- #
# strip_markdown_artifacts
# --------------------------------------------------------------------------- #
def test_strip_markdown_artifacts_bold_and_heading():
    text = "## Heading\n**bold** and __more__"
    out = strip_markdown_artifacts(text)
    assert "##" not in out
    assert "**" not in out
    assert "bold" in out
    assert "more" in out


def test_strip_markdown_artifacts_preserves_embedded_underscores():
    # Word-character boundaries protect underscores embedded inside identifiers.
    text = "the file__name stays intact"
    out = strip_markdown_artifacts(text)
    assert "file__name" in out


# --------------------------------------------------------------------------- #
# _safe_working_subdir
# --------------------------------------------------------------------------- #
def test_safe_working_subdir_valid(tmp_path):
    with patch.object(output_utils, "WORKING_DIR", str(tmp_path)):
        result = _safe_working_subdir("Acme Corp")
    assert result.name == "Acme_Corp"
    assert result.parent == tmp_path.resolve()


def test_safe_working_subdir_normalizes_trailing_period(tmp_path):
    with patch.object(output_utils, "WORKING_DIR", str(tmp_path)):
        result = _safe_working_subdir("Acme, Inc.")

    assert result.name == "Acme,_Inc"
    assert result.parent == tmp_path.resolve()


def test_safe_working_subdir_rejects_traversal(tmp_path):
    import pytest

    with patch.object(output_utils, "WORKING_DIR", str(tmp_path)), pytest.raises(ValueError):
        _safe_working_subdir("..")


def test_safe_working_subdir_rejects_empty(tmp_path):
    import pytest

    with patch.object(output_utils, "WORKING_DIR", str(tmp_path)), pytest.raises(ValueError):
        _safe_working_subdir("")


# --------------------------------------------------------------------------- #
# load_section_results
# --------------------------------------------------------------------------- #
def test_load_section_results_missing_dir(tmp_path):
    with patch.object(output_utils, "WORKING_DIR", str(tmp_path)):
        result = load_section_results("DoesNotExist")
    assert result == {}


def test_load_section_results_reads_files_from_flat_dir(tmp_path):
    from primr.config.sections_config import SECTION_KEY_MAP

    section_key = next(iter(SECTION_KEY_MAP.values()))
    company_dir = tmp_path / "Acme_Corp"
    company_dir.mkdir()
    (company_dir / f"{section_key}.txt").write_text("  some content  ", encoding="utf-8")

    with patch.object(output_utils, "WORKING_DIR", str(tmp_path)):
        result = load_section_results("Acme Corp")

    assert result.get(section_key) == "some content"


def test_load_section_results_prefers_latest_subdir(tmp_path):
    from primr.config.sections_config import SECTION_KEY_MAP

    section_key = next(iter(SECTION_KEY_MAP.values()))
    company_dir = tmp_path / "Acme_Corp"
    company_dir.mkdir()
    older = company_dir / "2026-01-01_0000"
    newer = company_dir / "2026-02-01_0000"
    older.mkdir()
    newer.mkdir()
    (older / f"{section_key}.txt").write_text("old", encoding="utf-8")
    (newer / f"{section_key}.txt").write_text("new", encoding="utf-8")

    with patch.object(output_utils, "WORKING_DIR", str(tmp_path)):
        result = load_section_results("Acme Corp")

    assert result.get(section_key) == "new"


def test_load_section_results_uses_explicit_folder_not_newest_sibling(tmp_path):
    from primr.config.sections_config import SECTION_KEY_MAP

    section_key = next(iter(SECTION_KEY_MAP.values()))
    company_dir = tmp_path / "Acme_Corp"
    older = company_dir / "2026-01-01_0000"
    newer = company_dir / "2026-02-01_0000"
    older.mkdir(parents=True)
    newer.mkdir()
    (older / f"{section_key}.txt").write_text("old", encoding="utf-8")
    (newer / f"{section_key}.txt").write_text("new", encoding="utf-8")

    result = load_section_results("Acme Corp", folder_path=older)
    assert result.get(section_key) == "old"


# --------------------------------------------------------------------------- #
# save_report_as_txt
# --------------------------------------------------------------------------- #
def test_save_report_as_txt_writes_file(tmp_path):
    from primr.config.sections_config import SECTION_KEY_MAP

    section_key = next(iter(SECTION_KEY_MAP.values()))
    section_results = {section_key: "* bullet item\nplain line\n\n"}

    path = save_report_as_txt(section_results, "Acme Corp", output_dir=tmp_path)

    assert path is not None
    out = Path(path)
    assert out.exists()
    text = out.read_text(encoding="utf-8")
    assert "Strategic Company Overview" in text
    assert "bullet item" in text


def test_save_report_as_txt_handles_write_failure(tmp_path):
    section_results = {"k": "content"}
    with patch("builtins.open", side_effect=OSError("disk full")):
        result = save_report_as_txt(section_results, "Acme Corp", output_dir=tmp_path)
        assert result is None


def test_save_incomplete_markdown_report_labels_and_publishes_partial(tmp_path):
    result = output_utils.save_incomplete_markdown_report(
        "## Partial\nGrounded evidence.",
        "Acme Corp",
        output_dir=tmp_path,
    )

    assert result is not None
    published = Path(result)
    assert published.suffix == ".md"
    assert "Incomplete_Overview" in published.name
    assert "Strategic_Overview" not in published.name
    assert "# Incomplete Report" in published.read_text(encoding="utf-8")
    assert published.with_suffix(".txt").is_file()
    from primr.output.artifact_inventory import infer_artifact_role

    assert infer_artifact_role(published) != "primary_report"


# --------------------------------------------------------------------------- #
# save_report_as_docx (legacy)
# --------------------------------------------------------------------------- #
def test_save_report_as_docx_converts_txt(tmp_path):
    txt = tmp_path / "in.txt"
    txt.write_text(
        "Title Line\nDate Line\n## Section\n* a bullet\n1. a number\n### Sub\nsome text\n",
        encoding="utf-8",
    )
    path = save_report_as_docx(str(txt), "Acme Corp", output_dir=tmp_path)
    assert path is not None
    assert Path(path).exists()


def test_save_report_as_docx_handles_missing_txt(tmp_path):
    result = save_report_as_docx(str(tmp_path / "nope.txt"), "Acme Corp", output_dir=tmp_path)
    assert result is None


# --------------------------------------------------------------------------- #
# save_report_as_docx_premium
# --------------------------------------------------------------------------- #
def test_save_report_as_docx_premium_numbered(tmp_path):
    from primr.config.sections_config import SECTION_KEY_MAP

    section_key = next(iter(SECTION_KEY_MAP.values()))
    section_results = {section_key: "Content with a [link](https://example.com)."}
    path = save_report_as_docx_premium(
        section_results, "Acme Corp", citation_style="numbered", output_dir=tmp_path
    )
    assert path is not None
    assert Path(path).exists()


def test_save_report_as_docx_premium_sidecar_writes_sources(tmp_path):
    from primr.config.sections_config import SECTION_KEY_MAP

    section_key = next(iter(SECTION_KEY_MAP.values()))
    section_results = {section_key: "See [Acme](https://acme.example) for details."}
    path = save_report_as_docx_premium(
        section_results, "Acme Corp", citation_style="sidecar", output_dir=tmp_path
    )
    assert path is not None
    # A sidecar sources file should have been written into the same dir.
    sidecars = list(Path(tmp_path).glob("*sources*"))
    assert sidecars


def test_save_report_as_docx_premium_handles_builder_failure(tmp_path):
    section_results = {"k": "content"}
    with patch.object(output_utils, "DocumentBuilder", side_effect=RuntimeError("boom")):
        result = save_report_as_docx_premium(section_results, "Acme Corp", output_dir=tmp_path)
    assert result is None


# --------------------------------------------------------------------------- #
# convert_docx_to_pdf — no converter available
# --------------------------------------------------------------------------- #
def test_convert_docx_to_pdf_no_converter(tmp_path):
    docx = tmp_path / "report.docx"
    docx.write_text("docx", encoding="utf-8")
    with (
        patch("docx2pdf.convert", side_effect=RuntimeError("no office")),
        patch("shutil.which", return_value=None),
    ):
        result = convert_docx_to_pdf(docx)
    assert result is None


def test_convert_docx_to_pdf_soffice_returns_without_pdf(tmp_path):
    docx = tmp_path / "report.docx"
    docx.write_text("docx", encoding="utf-8")

    def _fake_run(*_args, **_kwargs):
        return None  # does not create a PDF

    with (
        patch("docx2pdf.convert", side_effect=RuntimeError("no office")),
        patch("shutil.which", return_value="/usr/bin/soffice"),
        patch("subprocess.run", side_effect=_fake_run),
    ):
        result = convert_docx_to_pdf(docx)
    assert result is None


# --------------------------------------------------------------------------- #
# zip_research_files / cleanup
# --------------------------------------------------------------------------- #
def test_zip_research_files_creates_archive(tmp_path):
    working = tmp_path / "working"
    output = tmp_path / "output"
    output.mkdir()
    company = working / "Acme_Corp"
    company.mkdir(parents=True)
    (company / "notes.txt").write_text("data", encoding="utf-8")

    with (
        patch.object(output_utils, "WORKING_DIR", str(working)),
        patch.object(output_utils, "OUTPUT_DIR", str(output)),
    ):
        output_utils.zip_research_files("Acme Corp")

    zips = list(output.glob("*_research_*.zip"))
    assert zips


def test_zip_research_files_handles_error_gracefully(tmp_path):
    working = tmp_path / "working"
    working.mkdir()
    # Company folder does not exist -> function returns without raising.
    with patch.object(output_utils, "WORKING_DIR", str(working)):
        output_utils.zip_research_files("Missing Co")  # no exception


def test_cleanup_removes_working_folder(tmp_path):
    working = tmp_path / "working"
    output = tmp_path / "output"
    output.mkdir()
    company = working / "Acme_Corp"
    company.mkdir(parents=True)
    (company / "f.txt").write_text("x", encoding="utf-8")

    with (
        patch.object(output_utils, "WORKING_DIR", str(working)),
        patch.object(output_utils, "OUTPUT_DIR", str(output)),
    ):
        output_utils.cleanup("Acme Corp")

    assert company.exists()
    assert (company / "f.txt").read_text(encoding="utf-8") == "x"


def test_cleanup_handles_error_gracefully(tmp_path):
    working = tmp_path / "working"
    working.mkdir()
    with (
        patch.object(output_utils, "WORKING_DIR", str(working)),
        patch.object(output_utils, "zip_research_files", side_effect=RuntimeError("boom")),
    ):
        output_utils.cleanup("Acme Corp")  # swallows the error


# --------------------------------------------------------------------------- #
# generate_final_report
# --------------------------------------------------------------------------- #
def test_generate_final_report_no_section_data(tmp_path):
    with patch.object(output_utils, "load_section_results", return_value={}):
        result = output_utils.generate_final_report("Acme Corp", output_dir=tmp_path)
    assert result is None


def test_generate_final_report_premium_success(tmp_path):
    sections = {"financial_overview": "Annual Revenue: $5M."}
    docx_target = tmp_path / "report.docx"

    def _fake_premium(*_a, **_k):
        docx_target.write_text("docx", encoding="utf-8")
        return str(docx_target)

    with (
        patch.object(output_utils, "load_section_results", return_value=sections),
        patch.object(output_utils, "save_report_as_docx_premium", side_effect=_fake_premium),
        patch.object(output_utils, "convert_docx_to_pdf", return_value=None) as conv,
        patch.object(output_utils, "cleanup") as clean,
        patch.object(output_utils, "save_report_as_txt", return_value=str(tmp_path / "r.txt")),
    ):
        result = output_utils.generate_final_report("Acme Corp", output_dir=tmp_path)

    assert result == str(docx_target)
    conv.assert_called_once()
    clean.assert_called_once_with("Acme Corp")


def test_generate_final_report_falls_back_to_legacy(tmp_path):
    sections = {"financial_overview": "content"}
    txt_path = tmp_path / "r.txt"
    legacy_docx = tmp_path / "legacy.docx"

    with (
        patch.object(output_utils, "load_section_results", return_value=sections),
        patch.object(output_utils, "save_report_as_txt", return_value=str(txt_path)),
        patch.object(output_utils, "save_report_as_docx_premium", return_value=None),
        patch.object(output_utils, "save_report_as_docx", return_value=str(legacy_docx)) as legacy,
        patch.object(output_utils, "convert_docx_to_pdf", return_value=None),
        patch.object(output_utils, "cleanup"),
    ):
        result = output_utils.generate_final_report("Acme Corp", output_dir=tmp_path)

    assert result == str(legacy_docx)
    legacy.assert_called_once()


def test_generate_final_report_skip_txt(tmp_path):
    sections = {"financial_overview": "content"}
    docx_target = tmp_path / "report.docx"

    with (
        patch.object(output_utils, "load_section_results", return_value=sections),
        patch.object(output_utils, "save_report_as_txt") as txt,
        patch.object(output_utils, "save_report_as_docx_premium", return_value=str(docx_target)),
        patch.object(output_utils, "convert_docx_to_pdf", return_value=None),
        patch.object(output_utils, "cleanup"),
    ):
        result = output_utils.generate_final_report(
            "Acme Corp", output_dir=tmp_path, write_txt=False
        )

    assert result == str(docx_target)
    # write_txt=False and no diagnostics_dir -> TXT writer not invoked.
    txt.assert_not_called()
