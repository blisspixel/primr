"""Coverage tests for primr.output.markdown_converter."""

from __future__ import annotations

from docx import Document

from primr.output.markdown_converter import (
    add_hyperlink,
    markdown_to_docx,
    parse_inline_markdown,
    render_section_content,
    render_table,
    sanitize_text,
    strip_heading_markers,
    strip_markdown_header_block,
)


# --------------------------------------------------------------------------- #
# sanitize_text
# --------------------------------------------------------------------------- #
def test_sanitize_text_empty():
    assert sanitize_text("") == ""


def test_sanitize_text_removes_control_chars():
    out = sanitize_text("a\x00b\x07c")
    assert out == "abc"


def test_sanitize_text_keeps_tab_newline():
    out = sanitize_text("a\tb\nc\rd")
    assert out == "a\tb\nc\rd"


# --------------------------------------------------------------------------- #
# add_hyperlink
# --------------------------------------------------------------------------- #
def test_add_hyperlink_basic():
    doc = Document()
    para = doc.add_paragraph()
    add_hyperlink(para, "Example", "https://example.com")
    # The hyperlink element should be appended to the paragraph XML.
    xml = para._p.xml
    assert "hyperlink" in xml
    assert "Example" in xml


def test_add_hyperlink_empty_skips():
    doc = Document()
    para = doc.add_paragraph()
    add_hyperlink(para, "", "https://example.com")
    assert "hyperlink" not in para._p.xml


# --------------------------------------------------------------------------- #
# parse_inline_markdown
# --------------------------------------------------------------------------- #
def test_parse_inline_markdown_empty():
    doc = Document()
    para = doc.add_paragraph()
    parse_inline_markdown(para, "")
    assert len(para.runs) == 0


def test_parse_inline_markdown_bold():
    doc = Document()
    para = doc.add_paragraph()
    parse_inline_markdown(para, "**bold**")
    assert para.runs[0].font.bold is True
    assert para.runs[0].text == "bold"


def test_parse_inline_markdown_italic():
    doc = Document()
    para = doc.add_paragraph()
    parse_inline_markdown(para, "*ital*")
    assert para.runs[0].font.italic is True


def test_parse_inline_markdown_code():
    doc = Document()
    para = doc.add_paragraph()
    parse_inline_markdown(para, "`code`")
    assert para.runs[0].text == "code"
    assert para.runs[0].font.name == "Consolas"


def test_parse_inline_markdown_link():
    doc = Document()
    para = doc.add_paragraph()
    parse_inline_markdown(para, "[text](https://example.com)")
    assert "hyperlink" in para._p.xml


def test_parse_inline_markdown_plain_mixed():
    doc = Document()
    para = doc.add_paragraph()
    parse_inline_markdown(para, "plain **bold** end")
    texts = [r.text for r in para.runs]
    assert "plain " in texts
    assert "bold" in texts
    assert " end" in texts


# --------------------------------------------------------------------------- #
# render_table
# --------------------------------------------------------------------------- #
def test_render_table_too_few_lines():
    doc = Document()
    render_table(doc, ["| only one |"])
    assert len(doc.tables) == 0


def test_render_table_basic():
    doc = Document()
    # Separator rows for a single-column table (no inner pipe) are detected/skipped.
    lines = [
        "| Name |",
        "|------|",
        "| Foo |",
        "| Bar |",
    ]
    render_table(doc, lines)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3  # header + 2 data rows (separator skipped)


def test_render_table_multicolumn_separator_is_skipped():
    doc = Document()
    # A multi-column separator (inner pipes) must be matched and skipped, not
    # rendered as a data row: header + two data rows = 3 rows.
    lines = [
        "| Name | Value |",
        "|------|-------|",
        "| Foo | 1 |",
        "| Bar | 2 |",
    ]
    render_table(doc, lines)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 3
    cell_texts = [c.text for row in doc.tables[0].rows for c in row.cells]
    assert not any("---" in t for t in cell_texts)


def test_render_table_only_separator_returns_empty():
    doc = Document()
    # Single-column separator lines match the separator regex and are skipped,
    # leaving no data rows -> no table created.
    render_table(doc, ["|---|", "|---|"])
    assert len(doc.tables) == 0


# --------------------------------------------------------------------------- #
# strip_heading_markers
# --------------------------------------------------------------------------- #
def test_strip_heading_markers_removes_formatting():
    assert strip_heading_markers("**Bold** *ital* `code`") == "Bold ital code"


def test_strip_heading_markers_empty():
    assert strip_heading_markers("") == ""
    assert strip_heading_markers("\x00") == ""


# --------------------------------------------------------------------------- #
# strip_markdown_header_block
# --------------------------------------------------------------------------- #
def test_strip_markdown_header_block_empty():
    assert strip_markdown_header_block("") == ""


def test_strip_markdown_header_block_removes_header():
    md = "# Title\n\n**Status:** Draft\n**Date:** Y\n\n---\n\nBody content here"
    out = strip_markdown_header_block(md)
    assert "Title" not in out
    assert "Body content here" in out


def test_strip_markdown_header_block_removes_fast_report_metadata():
    md = (
        "# Strategic Company Overview: Acme\n\n"
        "*July 07, 2026* | [https://acme.example](https://acme.example)\n\n"
        "---\n\n"
        "## Executive Summary\n\n"
        "Body content here"
    )
    out = strip_markdown_header_block(md)
    assert "Strategic Company Overview" not in out
    assert "July 07, 2026" not in out
    assert "https://acme.example" not in out
    assert "## Executive Summary" in out


def test_strip_markdown_header_block_no_heading_returns_same():
    md = "Just body text\nMore text"
    out = strip_markdown_header_block(md)
    assert "Just body text" in out


# --------------------------------------------------------------------------- #
# markdown_to_docx
# --------------------------------------------------------------------------- #
def test_markdown_to_docx_full(tmp_path):
    md = (
        "# Doc Title\n\n"
        "## Section A\n"
        "### Sub\n"
        "#### Minor\n"
        "Some paragraph text.\n\n"
        "- bullet one\n"
        "* bullet two\n"
        "1. numbered one\n"
        "2) numbered two\n\n"
        "> a quote line\n"
        "> second quote line\n\n"
        "---\n\n"
        "| H1 | H2 |\n"
        "|----|----|\n"
        "| a | b |\n"
    )
    out = tmp_path / "out.docx"
    result = markdown_to_docx(md, out, title="My Title", subtitle="My Subtitle")
    assert result == out
    assert out.exists()


def test_markdown_to_docx_does_not_duplicate_fast_report_subtitle(tmp_path):
    md = (
        "# Strategic Company Overview: Acme\n\n"
        "*July 07, 2026* | [https://acme.example](https://acme.example)\n\n"
        "---\n\n"
        "## Executive Summary\n\n"
        "Body content here"
    )
    out = tmp_path / "subtitle.docx"
    markdown_to_docx(
        md,
        out,
        title="Strategic Company Overview: Acme",
        subtitle="July 07, 2026 | https://acme.example",
    )
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert full_text.count("July 07, 2026") == 1
    assert full_text.count("https://acme.example") == 1
    assert "Executive Summary" in full_text


def test_markdown_to_docx_table_at_end(tmp_path):
    md = "Intro\n\n| A | B |\n|---|---|\n| 1 | 2 |"
    out = tmp_path / "tbl.docx"
    markdown_to_docx(md, out, title=None)
    assert out.exists()


def test_markdown_to_docx_no_title(tmp_path):
    md = "## Heading\n\nbody"
    out = tmp_path / "nt.docx"
    markdown_to_docx(md, out)
    assert out.exists()


# --------------------------------------------------------------------------- #
# render_section_content
# --------------------------------------------------------------------------- #
def test_render_section_content_all_elements():
    doc = Document()
    content = (
        "## Heading\n"
        "### Sub\n"
        "#### Minor\n"
        "paragraph\n\n"
        "- bullet\n"
        "* star bullet\n"
        "1. num\n"
        "> quote\n"
        "---\n"
        "| X | Y |\n"
        "|---|---|\n"
        "| 1 | 2 |\n"
    )
    render_section_content(doc, content)
    assert len(doc.tables) == 1
    assert len(doc.paragraphs) > 0


# --------------------------------------------------------------------------- #
# table-detection content loss (regression)
# --------------------------------------------------------------------------- #
def test_single_pipe_prose_not_dropped():
    """A prose line containing a single '|' but no table separator must be
    rendered as a paragraph, not silently dropped as a degenerate table."""
    doc = Document()
    render_section_content(doc, "Strengths | Weaknesses\n\nFollowing paragraph text")
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Strengths" in texts
    assert "Weaknesses" in texts
    assert "Following paragraph text" in texts


def test_real_table_still_renders_as_table():
    """A genuine markdown table (with a |---| separator row) renders as a DOCX table."""
    doc = Document()
    render_section_content(doc, "| A | B |\n| --- | --- |\n| 1 | 2 |")
    assert len(doc.tables) == 1


def test_pipe_in_heading_not_treated_as_table():
    """Bug-hunt round 3: a heading that merely contains "|" must not be
    mis-detected as a table and flushed as a plain paragraph leaking "## "."""
    doc = Document()
    render_section_content(doc, "## Market | Share\n\nBody text here.\n")
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Market | Share" in heading_texts
    assert not any(p.text.strip().startswith("## ") for p in doc.paragraphs)
    assert len(doc.tables) == 0


def test_parenthesized_url_link_not_truncated():
    """Bug-hunt round 3: a link whose URL has balanced parens must not be cut at
    the first ")" (which left a stray ")" in the prose)."""
    doc = Document()
    para = doc.add_paragraph()
    parse_inline_markdown(para, "See [Acme](https://en.wikipedia.org/wiki/Acme_(company)) here.")
    run_text = "".join(r.text for r in para.runs)
    assert ")" not in run_text  # no stray close-paren from a truncated URL
    assert "hyperlink" in para._p.xml


def test_math_asterisk_not_mis_italicized():
    """Bug-hunt round 3: "5*3" math must not be consumed as italic when a real
    *italic* appears later on the line."""
    doc = Document()
    para = doc.add_paragraph()
    parse_inline_markdown(para, "Revenue grew 5*3 and *real* emphasis.")
    run_text = "".join(r.text for r in para.runs)
    assert "5*3" in run_text
    italic_runs = [r.text for r in para.runs if r.font.italic]
    assert italic_runs == ["real"]
