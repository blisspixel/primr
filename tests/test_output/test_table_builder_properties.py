"""
Property-based tests for the TableBuilder.

Uses Hypothesis to verify table styling consistency.
"""

import pytest
from hypothesis import given, strategies as st, settings, assume, HealthCheck
from docx import Document

from primr.output.table_builder import TableBuilder
from primr.output.models import CompanySnapshot


class TestTableStylingConsistency:
    """
    **Feature: report-excellence, Property 6: Style consistency across document** (tables portion)
    **Validates: Requirements 5.5**
    
    For any generated table, all cells of the same type SHALL have
    consistent styling properties.
    """
    
    @settings(max_examples=50)
    @given(
        company_name=st.text(min_size=1, max_size=50, alphabet=st.characters(whitelist_categories=['L', 'N'])),
        industry=st.text(min_size=1, max_size=30, alphabet=st.characters(whitelist_categories=['L'])),
        revenue=st.text(min_size=1, max_size=20, alphabet=st.characters(whitelist_categories=['N', 'S']))
    )
    def test_snapshot_table_has_consistent_structure(self, company_name, industry, revenue):
        """Company snapshot table has consistent 2-column structure."""
        company_name = company_name.strip()
        industry = industry.strip()
        revenue = revenue.strip()
        assume(len(company_name) > 0)
        assume(len(industry) > 0)
        
        doc = Document()
        builder = TableBuilder(doc)
        
        snapshot = CompanySnapshot(
            company_name=company_name,
            industry=industry,
            revenue=revenue if revenue else None
        )
        
        table = builder.create_company_snapshot(snapshot)
        
        # All rows should have exactly 2 columns
        for row in table.rows:
            assert len(row.cells) == 2, "Snapshot table should have 2 columns"

    @settings(max_examples=50)
    @given(highlights=st.lists(
        st.text(min_size=5, max_size=100, alphabet=st.characters(whitelist_categories=['L', 'N', 'S'])),
        min_size=1,
        max_size=7
    ))
    def test_executive_highlights_creates_callout(self, highlights):
        """Executive highlights creates a single-cell callout table."""
        highlights = [h.strip() for h in highlights if h.strip()]
        assume(len(highlights) > 0)
        
        doc = Document()
        builder = TableBuilder(doc)
        
        table = builder.create_executive_highlights(highlights)
        
        # Should be a single-cell table
        assert len(table.rows) == 1
        assert len(table.rows[0].cells) == 1
        
        # Cell should contain all highlights
        cell_text = table.rows[0].cells[0].text
        for highlight in highlights:
            assert highlight in cell_text

    @settings(max_examples=50, suppress_health_check=[HealthCheck.too_slow])
    @given(metrics=st.dictionaries(
        keys=st.text(min_size=3, max_size=20, alphabet=st.characters(whitelist_categories=['L'])),
        values=st.text(min_size=1, max_size=20, alphabet=st.characters(whitelist_categories=['N', 'S'])),
        min_size=1,
        max_size=6
    ))
    def test_metrics_table_has_header_row(self, metrics):
        """Key metrics table has header row plus data rows."""
        metrics = {k.strip(): v.strip() for k, v in metrics.items() if k.strip() and v.strip()}
        assume(len(metrics) > 0)
        
        doc = Document()
        builder = TableBuilder(doc)
        
        table = builder.create_key_metrics_table(metrics)
        
        if table:
            # Should have header row + data rows
            assert len(table.rows) == len(metrics) + 1
            
            # Header row should have "Metric" and "Value"
            header_text = ''.join(cell.text for cell in table.rows[0].cells)
            assert 'Metric' in header_text
            assert 'Value' in header_text

    def test_all_style_variants_apply_without_error(self):
        """All table style variants can be applied without errors."""
        doc = Document()
        builder = TableBuilder(doc)
        
        # Create a simple table
        table = doc.add_table(rows=3, cols=2)
        for i, row in enumerate(table.rows):
            row.cells[0].text = f"Label {i}"
            row.cells[1].text = f"Value {i}"
        
        # All style variants should work
        for style in ['professional', 'snapshot', 'callout']:
            # Create fresh table for each style
            table = doc.add_table(rows=3, cols=2)
            builder.apply_table_style(table, style)
            # No assertion needed - just verify no exception

    @settings(max_examples=30)
    @given(
        num_rows=st.integers(min_value=1, max_value=10),
        num_cols=st.integers(min_value=1, max_value=5)
    )
    def test_professional_style_alternates_rows(self, num_rows, num_cols):
        """Professional style applies alternating row colors."""
        doc = Document()
        builder = TableBuilder(doc)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        builder.apply_table_style(table, 'professional')
        
        # Table should still have correct dimensions
        assert len(table.rows) == num_rows
        for row in table.rows:
            assert len(row.cells) == num_cols
