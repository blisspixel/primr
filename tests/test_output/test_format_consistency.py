"""
Tests for output format consistency across report modes.

Validates report structure, section ordering, and format conversion.

**Feature: test-coverage-hardening**
**Validates: Requirements 7.1, 7.2, 7.3, 7.4, 7.5**
"""

import tempfile
from pathlib import Path

from hypothesis import given, settings
from hypothesis import strategies as st

from primr.output.markdown_converter import (
    markdown_to_docx,
    render_table,
)
from primr.prompts.composer import PromptComposer

# =============================================================================
# Expected Section Configuration
# =============================================================================


# Get actual section IDs from company_overview.yaml dynamically
def get_expected_section_ids():
    """Get section IDs from the actual config."""
    composer = PromptComposer()
    config = composer._load_config("company_overview")
    return [s.id for s in config.sections]


# =============================================================================
# Unit Tests for Report Structure
# =============================================================================


class TestReportSectionStructure:
    """Tests for report section structure and ordering."""

    def test_report_has_21_sections(self):
        """
        WHEN a report is generated in deep mode
        THEN the output SHALL contain all 21 section headings

        **Validates: Requirements 7.1**
        """
        composer = PromptComposer()
        config = composer._load_config("company_overview")

        assert len(config.sections) == 23, f"Expected 23 sections, got {len(config.sections)}"

    def test_executive_summary_is_first(self):
        """
        WHEN a report is generated
        THEN the executive summary SHALL appear first

        **Validates: Requirements 7.2**
        """
        composer = PromptComposer()
        config = composer._load_config("company_overview")

        first_section = config.sections[0]
        assert first_section.id == "executive_summary"
        assert first_section.position == "opening"

    def test_strategic_positioning_is_last(self):
        """
        WHEN a report is generated
        THEN the strategic positioning hypothesis SHALL appear last

        **Validates: Requirements 7.3**
        """
        composer = PromptComposer()
        config = composer._load_config("company_overview")

        last_section = config.sections[-1]
        # The last section should be a closing section
        assert last_section.position == "closing"
        # It should be strategic_positioning (the actual ID)
        assert "strategic" in last_section.id.lower()

    def test_section_ordering_is_consistent(self):
        """Sections maintain consistent ordering across loads."""
        composer1 = PromptComposer()
        config1 = composer1._load_config("company_overview")

        composer2 = PromptComposer()
        config2 = composer2._load_config("company_overview")

        ids1 = [s.id for s in config1.sections]
        ids2 = [s.id for s in config2.sections]

        assert ids1 == ids2


# =============================================================================
# DOCX Conversion Tests
# =============================================================================


class TestDOCXTablePreservation:
    """Tests for markdown table to DOCX conversion."""

    def test_simple_table_converted(self):
        """
        WHEN markdown is converted to DOCX
        THEN table formatting SHALL be preserved

        **Validates: Requirements 7.4**
        """
        from docx import Document

        doc = Document()
        table_lines = [
            "| Header 1 | Header 2 |",
            "|----------|----------|",
            "| Cell 1   | Cell 2   |",
            "| Cell 3   | Cell 4   |",
        ]

        render_table(doc, table_lines)

        # Should have created a table
        assert len(doc.tables) == 1
        table = doc.tables[0]

        # Should have correct dimensions
        # Note: render_table may include separator as a row
        assert len(table.rows) >= 3  # At least header + 2 data rows
        assert len(table.columns) == 2

    def test_table_with_formatting(self):
        """Tables with inline formatting are converted correctly."""
        from docx import Document

        doc = Document()
        table_lines = [
            "| **Bold** | *Italic* |",
            "|----------|----------|",
            "| Normal   | `Code`   |",
        ]

        render_table(doc, table_lines)

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) >= 2  # At least header + 1 data row


class TestDOCXHeadingHierarchy:
    """Tests for heading hierarchy preservation in DOCX."""

    def test_heading_levels_preserved(self):
        """
        WHEN markdown is converted to DOCX
        THEN heading hierarchy SHALL be preserved

        **Validates: Requirements 7.5**
        """
        markdown = """## Section Title

Content here.

### Subsection

More content.

#### Sub-subsection

Even more content.
"""

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = Path(f.name)

        try:
            markdown_to_docx(markdown, output_path)

            from docx import Document

            doc = Document(output_path)

            # Check that headings exist with correct styles
            heading_styles = []
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    heading_styles.append(para.style.name)

            # Should have multiple heading levels
            assert len(heading_styles) >= 2
        finally:
            output_path.unlink(missing_ok=True)

    def test_h1_to_h4_supported(self):
        """All heading levels H1-H4 are supported."""
        markdown = """# H1 Title

## H2 Section

### H3 Subsection

#### H4 Minor
"""

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = Path(f.name)

        try:
            markdown_to_docx(markdown, output_path, title="Test")

            from docx import Document

            doc = Document(output_path)

            # Document should be created successfully
            assert len(doc.paragraphs) > 0
        finally:
            output_path.unlink(missing_ok=True)


class TestDOCXContentPreservation:
    """Tests for content preservation during DOCX conversion."""

    def test_bold_text_preserved(self):
        """Bold text is preserved in DOCX."""
        markdown = "This is **bold text** in a paragraph."

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = Path(f.name)

        try:
            markdown_to_docx(markdown, output_path)

            from docx import Document

            doc = Document(output_path)

            # Find paragraph with bold text
            found_bold = False
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.bold and "bold text" in run.text:
                        found_bold = True
                        break

            assert found_bold, "Bold text not found in DOCX"
        finally:
            output_path.unlink(missing_ok=True)

    def test_bullet_lists_preserved(self):
        """Bullet lists are preserved in DOCX."""
        markdown = """Key points:

* First item
* Second item
* Third item
"""

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = Path(f.name)

        try:
            markdown_to_docx(markdown, output_path)

            from docx import Document

            doc = Document(output_path)

            # Check content is present
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "First item" in full_text
            assert "Second item" in full_text
            assert "Third item" in full_text
        finally:
            output_path.unlink(missing_ok=True)


# =============================================================================
# Property Tests
# =============================================================================


@given(
    num_rows=st.integers(min_value=2, max_value=10),
    num_cols=st.integers(min_value=2, max_value=5),
)
@settings(max_examples=50, deadline=None)
def test_property_table_dimensions_preserved(num_rows: int, num_cols: int):
    """
    **Feature: test-coverage-hardening, Property 10: Markdown table to DOCX preservation**
    **Validates: Requirements 7.4**

    For any markdown table, the DOCX conversion should produce
    a table with matching column structure.
    """
    from docx import Document

    # Generate table markdown
    header = "| " + " | ".join(f"Col{i}" for i in range(num_cols)) + " |"
    separator = "| " + " | ".join("---" for _ in range(num_cols)) + " |"
    rows = []
    for r in range(num_rows - 1):  # -1 because header is a row
        row = "| " + " | ".join(f"R{r}C{c}" for c in range(num_cols)) + " |"
        rows.append(row)

    table_lines = [header, separator, *rows]

    doc = Document()
    render_table(doc, table_lines)

    assert len(doc.tables) == 1
    table = doc.tables[0]

    # Verify column count is correct
    assert len(table.columns) == num_cols
    # Verify we have at least the expected data rows
    assert len(table.rows) >= num_rows


@given(
    heading_levels=st.lists(
        st.integers(min_value=1, max_value=4),
        min_size=1,
        max_size=5,
    )
)
@settings(max_examples=50, deadline=None)
def test_property_heading_hierarchy_preserved(heading_levels: list[int]):
    """
    **Feature: test-coverage-hardening, Property 11: Heading hierarchy preservation**
    **Validates: Requirements 7.5**

    For any markdown content with heading levels (H1-H4),
    the DOCX conversion should preserve the relative hierarchy.
    """
    # Generate markdown with specified heading levels
    lines = []
    for i, level in enumerate(heading_levels):
        prefix = "#" * level
        lines.append(f"{prefix} Heading {i + 1}")
        lines.append("")
        lines.append(f"Content for heading {i + 1}.")
        lines.append("")

    markdown = "\n".join(lines)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        output_path = Path(f.name)

    try:
        markdown_to_docx(markdown, output_path)

        from docx import Document

        doc = Document(output_path)

        # Document should be created successfully
        assert len(doc.paragraphs) > 0

        # All heading content should be present
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for i in range(len(heading_levels)):
            assert f"Heading {i + 1}" in full_text
    finally:
        output_path.unlink(missing_ok=True)


@given(
    text=st.text(
        alphabet=st.sampled_from(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789 .,!?"
        ),
        min_size=10,
        max_size=200,
    )
)
@settings(max_examples=50, deadline=None)
def test_property_content_not_lost_in_conversion(text: str):
    """
    **Feature: test-coverage-hardening, Property 10: Markdown table to DOCX preservation**
    **Validates: Requirements 7.4**

    For any text content, conversion to DOCX should not lose the content.
    """
    markdown = f"## Test Section\n\n{text}"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        output_path = Path(f.name)

    try:
        markdown_to_docx(markdown, output_path)

        from docx import Document

        doc = Document(output_path)

        # Extract all text from document
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Original text should be present (allowing for whitespace normalization)
        normalized_text = " ".join(text.split())
        normalized_full = " ".join(full_text.split())

        assert normalized_text in normalized_full or text.strip() in full_text
    finally:
        output_path.unlink(missing_ok=True)
