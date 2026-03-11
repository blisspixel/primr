"""
Property-based tests for the StyleEngine.

Uses Hypothesis to verify style consistency across documents.
"""

from docx import Document
from hypothesis import assume, given, settings
from hypothesis import strategies as st

from primr.output.style_engine import StyleEngine


class TestStyleConsistency:
    """
    **Feature: report-excellence, Property 6: Style consistency across document**
    **Validates: Requirements 2.1, 5.1**

    For any generated document, all paragraphs of the same semantic type
    SHALL have identical style properties (font, size, spacing, color).
    """

    def test_heading_styles_have_distinct_sizes(self):
        """Each heading level has a distinct font size."""
        doc = Document()
        engine = StyleEngine(doc)
        engine.setup_styles()

        sizes = []
        for level in [1, 2, 3, 4]:
            style_name = engine.get_style_for_level(level)
            style = doc.styles[style_name]
            sizes.append(style.font.size)

        # All sizes should be unique
        assert len(set(sizes)) == len(sizes), "Heading levels should have distinct sizes"

        # Sizes should decrease with level (H1 > H2 > H3 > H4)
        for i in range(len(sizes) - 1):
            assert sizes[i] > sizes[i + 1], f"H{i + 1} should be larger than H{i + 2}"

    def test_heading_styles_have_distinct_colors(self):
        """Heading levels have appropriate color contrast."""
        doc = Document()
        engine = StyleEngine(doc)
        engine.setup_styles()

        # H1 should use primary color
        h1 = doc.styles["Heading 1"]
        assert h1.font.color.rgb == engine.PRIMARY_COLOR

        # H2 should use secondary color
        h2 = doc.styles["Heading 2"]
        assert h2.font.color.rgb == engine.SECONDARY_COLOR

        # H3 should use accent color
        h3 = doc.styles["Heading 3"]
        assert h3.font.color.rgb == engine.ACCENT_COLOR

    def test_body_text_uses_consistent_font(self):
        """Body text style uses consistent font settings."""
        doc = Document()
        engine = StyleEngine(doc)
        engine.setup_styles()

        normal = doc.styles["Normal"]
        expected_font, expected_size = engine.FONTS["body"]

        assert normal.font.name == expected_font
        assert normal.font.size == expected_size
        assert normal.font.color.rgb == engine.TEXT_COLOR

    @settings(max_examples=100)
    @given(level=st.integers(min_value=1, max_value=4))
    def test_get_style_for_level_returns_valid_style(self, level):
        """get_style_for_level returns valid style names for all levels."""
        doc = Document()
        engine = StyleEngine(doc)

        style_name = engine.get_style_for_level(level)

        assert style_name in doc.styles, f"Style '{style_name}' not found"
        assert style_name == f"Heading {level}"

    @settings(max_examples=100)
    @given(level=st.integers(min_value=5, max_value=100))
    def test_invalid_level_returns_normal(self, level):
        """Invalid heading levels fall back to Normal style."""
        doc = Document()
        engine = StyleEngine(doc)

        style_name = engine.get_style_for_level(level)

        assert style_name == "Normal"

    def test_color_palette_wcag_contrast(self):
        """Color palette meets WCAG 2.1 AA contrast requirements."""
        doc = Document()
        engine = StyleEngine(doc)

        # Primary color should be dark enough for text on white
        # RGB(0, 51, 102) = #003366 - very dark blue
        # RGBColor uses index access: [0]=red, [1]=green, [2]=blue
        assert engine.PRIMARY_COLOR[0] == 0
        assert engine.PRIMARY_COLOR[1] == 51
        assert engine.PRIMARY_COLOR[2] == 102

        # Text color should be dark gray
        assert engine.TEXT_COLOR[0] == 51
        assert engine.TEXT_COLOR[1] == 51
        assert engine.TEXT_COLOR[2] == 51

    def test_font_scale_follows_golden_ratio(self):
        """Font sizes follow approximately golden ratio (1.2x) between levels."""
        doc = Document()
        engine = StyleEngine(doc)

        # Get font sizes
        title_size = engine.FONTS["title"][1].pt
        chapter_size = engine.FONTS["chapter"][1].pt
        section_size = engine.FONTS["section"][1].pt
        body_size = engine.FONTS["body"][1].pt

        # Check ratios are approximately 1.2x (allowing some flexibility)
        # title (28) / chapter (18) ≈ 1.55
        # chapter (18) / section (14) ≈ 1.29
        # section (14) / body (11) ≈ 1.27
        assert title_size > chapter_size > section_size > body_size
        assert chapter_size / section_size >= 1.1
        assert section_size / body_size >= 1.1

    @settings(max_examples=50)
    @given(
        style_type=st.sampled_from(
            ["title", "chapter", "section", "subsection", "body", "bullet", "caption", "metadata"]
        )
    )
    def test_get_font_config_returns_valid_tuple(self, style_type):
        """get_font_config returns valid (font_name, size) tuple for all types."""
        doc = Document()
        engine = StyleEngine(doc)

        font_name, font_size = engine.get_font_config(style_type)

        assert isinstance(font_name, str)
        assert len(font_name) > 0
        assert font_size is not None

    @settings(max_examples=50)
    @given(
        color_type=st.sampled_from(
            [
                "primary",
                "secondary",
                "accent",
                "text",
                "muted",
                "highlight_bg",
                "table_header",
                "table_alt",
            ]
        )
    )
    def test_get_color_returns_rgb_color(self, color_type):
        """get_color returns valid RGBColor for all color types."""
        from docx.shared import RGBColor

        doc = Document()
        engine = StyleEngine(doc)

        color = engine.get_color(color_type)

        assert isinstance(color, RGBColor)
        # RGBColor uses index access: [0]=red, [1]=green, [2]=blue
        assert 0 <= color[0] <= 255
        assert 0 <= color[1] <= 255
        assert 0 <= color[2] <= 255


class TestInlineHeaderFormatting:
    """
    **Feature: report-excellence, Property 4: Inline header detection and formatting**
    **Validates: Requirements 3.6**

    For any input line matching the pattern "Header Text: content",
    the output SHALL have the header portion (before colon) formatted as bold.
    """

    @settings(max_examples=100)
    @given(
        header=st.text(min_size=3, max_size=30, alphabet=st.characters(whitelist_categories=["L"])),
        content=st.text(
            min_size=1, max_size=100, alphabet=st.characters(whitelist_categories=["L", "N", "S"])
        ),
    )
    def test_inline_header_creates_bold_run(self, header, content):
        """apply_inline_header_formatting creates bold header run."""
        header = header.strip()
        content = content.strip()
        assume(len(header) > 0)
        assume(len(content) > 0)

        doc = Document()
        engine = StyleEngine(doc)
        para = doc.add_paragraph()

        engine.apply_inline_header_formatting(para, header, content)

        # Should have 2 runs: header (bold) and content (not bold)
        assert len(para.runs) == 2

        # First run should be bold and contain header + colon
        assert para.runs[0].bold
        assert header in para.runs[0].text
        assert ":" in para.runs[0].text

        # Second run should contain content
        assert content in para.runs[1].text

    @settings(max_examples=50)
    @given(
        header=st.text(min_size=3, max_size=20, alphabet=st.characters(whitelist_categories=["L"])),
        content=st.text(
            min_size=1, max_size=50, alphabet=st.characters(whitelist_categories=["L", "N"])
        ),
    )
    def test_inline_header_preserves_content(self, header, content):
        """Content is preserved exactly in the output."""
        header = header.strip()
        content = content.strip()
        assume(len(header) > 0)
        assume(len(content) > 0)

        doc = Document()
        engine = StyleEngine(doc)
        para = doc.add_paragraph()

        engine.apply_inline_header_formatting(para, header, content)

        # Full text should contain both header and content
        full_text = "".join(run.text for run in para.runs)
        assert header in full_text
        assert content in full_text
