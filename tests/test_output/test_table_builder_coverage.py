"""Coverage tests for primr.output.table_builder.TableBuilder."""

from __future__ import annotations

from docx import Document

from primr.output.models import CompanySnapshot
from primr.output.table_builder import TableBuilder


def _builder():
    return TableBuilder(Document())


# --------------------------------------------------------------------------- #
# create_company_snapshot
# --------------------------------------------------------------------------- #
def test_create_company_snapshot_full():
    tb = _builder()
    snap = CompanySnapshot(
        company_name="Acme",
        website="https://acme.example",
        industry="Tech",
        founded="2001",
        headquarters="Anytown",
        revenue="$10M",
        employees="500",
        ticker="ACME",
    )
    table = tb.create_company_snapshot(snap)
    assert len(table.rows) == 7  # all optional fields present


def test_create_company_snapshot_minimal_falls_back_to_company_row():
    tb = _builder()
    snap = CompanySnapshot(company_name="Acme")
    table = tb.create_company_snapshot(snap)
    # Only the company name fallback row.
    assert len(table.rows) == 1
    assert "Company" in table.rows[0].cells[0].text


# --------------------------------------------------------------------------- #
# create_executive_highlights (callout style)
# --------------------------------------------------------------------------- #
def test_create_executive_highlights():
    tb = _builder()
    table = tb.create_executive_highlights(["one", "two"], title="TAKEAWAYS")
    cell = table.rows[0].cells[0]
    assert "TAKEAWAYS" in cell.text
    assert "one" in cell.text
    assert "two" in cell.text


# --------------------------------------------------------------------------- #
# create_key_metrics_table (professional style)
# --------------------------------------------------------------------------- #
def test_create_key_metrics_table_none_when_empty():
    tb = _builder()
    assert tb.create_key_metrics_table({}) is None


def test_create_key_metrics_table_populated():
    tb = _builder()
    table = tb.create_key_metrics_table({"annual_revenue": "$5M", "growth_rate": "10%"})
    assert len(table.rows) == 3  # header + 2 metrics
    # Header row text
    assert "Metric" in table.rows[0].cells[0].text
    # Underscore replaced with space and title-cased
    assert "Annual Revenue" in table.rows[1].cells[0].text


# --------------------------------------------------------------------------- #
# apply_table_style dispatch / unknown style
# --------------------------------------------------------------------------- #
def test_apply_table_style_unknown_is_noop():
    tb = _builder()
    table = tb.document.add_table(rows=1, cols=1)
    # Should not raise even when an unsupported style is requested.
    tb.apply_table_style(table, "does-not-exist")


def test_apply_professional_style_alternating_rows():
    tb = _builder()
    table = tb.document.add_table(rows=4, cols=2)
    tb.apply_table_style(table, "professional")
    # Even-indexed data rows get shading; ensure no exception and shading XML present.
    assert "shd" in table.rows[2].cells[0]._tc.get_or_add_tcPr().xml


def test_apply_snapshot_style_alternating_rows():
    tb = _builder()
    table = tb.document.add_table(rows=3, cols=2)
    tb.apply_table_style(table, "snapshot")
    assert "shd" in table.rows[1].cells[0]._tc.get_or_add_tcPr().xml


def test_apply_callout_style():
    tb = _builder()
    table = tb.document.add_table(rows=1, cols=1)
    tb.apply_table_style(table, "callout")
    assert "shd" in table.rows[0].cells[0]._tc.get_or_add_tcPr().xml


# --------------------------------------------------------------------------- #
# create_from_markdown
# --------------------------------------------------------------------------- #
def test_create_from_markdown_empty_returns_none():
    tb = _builder()
    assert tb.create_from_markdown([], []) is None


def test_create_from_markdown_zero_columns_returns_none():
    tb = _builder()
    # No headers and a row that yields zero columns.
    assert tb.create_from_markdown([], [[]]) is None


def test_create_from_markdown_with_headers_and_rows():
    tb = _builder()
    table = tb.create_from_markdown(["A", "B"], [["1", "2"], ["3", "4"]])
    assert len(table.rows) == 3  # header + 2
    assert "A" in table.rows[0].cells[0].text
    assert "1" in table.rows[1].cells[0].text


def test_create_from_markdown_rows_only_no_headers():
    tb = _builder()
    table = tb.create_from_markdown([], [["x", "y"], ["z", "w"]])
    assert len(table.rows) == 2


def test_create_from_markdown_ragged_rows_truncated():
    tb = _builder()
    # Extra cells beyond num_cols are ignored.
    table = tb.create_from_markdown(["A", "B"], [["1", "2", "3"]])
    assert len(table.columns) == 2
