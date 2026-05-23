"""Coverage tests for primr.output.markdown_parser (MarkdownParser + ArtifactDetector)."""

from __future__ import annotations

from docx import Document

from primr.output.markdown_parser import ArtifactDetector, MarkdownParser
from primr.output.models import ContentBlock, ParsedLine


# --------------------------------------------------------------------------- #
# MarkdownParser.parse_line
# --------------------------------------------------------------------------- #
def test_parse_line_empty():
    p = MarkdownParser()
    pl = p.parse_line("   ")
    assert pl.type == "empty"


def test_parse_line_heading_levels():
    p = MarkdownParser()
    assert p.parse_line("# H1").level == 1
    assert p.parse_line("#### H4").type == "heading"


def test_parse_line_bullet_with_indent():
    p = MarkdownParser()
    pl = p.parse_line("    - item")
    assert pl.type == "bullet"
    assert pl.level == 1
    assert pl.metadata["bullet_char"] == "-"


def test_parse_line_numbered():
    p = MarkdownParser()
    pl = p.parse_line("3. third")
    assert pl.type == "numbered"
    assert pl.metadata["number"] == "3"


def test_parse_line_inline_header():
    p = MarkdownParser()
    pl = p.parse_line("Revenue: $5M")
    assert pl.type == "inline_header"
    assert pl.metadata["header_text"] == "Revenue"


def test_parse_line_inline_header_excludes_url():
    p = MarkdownParser()
    # "Https" matches header regex but is excluded by URL_PREFIXES.
    pl = p.parse_line("Https: something")
    assert pl.type != "inline_header"


def test_parse_line_table_row_and_separator():
    p = MarkdownParser()
    sep = p.parse_line("|---|---|")
    assert sep.type == "table_separator"
    row = p.parse_line("| a | b |")
    assert row.type == "table_row"
    assert row.metadata["cells"] == ["a", "b"]


def test_parse_line_plain_text():
    p = MarkdownParser()
    assert p.parse_line("just words").type == "text"


# --------------------------------------------------------------------------- #
# parse_content block grouping
# --------------------------------------------------------------------------- #
def test_parse_content_groups_blocks():
    p = MarkdownParser()
    content = (
        "# Heading\n"
        "para line one\n"
        "para line two\n"
        "\n"
        "- bullet a\n"
        "- bullet b\n"
        "\n"
        "1. num a\n"
        "2. num b\n"
    )
    blocks = p.parse_content(content)
    types = [b.type for b in blocks]
    assert "heading" in types
    assert "bullet_list" in types
    assert "numbered_list" in types


def test_parse_content_detects_subheading_before_bullet():
    p = MarkdownParser()
    content = "Capabilities\n- item one\n- item two\n"
    blocks = p.parse_content(content)
    # The plain text "Capabilities" followed by bullets becomes a heading block.
    assert blocks[0].type == "heading"
    assert blocks[0].lines[0].type == "subheading"


def test_parse_content_block_type_switch():
    p = MarkdownParser()
    content = "text para\n- bullet\n"
    blocks = p.parse_content(content)
    # text followed directly by bullet -> subheading then bullet list
    assert any(b.type == "bullet_list" for b in blocks)


def test_parse_content_table_block():
    p = MarkdownParser()
    content = "| H1 | H2 |\n|----|----|\n| 1 | 2 |\n"
    blocks = p.parse_content(content)
    assert any(b.type == "table" for b in blocks)


# --------------------------------------------------------------------------- #
# helper predicate methods
# --------------------------------------------------------------------------- #
def test_same_block_type_matrix():
    p = MarkdownParser()
    pl_text = ParsedLine("text", "x", 0, "x")
    pl_bullet = ParsedLine("bullet", "x", 0, "x")
    pl_num = ParsedLine("numbered", "x", 0, "x")
    pl_tbl = ParsedLine("table_row", "x", 0, "x")
    assert p._same_block_type("paragraph", pl_text)
    assert p._same_block_type("bullet_list", pl_bullet)
    assert p._same_block_type("numbered_list", pl_num)
    assert p._same_block_type("table", pl_tbl)
    assert not p._same_block_type("paragraph", pl_bullet)


def test_get_block_type_mapping():
    p = MarkdownParser()
    assert p._get_block_type(ParsedLine("text", "", 0, "")) == "paragraph"
    assert p._get_block_type(ParsedLine("bullet", "", 0, "")) == "bullet_list"
    assert p._get_block_type(ParsedLine("numbered", "", 0, "")) == "numbered_list"
    assert p._get_block_type(ParsedLine("heading", "", 0, "")) == "heading"
    assert p._get_block_type(ParsedLine("inline_header", "", 0, "")) == "paragraph"
    assert p._get_block_type(ParsedLine("table_row", "", 0, "")) == "table"
    assert p._get_block_type(ParsedLine("unknown", "", 0, "")) == "paragraph"


# --------------------------------------------------------------------------- #
# parse_table_block
# --------------------------------------------------------------------------- #
def test_parse_table_block():
    p = MarkdownParser()
    block = ContentBlock(
        "table",
        [
            ParsedLine("table_row", "", 0, "", {"cells": ["H1", "H2"]}),
            ParsedLine("table_separator", "", 0, ""),
            ParsedLine("table_row", "", 0, "", {"cells": ["a", "b"]}),
        ],
    )
    data = p.parse_table_block(block)
    assert data["headers"] == ["H1", "H2"]
    assert data["rows"] == [["a", "b"]]


# --------------------------------------------------------------------------- #
# apply_inline_formatting
# --------------------------------------------------------------------------- #
def test_apply_inline_formatting_no_bold():
    p = MarkdownParser()
    doc = Document()
    para = doc.add_paragraph()
    p.apply_inline_formatting(para, "plain text")
    assert para.runs[0].text == "plain text"


def test_apply_inline_formatting_with_bold():
    p = MarkdownParser()
    doc = Document()
    para = doc.add_paragraph()
    p.apply_inline_formatting(para, "a **b** c")
    assert any(r.bold for r in para.runs)


# --------------------------------------------------------------------------- #
# strip_markdown_formatting / extract_bold_segments
# --------------------------------------------------------------------------- #
def test_strip_markdown_formatting():
    p = MarkdownParser()
    out = p.strip_markdown_formatting("**bold** and *ital*")
    assert "**" not in out
    assert "bold" in out
    assert "ital" in out


def test_extract_bold_segments_no_bold():
    p = MarkdownParser()
    assert p.extract_bold_segments("plain") == [("plain", False)]


def test_extract_bold_segments_with_bold():
    p = MarkdownParser()
    segs = p.extract_bold_segments("a **b** c")
    assert ("b", True) in segs
    assert ("a ", False) in segs
    assert (" c", False) in segs


# --------------------------------------------------------------------------- #
# ArtifactDetector
# --------------------------------------------------------------------------- #
def test_scan_text_finds_artifacts():
    det = ArtifactDetector()
    arts = det.scan_text("## Heading\n**bold**\n- bullet text", context="ctx")
    types = {a["type"] for a in arts}
    assert "heading" in types
    assert "bold" in types
    assert "bullet" in types


def test_scan_text_clean_returns_empty():
    det = ArtifactDetector()
    assert det.scan_text("plain clean text") == []


def test_scan_document_paragraphs_and_tables():
    det = ArtifactDetector()
    doc = Document()
    doc.add_paragraph("**bold artifact**")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "## heading artifact"
    arts = det.scan_document(doc)
    assert len(arts) >= 2
    assert det.artifacts_found == arts


def test_has_artifacts_true_and_false():
    det = ArtifactDetector()
    doc_clean = Document()
    doc_clean.add_paragraph("clean text")
    assert det.has_artifacts(doc_clean) is False

    doc_dirty = Document()
    doc_dirty.add_paragraph("**bold**")
    assert det.has_artifacts(doc_dirty) is True


def test_get_artifact_summary_empty():
    det = ArtifactDetector()
    assert det.get_artifact_summary() == "No markdown artifacts found."


def test_get_artifact_summary_with_many():
    det = ArtifactDetector()
    det.artifacts_found = [
        {"type": "bold", "match": f"**x{i}**", "context": "c"} for i in range(15)
    ]
    summary = det.get_artifact_summary()
    assert "15 markdown artifact" in summary
    assert "and 5 more" in summary
