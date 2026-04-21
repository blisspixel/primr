from pathlib import Path

from docx import Document

import primr.output.output_utils as output_utils
from primr.core.research_agent import (
    _clean_strategy_output,
    _compute_strategy_qa_metrics,
    _convert_deep_research_to_docx,
    _ensure_strategy_source_inventory,
    _prepare_strategy_for_output,
    _prepare_strategy_markdown_for_shipping,
    _save_strategy_output,
    _validate_output_docx,
    _validate_output_markdown,
    improve_output_file,
)
from primr.output.final_artifact import (
    GeneratedSection,
    canonicalize_final_markdown,
    parse_final_markdown,
)


def test_compute_strategy_qa_metrics_flags_budget_inconsistency():
    content = (
        "# AI Strategy: Demo\n\n"
        "We recommend $5-8M Year 1 investment.\n\n"
        "## Investment Framework\n\n"
        "**Year 1**: Tech $1.5M, People $2M, Impl $1.5M. **Total $5M**.\n\n"
        "## BOARD SUMMARY\n\n"
        "Quick Wins: $1M\n"
        "Bigger Bets: $4M\n"
        "Org/People: $2.5M\n"
        "**Total: $7.5M**\n"
        "[Source: https://example.com/a]\n"
        "[Source: https://example.com/b]\n"
    )
    metrics = _compute_strategy_qa_metrics(content)
    assert metrics["budget_inconsistent"] is True
    assert metrics["source_urls"] >= 2


def test_compute_strategy_qa_metrics_allows_reasonable_budget_variance():
    content = (
        "# AI Strategy: Demo\n\n"
        "Year 1 investment: $1.2-1.8M.\n\n"
        "## BOARD SUMMARY\n\n"
        "**Total: $1.7M**\n"
        "[Source: https://example.com/a]\n"
        "[Source: https://example.com/b]\n"
    )
    metrics = _compute_strategy_qa_metrics(content)
    assert metrics["budget_inconsistent"] is False
    assert metrics["qa_gate_passed"] is True


def test_clean_strategy_output_strips_internal_placeholders():
    content = "Claim [Reported: Analysis Context].\n\n[citation inventory: 1=example.com]\n"
    cleaned = _clean_strategy_output(content)
    assert "Analysis Context" not in cleaned
    assert "citation inventory" not in cleaned.lower()


def test_clean_strategy_output_normalizes_source_tags():
    content = "## Strategy\n\nClaim [Source: https://example.com/a].\n"
    cleaned = _clean_strategy_output(content)
    assert "[Source:" not in cleaned
    assert "[cite: 1]" in cleaned
    assert "## Sources" in cleaned


def test_improve_output_file_writes_improved_report(tmp_path: Path):
    p = tmp_path / "demo.md"
    p.write_text(
        "## Executive Summary\n\n"
        "Claim [Source: https://example.com/a].\n\n"
        "[citation inventory: 1=example.com]\n",
        encoding="utf-8",
    )

    out = improve_output_file(str(p), in_place=False, use_agentic=False)
    assert out is not None
    out_path = Path(out)
    assert out_path.exists()
    text = out_path.read_text(encoding="utf-8")
    assert "## Sources" in text
    assert "citation inventory" not in text.lower()


def test_validate_output_markdown_flags_internal_artifacts():
    result = _validate_output_markdown(
        "## Report\n\nClaim [Workbook: Financial Profile] [Source: https://example.com/a].\n"
    )
    assert result["passed"] is False
    assert any("workbook_ref" in issue for issue in result["issues"])
    assert any("raw_source_tag" in issue for issue in result["issues"])


def test_validate_output_markdown_flags_internal_analysis_terms():
    result = _validate_output_markdown(
        "## Strategy\n\nBuilt from vendor-research-aws-2026-03.txt with Internal ROI Model and [Analysis: 4].\n"
    )
    assert result["passed"] is False
    assert any("vendor_research_file" in issue for issue in result["issues"])
    assert any("internal_roi_model" in issue for issue in result["issues"])


def test_validate_output_docx_flags_internal_artifacts(tmp_path: Path):
    doc_path = tmp_path / "bad.docx"
    doc = Document()
    doc.add_paragraph("Claim [Workbook: Financial Profile]")
    doc.save(doc_path)

    result = _validate_output_docx(doc_path)
    assert result["passed"] is False
    assert any("workbook_ref" in issue for issue in result["issues"])


def test_validate_output_docx_accepts_clean_document(tmp_path: Path):
    doc_path = tmp_path / "clean.docx"
    doc = Document()
    doc.add_paragraph("Strategic Thesis")
    doc.add_paragraph("Claim [cite: 1]")
    doc.save(doc_path)

    result = _validate_output_docx(doc_path)
    assert result["passed"] is True
    assert result["issues"] == []


def test_convert_deep_research_to_docx_salvages_recoverable_markdown(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(output_utils, "OUTPUT_DIR", str(tmp_path))
    result = _convert_deep_research_to_docx(
        "## Report\n\nClaim [Workbook: Financial Profile].\n",
        "DemoCo",
        "https://example.com",
    )
    assert result is not None
    assert result.endswith(".docx")
    assert not list(tmp_path.glob("*markdown_validation.txt"))
    md_files = list(tmp_path.glob("DemoCo_Strategic_Overview_*.md"))
    assert md_files
    md_text = md_files[0].read_text(encoding="utf-8")
    assert "[Workbook:" not in md_text


def test_convert_deep_research_to_docx_allows_clean_markdown(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(output_utils, "OUTPUT_DIR", str(tmp_path))
    markdown = (
        "## Executive Summary\n\n"
        "Claim (Reported) [cite: 1].\n\n"
        "What to validate: Confirm the primary claim.\n\n"
        "## Sources\n\n"
        "[cite: 1] https://example.com/source\n"
    )
    result = _convert_deep_research_to_docx(markdown, "DemoCo", "https://example.com")
    assert result is not None
    assert result.endswith(".docx")
    assert Path(result).exists()


def test_save_strategy_output_ships_docx_with_budget_warning(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(output_utils, "OUTPUT_DIR", str(tmp_path))
    content = (
        "## AI Strategy\n\n"
        "Recommended Year 1 investment: $1.2-1.8M.\n\n"
        "## BOARD SUMMARY\n\n"
        "**Total: $2.5M**\n"
    )
    result = _save_strategy_output(content, "DemoCo", "azure", strategy_label="AI_Strategy")
    assert result is not None
    assert result.endswith(".docx")
    assert list(tmp_path.glob("*markdown_validation.txt"))
    assert list(tmp_path.glob("*.docx"))


def test_compute_strategy_qa_metrics_counts_numeric_citations_as_sources():
    content = (
        "## Strategy\n\n"
        "Claim [cite: 1].\n\n"
        "Another claim [cite: 2].\n\n"
        "## Sources\n\n"
        "[cite: 1] https://example.com/a\n"
        "[cite: 2] https://example.com/b\n"
    )
    metrics = _compute_strategy_qa_metrics(content)
    assert metrics["source_urls"] >= 2


def test_convert_deep_research_to_docx_salvages_raw_source_tags(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(output_utils, "OUTPUT_DIR", str(tmp_path))
    markdown = (
        "## Executive Summary\n\n"
        "Claim [Source: https://example.com/source].\n\n"
        "What to validate: Confirm the primary claim.\n"
    )
    result = _convert_deep_research_to_docx(markdown, "DemoCoSalvage", "https://example.com")
    assert result is not None
    assert result.endswith(".docx")
    md_files = list(tmp_path.glob("DemoCoSalvage_Strategic_Overview_*.md"))
    assert md_files
    md_text = md_files[0].read_text(encoding="utf-8")
    assert "[Source:" not in md_text
    assert "## Sources" in md_text


def test_save_strategy_output_salvages_internal_vendor_reference(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(output_utils, "OUTPUT_DIR", str(tmp_path))
    content = (
        "## AI Strategy\n\n"
        "Use vendor-research-aws-2026-03.txt with Claim [cite: 1].\n\n"
        "## Sources\n\n"
        "[cite: 1] https://example.com/a\n"
        "[cite: 2] https://example.com/b\n"
    )
    result = _save_strategy_output(content, "DemoCoAdaptive", "aws", strategy_label="AI_Strategy")
    assert result is not None
    assert result.endswith(".docx")
    md_files = list(tmp_path.glob("DemoCoAdaptive_AI_Strategy_AWS_*.md"))
    assert md_files
    md_text = md_files[0].read_text(encoding="utf-8")
    assert "vendor-research-aws-2026-03.txt" not in md_text


def test_ensure_strategy_source_inventory_appends_sources_when_missing():
    content = "## AI Strategy\n\nRecommendation without citations.\n"
    improved = _ensure_strategy_source_inventory(
        content,
        ["https://example.com/a", "https://example.com/b"],
    )
    assert "## Sources" in improved
    assert "[cite: 1] https://example.com/a" in improved
    assert "[cite: 2] https://example.com/b" in improved


def test_ensure_strategy_source_inventory_preserves_existing_citations():
    content = (
        "## AI Strategy\n\nClaim [cite: 1].\n\n## Sources\n\n[cite: 1] https://example.com/a\n"
    )
    improved = _ensure_strategy_source_inventory(
        content, ["https://example.com/a", "https://example.com/b"]
    )
    assert improved.count("[cite: 1] https://example.com/a") == 1
    assert "[cite: 2] https://example.com/b" in improved


def test_compute_strategy_qa_metrics_flags_invalid_source_urls():
    content = (
        "## Strategy\n\n"
        "Claim [cite: 1].\n\n"
        "## Sources\n\n"
        "[cite: 1] https://-aws-2026-03\n"
        "[cite: 2] https://example.com/b\n"
    )
    metrics = _compute_strategy_qa_metrics(content)
    assert metrics["invalid_source_urls"] == 1
    assert metrics["qa_gate_passed"] is False


def test_ensure_strategy_source_inventory_skips_invalid_urls():
    content = "## AI Strategy\n\nRecommendation without citations.\n"
    improved = _ensure_strategy_source_inventory(
        content,
        ["https://example.com/a", "https://-aws-2026-03", "not-a-url", "https://example.com/b"],
    )
    assert "https://example.com/a" in improved
    assert "https://example.com/b" in improved
    assert "https://-aws-2026-03" not in improved
    assert "not-a-url" not in improved


def test_prepare_strategy_for_output_repairs_budget_and_sources(monkeypatch):
    def fake_grok_llm(prompt, **kwargs):
        return (
            "## AI Strategy\n\n"
            "Recommended Year 1 investment: $1.2M.\n\n"
            "## BOARD SUMMARY\n\n"
            "**Total: $1.2M**\n\n"
            "## Sources\n\n"
            "[cite: 1] https://example.com/a\n"
            "[cite: 2] https://example.com/b\n"
        )

    monkeypatch.setattr("primr.ai.grok_client.grok_llm", fake_grok_llm)

    prepared, qa, rejected = _prepare_strategy_for_output(
        "## AI Strategy\n\n"
        "Recommended Year 1 investment: $1.2-1.8M.\n\n"
        "## BOARD SUMMARY\n\n"
        "**Total: $1.9M**\n\n"
        "## Sources\n\n"
        "[cite: 1] https://-aws-2026-03\n",
        "DemoCo",
        "aws",
        "AI Strategy",
        ["https://example.com/a", "https://example.com/b", "https://-aws-2026-03"],
    )

    assert rejected == ["https://-aws-2026-03"]
    assert qa["budget_inconsistent"] is False
    assert qa["invalid_source_urls"] == 0
    assert qa["missing_citations"] == 0
    assert qa["qa_gate_passed"] is True
    assert "https://-aws-2026-03" not in prepared


def test_validate_output_docx_allows_literal_hash_table_headers(tmp_path: Path):
    doc_path = tmp_path / "hash-table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "# Items"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "42"
    table.rows[1].cells[1].text = "OK"
    doc.save(doc_path)

    result = _validate_output_docx(doc_path)
    assert result["passed"] is True
    assert result["issues"] == []


def test_canonicalize_final_markdown_merges_reference_sections_at_end():
    content = (
        "# Report\n\n"
        "## Executive Summary\n\n"
        "Intro (Reported).\n\n"
        "## References\n\n"
        "[cite: 2] https://example.com/b\n\n"
        "## Products and Services\n\n"
        "Detail (Estimated).\n\n"
        "## Sources\n\n"
        "[cite: 1] https://example.com/a\n"
    )

    normalized = canonicalize_final_markdown(content)

    assert normalized.count("## Sources") == 1
    assert normalized.rstrip().endswith("[cite: 1] https://example.com/a")
    assert "## References" not in normalized
    assert normalized.index("## Products and Services") < normalized.index("## Sources")


def test_parse_final_markdown_preserves_preamble_and_content_section_order():
    content = (
        "# Strategic Company Overview: DemoCo\n\n"
        "*April 12, 2026*\n\n"
        "## Executive Summary\n\n"
        "Summary.\n\n"
        "## SWOT Analysis\n\n"
        "SWOT body.\n"
    )

    parsed = parse_final_markdown(content)

    assert parsed.preamble.startswith("# Strategic Company Overview")
    assert [section.heading for section in parsed.sections] == [
        "Executive Summary",
        "SWOT Analysis",
    ]
    assert parsed.sources_body == ""


def test_generated_section_to_markdown_formats_section():
    section = GeneratedSection(
        title="Executive Summary",
        content="Deep content.\n\nWhat to validate: Confirm the key claim.",
        words=8,
        validate_line="What to validate: Confirm the key claim.",
        citation_numbers=[1, 2],
    )

    assert (
        section.to_markdown()
        == "## Executive Summary\n\nDeep content.\n\nWhat to validate: Confirm the key claim."
    )


def test_prepare_strategy_markdown_for_shipping_merges_reference_sections():
    content = (
        "## AI Strategy\n\n"
        "Recommendation [cite: 1].\n\n"
        "## References\n\n"
        "[cite: 2] https://example.com/b\n\n"
        "## Recommended AI Architecture Posture\n\n"
        "Architecture detail [cite: 2].\n\n"
        "## Sources\n\n"
        "[cite: 1] https://example.com/a\n"
    )

    prepared = _prepare_strategy_markdown_for_shipping(content)

    assert prepared.count("## Sources") == 1
    assert "## References" not in prepared
    assert prepared.index("## Recommended AI Architecture Posture") < prepared.index("## Sources")


def test_clean_strategy_output_strips_unresolved_section_cross_refs():
    content = (
        "## Executive Summary\n\n"
        "Quick wins are prioritized [see ## Five Quick Wins and ## Five Bigger Bets].\n"
    )
    cleaned = _clean_strategy_output(content)
    assert "[see ##" not in cleaned
    assert "Quick wins are prioritized" in cleaned


def test_compute_strategy_qa_metrics_flags_unresolved_section_cross_refs():
    content = (
        "# AI Strategy: Demo\n\n"
        "Reference [see ## Five Quick Wins].\n\n"
        "## Sources\n\n"
        "[cite: 1] https://example.com/a\n"
        "[cite: 2] https://example.com/b\n"
    )
    metrics = _compute_strategy_qa_metrics(content)
    assert metrics["placeholder_refs"] >= 1
    assert metrics["qa_gate_passed"] is False
